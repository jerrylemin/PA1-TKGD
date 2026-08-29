"""Build the PA4 editable DOCX sources and selectable-text PDFs.

The DOCX and PDF renderers consume the same report-data models. The summative
model reads the canonical ``analysis-result.json`` state so the report changes
when evidence changes, without source edits or invented findings.
"""

from __future__ import annotations

import html
import json
import os
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as PdfImage,
    KeepTogether,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


PA4 = Path(__file__).resolve().parents[1]
SOURCE = PA4 / "source"
FINAL = PA4 / "final"
SCREENSHOTS = PA4 / "evidence" / "prototype-screenshots"
ANALYSIS_RESULT = PA4 / "study" / "analysis" / "analysis-result.json"

NAVY = "081D2B"
TEAL = "1D7387"
MINT = "DFF3F2"
GOLD = "E8B765"
CORAL = "F37763"
GREY = "6D7F84"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths: list[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")


def setup_docx(doc: Document, short_title: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12
    for name, size, color, before, after in [
        ("Title", 28, NAVY, 0, 8),
        ("Heading 1", 17, NAVY, 15, 7),
        ("Heading 2", 12, TEAL, 10, 4),
        ("Heading 3", 10, NAVY, 7, 3),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    header = section.header.paragraphs[0]
    header.text = f"GROUP10 / PA4     {short_title.upper()}"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.name = "Arial"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.bold = True
    header.runs[0].font.color.rgb = RGBColor.from_string(TEAL)
    footer = section.footer.paragraphs[0]
    footer.text = "CSC13112 UI/UX Design · 2026 · Evidence state must remain truthful"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.name = "Arial"
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor.from_string(GREY)


def docx_callout(doc: Document, label: str, text: str, fill: str = MINT, color: str = TEAL) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 180, 180, 180, 180)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    label_run = paragraph.add_run(f"{label}\n")
    label_run.bold = True
    label_run.font.size = Pt(8)
    label_run.font.color.rgb = RGBColor.from_string(color)
    text_run = paragraph.add_run(text)
    text_run.font.size = Pt(10)
    text_run.font.color.rgb = RGBColor.from_string(NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def docx_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.add_run(item)


def docx_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.add_run(item)


def docx_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int] | None = None, font_size: float = 8) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    widths = widths or [9360 // len(headers)] * len(headers)
    set_table_widths(table, widths)
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(255, 255, 255)
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for index, text in enumerate(row):
            cell = cells[index]
            set_cell_margins(cell)
            if row_index % 2 == 0:
                set_cell_shading(cell, "F1F6F7")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(str(text))
            run.font.name = "Arial"
            run.font.size = Pt(font_size)
            run.font.color.rgb = RGBColor.from_string(NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def docx_image(doc: Document, filename: str, width: float = 6.6) -> None:
    path = SCREENSHOTS / filename
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(width))
    caption = doc.add_paragraph(f"Figure: {filename.replace('-', ' ').replace('.png', '')}")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].italic = True
    caption.runs[0].font.size = Pt(8)
    caption.runs[0].font.color.rgb = RGBColor.from_string(GREY)


def docx_title(doc: Document, title: str, subtitle: str) -> None:
    paragraph = doc.add_paragraph(style="Title")
    paragraph.add_run(title)
    subtitle_paragraph = doc.add_paragraph(subtitle)
    subtitle_paragraph.paragraph_format.space_after = Pt(13)
    subtitle_paragraph.runs[0].font.size = Pt(11)
    subtitle_paragraph.runs[0].font.color.rgb = RGBColor.from_string(TEAL)


def pdf_styles() -> dict[str, ParagraphStyle]:
    styles = getSampleStyleSheet()
    return {
        "title": ParagraphStyle("pa4Title", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=27, leading=28, textColor=colors.HexColor(f"#{NAVY}"), spaceAfter=8),
        "subtitle": ParagraphStyle("pa4Subtitle", parent=styles["Normal"], fontName="Helvetica", fontSize=10.5, leading=15, textColor=colors.HexColor(f"#{TEAL}"), spaceAfter=13),
        "h1": ParagraphStyle("pa4H1", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=16, leading=19, textColor=colors.HexColor(f"#{NAVY}"), spaceBefore=15, spaceAfter=7, keepWithNext=True),
        "h2": ParagraphStyle("pa4H2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=11.5, leading=14, textColor=colors.HexColor(f"#{TEAL}"), spaceBefore=9, spaceAfter=4, keepWithNext=True),
        "body": ParagraphStyle("pa4Body", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.2, leading=13.2, textColor=colors.HexColor(f"#{NAVY}"), spaceAfter=5),
        "small": ParagraphStyle("pa4Small", parent=styles["BodyText"], fontName="Helvetica", fontSize=7.5, leading=10, textColor=colors.HexColor(f"#{GREY}"), spaceAfter=3),
        "caption": ParagraphStyle("pa4Caption", parent=styles["BodyText"], fontName="Helvetica-Oblique", fontSize=7.5, leading=9, alignment=TA_CENTER, textColor=colors.HexColor(f"#{GREY}"), spaceAfter=8),
        "coverKicker": ParagraphStyle("pa4Kicker", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=8, leading=10, textColor=colors.HexColor(f"#{TEAL}"), spaceAfter=8),
    }


def pdf_p(text: str, style: ParagraphStyle) -> Paragraph:
    escaped = html.escape(str(text)).replace("\n", "<br/>")
    return Paragraph(escaped, style)


def pdf_bullets(items: Iterable[str], style: ParagraphStyle) -> ListFlowable:
    return ListFlowable(
        [ListItem(pdf_p(item, style), leftIndent=8) for item in items],
        bulletType="bullet",
        start="circle",
        leftIndent=15,
        bulletFontName="Helvetica",
        bulletFontSize=6,
        bulletOffsetY=2,
    )


def pdf_table(headers: list[str], rows: list[list[str]], widths: list[float], styles: dict[str, ParagraphStyle]) -> Table:
    data = [[Paragraph(f"<b>{html.escape(str(item))}</b>", styles["small"]) for item in headers]]
    data.extend([[pdf_p(str(item), styles["small"]) for item in row] for row in rows])
    table = Table(data, colWidths=[width * inch for width in widths], repeatRows=1, hAlign="LEFT")
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{NAVY}")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#D5E1E3")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        *[("BACKGROUND", (0, index), (-1, index), colors.HexColor("#F1F6F7")) for index in range(2, len(data), 2)],
    ]))
    return table


def pdf_callout(label: str, text: str, styles: dict[str, ParagraphStyle], fill: str = MINT) -> Table:
    content = Paragraph(f"<b>{html.escape(label)}</b><br/>{html.escape(text)}", styles["body"])
    table = Table([[content]], colWidths=[7.1 * inch])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(f"#{fill}")),
        ("BOX", (0, 0), (-1, -1), 0.55, colors.HexColor(f"#{TEAL}")),
        ("LEFTPADDING", (0, 0), (-1, -1), 12),
        ("RIGHTPADDING", (0, 0), (-1, -1), 12),
        ("TOPPADDING", (0, 0), (-1, -1), 10),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
    ]))
    return table


def pdf_image(filename: str, width: float = 6.5):
    path = SCREENSHOTS / filename
    if not path.exists():
        return Spacer(1, 1)
    from PIL import Image as PilImage

    with PilImage.open(path) as image:
        ratio = image.height / image.width
    styles = pdf_styles()
    return KeepTogether([
        PdfImage(str(path), width=width * inch, height=width * ratio * inch),
        Paragraph(f"Figure: {filename.replace('-', ' ').replace('.png', '')}", styles["caption"]),
    ])


def page_decor(canvas, doc, label: str, cover_color: str = "F4F7F8") -> None:
    canvas.saveState()
    width, height = letter
    if doc.page == 1:
        canvas.setFillColor(colors.HexColor(f"#{cover_color}"))
        canvas.rect(0, 0, width, height, stroke=0, fill=1)
        canvas.setFillColor(colors.HexColor(f"#{NAVY}"))
        canvas.rect(0, height - 0.16 * inch, width, 0.16 * inch, stroke=0, fill=1)
        canvas.setFillColor(colors.HexColor(f"#{GOLD}"))
        canvas.circle(width - 0.7 * inch, height - 0.85 * inch, 0.28 * inch, stroke=0, fill=1)
    else:
        canvas.setStrokeColor(colors.HexColor("#D5E1E3"))
        canvas.line(0.7 * inch, height - 0.48 * inch, width - 0.7 * inch, height - 0.48 * inch)
        canvas.setFont("Helvetica-Bold", 7.5)
        canvas.setFillColor(colors.HexColor(f"#{TEAL}"))
        canvas.drawString(0.7 * inch, height - 0.36 * inch, f"GROUP10 / PA4 · {label.upper()}")
    canvas.setStrokeColor(colors.HexColor("#D5E1E3"))
    canvas.line(0.7 * inch, 0.48 * inch, width - 0.7 * inch, 0.48 * inch)
    canvas.setFont("Helvetica", 7.5)
    canvas.setFillColor(colors.HexColor(f"#{GREY}"))
    canvas.drawCentredString(width / 2, 0.29 * inch, f"CSC13112 UI/UX Design · {doc.page}")
    canvas.restoreState()


def _read_json(path: Path) -> dict[str, Any]:
    try:
        with path.open("r", encoding="utf-8") as handle:
            value = json.load(handle)
    except (OSError, json.JSONDecodeError):
        return {}
    return value if isinstance(value, dict) else {}


def _int_value(value: Any) -> int:
    try:
        return int(value)
    except (TypeError, ValueError):
        return 0


def build_hifi_report_data() -> dict[str, Any]:
    return {
        "urls": {
            "fifa": "?mode=study&product=fifa#fifa",
            "chess": "?mode=study&product=chess#chess",
        },
        "chess": {
            "mistake": "Qh5",
            "consequence": "Nxh5 captures the queen",
            "better": "Qe2",
            "practice": "Qd3",
            "concept": "Check an opponent's attack before moving a valuable piece.",
        },
    }


def build_summative_report_data(analysis_path: Path = ANALYSIS_RESULT) -> dict[str, Any]:
    """Return the single substantive data model consumed by DOCX and PDF."""

    result = _read_json(analysis_path)
    participant_gate = result.get("gate", {}) if isinstance(result.get("gate"), dict) else {}
    gates = result.get("gates", {}) if isinstance(result.get("gates"), dict) else {}
    gate_status = lambda key, fallback: str((gates.get(key) or {}).get("status", fallback))
    verified_count = _int_value(participant_gate.get("verified_participant_count", result.get("verified_participant_count")))
    participant_rows = _int_value(participant_gate.get("participant_rows", result.get("participant_rows")))
    task_rows = _int_value(result.get("task_row_count", result.get("task_rows")))
    questionnaire_rows = _int_value(result.get("questionnaire_row_count", result.get("questionnaire_rows")))
    interview_rows = _int_value(result.get("interview_row_count", result.get("interview_rows")))
    analysis_status = str(result.get("status") or "BLOCKED_EXTERNAL_EVIDENCE")
    participant_state = f"{verified_count} verified participant sessions" if verified_count else "No verified participant sessions"
    recording_state = f"{verified_count} verified video recordings" if verified_count else "No verified video recordings"
    task_state = f"{task_rows} task rows" if task_rows else "No task rows (header-only template)"
    questionnaire_state = f"{questionnaire_rows} questionnaire rows" if questionnaire_rows else "No questionnaire rows (header-only template)"
    interview_state = f"{interview_rows} interview/feedback rows" if interview_rows else "No interview/feedback rows (header-only template)"
    current_states = [
        ["Verified participants", participant_state, "At least five verified, anonymized participant sessions"],
        ["Recordings", recording_state, "One consented MP4 with positive duration and a video stream per session"],
        ["Task results", task_state, "All assigned task rows with timestamps, 2/1/0 score, recovery, and hesitation fields"],
        ["Questionnaire", questionnaire_state, "Five raw items per product and verified participant"],
        ["Interview", interview_state, "At least one timestamped non-empty feedback row per verified participant"],
    ]
    gate_rows = [[key, str((gates.get(key) or {}).get("name", "")), gate_status(key, "BLOCKED_EXTERNAL_EVIDENCE")] for key in ("G01", "G02", "G03", "G04", "G05", "G06", "G07", "G08")]
    tasks = [
        ["FIFA-T1", "Status, owner, timing, and next action explained", "2/1/0; time; errors; wrong paths; assistance; hesitation"],
        ["FIFA-T2", "Order or ticket detail opened and interpreted", "Same task measures"],
        ["FIFA-T3", "Partner destination recognized before leaving", "Same measures; boundary recognition"],
        ["FIFA-T4", "Return restores status context", "Same measures; orientation"],
        ["CHESS-T1", "Beginner Review started independently", "Same task measures"],
        ["CHESS-T2", "Mistake and immediate consequence explained", "Same measures; own-words explanation"],
        ["CHESS-T3", "Safer move performed and feedback understood", "Same measures; source-to-destination interaction; retry/recovery"],
        ["CHESS-T4", "New practice position completed and context retained", "Same measures; practice input and continuation"],
    ]
    return {
        "analysis_status": analysis_status,
        "participant_rows": participant_rows,
        "verified_count": verified_count,
        "recording_status": gate_status("G03", "BLOCKED_RECORDINGS"),
        "task_status": gate_status("G04", "BLOCKED_TASK_DATA"),
        "questionnaire_status": gate_status("G05", "BLOCKED_QUESTIONNAIRE"),
        "interview_status": gate_status("G06", "BLOCKED_INTERVIEW_DATA"),
        "participant_state": participant_state,
        "recording_state": recording_state,
        "task_state": task_state,
        "questionnaire_state": questionnaire_state,
        "interview_state": interview_state,
        "current_states": current_states,
        "gate_rows": gate_rows,
        "tasks": tasks,
        "procedure": [
            "Greeting, purpose, consent, and recording confirmation.",
            "Background questions and think-aloud practice.",
            "Run the first assigned product tasks, then that product's questionnaire.",
            "Reset the prototype.",
            "Run the second assigned product tasks, then that product's questionnaire.",
            "Neutral post-test interview, closing, and recording filename confirmation.",
        ],
        "orders": [
            "A_FIFA_FIRST: FIFA tasks → FIFA questionnaire → reset → Chess tasks → Chess questionnaire.",
            "B_CHESS_FIRST: Chess tasks → Chess questionnaire → reset → FIFA tasks → FIFA questionnaire.",
        ],
        "scoring": "2 = independent completion without facilitator assistance; 1 = completion after one neutral prompt or a defined recoverable wrong path; 0 = failure, abandonment, or solution-revealing/direct assistance.",
        "hesitation": "Track hesitation separately: one event is at least 5 consecutive seconds without task-progress action while the participant is visibly attending, excluding expected reading. Store hesitation_count, timestamp, and observable behavior; hesitation alone does not lower success_score.",
        "recording_rule": "A verified recording is an existing MP4 with a positive duration and at least one codec_type=video stream. A present file without a probe is RECORDING_PRESENT_UNVERIFIED and never counts as verified.",
        "results_summary": f"The report reflects the canonical analysis result: {analysis_status}. It changes when evidence is ingested and the analyzer is rerun; no participant findings are inferred from empty or synthetic templates.",
        "limitations": [
            "The prototype uses static fictional data and does not measure live service reliability or partner completion.",
            "Participant-dependent findings remain subject to the canonical analysis gates shown above.",
            "PA3 formative evidence is continuity evidence and is not substituted for PA4 human sessions.",
            "Concurrent think-aloud timings are descriptive and are not a natural unmoderated baseline.",
        ],
    }


def build_hifi_docx(data: dict[str, Any] | None = None) -> Path:
    data = data or build_hifi_report_data()
    doc = Document()
    setup_docx(doc, "Hi-fi Prototype")
    docx_title(doc, "Group10 PA4 · Hi-fi Prototype", "FIFA Status Dashboard + Chess Beginner Review Flow")
    docx_callout(doc, "YOUTUBE DEMO LINK", "YouTube demo link: REQUIRED EXTERNAL EVIDENCE BEFORE SUBMISSION", fill="FFF0D8", color="97551D")
    doc.add_paragraph("This editable report records the locally achievable hi-fi implementation. No genuine external demo URL is invented while that evidence is missing.")
    doc.add_heading("Project continuity", level=1)
    doc.add_paragraph("PA1 established FIFA.com as a browse-first official football ecosystem and Chess.com as an action-first play, review, and learning platform. PA2 translated those findings into the Status Dashboard and Beginner Review concepts. PA3 selected FIFA Alt 1 and Chess Alt 1 for the PA4 hi-fi implementation.")
    doc.add_heading("Hi-fi design objectives", level=1)
    docx_bullets(doc, ["Make the current state or learning moment visible before secondary detail.", "Place the next action beside the explanation that makes it meaningful.", "Use text, icon, and state treatment together so color is not the only cue.", "Make external transitions and recovery states testable in a deterministic offline demo.", "Keep the responsive surface usable at 1440 × 900 and 390 × 844."])
    doc.add_heading("FIFA Status Dashboard", level=1)
    doc.add_paragraph("The FIFA flow uses a status-first editorial surface with fictional Mexico City pending and Toronto confirmed events. Official source, freshness, ownership, action adjacency, and the partner boundary remain visible without claiming a live transaction.")
    docx_image(doc, "fifa-desktop-overview.png", 6.35)
    doc.add_heading("FIFA key interactions", level=2)
    docx_numbered(doc, ["Select Pending or Confirmed event rows.", "Reveal the status definition and next owner.", "Open order or ticket detail.", "Save the confirmed event to the calendar.", "Preview the partner destination before transfer.", "Stay or continue, then return with context.", "Presenter mode can preview unavailable and reset states; study mode removes those researcher controls."])
    docx_image(doc, "fifa-desktop-handoff.png", 5.7)
    doc.add_heading("Chess Beginner Review Flow", level=1)
    doc.add_paragraph(f"The Chess flow uses a guided one-mistake-at-a-time route. The validated scenario checks {data['chess']['mistake']}: {data['chess']['consequence']}. The revealed alternative is {data['chess']['better']}, and the separate practice position uses {data['chess']['practice']}. {data['chess']['concept']}")
    docx_image(doc, "chess-desktop-mistake.png", 6.35)
    docx_numbered(doc, ["Start from an intro state that does not reveal the answer.", "Read the mistake and immediate consequence in plain language.", "Reveal the better move only at the appropriate stage.", "Select the source piece and destination on the trial board; wrong moves give feedback and can be retried.", "Complete the separate practice position with the same source-to-destination interaction.", "Finish the route or return to the review without losing context."])
    docx_image(doc, "chess-desktop-practice.png", 5.7)
    doc.add_heading("Study routes and responsive validation", level=1)
    docx_table(doc, ["Route", "URL"], [["FIFA study", data["urls"]["fifa"]], ["Chess study", data["urls"]["chess"]]], [1900, 7460])
    docx_table(doc, ["Viewport", "Validated behavior"], [["1440 × 900", "Product hierarchy, rail guidance, board interaction, and primary controls remain usable."], ["390 × 844", "Columns stack, board remains readable, controls wrap, and no horizontal overflow is expected."], ["Study mode", "Researcher branding, launcher, reset/preview controls, and demo help are removed from the participant DOM."], ["Presenter mode", "Launcher, demo labels, help, reset, and unavailable-state preview remain available for demonstration and QA."]], [1800, 7560])
    doc.add_heading("Limitations and external evidence", level=1)
    docx_callout(doc, "BLOCKED EXTERNALLY", "A genuine YouTube demo URL, real participant sessions, verified recordings, questionnaire/interview evidence, and measured task timings are not local facts. The prototype is locally runnable and browser-tested, but it does not use live FIFA, ticketing, partner, or Chess.com backends.", fill="FFF0D8", color="97551D")
    path = SOURCE / "Group10-PA4-HifiProtype.docx"
    doc.save(path)
    return path


def build_study_docx(data: dict[str, Any] | None = None) -> Path:
    data = data or build_summative_report_data()
    doc = Document()
    setup_docx(doc, "Summative User Study")
    docx_title(doc, "Group10 PA4 · Summative User Study", "Study design, instruments, evidence gate, and analysis plan")
    docx_callout(doc, "STUDY STATUS", f"{data['analysis_status']} · {data['participant_state']}. This report contains no inferred participant findings.", fill="FFF0D8", color="97551D")
    doc.add_heading("Study objective and research questions", level=1)
    doc.add_paragraph("Measure whether the selected hi-fi flows support effective, efficient, low-error, and satisfactory completion of the core tasks. The study evaluates interface behavior, not the participant.")
    docx_bullets(doc, ["Can users identify FIFA status, ownership, and the next action without coaching?", "Can users recognize the partner boundary and return with context?", "Can users identify and explain the Chess mistake before advanced analysis?", "Can users perform a safer Chess move, recover from an incorrect trial, and complete related practice?", "What errors, pauses, wording problems, and recovery behaviors remain?"])
    doc.add_heading("Canonical analysis state", level=1)
    doc.add_paragraph(data["results_summary"])
    docx_table(doc, ["Evidence item", "Current state", "Required before PASS"], data["current_states"], [1900, 2600, 4860], font_size=7.5)
    docx_table(doc, ["Gate", "Name", "Status"], data["gate_rows"], [850, 4500, 4010], font_size=7.5)
    doc.add_heading("Setup, roles, and order-neutral procedure", level=1)
    docx_bullets(doc, ["Quiet room or call; desktop reference 1440 × 900 and mobile reference 390 × 844 where appropriate.", "Moderated concurrent think-aloud with descriptive task time; use neutral prompts only.", "Observer records first action, timestamps, errors, wrong paths, assistance, recovery, and hesitation.", "Recording begins after consent and is verified using the media rule below.", "Target session length is 25–35 minutes; stop for withdrawal, discomfort, recording failure, or coaching risk."])
    docx_numbered(doc, data["procedure"])
    doc.add_heading("Counterbalancing", level=2)
    docx_bullets(doc, data["orders"])
    doc.add_heading("Tasks and measures", level=1)
    docx_table(doc, ["Task", "Success endpoint", "Measures"], data["tasks"], [1200, 3900, 4260], font_size=7.5)
    doc.add_paragraph(data["scoring"])
    doc.add_paragraph(data["hesitation"])
    doc.add_heading("Questionnaire and interview", level=1)
    doc.add_paragraph("Use the custom five-item Likert questionnaire after each assigned product flow. Store raw 1–5 responses and do not replace missing answers. Ask the neutral interview questions in `PA4/study/post-test-interview.md` and create only timestamped, non-empty evidence-backed feedback rows.")
    docx_table(doc, ["ID", "Statement"], [["Q1", "I could find the first useful action without help."], ["Q2", "The status or explanation was clear."], ["Q3", "I felt confident about what to do next."], ["Q4", "The feedback after my action was useful."], ["Q5", "I was satisfied with this flow."]], [1000, 8360])
    doc.add_heading("Recording and analysis rule", level=1)
    doc.add_paragraph(data["recording_rule"])
    doc.add_heading("Limitations", level=1)
    docx_bullets(doc, data["limitations"])
    doc.add_heading("Evidence index", level=1)
    docx_table(doc, ["Artifact", "Purpose", "State"], [["study-plan.md", "Design, criteria, procedure, tasks, measures, privacy", "Prepared"], ["facilitator-script.md", "Neutral moderated session script", "Prepared"], ["data/*.csv", "Raw evidence capture schema", "Header-only until real evidence is ingested"], ["evidence/recordings/README.md", "Recording convention and video verification rule", "Prepared; no recordings claimed"], ["study/analysis/", "Canonical gates, result, metrics, and synthetic-only tests", "Generated from current evidence state"]], [2400, 4800, 2160], font_size=7.5)
    docx_callout(doc, "NEXT REQUIRED ACTION", "Run the prepared study, verify consented video recordings and timestamped feedback, ingest real evidence, rerun the analyzer, and rebuild the reports before any submission package is generated.", fill="DFF3F2")
    path = SOURCE / "Group10-PA4-SummativeUserStudy.docx"
    doc.save(path)
    return path


def build_weekly_docx() -> Path:
    doc = Document()
    setup_docx(doc, "Weekly Report")
    docx_title(doc, "Group10 PA4 · Weekly Report", "Artifact snapshot and team continuity tracker")
    docx_callout(doc, "REPORT STATUS", "The local snapshot tracks the recovered Group10 roster and PA4 ownership. Participant sessions and outcome evidence remain externally blocked; future work is marked Planned.", fill="FFF0D8", color="97551D")
    doc.add_heading("Group and artifact snapshot", level=1)
    docx_table(doc, ["Field", "Value"], [["Group", "Group10"], ["Course", "CSC13112 UI/UX Design · FIT-HCMUS"], ["Snapshot date", "2026-08-23 · local artifact audit/build snapshot, not a recorded team meeting"], ["Selected directions", "FIFA Alt 1 Status Dashboard; Chess Alt 1 Beginner Review Flow"], ["External gate", "YouTube URL; five real participants; recordings; questionnaire/interview/timing evidence"], ["Template gate", "Official Weekly Report template is not present in the local workspace; this is a labeled local snapshot"]], [1800, 7560])
    doc.add_heading("Real Group10 roster", level=1)
    docx_table(doc, ["Member", "ID", "Prior continuity role"], [["Le Minh", "21127645", "Coordination, integration, peer review, weekly report, packaging"], ["Nguyen Vu Bach", "21127224", "FIFA research and evidence lead"], ["Pham Nguyen Gia Bao", "20127119", "Chess research and evidence lead"], ["Trang Minh Nhut", "22127318", "HCI analysis, solutions, visual QA"]], [1900, 1300, 6160])
    doc.add_heading("Evidence and blockers", level=1)
    docx_bullets(doc, ["Local prototype, browser QA, scenario validation, study materials, analysis gates, reports, and working-package separation are maintained as artifacts.", "External blockers remain: genuine YouTube demo URL, real participants, consented recordings, task/questionnaire/interview evidence, and any comparable current-practice baseline.", "The official Weekly Report template is not present locally and is not claimed as satisfied.", "No participant identities, demographics, quotes, recordings, timing values, or questionnaire answers are invented."])
    doc.add_heading("Next weekly checkpoint", level=1)
    doc.add_paragraph("Confirm ownership, run sessions, verify recordings and timestamped feedback, ingest real evidence, rerun analysis, and replace only evidence-backed gates.")
    path = SOURCE / "Group10-PA4-WeeklyReport.docx"
    doc.save(path)
    return path


def build_hifi_pdf(data: dict[str, Any] | None = None) -> Path:
    data = data or build_hifi_report_data()
    styles = pdf_styles()
    path = FINAL / "Group10-PA4-HifiProtype.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=letter, leftMargin=0.7 * inch, rightMargin=0.7 * inch, topMargin=0.7 * inch, bottomMargin=0.65 * inch, title="Group10 PA4 Hi-fi Prototype", author="Group10")
    story: list[Any] = [
        pdf_p("GROUP10 / PA4 · HI-FI PROTOTYPE", styles["coverKicker"]),
        pdf_p("Hi-fi Prototype", styles["title"]),
        pdf_p("FIFA Status Dashboard + Chess Beginner Review Flow", styles["subtitle"]),
        pdf_callout("YOUTUBE DEMO LINK", "YouTube demo link: REQUIRED EXTERNAL EVIDENCE BEFORE SUBMISSION", styles, fill="FFF0D8"),
        Spacer(1, 0.18 * inch),
        pdf_p("A coded, offline demonstrator carrying the selected PA3 directions into a polished, testable interface. Fictional data is labeled and no live service completion is claimed.", styles["body"]),
        pdf_table(["Product", "Selected PA3 base", "PA4 focus"], [["FIFA", "Alt 1 · Status Dashboard", "Status → owner → action → handoff"], ["Chess", "Alt 1 · Beginner Review Flow", "Mistake → consequence → try → practice"]], [1.1, 2.35, 3.65], styles),
        PageBreak(),
        pdf_p("FIFA Status Dashboard", styles["h1"]),
        pdf_p("The FIFA surface keeps fictional status, owner, freshness, action adjacency, and the partner boundary visible.", styles["body"]),
        pdf_image("fifa-desktop-overview.png", 6.35),
        pdf_bullets(["Select Pending or Confirmed event rows.", "Reveal the Pending definition and next owner.", "Open order or ticket detail and save the confirmed event.", "Preview the partner destination, stay or continue, and return with context.", "Presenter mode retains reset and unavailable-state preview for QA; study mode removes researcher controls."], styles["body"]),
        pdf_image("fifa-desktop-handoff.png", 5.7),
        pdf_p("Chess Beginner Review Flow", styles["h1"]),
        pdf_p(f"The validated Chess scenario checks {data['chess']['mistake']}: {data['chess']['consequence']}. The revealed alternative is {data['chess']['better']}; the related practice move is {data['chess']['practice']}. {data['chess']['concept']}", styles["body"]),
        pdf_image("chess-desktop-mistake.png", 6.35),
        pdf_bullets(["Intro hides the exact mistake and answer.", "Mistake state reveals the immediate consequence before the better move.", "Trial requires selecting a source piece and destination; wrong moves provide feedback and retry.", "Practice uses a new validated micro-position and requires the same two-step input.", "Completion provides a return path without claiming engine or live-service behavior."], styles["body"]),
        pdf_image("chess-desktop-practice.png", 5.7),
        pdf_p("Study routes and responsive validation", styles["h1"]),
        pdf_table(["Route", "URL"], [["FIFA study", data["urls"]["fifa"]], ["Chess study", data["urls"]["chess"]]], [1.35, 5.75], styles),
        pdf_table(["Viewport", "Validated behavior"], [["1440 × 900", "Product hierarchy, rail guidance, board interaction, and primary controls remain usable."], ["390 × 844", "Columns stack, the board remains readable, controls wrap, and no horizontal overflow is expected."], ["Study mode", "Researcher branding, launcher, reset/preview controls, and demo help are removed from the participant DOM."], ["Presenter mode", "Launcher, demo labels, help, reset, and unavailable-state preview remain available for QA."]], [1.35, 5.75], styles),
        pdf_p("Limitations and external evidence", styles["h1"]),
        pdf_callout("BLOCKED EXTERNALLY", "A genuine YouTube demo URL, real participant sessions, verified recordings, questionnaire/interview evidence, and measured task timings are not local facts. The prototype is locally runnable and browser-tested, but it does not use live backends.", styles, fill="FFF0D8"),
    ]
    doc.build(story, onFirstPage=lambda canvas, document: page_decor(canvas, document, "Hi-fi Prototype", "F4F7F8"), onLaterPages=lambda canvas, document: page_decor(canvas, document, "Hi-fi Prototype"))
    return path


def build_study_pdf(data: dict[str, Any] | None = None) -> Path:
    data = data or build_summative_report_data()
    styles = pdf_styles()
    path = FINAL / "Group10-PA4-SummativeUserStudy.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=letter, leftMargin=0.7 * inch, rightMargin=0.7 * inch, topMargin=0.7 * inch, bottomMargin=0.65 * inch, title="Group10 PA4 Summative User Study", author="Group10")
    story: list[Any] = [
        pdf_p("GROUP10 / PA4 · SUMMATIVE USER STUDY", styles["coverKicker"]),
        pdf_p("Summative User Study", styles["title"]),
        pdf_p("Study design, instruments, evidence gate, and analysis plan", styles["subtitle"]),
        pdf_callout("STUDY STATUS", f"{data['analysis_status']} · {data['participant_state']}. No participant findings are inferred.", styles, fill="FFF0D8"),
        Spacer(1, 0.16 * inch),
        pdf_p(data["results_summary"], styles["body"]),
        PageBreak(),
        pdf_p("Study objective and research questions", styles["h1"]),
        pdf_p("Measure whether the selected hi-fi flows support effective, efficient, low-error, and satisfactory completion of the core tasks. The study evaluates interface behavior, not the participant.", styles["body"]),
        pdf_bullets(["Can users identify FIFA status, ownership, and the next action without coaching?", "Can users recognize the partner boundary and return with context?", "Can users identify and explain the Chess mistake before advanced analysis?", "Can users perform a safer Chess move, recover, and complete related practice?", "What errors, pauses, wording problems, and recovery behaviors remain?"], styles["body"]),
        pdf_p("Canonical analysis state", styles["h1"]),
        pdf_table(["Evidence item", "Current state", "Required before PASS"], data["current_states"], [1.55, 2.0, 3.55], styles),
        pdf_table(["Gate", "Name", "Status"], data["gate_rows"], [0.7, 4.4, 2.0], styles),
        pdf_p("Setup and order-neutral procedure", styles["h1"]),
        pdf_bullets(["Quiet room or call; desktop 1440 × 900 and mobile 390 × 844 where appropriate.", "Moderated concurrent think-aloud with descriptive task time and neutral prompts only.", "Observer records first action, timestamps, errors, wrong paths, assistance, recovery, and hesitation.", "Recording begins after consent and is verified using the media rule below.", "Stop for withdrawal, discomfort, recording failure, or coaching risk."], styles["body"]),
        pdf_table(["Step", "Activity"], [[str(index + 1), item] for index, item in enumerate(data["procedure"])], [0.55, 6.55], styles),
        pdf_p("Counterbalancing", styles["h2"]),
        pdf_bullets(data["orders"], styles["body"]),
        pdf_p("Tasks and measures", styles["h1"]),
        pdf_table(["Task", "Success endpoint", "Measures"], data["tasks"], [1.0, 3.2, 2.9], styles),
        pdf_p(data["scoring"], styles["body"]),
        pdf_p(data["hesitation"], styles["body"]),
        pdf_p("Questionnaire and interview", styles["h1"]),
        pdf_p("Use the custom five-item Likert questionnaire after each assigned product flow. Store raw 1–5 responses and do not replace missing answers. Ask the neutral interview questions and create only timestamped, non-empty evidence-backed feedback rows.", styles["body"]),
        pdf_table(["ID", "Statement"], [["Q1", "I could find the first useful action without help."], ["Q2", "The status or explanation was clear."], ["Q3", "I felt confident about what to do next."], ["Q4", "The feedback after my action was useful."], ["Q5", "I was satisfied with this flow."]], [0.65, 6.35], styles),
        pdf_p("Recording and analysis rule", styles["h1"]),
        pdf_p(data["recording_rule"], styles["body"]),
        pdf_p("Limitations", styles["h1"]),
        pdf_bullets(data["limitations"], styles["body"]),
        pdf_p("Evidence index", styles["h1"]),
        pdf_table(["Artifact", "Purpose", "State"], [["study-plan.md", "Design, criteria, procedure, tasks, measures, privacy", "Prepared"], ["facilitator-script.md", "Neutral moderated session script", "Prepared"], ["data/*.csv", "Raw evidence capture schema", "Header-only until real evidence is ingested"], ["evidence/recordings/README.md", "Recording convention and video verification rule", "Prepared; no recordings claimed"], ["study/analysis/", "Canonical gates, result, metrics, and synthetic-only tests", "Generated from current evidence state"]], [2.0, 3.8, 1.3], styles),
        pdf_callout("NEXT REQUIRED ACTION", "Run the prepared study, verify consented video recordings and timestamped feedback, ingest real evidence, rerun the analyzer, and rebuild the reports before any submission package is generated.", styles, fill="DFF3F2"),
    ]
    doc.build(story, onFirstPage=lambda canvas, document: page_decor(canvas, document, "Summative User Study", "FFF9EE"), onLaterPages=lambda canvas, document: page_decor(canvas, document, "Summative User Study"))
    return path


def build_weekly_pdf() -> Path:
    styles = pdf_styles()
    path = FINAL / "Group10-PA4-WeeklyReport.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=letter, leftMargin=0.7 * inch, rightMargin=0.7 * inch, topMargin=0.7 * inch, bottomMargin=0.65 * inch, title="Group10 PA4 Weekly Report", author="Group10")
    story: list[Any] = [
        pdf_p("GROUP10 / PA4 · WEEKLY REPORT", styles["coverKicker"]),
        pdf_p("Weekly Report", styles["title"]),
        pdf_p("Artifact snapshot and team continuity tracker", styles["subtitle"]),
        pdf_callout("REPORT STATUS", "The local snapshot tracks the recovered Group10 roster and PA4 ownership. Participant sessions and outcome evidence remain externally blocked; future work is marked Planned.", styles, fill="FFF0D8"),
        pdf_p("Group and artifact snapshot", styles["h1"]),
        pdf_table(["Field", "Value"], [["Group", "Group10"], ["Course", "CSC13112 UI/UX Design · FIT-HCMUS"], ["Snapshot date", "2026-08-23 · local artifact audit/build snapshot, not a recorded team meeting"], ["Selected directions", "FIFA Alt 1 Status Dashboard; Chess Alt 1 Beginner Review Flow"], ["External gate", "YouTube URL; five real participants; recordings; questionnaire/interview/timing evidence"], ["Template gate", "Official Weekly Report template is not present locally; this snapshot does not claim template compliance"]], [1.35, 5.75], styles),
        pdf_p("Real Group10 roster", styles["h1"]),
        pdf_table(["Member", "ID", "Prior continuity role"], [["Le Minh", "21127645", "Coordination, integration, peer review, weekly report, packaging"], ["Nguyen Vu Bach", "21127224", "FIFA research and evidence lead"], ["Pham Nguyen Gia Bao", "20127119", "Chess research and evidence lead"], ["Trang Minh Nhut", "22127318", "HCI analysis, solutions, visual QA"]], [1.45, 1.0, 4.65], styles),
        pdf_p("Evidence and blockers", styles["h1"]),
        pdf_bullets(["Local prototype, browser QA, scenario validation, study materials, analysis gates, reports, and working-package separation are maintained as artifacts.", "External blockers remain: genuine YouTube demo URL, real participants, consented recordings, task/questionnaire/interview evidence, and any comparable current-practice baseline.", "The official Weekly Report template is not present locally and is not claimed as satisfied.", "No participant identities, demographics, quotes, recordings, timing values, or questionnaire answers are invented."], styles["body"]),
        pdf_p("Next weekly checkpoint", styles["h1"]),
        pdf_p("Confirm ownership, run sessions, verify recordings and timestamped feedback, ingest real evidence, rerun analysis, and replace only evidence-backed gates.", styles["body"]),
    ]
    doc.build(story, onFirstPage=lambda canvas, document: page_decor(canvas, document, "Weekly Report", "F4F7F8"), onLaterPages=lambda canvas, document: page_decor(canvas, document, "Weekly Report"))
    return path


def main() -> None:
    SOURCE.mkdir(parents=True, exist_ok=True)
    FINAL.mkdir(parents=True, exist_ok=True)
    hifi_data = build_hifi_report_data()
    study_data = build_summative_report_data()
    docx_paths = [build_hifi_docx(hifi_data), build_study_docx(study_data), build_weekly_docx()]
    pdf_paths = [build_hifi_pdf(hifi_data), build_study_pdf(study_data), build_weekly_pdf()]
    for path in docx_paths + pdf_paths:
        print(path)


if __name__ == "__main__":
    main()
