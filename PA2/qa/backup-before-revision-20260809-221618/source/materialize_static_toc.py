from __future__ import annotations

import re
from pathlib import Path

import pdfplumber
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(r"C:\Users\Administrator\Documents\MEGA\tkgd\PA2")
SOURCE = ROOT / "source" / "rebuilt"
FINAL = ROOT / "final"
STEMS = [
    "Group10-PA2-UserResearch",
    "Group10-PA2-UserAnalysis",
    "Group10-PA2-ProjectProposal",
    "Group10-PA2-UseCaseDocument",
]


def normalize(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip().casefold()


def page_texts(pdf: Path) -> list[str]:
    with pdfplumber.open(pdf) as document:
        return [normalize(page.extract_text() or "") for page in document.pages]


def locate(heading: str, pages: list[str]) -> int:
    needle = normalize(heading)
    candidates = [needle, re.sub(r"^\d+\.\s*", "", needle)]
    for candidate in candidates:
        for page_number, page in enumerate(pages[2:], start=3):
            if candidate and candidate in page:
                return page_number
    # Match stable first words when PDF extraction changes punctuation.
    short = " ".join(re.sub(r"^\d+\.\s*", "", needle).split()[:5])
    for page_number, page in enumerate(pages[2:], start=3):
        if short and short in page:
            return page_number
    raise ValueError(f"Heading not found in PDF text: {heading}")


def clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        p.remove(child)


def main() -> None:
    for stem in STEMS:
        docx = SOURCE / f"{stem}.docx"
        pdf = FINAL / f"{stem}.pdf"
        pages = page_texts(pdf)
        doc = Document(docx)
        headings = [
            paragraph.text.strip()
            for paragraph in doc.paragraphs
            if paragraph.style.name == "Heading 1"
            and paragraph.text.strip()
            and paragraph.text.strip() != "Contents"
        ]
        entries = [(heading, locate(heading, pages)) for heading in headings]
        contents_index = next(
            i for i, paragraph in enumerate(doc.paragraphs) if paragraph.text.strip() == "Contents"
        )
        toc = doc.paragraphs[contents_index + 1]
        clear_paragraph(toc)
        toc.alignment = WD_ALIGN_PARAGRAPH.LEFT
        toc.paragraph_format.space_before = Pt(2)
        toc.paragraph_format.space_after = Pt(0)
        toc.paragraph_format.line_spacing = 1.08
        for index, (heading, page) in enumerate(entries):
            run = toc.add_run(f"{heading}  {'.' * max(4, 54 - len(heading))}  {page}")
            run.font.name = "Arial"
            run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
            run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string("1F2937")
            if index < len(entries) - 1:
                run.add_break()
        doc.save(docx)
        print(stem, entries)


if __name__ == "__main__":
    main()
