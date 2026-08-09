from __future__ import annotations

import csv
import math
import os
import textwrap
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\Administrator\Documents\MEGA\tkgd\PA2")
PA1 = Path(r"C:\Users\Administrator\Documents\MEGA\tkgd\PA1")
SOURCE = ROOT / "source" / "rebuilt"
DIAGRAMS = ROOT / "generated-diagrams"
FINAL = ROOT / "final"
CAP = ROOT / "capture-work"

TEAM = [
    ("Le Minh", "21127645"),
    ("Nguyen Vu Bach", "21127224"),
    ("Pham Nguyen Gia Bao", "20127119"),
    ("Trang Minh Nhut", "22127318"),
]

NAVY = "123B65"
BLUE = "1D70A2"
GREEN = "4F7F35"
INK = "1F2937"
MUTED = "5B6573"
LIGHT = "F3F6FA"
CHESS_LIGHT = "EEF5E9"
FIFA_LIGHT = "EAF2F8"
ORANGE = "C66A1B"
OBS_FILL = "#DCEEFF"
DOC_FILL = "#FDE8C8"

FONT_REG = Path(r"C:\Windows\Fonts\arial.ttf")
FONT_BOLD = Path(r"C:\Windows\Fonts\arialbd.ttf")


def font(size: int, bold: bool = False):
    return ImageFont.truetype(str(FONT_BOLD if bold else FONT_REG), size)


def add_field(paragraph, instruction: str, placeholder: str = "1"):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, sep, text, end])


def shade(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_table_header_and_geometry(table, widths_cm: list[float] | None = None):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if table.rows:
        repeat_table_header(table.rows[0])
    if widths_cm:
        for row in table.rows:
            for i, width in enumerate(widths_cm):
                row.cells[i].width = Cm(width)
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.space_before = Pt(1)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
                    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
                    run.font.size = Pt(8.5 if len(table.columns) >= 4 else 9)
                    run.font.color.rgb = RGBColor.from_string(INK)
            if ri == 0:
                shade(cell, NAVY)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)


def base_document(title: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, color, before, after in [
        ("Title", 28, NAVY, 0, 8),
        ("Subtitle", 14, MUTED, 0, 12),
        ("Heading 1", 17, NAVY, 14, 7),
        ("Heading 2", 13, BLUE, 10, 5),
        ("Heading 3", 11, GREEN, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Cm(0.65)
        style.paragraph_format.first_line_indent = Cm(-0.3)
        style.paragraph_format.space_after = Pt(3)
    doc.core_properties.title = f"Group10 PA2 - {title}"
    doc.core_properties.subject = "CSC13112 UI/UX Design Project Assignment 2"
    doc.core_properties.author = "Group10"
    return doc


def add_header_footer(doc: Document, short_title: str):
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.text = f"GROUP10  |  CSC13112 UI/UX DESIGN  |  {short_title.upper()}"
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(8)
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(NAVY)
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.add_run(f"{short_title}  |  Group10  |  ")
        add_field(fp, "PAGE")
        for run in fp.runs:
            run.font.name = "Arial"
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor.from_string(MUTED)


def add_cover(doc: Document, title: str, subtitle: str, status: str | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(70)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PROJECT ASSIGNMENT 2")
    r.font.name = "Arial"
    r.font.size = Pt(11)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(ORANGE)
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(title)
    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(subtitle)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("FIFA.com and Chess.com")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    table = doc.add_table(rows=5 + (1 if status else 0), cols=2)
    table.style = "Table Grid"
    values = [
        ("Course", "CSC13112 - UI/UX Design, FIT-HCMUS"),
        ("Group", "Group10"),
        ("Team", "; ".join(f"{n} ({sid})" for n, sid in TEAM)),
        ("Evidence date", "Local evidence audited through 30 July 2026"),
        ("Document type", "Evidence-based redesign report"),
    ]
    if status:
        values.append(("Status", status))
    for i, (label, value) in enumerate(values):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value
        shade(table.cell(i, 0), FIFA_LIGHT if i % 2 == 0 else CHESS_LIGHT)
        for run in table.cell(i, 0).paragraphs[0].runs:
            run.bold = True
    set_repeat_table_header_and_geometry(table, [3.4, 12.2])
    for cell in table.rows[0].cells:
        # This is metadata, not a header row.
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor.from_string(INK)
    doc.add_page_break()


def add_toc_placeholder(doc: Document):
    doc.add_heading("Contents", level=1)
    doc.add_paragraph("[[TOC]]")
    doc.add_page_break()


def callout(doc: Document, label: str, text: str, fill: str = FIFA_LIGHT):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(f"{label}: ")
    r.bold = True
    p.add_run(text)
    set_cell_margins(cell, 130, 160, 130, 160)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered_steps(doc: Document, items: list[str]):
    """Add a real Word-numbered list that restarts at 1 for this call."""
    numbering = doc.part.numbering_part.element
    abstract_id = None
    for abstract in numbering.findall(qn("w:abstractNum")):
        for style in abstract.findall(".//" + qn("w:pStyle")):
            if style.get(qn("w:val")) == "ListNumber":
                abstract_id = abstract.get(qn("w:abstractNumId"))
                break
        if abstract_id is not None:
            break
    if abstract_id is None:
        abstract_id = "0"
    existing = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = str(max(existing, default=0) + 1)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), num_id)
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), abstract_id)
    num.append(abstract_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.add_run(item)
        properties = paragraph._p.get_or_add_pPr()
        num_pr = properties.get_or_add_numPr()
        num_pr.get_or_add_ilvl().set(qn("w:val"), "0")
        num_pr.get_or_add_numId().set(qn("w:val"), num_id)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.cell(0, i).text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    set_repeat_table_header_and_geometry(table, widths)
    return table


def add_figure(
    doc: Document,
    path: Path,
    figure_id: str,
    caption: str,
    source: str,
    related: str,
    width_cm=16.0,
    max_height_cm=17.5,
):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    with Image.open(path) as image:
        projected_height_cm = width_cm * image.height / image.width
    if projected_height_cm > max_height_cm:
        run.add_picture(str(path), height=Cm(max_height_cm))
    else:
        run.add_picture(str(path), width=Cm(width_cm))
    cp = doc.add_paragraph()
    cp.paragraph_format.keep_together = True
    cp.paragraph_format.space_after = Pt(8)
    r = cp.add_run(f"Figure {figure_id}. ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    cp.add_run(caption)
    sr = cp.add_run(f"\nSource: {source}\nRelated: {related}")
    sr.italic = True
    sr.font.size = Pt(8)
    sr.font.color.rgb = RGBColor.from_string(MUTED)


def wrap(text: str, width: int):
    return textwrap.wrap(text, width=width, break_long_words=False, break_on_hyphens=False)


def rounded_box(draw, box, fill, outline, radius=18, width=3):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def flow_diagram(name: str, title: str, nodes: list[str], breakdowns: list[str], accent="#1D70A2"):
    DIAGRAMS.mkdir(parents=True, exist_ok=True)
    w, h = 2200, 1250
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    d.text((80, 55), title, font=font(50, True), fill="#123B65")
    n = len(nodes)
    box_w = min(340, int((w - 180) / max(3, min(n, 5))) - 30)
    cols = min(5, n)
    rows = math.ceil(n / cols)
    positions = []
    for i, node in enumerate(nodes):
        row, col = divmod(i, cols)
        x = 90 + col * ((w - 180) / cols)
        y = 210 + row * 320
        box = (int(x), int(y), int(x + box_w), int(y + 180))
        rounded_box(d, box, "#F3F6FA", accent)
        lines = wrap(node, 22)
        yy = y + 38
        for line in lines[:4]:
            bbox = d.textbbox((0, 0), line, font=font(28, True))
            d.text((x + (box_w - (bbox[2] - bbox[0])) / 2, yy), line, font=font(28, True), fill="#1F2937")
            yy += 36
        positions.append((box[2], (box[1] + box[3]) // 2, box[0], (box[1] + box[3]) // 2))
        if i > 0:
            px = positions[i - 1][0]
            py = positions[i - 1][1]
            cx = box[0]
            cy = (box[1] + box[3]) // 2
            if row == (i - 1) // cols:
                d.line((px + 8, py, cx - 12, cy), fill=accent, width=6)
                d.polygon([(cx - 12, cy), (cx - 30, cy - 10), (cx - 30, cy + 10)], fill=accent)
    y0 = 210 + rows * 320
    d.text((90, y0), "Breakdowns / decision risks", font=font(34, True), fill="#9B1C1C")
    x = 90
    for item in breakdowns:
        lines = wrap(item, 28)
        height = 52 + len(lines) * 30
        box = (x, y0 + 65, min(x + 430, w - 90), y0 + 65 + height)
        rounded_box(d, box, "#FFF1F1", "#9B1C1C", radius=12, width=2)
        yy = box[1] + 18
        for line in lines:
            d.text((box[0] + 18, yy), line, font=font(24), fill="#5E1B1B")
            yy += 30
        x += 450
    png = DIAGRAMS / f"{name}.png"
    img.save(png, dpi=(220, 220))
    svg = DIAGRAMS / f"{name}.svg"
    svg_nodes = []
    for i, node in enumerate(nodes):
        row, col = divmod(i, cols)
        x = 90 + col * ((w - 180) / cols)
        y = 210 + row * 320
        svg_nodes.append(
            f'<rect x="{x:.0f}" y="{y}" width="{box_w}" height="180" rx="18" fill="#F3F6FA" stroke="{accent}" stroke-width="3"/>'
            f'<text x="{x + box_w/2:.0f}" y="{y+92}" text-anchor="middle" font-family="Arial" font-size="26" font-weight="700">{node.replace("&","&amp;")}</text>'
        )
    svg.write_text(
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{w}" height="{h}" viewBox="0 0 {w} {h}">'
        f'<rect width="100%" height="100%" fill="white"/><text x="80" y="90" font-family="Arial" font-size="50" font-weight="700" fill="#123B65">{title}</text>'
        + "".join(svg_nodes)
        + "</svg>",
        encoding="utf-8",
    )
    return png


def affinity_diagram():
    clusters = [
        ("FIFA-A1", "Quick task entry and navigation"),
        ("FIFA-A2", "Match and search scanning"),
        ("FIFA-A3", "Ticket status, planning, and trust"),
        ("FIFA-A4", "Partner and sibling-property continuity"),
        ("FIFA-A5", "Mobile interruption and long pages"),
        ("CHESS-A1", "Beginner task selection"),
        ("CHESS-A2", "Play setup and pre-task choices"),
        ("CHESS-A3", "Puzzle and feedback flow"),
        ("CHESS-A4", "Learning path and content priority"),
        ("CHESS-A5", "Analysis comprehension and disclosure"),
        ("CHESS-A6", "Promotions and mobile attention"),
    ]
    notes = {
        "FIFA-A1":["[OBS] Modal covers first task view","[OBS] Compact header hides full map","[PA1] Ecosystem competes at top level","[OBS] Rankings and tickets are separate routes","[IDEA] Task-first intent bar"],
        "FIFA-A2":["[OBS] Date rail supports time model","[OBS] Filters add selection work","[OBS] Search mixes content types","[OBS] Pagination extends comparison","[IDEA] Preserve filters while narrowing"],
        "FIFA-A3":["[OBS] Cards organize by tournament","[OBS] Register interest and Buy now coexist","[OBS] No consolidated status view","[OBS] Mobile cards continue vertically","[PA1] F-D4 status uncertainty"],
        "FIFA-A4":["[OBS] FIFA+ route can end in error state","[OBS] DAZN changes brand context","[OBS] Consent adds a step","[OBS] Store/Collect/Rewards use new vocabularies","[PA1] F-D2 continuity break"],
        "FIFA-A5":["[OBS] Mobile home is very long","[OBS] Mobile Match Centre stacks groups","[OBS] Ticket cards stack vertically","[OBS] Article has a long sparse interval","[IDEA] Preserve orientation on return"],
        "CHESS-A1":["[OBS] Home exposes many task families","[OBS] Mobile home stacks all task families","[PA1] C-D1 feature overload","[OBS] Clear labels aid recognition","[IDEA] One recommended next action"],
        "CHESS-A2":["[OBS] Start page has four play routes","[OBS] Skill modal adds a gate","[OBS] Bot page exposes many levels","[OBS] Ads sit beside board on desktop","[IDEA] Beginner preset before options"],
        "CHESS-A3":["[OBS] Board and progression split attention","[OBS] Solve Puzzles is a clear CTA","[OBS] Candidate state does not prove outcome","[OBS] Mobile removes the board at entry","[IDEA] Feedback close to the move"],
        "CHESS-A4":["[OBS] Lessons page has a long catalog","[OBS] Learn-to-Play shows clear Next Lesson","[OBS] Study plan is article content","[OBS] Mobile list is extremely long","[IDEA] Continue one path first"],
        "CHESS-A5":["[OBS] Analysis entry offers many setup paths","[OBS] Start Analysis is visible","[OBS] No beginner explanation at entry","[PA1] C-D2 analysis overload","[IDEA] Reveal advanced detail on demand"],
        "CHESS-A6":["[OBS] Promotions appear near lessons","[OBS] Ads appear beside play and puzzles","[OBS] Mobile stacking extends attention cost","[PA1] Purposeful board whitespace can help","[IDEA] Protect the primary task region"],
    }
    w, h = 3200, 2500
    img = Image.new("RGB", (w, h), "#F7F8FA")
    d = ImageDraw.Draw(img)
    d.text((80, 50), "PA2 Affinity Diagram - 55 Evidence and Idea Notes", font=font(54, True), fill="#123B65")
    d.text((80, 125), "Blue [OBS] = screenshot evidence  |  Sand [PA1] = PA1/course evidence  |  Gray [IDEA] = design inference", font=font(28), fill="#374151")
    cols = 3
    card_w = 980
    card_h = 520
    for i, (cid, title) in enumerate(clusters):
        row, col = divmod(i, cols)
        x, y = 80 + col * 1035, 210 + row * 555
        rounded_box(d, (x, y, x + card_w, y + card_h), "white", "#CBD5E1", 16, 3)
        d.text((x + 22, y + 18), cid, font=font(29, True), fill="#123B65" if "FIFA" in cid else "#4F7F35")
        d.text((x + 160, y + 18), title, font=font(27, True), fill="#1F2937")
        yy = y + 75
        for note in notes[cid]:
            fill = OBS_FILL if note.startswith("[OBS]") else DOC_FILL if note.startswith("[PA1]") else "#E5E7EB"
            rounded_box(d, (x + 25, yy, x + card_w - 25, yy + 72), fill, "#94A3B8", 9, 2)
            d.text((x + 42, yy + 19), note, font=font(24), fill="#1F2937")
            yy += 82
    png = DIAGRAMS / "affinity-diagram.png"
    img.save(png, dpi=(220, 220))
    svg = DIAGRAMS / "affinity-diagram.svg"
    svg.write_text(
        '<svg xmlns="http://www.w3.org/2000/svg" width="3200" height="2500"><rect width="100%" height="100%" fill="#F7F8FA"/>'
        '<text x="80" y="90" font-family="Arial" font-size="54" font-weight="700" fill="#123B65">PA2 Affinity Diagram - 55 Evidence and Idea Notes</text>'
        + "".join(
            f'<text x="{100+(i%3)*1035}" y="{270+(i//3)*555}" font-family="Arial" font-size="28" font-weight="700">{cid}: {title.replace("&","&amp;")}</text>'
            for i, (cid, title) in enumerate(clusters)
        )
        + "</svg>",
        encoding="utf-8",
    )
    return png


def affinity_product_view(source: Path, name: str, title: str, indices: list[int]):
    """Assemble readable product-specific views from the unchanged full-note cards."""
    with Image.open(source) as full:
        card_w, card_h = 980, 520
        canvas = Image.new("RGB", (2200, 350 + math.ceil(len(indices) / 2) * 570), "#F7F8FA")
        draw = ImageDraw.Draw(canvas)
        draw.text((70, 55), title, font=font(48, True), fill="#123B65")
        draw.text((70, 125), "Exact note-card crops from the full 55-note artifact; no note content changed.", font=font(26), fill="#374151")
        for out_i, source_i in enumerate(indices):
            source_row, source_col = divmod(source_i, 3)
            sx, sy = 80 + source_col * 1035, 210 + source_row * 555
            card = full.crop((sx, sy, sx + card_w, sy + card_h))
            row, col = divmod(out_i, 2)
            x, y = 70 + col * 1050, 220 + row * 570
            canvas.paste(card, (x, y))
    path = DIAGRAMS / f"{name}.png"
    canvas.save(path, dpi=(220, 220))
    return path


def evidence_crop(figure_id: str, path: Path) -> Path:
    """Create a readable first-decision-surface crop while preserving source pixels."""
    crop_dir = DIAGRAMS / "evidence-crops"
    crop_dir.mkdir(parents=True, exist_ok=True)
    output = crop_dir / f"{figure_id.lower()}-decision-surface.png"
    with Image.open(path) as image:
        width, height = image.size
        crop_height = min(height, int(width * (1.25 if width < 600 else 0.78)))
        cropped = image.crop((0, 0, width, crop_height))
        if cropped.width < 1200:
            scale = 1200 / cropped.width
            cropped = cropped.resize((1200, round(cropped.height * scale)), Image.Resampling.LANCZOS)
        cropped.save(output)
    return output


def lowfi_screen_map(name: str, title: str, panels: list[tuple[str, list[str]]], accent="#1D70A2"):
    """Generate a conceptual low-fidelity screen map from shapes and labels only."""
    w, h = 2200, 1300
    image = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(image)
    draw.text((80, 55), title, font=font(50, True), fill="#123B65")
    panel_w = int((w - 180 - (len(panels) - 1) * 45) / len(panels))
    for index, (panel_title, components) in enumerate(panels):
        x = 90 + index * (panel_w + 45)
        rounded_box(draw, (x, 190, x + panel_w, 1150), "#FAFAFA", accent, 18, 4)
        draw.rectangle((x, 190, x + panel_w, 280), fill=accent)
        draw.text((x + 25, 215), panel_title, font=font(30, True), fill="white")
        y = 330
        for component in components:
            box_h = 105 if len(component) < 35 else 145
            rounded_box(draw, (x + 35, y, x + panel_w - 35, y + box_h), "#F3F6FA", "#94A3B8", 10, 2)
            lines = wrap(component, 28)
            yy = y + 25
            for line in lines[:3]:
                draw.text((x + 55, yy), line, font=font(25, True), fill="#1F2937")
                yy += 31
            y += box_h + 32
    output = DIAGRAMS / f"{name}.png"
    image.save(output, dpi=(220, 220))
    return output


def uml_use_case_diagram(name: str, title: str, system_name: str, actor: str, external: list[str], use_cases: list[str], accent="#1D70A2"):
    """Draw a UML-style system boundary with actors outside and use-case ellipses inside."""
    w, h = 2400, 1700
    image = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(image)
    draw.text((80, 45), title, font=font(50, True), fill="#123B65")
    boundary = (520, 170, 1880, 1580)
    draw.rectangle(boundary, outline=accent, width=6)
    draw.text((560, 195), system_name, font=font(32, True), fill=accent)

    def stick_figure(cx, cy, label):
        draw.ellipse((cx - 38, cy - 120, cx + 38, cy - 44), outline="#1F2937", width=5)
        draw.line((cx, cy - 44, cx, cy + 70), fill="#1F2937", width=5)
        draw.line((cx - 70, cy, cx + 70, cy), fill="#1F2937", width=5)
        draw.line((cx, cy + 70, cx - 65, cy + 160), fill="#1F2937", width=5)
        draw.line((cx, cy + 70, cx + 65, cy + 160), fill="#1F2937", width=5)
        for line_i, line in enumerate(wrap(label, 18)):
            bbox = draw.textbbox((0, 0), line, font=font(26, True))
            draw.text((cx - (bbox[2] - bbox[0]) / 2, cy + 185 + line_i * 30), line, font=font(26, True), fill="#1F2937")

    stick_figure(245, 760, actor)
    ext_positions = [(2160, 520 + i * 430) for i in range(len(external))]
    for (cx, cy), label in zip(ext_positions, external):
        stick_figure(cx, cy, label)

    centers = []
    for i, label in enumerate(use_cases):
        row, col = divmod(i, 2)
        cx, cy = 870 + col * 640, 420 + row * 350
        draw.ellipse((cx - 260, cy - 90, cx + 260, cy + 90), fill="#F3F6FA", outline=accent, width=4)
        lines = wrap(label, 32)
        yy = cy - (len(lines) * 18)
        for line in lines:
            bbox = draw.textbbox((0, 0), line, font=font(25, True))
            draw.text((cx - (bbox[2] - bbox[0]) / 2, yy), line, font=font(25, True), fill="#1F2937")
            yy += 34
        draw.line((315, 760, cx - 260, cy), fill="#64748B", width=3)
        centers.append((cx, cy))
    for ext_i, (cx, cy) in enumerate(ext_positions):
        target = centers[min(len(centers) - 1, 1 + ext_i * 3)]
        draw.line((1880, target[1], cx - 75, cy), fill="#64748B", width=3)
    output = DIAGRAMS / f"{name}.png"
    image.save(output, dpi=(220, 220))
    return output


def generate_diagrams():
    affinity = affinity_diagram()
    diagrams = {
        "wm_f1": flow_diagram("wm-f1-ticket-planning","WM-F1 FIFA Ticket-Planning Flow",["Ticket buyer","FIFA.com","Tournament card","Tickets & Hospitality","Official partner","Identity boundary","Notification channel","Payment context"],["Status uncertainty","Seat availability not consolidated","Partner boundary","Return-path uncertainty"]),
        "wm_f2": flow_diagram("wm-f2-ticket-verification","WM-F2 FIFA Ticket-Verification Sequence",["Intent / trigger","Select tournament","Read visible card","Status sufficient?","Review destination","Continue or stay","Return and verify"],["Loop: status insufficient","Branch: act now or wait","Breakdown: destination uncertainty"], "#1D70A2"),
        "wm_c1": flow_diagram("wm-c1-beginner-learning","WM-C1 Chess Beginner-Learning Flow",["Beginner","Homepage","Play","Puzzles","Lessons","Learn-to-Play path","Account / access boundary"],["Feature choice","Learning-path selection","Promotions compete","Loss of task priority"], "#4F7F35"),
        "wm_c2": flow_diagram("wm-c2-analysis-entry","WM-C2 Chess Analysis-Entry Sequence",["Trigger after game","Open analysis","Choose setup path","Import or position","Start Analysis","Interpret next action","Reveal advanced detail"],["Many setup paths","No beginner explanation at entry","Decision priority unclear"], "#4F7F35"),
        "wm_f_artifact": flow_diagram("wm-f-artifacts","WM-F3 FIFA Artifact Model",["Tournament card | identity + CTA","Status summary | state + timestamp","Partner preview | name + domain","Saved plan | preferences + alert state"],["User action: compare or save","Downstream decision: act, wait, or stay","Breakdown: missing/stale governed data"], "#1D70A2"),
        "wm_c_artifact": flow_diagram("wm-c-artifacts","WM-C3 Chess Artifact Model",["Board position | pieces + turn","Review card | main mistake + move","Explanation | plain-language reason","Learning link | lesson or puzzle"],["User action: inspect or try move","Downstream decision: practice or reveal depth","Breakdown: entry state does not prove review output"], "#4F7F35"),
        "affinity": affinity,
        "affinity_fifa": affinity_product_view(affinity, "affinity-fifa", "FIFA Evidence-Derived Affinity View", list(range(5))),
        "affinity_chess": affinity_product_view(affinity, "affinity-chess", "Chess Evidence-Derived Affinity View", list(range(5, 11))),
        "fifa_uc": uml_use_case_diagram("fifa-system-use-cases", "FIFA Ticket Confidence Layer - UML Use-Case Diagram", "FIFA Ticket Confidence Layer", "Ticket Buyer", ["Official Ticketing Partner", "Notification Service"], ["Select tournament/team", "Inspect status/options", "Compare sale/resale/hospitality", "Save plan + subscribe", "Review partner; continue/stay", "Return to saved plan"]),
        "chess_uc": uml_use_case_diagram("chess-system-use-cases", "Chess Guided Beginner Review - UML Use-Case Diagram", "Chess Guided Beginner Review", "Beginner / Returning Player", ["Analysis Engine", "Learning Catalog"], ["Open game/position", "Start Beginner Review", "Understand mistake + move", "Try recommended move", "Reveal advanced analysis", "Continue lesson/puzzle"], "#4F7F35"),
    } | {
        key: flow_diagram(key, title, nodes, breaks, accent)
        for key, title, nodes, breaks, accent in [
            ("f-a1-flow","F-A1 Status Dashboard Flow",["1 Choose tournament","2 Read current state","3 Compare options","4 Inspect summary","5 Act or stay"],["Branch at 2: stale status -> safe retry","Branch at 3: no matching category -> save plan"],"#1D70A2"),
            ("f-a2-flow","F-A2 Guided Ticket Concierge Flow",["1 State intent","2 Answer preferences","3 Check eligibility","4 Explain official path","5 Confirm handoff"],["Branch at 2: incomplete answer -> explain field","Branch at 3: unsupported market -> stay"],"#1D70A2"),
            ("f-a3-flow","F-A3 Alert-First Ticket Planner Flow",["1 Follow context","2 Save requirements","3 Wait for change","4 Receive alert","5 Open destination"],["Branch at 2: decline consent -> local plan only","Loop at 4: irrelevant alert -> adjust preferences"],"#1D70A2"),
            ("c-a1-flow","C-A1 Beginner Analysis Preset Flow",["1 Open position","2 Select Beginner Review","3 See main mistake","4 Try best move","5 Open next lesson"],["Branch at 1: no game -> import/setup","Branch at 3: request depth -> advanced disclosure"],"#4F7F35"),
            ("c-a2-flow","C-A2 Conversational Coach Flow",["1 Open review","2 State intention","3 Coach clarifies","4 Explain position","5 Link practice"],["Branch at 2: skip question -> generic explanation","Branch at 3: ambiguous intent -> choose from options"],"#4F7F35"),
            ("c-a3-flow","C-A3 Visual Game Story Flow",["1 Open review","2 View chapters","3 Inspect turning point","4 Replay chance","5 Choose next step"],["Branch at 2: long game -> collapse chapters","Loop at 3: multiple turning points -> compare"],"#4F7F35"),
        ]
    }
    lowfi_specs = {
        "F-A1": [("Dashboard", ["Tournament selector", "Status cards + timestamp", "Compare options table", "Official action"]), ("Status detail", ["Evidence/provenance", "Availability summary", "Partner preview"])],
        "F-A2": [("Concierge", ["Goal + tournament", "Party/eligibility questions", "Back + progress"]), ("Recommendation", ["Plain-language path", "Why this route", "Continue or stay"])],
        "F-A3": [("Planner", ["Follow team/tournament", "Saved requirements", "Alert channel + consent"]), ("Alert detail", ["Verified change", "Timestamp", "Official destination"])],
        "C-A1": [("Beginner Review", ["Main mistake", "Recommended move", "Plain-language reason", "Try move"]), ("More detail", ["Advanced analysis disclosure", "Lesson/puzzle link"])],
        "C-A2": [("Coach", ["What were you trying?", "Response choices", "Critical position"]), ("Explanation", ["Coach response", "Try move", "Related practice"])],
        "C-A3": [("Game Story", ["Chapter timeline", "Turning point", "Board snapshot"]), ("Chapter detail", ["What changed", "Replay chance", "Next learning step"])],
    }
    for concept_id, panels in lowfi_specs.items():
        diagrams[f"{concept_id.lower()}-lowfi"] = lowfi_screen_map(
            f"{concept_id.lower()}-lowfi",
            f"{concept_id} Conceptual Low-Fidelity Screen Map",
            panels,
            "#1D70A2" if concept_id.startswith("F") else "#4F7F35",
        )
    return diagrams


def read_evidence():
    with (ROOT / "evidence-index.csv").open(encoding="utf-8-sig", newline="") as stream:
        return list(csv.DictReader(stream))


def read_capture_manifest():
    with (CAP / "capture-manifest.csv").open(encoding="utf-8-sig", newline="") as stream:
        return {row["filename"]: row for row in csv.DictReader(stream)}


def evidence_source(row: dict, manifest: dict) -> str:
    record = manifest.get(Path(row["local_path"]).name, {})
    parts = [row["local_path"].replace("\\", "/")]
    if record.get("source_url"):
        parts.append(f"URL: {record['source_url']}")
    if record.get("captured_at_local"):
        parts.append(f"captured: {record['captured_at_local']}")
    return "; ".join(parts)


def build_user_research(diagrams, evidence):
    doc = base_document("User Research")
    add_cover(doc, "Group10-PA2 User Research", "Source analysis and structured try-it-yourself artifact walkthrough")
    add_toc_placeholder(doc)
    manifest = read_capture_manifest()
    by_id = {row["figure_id"]: row for row in evidence}

    doc.add_heading("1. Scope and evidence boundary", level=1)
    callout(doc, "Observed evidence", "The approved corpus contains 105 local website captures (54 FIFA and 51 Chess.com). Captures prove interface states only. No participant behavior, quotation, preference, completion time, or success rate is inferred.")
    doc.add_paragraph("The study continues the PA1 product pair and uses two permitted methods: source analysis and a structured try-it-yourself artifact walkthrough (PA2-LKDuy-2026-Public.pdf, p. 2). No external participant session is represented. User experience is defined relative to specified users, goals, and context; those contextual fields are therefore stated explicitly and kept provisional where direct research is absent (LN01 - Introduction - v2, pp. 20, 24).")

    doc.add_heading("2. Research questions", level=1)
    add_bullets(doc, [
        "RQ-F1/F2: What visible structures support fixture checking and official information lookup?",
        "RQ-F3/F4/F5: What ticket status, destination boundary, and future-action support is visible before action?",
        "RQ-C1/C2: What choices precede play, and how is an immediate puzzle action signposted?",
        "RQ-C3/C4/C5: How is a learning path continued, what is visible at analysis entry, and what evidence could support a comprehensible review narrative?",
        "Across products: which goals, contexts, capabilities, artifacts, and breakdowns can be documented without behavioral inference?",
    ])

    doc.add_heading("3. Methods", level=1)
    add_table(doc, ["Method", "Material", "Procedure", "Validity boundary"], [
        ["Source analysis", "PA1 ProductResearch, PotentialSolutions, PeerReview, WeeklyReport; PA1/PA2 briefs; LN01-LN04.", "Preserve canonical IDs; compare claims with the current approved captures; cite report and page.", "Prior PA1 statements remain documentary evidence, not new PA2 participant evidence."],
        ["Structured artifact walkthrough", "evidence-index.csv, capture-manifest.csv, approved images, evidence-validation.md.", "Trace each focal task from a named start state to a visible decision/end state; record controls, branches, missing information, and artifacts.", "Researcher inspection of artifacts only; screenshots do not establish what a user did."],
        ["Evidence-gap audit", "Workspace files and supplied attachments.", "Search for session notes, recordings, surveys, brainstorming records, feedback, meetings, and ownership.", "Absence is reported as a blocker; it is never filled with synthetic data."],
    ], [3.0, 4.4, 5.1, 3.3])
    doc.add_paragraph("Task analysis documents goals, preconditions, subtasks, exceptions, context, frequency assumptions, and other involved parties (LN04 - Task Analysis, pp. 6, 39-40, 48). Walk-through is an accepted refinement technique (LN04, p. 44).")

    doc.add_heading("4. Walkthrough-session protocol and verified record", level=1)
    callout(doc, "Member-specific limitation", "No verified PA2 record names the researcher/member who performed a walkthrough, its duration, or a contemporaneous recording method. The table below is the reproducible artifact protocol used for this report, not a reconstructed human-session claim.")
    add_table(doc, ["Protocol field", "Verified value or limitation"], [
        ["Researcher/member", "Not recorded in a valid PA2 walkthrough log."],
        ["Date", "Capture-manifest records approved captures on 30 July 2026; this is capture metadata, not a participant-session date."],
        ["Duration", "Not recorded; no duration is reported."],
        ["Device", "Manifest viewport labels and pixel dimensions are used only for the associated capture (desktop or mobile)."],
        ["Task", "FR-T1-FR-T4 and CR-T1-CR-T4, defined below."],
        ["Start/end state", "Named visible page/state from evidence-index.csv; no unobserved transition is asserted."],
        ["Success condition", "Information/action that must be visible at the decision point; this is a task criterion, not a measured success."],
        ["Recording", "Approved PNG capture, evidence-index row, manifest metadata, and traceability-matrix row."],
        ["Observation wording", "Use 'researcher observed in the artifact' for pixels/controls; never 'the user did'."],
    ], [4.2, 11.4])

    doc.add_heading("5. User and task characteristics", level=1)
    task_rows = [
        ["FR-T1","Check today's fixtures/results","Short mobile or desktop check","Today/date anchor; visible match state","F2-E03, E04, E17","F-D5"],
        ["FR-T2","Search official World Cup information","Research or tournament following","Find an official, relevant result","F2-E05, E06, E07, E08","F-D5"],
        ["FR-T3","Verify ticket availability and next official action","Trust-sensitive planning","Understand status before leaving FIFA.com","F2-E09, E18","F-D4"],
        ["FR-T4","Move to FIFA+ or sibling property with orientation","Watch or ecosystem transition","Understand destination and return path","F2-E10-E16, E20","F-D1, F-D2, F-D3"],
        ["CR-T1","Start a game or bot session","Short or focused play","Reach an appropriate start state","C2-E02-E04, E13, E16","C-D1"],
        ["CR-T2","Solve a puzzle","Practice session","Identify the immediate puzzle action","C2-E05, E06, E14","C-D1"],
        ["CR-T3","Find beginner lesson and continue path","Home/school learning","Select and continue one coherent path","C2-E07-E09, E12, E15","C-D1"],
        ["CR-T4","Open analysis and identify next action","Post-game reflection","Know what to do before advanced detail","C2-E10","C-D2"],
    ]
    add_table(doc, ["Task","Goal","Context","Success condition","Evidence","PA1"],task_rows,[1.2,3.1,3.0,3.7,2.6,1.3])
    add_table(doc,["Persona","Domain experience","Application experience","Capabilities","Constraints / risk"],[
        ["F-P1 Casual football fan","Basic to moderate","Occasional FIFA.com","Recognizes teams, dates, scores","Interrupted attention; mobile scanning"],
        ["F-P2 Tournament follower","Moderate to high","Repeat tournament browsing","Compares fixtures, rankings, stories","Many parallel information needs"],
        ["F-P3 Trust-sensitive planner","Moderate","Cross-property research","Checks official source and destination","Ticket risk; identity/payment boundary"],
        ["C-P1 Beginner learner","Low","New or early-stage","Can follow clear guided steps","Low chess vocabulary; choice overload"],
        ["C-P2 Competitive player","High","Frequent play","Fast board manipulation","Time pressure; motor accuracy"],
        ["C-P3 Returning player","Moderate","Intermittent","Recognizes board and common tasks","Forgets feature location and analysis meaning"],
    ],[2.6,2.6,2.8,3.5,4.1])

    doc.add_heading("6. Focal task findings", level=1)
    focal_figs = ["F2-E03", "F2-E09", "F2-E11", "F2-E18", "C2-E02", "C2-E07", "C2-E08", "C2-E10"]
    for index, eid in enumerate(focal_figs, 1):
        row = by_id[eid]
        full_path = ROOT / row["local_path"]
        crop = evidence_crop(eid, full_path)
        caption = (
            f"Observed state: {row['visible_page_or_state']} "
            f"Supported claim: {row['supported_claims']} "
            f"Limitation: {row['unsupported_claims']}"
        )
        add_figure(doc, crop, f"UR-{index:02}", caption, evidence_source(row, manifest), f"{row['related_persona']}; {row['related_task']}")
        if eid == "F2-E09":
            doc.add_paragraph("Researcher observed a ticket landing decision surface organized by tournament cards. The captured state supports a need to compare visible options; it does not prove universal absence across every FIFA ticket channel.")
        elif eid == "C2-E10":
            doc.add_paragraph("Researcher observed multiple analysis-entry and setup paths. The artifact does not show a completed Game Review, engine lines, move classifications, or demonstrated user overload.")
        else:
            doc.add_paragraph("Researcher observed the listed controls and content hierarchy in this artifact. Any expected redesign benefit remains a PA3 hypothesis.")

    doc.add_heading("7. Provisional role-based proto-personas", level=1)
    callout(doc, "Status", "These are role-based proto-personas derived from PA1 and artifact analysis. They have no invented names, ages, quotations, or direct-research claims.")
    personas = [
        ("F-P1","Casual fixture checker","Phone during a short, interruptible check","Check today's fixtures or results","Recognizes dates, teams, score states, and common navigation labels","Low attention; long vertical surfaces","Match rows and date/status labels","Scenario: opens Match Centre, anchors on the current date, scans a relevant match state, and exits once the result is visible.","Visible date and match state without claiming measured speed.","F2-E03/E17; Group10-PA1-ProductResearch, pp. 3, 5"),
        ("F-P2","Tournament information follower","Laptop during a research or follow-up session","Compare official tournament, ranking, and article information","Moderate football vocabulary; repeat web browsing","Mixed result types; modal or article interruption","Search cards, tournament identity, ranking table","Scenario: searches an official tournament topic, distinguishes result types, opens one source, then relocates the tournament context.","Official source and context remain identifiable.","F2-E05-E08; Group10-PA1-ProductResearch, pp. 3, 6-7"),
        ("F-P3","Trust-sensitive ticket planner","Laptop plus phone across repeated planning checks","Determine current visible option and official destination before leaving FIFA.com","Moderate domain knowledge; checks source identity and action labels","Money/identity boundary; incomplete visible status","Tournament card, status summary, partner preview, saved plan","Scenario: selects a tournament, compares sale/hospitality/waiting labels, reviews the named destination, and chooses act, wait, or stay.","Decision can be made without inventing unavailable seat or resale data.","F2-E09-E13/E18; Group10-PA1-PotentialSolutions, pp. 6, 8"),
        ("C-P1","Beginner learner","Home or school on laptop/phone","Choose one comprehensible learning or review action","General web navigation; limited chess/application vocabulary","Many first-level routes; analysis entry lacks completed review explanation","Board, lesson card, review card, explanation","Scenario: opens a completed game or position, selects Beginner Review, reads one critical point, tries the move, and continues to one relevant practice item.","One next action is understandable before advanced depth.","C2-E01/E07/E08/E10/E12/E15; Group10-PA1-ProductResearch, pp. 3, 10-12"),
        ("C-P2","Focused competitive player","Desktop during time-sensitive play","Start an appropriate game efficiently","High chess and application experience; fast board manipulation","Time pressure and motor accuracy","Board, clock, time control, start action","Scenario: opens Play, confirms time control, starts the intended game route, and retains a stable board focus.","Chosen start state is visible and reversible before play.","C2-E02/E13; Group10-PA1-ProductResearch, pp. 3, 10-11"),
        ("C-P3","Returning learning player","Intermittent laptop/phone sessions","Resume a coherent learning path and regain analysis orientation","Moderate chess knowledge; incomplete memory of feature locations","Reorientation after gaps; advanced terminology","Lesson path, board position, analysis entry, learning link","Scenario: returns after a gap, recognizes a prior learning artifact, opens a review, and selects one relevant lesson or puzzle.","Continuation is identifiable without removing optional expert controls.","C2-E04/E08-E10/E16; Group10-PA1-ProductResearch, pp. 3, 12"),
    ]
    for pid, role, environment, goal, capabilities, constraints, artifacts, scenario, success, basis in personas:
        doc.add_heading(f"{pid} - {role} (provisional)", level=2)
        add_table(doc,["Attribute","Provisional profile"],[
            ["Role", role], ["Environment", environment], ["Goals and tasks", goal], ["Capabilities", capabilities],
            ["Domain/application experience", "As stated in Section 5; provisional and not measured."], ["Constraints", constraints],
            ["Artifacts", artifacts], ["Distinct scenario", scenario], ["Success criteria", success], ["Evidence basis", basis],
        ],[4.2,11.4])

    doc.add_heading("8. Work models", level=1)
    doc.add_paragraph("The models separate people/places/artifacts and coordination (flow), intent-trigger-steps-branches-loops-breakdowns (sequence), and artifact-information-action-downstream decision-breakdown (artifact). They are provisional because no direct observation session exists.")
    for fid,key,caption in [
        ("WM-F1","wm_f1","FIFA flow model: people/groups, virtual places, artifacts, directed coordination, and ticket-planning breakdowns."),
        ("WM-F2","wm_f2","FIFA sequence model: intent, trigger, steps, status-sufficient loop, act/wait branch, hesitation, and destination breakdown."),
        ("WM-F3","wm_f_artifact","FIFA artifact model: ticket artifacts, information carried, user action, downstream decision, and missing/stale-data breakdown."),
        ("WM-C1","wm_c1","Chess flow model: beginner/returning roles, home/play/puzzles/lessons places, artifacts, coordination, and priority breakdowns."),
        ("WM-C2","wm_c2","Chess sequence model: post-game trigger, setup branches, interpretation, optional detail, hesitation, and entry-state limitation."),
        ("WM-C3","wm_c_artifact","Chess artifact model: board/review/explanation/learning artifacts, information, action, downstream practice decision, and evidence limits."),
    ]:
        add_figure(doc, diagrams[key], fid, caption, f"generated-diagrams/{diagrams[key].name}", "Product-specific provisional roles and focal tasks")

    doc.add_heading("9. Insights and limitations", level=1)
    add_bullets(doc, ["Both products expose recognizable domain artifacts, but the decision burden differs: FIFA asks users to verify information and boundaries; Chess.com asks users to prioritize actions and interpret depth.","Mobile layouts preserve major tasks through stacking, while increasing vertical continuation and reducing simultaneous overview.","A positive baseline exists in Chess.com's Learn-to-Play path: a visible Next Lesson action demonstrates how one clear continuation can coexist with a rich product.","The selected problems are TP-FIFA ticket confidence and TP-CHESS beginner learning/review priority."])
    add_bullets(doc, ["No external participant behavior, interview, survey, quotation, observation transcript, or measured outcome exists.","No verified member/date/duration record exists for a human walkthrough.","Screenshots are time-bound states and cannot establish universal product behavior.","The Chess analysis capture is entry-only; it does not validate completed Game Review overload.","Role-based proto-personas and work models require future validation with representative users."])

    doc.add_heading("10. Traceability", level=1)
    add_table(doc,["Problem","Personas","Tasks","Evidence","PA1 continuity","PA3 direction"],[
        ["TP-FIFA","F-P3","FR-T3, FR-T4","F2-E09, E10, E11, E18","F-D2/F-D4; PA1 PotentialSolutions pp. 6, 8","F-A1/F-A2/F-A3; F-UC01-F-UC06"],
        ["TP-CHESS","C-P1, C-P3","CR-T3, CR-T4","C2-E01, E07, E08, E10, E12, E15","C-D1/C-D2; PA1 PotentialSolutions pp. 11-12","C-A1/C-A2/C-A3; C-UC01-C-UC06"],
    ],[2.2,2.4,2.2,3.6,2.5,3.1])

    doc.add_heading("11. References", level=1)
    add_bullets(doc, [
        "Group10-PA1-ProductResearch.pdf, pp. 2-12.",
        "Group10-PA1-PotentialSolutions.pdf, pp. 2, 6, 8, 11-12.",
        "PA2-LKDuy-2026-Public.pdf, pp. 1-3.",
        "LN01 - Introduction - v2.pdf, pp. 20, 22, 24.",
        "LN02 - Fundamental Concepts - Usability Dimensions_2.pdf, pp. 62-66, 116-118.",
        "LN03 - UI Design Process.pdf, pp. 16, 23-24, 28, 30, 36.",
        "LN04 - Task Analysis.pdf, pp. 3, 6-7, 39-40, 44, 47-49.",
    ])

    doc.add_heading("12. Evidence appendix", level=1)
    doc.add_paragraph("The exact approved full-height images used for the main focal figures follow. Main-body figures are pixel-preserving crops for legibility; these appendix images retain the complete capture.")
    for index, eid in enumerate(focal_figs, 1):
        row = by_id[eid]
        add_figure(
            doc,
            ROOT / row["local_path"],
            f"UR-A{index:02}",
            f"Full approved image for {eid}. Observed state: {row['visible_page_or_state']} Supported claim: {row['supported_claims']} Limitation: {row['unsupported_claims']}",
            evidence_source(row, manifest),
            f"{row['related_persona']}; {row['related_task']}",
            width_cm=15.2,
            max_height_cm=18.0,
        )
    add_table(doc,["ID","State","Supported claim","Unsupported claim"],[[r["figure_id"],r["visible_page_or_state"],r["supported_claims"],r["unsupported_claims"]] for r in evidence],[1.4,5.1,5.1,4.0])
    add_header_footer(doc,"User Research")
    path=SOURCE/"Group10-PA2-UserResearch.docx"; doc.save(path); return path


def build_user_analysis(diagrams):
    doc=base_document("User Analysis")
    add_cover(doc,"Group10-PA2 User Analysis","Evidence-derived prework for group brainstorming and prioritization","INCOMPLETE - verified group brainstorming evidence is absent")
    add_toc_placeholder(doc)
    doc.add_heading("1. Requirement status and evidence boundary",level=1)
    callout(doc,"Requirement 2 blocker","PA2 requires real group brainstorming. No verified date, duration, participant list, tool/export, raw-note generation record, clustering record, votes, scoring sheet, consensus record, or board screenshot/export was found. This report therefore remains incomplete.")
    doc.add_paragraph("The material below is analyst prework derived from validated screenshots, PA1 findings, and the User Research task analysis. It must not be presented as a group meeting or group consensus.")

    doc.add_heading("2. Required real-session protocol",level=1)
    add_table(doc,["Required field","Current status","Completion evidence needed"],[
        ["Date and duration","Missing","Dated session record with start/end or duration."],
        ["Participants","Missing","Verified attendee list; do not infer from roster."],
        ["Tool and board","Missing","Real board screenshot/export from the session."],
        ["Raw note generation","Missing","Unclustered notes with authorship or session provenance."],
        ["Clustering","Missing","Record of how notes moved into FIFA and Chess themes."],
        ["Voting/scoring","Missing","Raw individual votes or an explicit consensus record."],
        ["Decisions and ownership","Missing","Decision log naming owner and affected requirement."],
    ],[4.2,3.0,8.4])
    doc.add_paragraph("When supplied, the session should diverge before converging, compare alternatives, and document review/evaluation as part of an iterative, user-centered design process (LN03 - UI Design Process, pp. 19, 23-24, 30).")

    doc.add_heading("3. Analyst prework legend and note inventory",level=1)
    add_bullets(doc,["[OBS] = visible screenshot state.","[PA1] = inherited PA1 report finding with report/page reference.","[IDEA] = design inference or proposal, never observed behavior.","No note label establishes frequency, severity, vote, or user preference."])
    add_table(doc,["Cluster","Representative notes","Evidence basis"],[
        ["FIFA-A1/A2","Task entry, date/filter scanning, mixed result types, interruption.","F2-E01-E05; PA1 ProductResearch pp. 4-6"],
        ["FIFA-A3","Ticket cards show tournament-specific actions but not a consolidated comparison surface in the capture.","F2-E09/E18; F-D4, PA1 PotentialSolutions p. 8"],
        ["FIFA-A4","FIFA+ entry, DAZN brand/account context, consent, and sibling-property vocabulary change.","F2-E10-E16; F-D2, PA1 PotentialSolutions p. 6"],
        ["FIFA-A5","Long mobile match, ticket, and article surfaces.","F2-E17-E20"],
        ["CHESS-A1-A4","Many first-level task/setup routes; puzzle and lesson continuation artifacts.","C2-E01-E09/E12-E16; C-D1, PA1 PotentialSolutions p. 11"],
        ["CHESS-A5","Analysis entry shows multiple setup/import paths and Start Analysis; no completed review is visible.","C2-E10; C-D2 continuity only, PA1 PotentialSolutions p. 12"],
        ["CHESS-A6","Promotional/long-page content near learning; ambiguous ad capture is restricted.","C2-E07/E11/E15; evidence-validation.md"],
    ],[3.0,8.0,4.6])

    doc.add_heading("4. Readable affinity views",level=1)
    add_figure(doc,diagrams["affinity_fifa"],"UA-F01","FIFA evidence-derived affinity prework. This is not a real brainstorming-board export and does not satisfy Requirement 2.",f"generated-diagrams/{diagrams['affinity_fifa'].name}","FIFA-A1-FIFA-A5; TP-FIFA",width_cm=16.0)
    add_figure(doc,diagrams["affinity_chess"],"UA-C01","Chess evidence-derived affinity prework. This is not a real brainstorming-board export and does not satisfy Requirement 2.",f"generated-diagrams/{diagrams['affinity_chess'].name}","CHESS-A1-CHESS-A6; TP-CHESS",width_cm=16.0)

    doc.add_heading("5. Candidate ideas retained for a future group session",level=1)
    add_table(doc,["Problem family","Concept candidates","Evidence/idea status"],[
        ["Ticket decision surface","F-A1 Status Dashboard; F-A2 Guided Ticket Concierge; F-A3 Alert-First Ticket Planner.","Proposal ideas linked to F2-E09/E18 and F-D4."],
        ["Partner handoff","Partner/domain preview; Continue or Stay; return-path preservation.","Proposal ideas linked to F2-E10/E11 and F-D2."],
        ["Beginner review entry","C-A1 Beginner Analysis Preset; C-A2 Conversational Coach; C-A3 Visual Game Story.","Proposal ideas linked narrowly to C2-E10 entry state plus C-D2."],
        ["Learning continuation","One recommended next action; relevant lesson/puzzle; advanced detail on demand.","Proposal ideas linked to C2-E07/E08/E15 and C-D1."],
    ],[4.0,7.5,4.1])

    doc.add_heading("6. Provisional tough problems",level=1)
    callout(doc,"TP-FIFA (provisional)","On the captured FIFA ticket landing decision surface, tournament cards expose different visible actions, while a consolidated comparison of sale, resale, hospitality, waiting, availability, freshness, and partner destination is not visible. This is not a universal claim about every FIFA ticket channel.",FIFA_LIGHT)
    callout(doc,"HMW-F","How might a responsive web redesign help a ticket planner compare governed visible states and review the official destination before choosing act, wait, continue, or stay?",FIFA_LIGHT)
    callout(doc,"TP-CHESS (provisional)","At the captured Chess.com analysis-entry state, a beginner or returning player sees multiple setup/import paths and Start Analysis, but the screenshot does not show a completed review or a plain-language next-learning explanation. The problem is entry orientation, not proven completed-review overload.",CHESS_LIGHT)
    callout(doc,"HMW-C","How might a responsive web redesign make one beginner review path explicit at entry, then connect a critical position to optional advanced detail and relevant practice?",CHESS_LIGHT)

    doc.add_heading("7. Scoring formula and missing consensus record",level=1)
    doc.add_paragraph("Proposed group formula: Total = Severity + Frequency + Risk + Reach + Evidence + PA3 relevance, with each dimension scored 1-5. A valid total requires the raw votes or a documented consensus for every dimension.")
    add_table(doc,["Dimension","1 anchor","3 anchor","5 anchor","Required rationale/owner"],[
        ["Severity","Minor inconvenience","Meaningful friction","Blocks confidence/comprehension","Named voter or consensus record"],
        ["Frequency","Rare assumption","Occasional assumption","Repeated/frequent evidence","Direct evidence or explicit assumption"],
        ["Risk","Negligible","Moderate","Trust/error critical","Evidence-linked rationale"],
        ["Reach","Narrow role","One role class","Multiple roles/devices","Scope rationale"],
        ["Evidence","Single ambiguous item","Multiple aligned items","Direct + PA1 triangulation","Exact evidence IDs/pages"],
        ["PA3 relevance","Weak","Useful","Two-scenario core","Concept/prototype rationale"],
    ],[2.6,2.4,2.6,2.8,5.2])
    add_table(doc,["Candidate","Severity","Frequency","Risk","Reach","Evidence","PA3","Total / owner"],[
        ["TP-FIFA","Not voted","Not voted","Not voted","Not voted","F2-E09/E18 + F-D4","Not voted","Not computed; owner unavailable"],
        ["TP-CHESS","Not voted","Not voted","Not voted","Not voted","C2-E10 + C-D2; entry-only limitation","Not voted","Not computed; owner unavailable"],
    ],[3.0,1.8,1.8,1.8,1.8,3.3,1.8,3.3])
    doc.add_paragraph("The two problems are retained provisionally to keep the Proposal and Use Case documents internally traceable. Their selection is an evidence-driven editorial decision, not a fabricated group vote.")

    doc.add_heading("8. Dimension-by-dimension rationale and ownership gap",level=1)
    add_table(doc,["Problem / dimension","Provisional rationale","Verified owner"],[
        ["TP-FIFA severity/risk","Ticket planning may cross identity/payment boundaries; PA1 documents trust sensitivity (PA1 PeerReview pp. 2-3).","None recorded"],
        ["TP-FIFA frequency/reach","Repeated checking and multi-device scope are PA1/design assumptions, not PA2 measurements.","None recorded"],
        ["TP-FIFA evidence","F2-E09/E18 show ticket landing/card states; F2-E10/E11 show destination/context changes.","None recorded"],
        ["TP-FIFA PA3 relevance","Three distinct dashboard, concierge, and alert concepts can test the same decision surface.","None recorded"],
        ["TP-CHESS severity/risk","C2-E10 supports entry-orientation uncertainty only; no completed-review error outcome is observed.","None recorded"],
        ["TP-CHESS frequency/reach","Beginner/returning scope is provisional; no PA2 participant frequency exists.","None recorded"],
        ["TP-CHESS evidence","C2-E10 plus C2-E07/E08/E15 and PA1 C-D1/C-D2 establish continuity, not measured overload.","None recorded"],
        ["TP-CHESS PA3 relevance","Three preset, conversational, and narrative concepts can be compared as parallel prototypes.","None recorded"],
    ],[4.1,9.2,2.3])

    doc.add_heading("9. Provisional design requirements",level=1)
    add_table(doc,["Req.","Design requirement","Evidence / rationale"],[
        ["F-R1","Show a plain-language current ticket state before any external exit.","F2-E09, E18; F-D4"],
        ["F-R2","Compare only governed sale, resale, hospitality, waiting, and availability fields; label unavailable fields rather than inventing them.","TP-FIFA"],
        ["F-R3","Disclose partner name, destination domain, account boundary, and return path before continuation.","F2-E10-E13; F-D2"],
        ["F-R4","Encode status with text and icon as well as color; expose last-updated time.","Trust and accessibility"],
        ["C-R1","Place one explicit Beginner Review entry action before advanced setup paths.","C2-E10; TP-CHESS"],
        ["C-R2","In the proposed review output, explain one critical position, recommended move, and next learning action in plain language.","Design hypothesis from C-D2; not captured current output"],
        ["C-R3","Keep engine detail and advanced setup available through explicit progressive disclosure.","C2-E10"],
        ["C-R4","Provide keyboard and screen-reader labels for every core action; never rely only on color.","Accessibility"],
    ],[1.3,8.8,5.5])

    doc.add_heading("10. Traceability, limitations, and completion gate",level=1)
    add_bullets(doc,["Every provisional tough problem links to approved screenshot evidence and a PA1 drawback page.","Every Proposal alternative and Use Case ID maps back to one provisional tough problem.","No measured improvement, group attendance, note ownership, vote, or consensus is claimed.","Completion requires the real group board export, session metadata, raw notes, clustering record, votes/consensus, decisions, and owners."])
    doc.add_heading("Appendix A. Full 55-note analyst prework",level=1)
    sec=doc.add_section(WD_SECTION.NEW_PAGE); sec.orientation=WD_ORIENT.LANDSCAPE; sec.page_width=Cm(29.7); sec.page_height=Cm(21.0); sec.left_margin=sec.right_margin=Cm(1.2)
    add_figure(doc,diagrams["affinity"],"UA-A01","Full 55-note evidence-derived analyst prework. It is not a real group brainstorming-board screenshot/export.",f"generated-diagrams/{diagrams['affinity'].name}","All focal tasks and provisional tough problems",width_cm=23.8,max_height_cm=14.2)
    add_header_footer(doc,"User Analysis")
    path=SOURCE/"Group10-PA2-UserAnalysis.docx"; doc.save(path); return path


CONCEPTS = [
    ("F-A1","Status Dashboard","F-P3","Compare ticket states before action","Comparison-first dashboard","Tournament, sale phase, ticket type, availability summary, resale, waiting state, milestone, last update, one official action.","Comparison table and status cards","Dashboard; status detail; destination preview","State -> evidence timestamp -> official action","Visibility of system status; recognition; error prevention","Hypothesis: improves confidence and comparison speed.","Depends on governed real-time status data.","Could appear authoritative when data is stale.","Text/icon status, keyboard comparison, screen-reader timestamps.","F2-E09, F2-E18; F-D4"),
    ("F-A2","Guided Ticket Concierge","F-P3","Resolve an individual ticket path","Step-by-step wizard","Ask tournament, date/team, ticket type, party size, and eligibility; then explain next official path.","Wizard, eligibility summary, handoff confirmation","Preference steps; result; partner preview","One question -> eligibility -> destination","Progressive disclosure; feedforward; error prevention","Hypothesis: reduces interpretation burden for complex eligibility.","Requires eligibility and market rules.","Collects preference data and must minimize retention.","Explicit labels, error summary, back navigation, no time limit.","F2-E09-E11; F-D2/F-D4"),
    ("F-A3","Alert-First Ticket Planner","F-P3","Avoid repeated manual checking","Planning workspace","Follow tournament/team, save requirements, receive verified state changes, and open partner only when actionable.","Watchlist, requirement card, alert receipt","Planner; consent; verified alert; destination preview","Plan -> wait -> verified change -> action","Prospective-memory support; user control; trust","Hypothesis: replaces repeated checking with governed alerts.","Requires notification service and consent state.","Notification privacy and fatigue.","Channel choice, quiet hours, accessible alert history.","F2-E09, F2-E18; F-D4"),
    ("C-A1","Beginner Analysis Preset","C-P1; C-P3","Understand one critical review point","Simplified review surface","Show one main mistake, one best move, one plain-language explanation, and one next-learning action; advanced data on demand.","Position card, move comparison, explanation, lesson link","Preset entry; critical position; try move; next action","Primary action -> explanation -> practice -> optional advanced","Progressive disclosure; feedback; recognition","Hypothesis: makes review interpretable before engine detail.","Requires review engine and explanation rules.","Explanations must state uncertainty and avoid false authority.","Keyboard board alternative, text move notation, non-color status.","C2-E08, C2-E10; C-D1/C-D2"),
    ("C-A2","Conversational Coach","C-P1; C-P3","Connect intention to feedback","Guided Q&A","Ask what the player intended, explain the critical position, and link one lesson or puzzle.","Question card, response options, coach explanation","Intent question; critical position; resource link","Intent -> clarification -> explanation -> next task","Conversational scaffolding; reflection; contextual help","Hypothesis: increases personal relevance of review.","Requires safe prompt/explanation logic.","Conversation data and incorrect inference require controls.","Skip option, concise language, screen-reader announcements.","C2-E08, C2-E10; C-D2"),
    ("C-A3","Visual Game Story","C-P1; C-P3","Understand the game as a sequence","Chapter/timeline review","Group opening, turning point, main error, recovery chance, and next step; collapse engine detail.","Timeline chapters, board snapshot, recovery action","Story overview; chapter detail; practice","Overview -> turning point -> replay -> next step","Chunking; external cognition; progressive disclosure","Hypothesis: supports recall through narrative structure.","Requires reliable event segmentation.","Story labels can oversimplify complex games.","Linear keyboard order, text chapter summary, scalable board.","C2-E10; C-D2"),
]


def build_project_proposal(diagrams):
    doc=base_document("Project Proposal")
    add_cover(doc,"Group10-PA2 Project Proposal","Task-first redesign concepts for ticket confidence and guided beginner review")
    add_toc_placeholder(doc)
    doc.add_heading("1. Proposal summary",level=1)
    callout(doc,"Overall concept","A task-first web redesign program with two product-specific modules: Module A, FIFA Ticket Confidence Layer; Module B, Chess Guided Beginner Review.")
    doc.add_paragraph("These are redesign concepts for existing web experiences. They are not implemented or tested products. Benefits are hypotheses for PA3 prototyping and informal testing.")
    doc.add_heading("2. Evidence-based problem selection",level=1)
    add_table(doc,["Module","Tough problem","Evidence","PA1 continuity"],[
        ["A. FIFA Ticket Confidence Layer","Status, availability, resale/waiting, and destination are not consolidated before external action.","F2-E09, E10, E11, E18","F-D2, F-D4"],
        ["B. Chess Guided Beginner Review","Learning and analysis entry expose competing choices and advanced controls before a clear beginner next action.","C2-E01, E07, E08, E10, E12, E15","C-D1, C-D2"],
    ],[3.1,7.4,3.1,2.0])
    doc.add_heading("3. Target users and contexts",level=1)
    add_bullets(doc,["F-P3 plans tickets across days or devices and needs official status before identity, partner, or payment boundaries.","C-P1 learns with limited terminology; C-P3 returns after gaps and needs recognition plus a clear continuation.","Secondary personas remain in scope for compatibility, but the initial prototypes focus on these scenarios."])
    doc.add_heading("4. Concept family A - FIFA",level=1)
    for cid,name,persona,scenario,summary,interaction,artifacts,screens,hierarchy,principles,benefit,dependency,trust,a11y,evidence in CONCEPTS[:3]:
        doc.add_heading(f"{cid} {name}",level=2)
        add_table(doc,["Field","Concept definition"],[
            ["Target proto-persona / scenario",f"{persona} (provisional) - {scenario}."],["Problem addressed","TP-FIFA."],["Form factor","Responsive web redesign."],["Concept summary",summary],["Interaction model",interaction],["Artifacts / data",artifacts],["Core screens / components",screens],["Information hierarchy",hierarchy],["HCI principles",f"{principles}. See LN02 - Fundamental Concepts, pp. 66, 116; LN03 - UI Design Process, p. 36."],["Expected benefit",benefit],["Tradeoff","Adds new status, guidance, or planning structure that must not obscure direct official routes."],["Implementation dependency",dependency],["Privacy / trust",trust],["Accessibility",a11y],["Evidence links",evidence],
        ],[4.0,11.6])
        add_figure(doc,diagrams[cid.lower()+"-lowfi"],f"PP-{cid}-L",f"{name} conceptual low-fidelity screen map made only from shapes and labels.",f"generated-diagrams/{diagrams[cid.lower()+'-lowfi'].name}",f"{persona}; TP-FIFA")
        add_figure(doc,diagrams[cid.lower()+"-flow"],f"PP-{cid}-F",f"{name} flow model. Numbered steps, branches, and breakdowns are proposal hypotheses, not existing or tested behavior.",f"generated-diagrams/{diagrams[cid.lower()+'-flow'].name}",f"{persona}; TP-FIFA")
    doc.add_heading("5. Concept family B - Chess.com",level=1)
    for cid,name,persona,scenario,summary,interaction,artifacts,screens,hierarchy,principles,benefit,dependency,trust,a11y,evidence in CONCEPTS[3:]:
        doc.add_heading(f"{cid} {name}",level=2)
        add_table(doc,["Field","Concept definition"],[
            ["Target proto-persona / scenario",f"{persona} (provisional) - {scenario}."],["Problem addressed","TP-CHESS."],["Form factor","Responsive web redesign."],["Concept summary",summary],["Interaction model",interaction],["Artifacts / data",artifacts],["Core screens / components",screens],["Information hierarchy",hierarchy],["HCI principles",f"{principles}. See LN02 - Fundamental Concepts, pp. 66, 116; LN03 - UI Design Process, p. 36."],["Expected benefit",benefit],["Tradeoff","Simplification can hide nuance; advanced analysis must remain explicitly available."],["Implementation dependency",dependency],["Privacy / trust",trust],["Accessibility",a11y],["Evidence links",evidence],
        ],[4.0,11.6])
        add_figure(doc,diagrams[cid.lower()+"-lowfi"],f"PP-{cid}-L",f"{name} conceptual low-fidelity screen map made only from shapes and labels.",f"generated-diagrams/{diagrams[cid.lower()+'-lowfi'].name}",f"{persona}; TP-CHESS")
        add_figure(doc,diagrams[cid.lower()+"-flow"],f"PP-{cid}-F",f"{name} flow model. Numbered steps, branches, and breakdowns are proposal hypotheses, not existing or tested behavior.",f"generated-diagrams/{diagrams[cid.lower()+'-flow'].name}",f"{persona}; TP-CHESS")
    doc.add_heading("6. Weighted comparison",level=1)
    weights={"Problem fit":25,"Learnability":15,"Error prevention / trust":15,"Task efficiency":10,"Feasibility":10,"Cross-device fit":10,"Accessibility":10,"Evidence strength":5}
    scores={
        "F-A1":[5,4,5,5,3,4,4,5],"F-A2":[5,5,5,3,3,4,5,4],"F-A3":[4,4,4,5,2,5,4,4],
        "C-A1":[5,5,4,5,4,4,5,5],"C-A2":[4,5,4,3,2,4,4,4],"C-A3":[4,4,3,4,3,4,5,4],
    }
    rows=[]
    for cid,vals in scores.items():
        total=sum(v*w for v,w in zip(vals,weights.values()))/5
        rows.append([cid,*map(str,vals),f"{total:.0f}/100"])
    add_table(doc,["Concept",*weights.keys(),"Weighted"],rows,[1.2,1.55,1.45,1.6,1.35,1.2,1.35,1.35,1.35,1.4])
    doc.add_paragraph("Weights total 100%. Scoring uses 1 (weak) to 5 (strong). Weighted total = Σ(score × criterion weight ÷ 5). These are design-team analytical comparison scores, not usability measurements or brainstorming votes.")
    add_table(doc,["Concept","Short score rationale"],[
        ["F-A1","Strongest problem/evidence fit and efficient comparison; feasibility is lower because governed, fresh status data is required."],
        ["F-A2","High learnability, error prevention, and accessibility through guided steps; slower for repeat users and dependent on eligibility rules."],
        ["F-A3","Strong repeated-checking and cross-device fit; feasibility is constrained by notifications, consent, and alert governance."],
        ["C-A1","Strongest entry-orientation, learnability, efficiency, accessibility, and evidence continuity; depends on reliable explanation rules."],
        ["C-A2","High reflective scaffolding; lower feasibility/efficiency because conversation logic may infer intent incorrectly."],
        ["C-A3","Strong accessibility and narrative chunking; segmentation may oversimplify games and has moderate implementation complexity."],
    ],[2.0,13.6])
    doc.add_heading("7. Recommended direction",level=1)
    callout(doc,"FIFA recommendation","Prototype F-A1 Status Dashboard first, while testing F-A2 and F-A3 as meaningfully different wizard and planning alternatives.",FIFA_LIGHT)
    callout(doc,"Chess recommendation","Prototype C-A1 Beginner Analysis Preset first, while retaining C-A2 conversational and C-A3 narrative alternatives.",CHESS_LIGHT)
    doc.add_heading("8. Scope and non-scope",level=1)
    add_table(doc,["In scope","Out of scope"],[
        ["Responsive web concept flows; task entry; information hierarchy; partner disclosure; beginner review; accessible interaction.","Implementation, live ticket inventory, payment processing, identity migration, engine accuracy validation, production integration."],
        ["Three alternatives per tough problem; PA3 hypothesis and test criteria.","Claimed measured improvement, fabricated participant results, premove or Focus Mode redesign as primary scenario."],
    ],[7.8,7.8])
    doc.add_heading("9. Design principles, risks, and success criteria",level=1)
    add_bullets(doc,["State before action: disclose current status and destination before commitment.","One clear next action: present the beginner's primary step before advanced depth.","Progressive disclosure without removal: advanced and direct routes remain available.","Trust through provenance: official destination, timestamp, and uncertainty are visible.","Accessible redundancy: text, icon, structure, and programmatic labels support color."])
    add_table(doc,["Design goal for PA3","Target (not a tested result)"],[
        ["Ticket state recognition","A first-time prototype participant identifies the current state without opening another property."],
        ["Destination feedforward","The partner name and destination domain are visible before exit."],
        ["Beginner review priority","The main review action appears in the initial viewport and keyboard order."],
        ["Advanced access","Advanced analysis remains available through a labeled disclosure control."],
        ["Accessibility","All core actions have keyboard and screen-reader definitions; color is never the only status channel."],
    ],[5.0,10.6])
    doc.add_heading("10. Traceability and preparation for PA3",level=1)
    add_bullets(doc,["TP-FIFA -> F-A1/F-A2/F-A3 -> FIFA use cases F-UC01-F-UC06.","TP-CHESS -> C-A1/C-A2/C-A3 -> Chess use cases C-UC01-C-UC06.","Create two PA3 scenarios, each with three parallel paper-prototype alternatives; test comprehension, next-action selection, error recognition, and return orientation.","Do not interpret expected benefits as measured outcomes."])
    add_header_footer(doc,"Project Proposal")
    path=SOURCE/"Group10-PA2-ProjectProposal.docx"; doc.save(path); return path


FIFA_UCS = [
    {
        "id":"F-UC01","name":"Select tournament or team","actor":"Ticket Buyer","support":"FIFA tournament catalog","persona":"F-P3","concept":"F-A1/F-A2/F-A3","evidence":"F2-E09, F2-E18","trigger":"The buyer opens the ticket-planning module with a tournament or team in mind.","pre":"The catalog has a governed identifier and display label; the module can disclose stale/unavailable catalog data.","frequency":"At the start of a planning session; assumption to validate.","priority":"High","stakeholders":"Buyer wants correct context; FIFA content owner wants canonical tournament/team identity.","main":["Open the tournament/team selector.","Search or browse the governed catalog.","Inspect the matching label, competition, and season.","Choose one tournament or team.","Confirm the selected context shown in the planning header.","Continue to ticket-status inspection."],"alts":["A1 at step 2: No result matches; show spelling guidance and browse categories without creating a fake item.","A2 at step 4: Multiple seasons share a label; require an explicit season choice.","E1 at step 2: Catalog is unavailable; retain the prior context and offer retry."],"data":"Tournament/team ID, canonical label, competition, season, locale.","rules":"Only governed catalog entries may be selected; ambiguous names require disambiguation.","output":"Selected ticket context and accessible confirmation.","special":"Search results expose type and season in text, not color alone.","open":"Who owns catalog freshness across FIFA properties?"},
    {
        "id":"F-UC02","name":"Inspect official ticket status and available options","actor":"Ticket Buyer","support":"Ticket Status Service","persona":"F-P3","concept":"F-A1/F-A2","evidence":"F2-E09, F2-E18; F-D4","trigger":"A ticket context is selected and the buyer asks what is currently visible/available.","pre":"F-UC01 is complete; each returned field includes provenance and freshness or an unavailable marker.","frequency":"Repeated during planning; frequency is an assumption, not measured.","priority":"Critical","stakeholders":"Buyer needs a bounded current state; data owner must avoid false precision.","main":["Request status for the selected context.","Read the plain-language overall state and last-updated time.","Review which official option categories have governed data.","Open the provenance explanation for a category.","Inspect the single next official action associated with that state.","Choose to compare options, save the plan, or stop."],"alts":["A1 at step 2: Status is stale; label it stale and suppress action language that implies currency.","A2 at step 3: A category has no governed data; show 'not provided' rather than 'unavailable'.","E1 at step 1: Status service fails; show no cached value unless its timestamp and source remain visible."],"data":"Context ID, state label, option-category presence, source, timestamp, action eligibility.","rules":"Unknown, not provided, closed, and sold out are distinct states.","output":"Readable status summary with provenance and one safe next action.","special":"Timestamp is localized and screen-reader text conveys the same state as icons.","open":"What official system is authoritative for each field?"},
    {
        "id":"F-UC03","name":"Compare sale, resale, hospitality, and waiting states","actor":"Ticket Buyer","support":"Ticket Status Service; Hospitality Catalog","persona":"F-P3","concept":"F-A1","evidence":"F2-E09; F-D4","trigger":"The buyer chooses Compare from a governed ticket-status summary.","pre":"At least one category has a governed value; unavailable fields are explicitly marked.","frequency":"Once or more per selected tournament; assumption to validate.","priority":"High","stakeholders":"Buyer wants comparable definitions; FIFA/partners need categories represented accurately.","main":["Open the option comparison for the selected context.","Read category definitions for sale, resale, hospitality, and waiting.","Compare state, eligibility, key date, and official destination by row.","Filter rows to categories relevant to the buyer.","Inspect why any row lacks data or action.","Select one row for details or return to the unfiltered comparison."],"alts":["A1 at step 2: A market does not support resale; explain the market rule instead of hiding the row.","A2 at step 4: Filtering would produce zero rows; keep filters visible and offer Clear filters.","E1 at step 3: Categories use incompatible freshness times; display each timestamp rather than a misleading shared timestamp."],"data":"Category taxonomy, state, eligibility summary, key date, destination, per-row timestamp.","rules":"Rows cannot be ranked as better without an explicit buyer criterion.","output":"Persisted comparison/filter state and selected option, if any.","special":"Table supports keyboard row navigation and a stacked mobile reading order.","open":"Which comparison fields are legally/operationally available per market?"},
    {
        "id":"F-UC04","name":"Save a ticket plan and subscribe to verified alerts","actor":"Ticket Buyer","support":"Identity Service; Notification Service","persona":"F-P3","concept":"F-A3","evidence":"F2-E09, F2-E18; F-D4","trigger":"The buyer wants to avoid manually repeating the same check.","pre":"A context exists; consent and authentication requirements are disclosed before collecting contact data.","frequency":"Occasional setup followed by event-driven alerts; assumption to validate.","priority":"High","stakeholders":"Buyer wants control and low alert noise; services need valid consent and deliverability.","main":["Open Save plan from the selected ticket context.","Choose tracked categories and optional team/date preferences.","Review which fields are stored locally or in an account.","Save the plan with a recognizable label.","Choose an alert channel and quiet-hours preference.","Confirm consent and view the verified-alert criteria.","Receive an accessible plan summary with edit/unsubscribe controls."],"alts":["A1 at step 3: Buyer declines account storage; offer a local-only plan if supported.","A2 at step 5: Buyer declines notifications; save the plan without a subscription.","E1 at step 6: Channel verification fails; keep the plan and mark alerts inactive."],"data":"Plan ID/label, context, categories, preferences, storage mode, channel, consent, quiet hours.","rules":"No marketing consent is bundled with status alerts; alerts fire only on governed changes.","output":"Saved plan, subscription state, consent receipt, and controls.","special":"Consent is reversible; notification content minimizes sensitive preference data.","open":"What retention period and channels are approved?"},
    {
        "id":"F-UC05","name":"Review partner identity and choose Continue or Stay","actor":"Ticket Buyer","support":"Official Ticketing Partner","persona":"F-P3","concept":"F-A1/F-A2/F-A3","evidence":"F2-E10, F2-E11; F-D2","trigger":"The buyer selects an action that leaves the FIFA-hosted module.","pre":"Destination is allow-listed and tied to the chosen option; the current plan/context can be preserved.","frequency":"Before every external ticket/hospitality handoff.","priority":"Critical","stakeholders":"Buyer needs informed control; FIFA and partner need a trustworthy boundary.","main":["Open the destination preview for the selected action.","Read the official partner name and destination domain.","Review whether sign-in, identity, or payment context may change.","Read what context will be preserved for return.","Choose Continue or Stay on FIFA.com.","If Continue, open the allow-listed destination and record a return token; if Stay, retain the current decision surface."],"alts":["A1 at step 2: Partner/domain verification is unavailable; disable Continue and offer retry.","A2 at step 5: Buyer chooses Stay; close the preview without external navigation.","E1 at step 6: Destination fails to open; keep the saved context and show a safe retry."],"data":"Partner ID/name, allow-listed domain, action type, account-boundary notice, return token.","rules":"No open redirect; Continue is unavailable when partner identity cannot be verified.","output":"Handoff decision, retained context, and optional return token.","special":"Focus enters the preview, remains trapped appropriately, then returns to the invoking control.","open":"Which partner/account disclosures are required for each route?"},
    {
        "id":"F-UC06","name":"Return to a saved plan after partner activity","actor":"Ticket Buyer","support":"Official Ticketing Partner; Identity Service","persona":"F-P3","concept":"F-A1/F-A2/F-A3","evidence":"F2-E11; F-D2","trigger":"The buyer follows a return link or reopens the planner after partner activity.","pre":"A valid local/account plan or time-limited return token exists; partner activity is not assumed to be available.","frequency":"After an external handoff or later planning session; assumption to validate.","priority":"High","stakeholders":"Buyer wants orientation; FIFA must not misrepresent partner-side completion.","main":["Open the return link or Saved plans list.","Validate the return token or authenticate for the stored plan.","Restore tournament/team, comparison filters, and alert state.","Label the restored data with current freshness/provenance.","Show that partner-side purchase/activity status is unknown unless explicitly returned by a governed integration.","Let the buyer refresh, edit the plan, inspect status, or remove it."],"alts":["A1 at step 2: Token expired; offer account lookup or safe plan recreation without exposing token data.","A2 at step 3: Only a partial plan is available; list omitted fields before continuing.","E1 at step 4: Refresh fails; keep the restored snapshot labeled with its prior timestamp."],"data":"Return token, plan ID, saved context/filter/alerts, current status timestamp, optional governed partner result.","rules":"Never infer purchase success from a return; expired tokens cannot be replayed.","output":"Restored plan with explicit freshness and next actions.","special":"Restoration is announced to assistive technology without moving focus unexpectedly.","open":"Can any partner return a signed completion state, and for how long?"},
]

CHESS_UCS = [
    {
        "id":"C-UC01","name":"Open a completed game or analysis position","actor":"Beginner / Returning Player","support":"Game Store; Analysis Engine","persona":"C-P1; C-P3","concept":"C-A1/C-A2/C-A3","evidence":"C2-E10","trigger":"The player chooses a completed game, paste/import, or setup route for review.","pre":"The source artifact is accessible; unsupported/invalid notation can be identified without losing input.","frequency":"Per review session; assumption to validate.","priority":"High","stakeholders":"Player wants the intended position; platform needs valid, permission-appropriate game data.","main":["Choose Completed game, Paste/import, or Set up position.","Select the game or enter the position data.","Review source, side to move, and orientation.","Correct any detected notation or setup issue.","Confirm the review artifact.","Open the review entry with Beginner Review and advanced-analysis choices."],"alts":["A1 at step 1: No completed game exists; keep import and setup routes available.","A2 at step 3: Orientation is wrong; switch sides before confirmation.","E1 at step 2: Notation is invalid; identify the first invalid segment and preserve the remaining input."],"data":"Game/position ID, notation, source, side to move, orientation, ownership/visibility.","rules":"Private games/positions follow account permissions; invalid data is never silently corrected.","output":"Validated review artifact and entry state.","special":"Board setup has a text-notation alternative and fully labeled pieces/squares.","open":"Which anonymous/import formats are supported?"},
    {
        "id":"C-UC02","name":"Start Beginner Review","actor":"Beginner / Returning Player","support":"Analysis Engine","persona":"C-P1; C-P3","concept":"C-A1/C-A2/C-A3","evidence":"C2-E10; C-D2","trigger":"A valid game/position is open and the player chooses Beginner Review.","pre":"C-UC01 is complete; the system can disclose analysis availability/entitlement.","frequency":"Per selected review artifact; assumption to validate.","priority":"Critical","stakeholders":"Player wants a low-vocabulary entry; engine owner must communicate uncertainty and availability.","main":["Choose Beginner Review from the review entry.","Read the mode summary: one critical point first, advanced detail optional.","Confirm whose perspective and skill-level language to use.","Start the requested analysis.","Wait with progress and cancellation controls.","Receive the Beginner Review overview or a clear unavailable explanation."],"alts":["A1 at step 3: Player skips skill-level choice; use neutral beginner language, not a hidden profile inference.","A2 at step 5: Player cancels; return to review entry with the artifact intact.","E1 at step 4: Analysis is unavailable or restricted; explain the condition before suggesting another route."],"data":"Artifact ID, perspective, language level, analysis request/status, entitlement disclosure.","rules":"Beginner mode changes presentation, not engine truth; uncertainty is retained.","output":"Beginner Review overview or actionable unavailable state.","special":"Progress is perceivable without relying on animation; cancellation is keyboard accessible.","open":"What entitlement and expected-duration disclosures are accurate?"},
    {
        "id":"C-UC03","name":"Understand the main mistake and recommended move","actor":"Beginner Player","support":"Analysis Engine; Explanation Service","persona":"C-P1","concept":"C-A1/C-A2/C-A3","evidence":"C2-E08, C2-E10; C-D2","trigger":"Beginner Review returns at least one explainable critical position.","pre":"C-UC02 succeeded; a move recommendation and confidence/uncertainty representation exist.","frequency":"One primary explanation per beginner overview; assumption to validate.","priority":"Critical","stakeholders":"Beginner wants meaning, not just a label; engine/explanation owners need calibrated language.","main":["Open the primary critical-position card.","See the played move and recommended move in notation and on the board.","Read the immediate tactical/positional consequence in plain language.","Inspect the evidence cue or uncertainty note behind the recommendation.","Move between before and after board states.","Choose Try this move, Next critical point, or More analysis."],"alts":["A1 at step 1: No single critical point is reliable; present a neutral summary and explain why.","A2 at step 3: A term is unfamiliar; open an inline definition without leaving the position.","E1 at step 4: Explanation cannot be generated reliably; show move comparison without synthetic prose."],"data":"Position, played/recommended moves, evaluation change, explanation concepts, uncertainty.","rules":"Do not state one move as universally best without the engine context; avoid invented player intent.","output":"Comprehended-position view and chosen next action.","special":"Every visual arrow/color has text notation and ordered board-state descriptions.","open":"Which explanation rules are safe at each language level?"},
    {
        "id":"C-UC04","name":"Try the recommended move","actor":"Beginner Player","support":"Move Validator; Analysis Engine","persona":"C-P1","concept":"C-A1/C-A2","evidence":"C2-E08, C2-E10","trigger":"The player selects Try this move from a critical-position explanation.","pre":"A legal starting position and recommended move are available; practice mode is non-rated.","frequency":"Zero or more times per critical position.","priority":"High","stakeholders":"Player wants immediate practice; validator must distinguish legal, recommended, and alternative moves.","main":["Open the critical position in practice mode.","Read the goal and whose turn it is.","Enter a move by board or notation.","Receive legal/illegal validation without revealing the answer prematurely.","If legal, compare the move with the recommendation and see the key consequence.","Retry, reveal the recommended move, or continue to practice."],"alts":["A1 at step 3: Player requests a hint; reveal one bounded cue and record no failure judgment.","A2 at step 5: Legal alternative is reasonable; explain the difference instead of marking it simply wrong.","E1 at step 4: Input method fails; preserve the position and offer notation entry/reset."],"data":"Practice position, entered move, legality, recommended/acceptable alternatives, hint state.","rules":"Practice never changes rating/game history; legality and recommendation are separate judgments.","output":"Practice attempt, feedback, and retry/continue choice.","special":"No time limit; keyboard move entry and full board reset are available.","open":"How many alternatives receive tailored feedback?"},
    {
        "id":"C-UC05","name":"Reveal advanced analysis on demand","actor":"Returning Player","support":"Analysis Engine","persona":"C-P3","concept":"C-A1/C-A3","evidence":"C2-E10","trigger":"The player chooses More analysis from a beginner explanation or overview.","pre":"A review artifact exists; advanced data availability and terminology level are disclosed.","frequency":"Optional per position; assumption to validate.","priority":"Medium","stakeholders":"Returning player wants depth; beginner flow must remain stable and recoverable.","main":["Activate the labeled More analysis disclosure.","Read a short preview of the additional data types.","Expand engine lines, evaluations, or setup controls intentionally.","Adjust one advanced setting if needed.","Inspect the updated advanced result with provenance.","Collapse advanced analysis and return to the same beginner position."],"alts":["A1 at step 3: Advanced data is unavailable; leave the beginner explanation visible and state why.","A2 at step 4: A setting changes the evaluated position; require confirmation before replacing the current context.","E1 at step 5: Engine refresh fails; retain the prior result with its timestamp/status."],"data":"Disclosure state, engine lines/evaluations, settings, provenance, refresh status.","rules":"Advanced detail is opt-in and reversible; collapsing never discards the beginner explanation.","output":"Advanced view state plus intact return location.","special":"Disclosure communicates expanded/collapsed state programmatically and restores focus.","open":"Which advanced controls belong in the first disclosure tier?"},
    {
        "id":"C-UC06","name":"Continue with a relevant lesson or puzzle","actor":"Beginner / Returning Player","support":"Learning Catalog","persona":"C-P1; C-P3","concept":"C-A1/C-A2/C-A3","evidence":"C2-E07, C2-E08, C2-E15","trigger":"The player finishes an explanation/practice attempt and chooses a learning follow-up.","pre":"A concept tag can be mapped to governed catalog items; entitlement is known or disclosed.","frequency":"Optional at the end of a review segment.","priority":"High","stakeholders":"Player wants relevant continuation; catalog owner needs accurate topic/level/entitlement metadata.","main":["Open the Recommended next step panel.","Read why each lesson or puzzle relates to the reviewed position.","Compare type, level, duration estimate, and entitlement.","Choose one lesson or puzzle.","Confirm whether progress will be saved and where the review return link appears.","Open the learning item with the review context preserved.","Complete/exit and return to the review or learning path."],"alts":["A1 at step 2: No strong match exists; say so and offer browse by concept instead of a false recommendation.","A2 at step 3: Item is restricted; disclose entitlement before selection and keep an accessible alternative.","E1 at step 6: Catalog item fails to open; retain review context and offer another item."],"data":"Concept tags, item ID/type/level/duration, relevance reason, entitlement, return context.","rules":"Recommendation rationale is visible; restricted items are not the only continuation path.","output":"Opened learning item, saved return context, and optional progress link.","special":"Cards have descriptive link names and predictable return navigation.","open":"How are relevance and duration validated for beginners?"},
]


def add_use_case_spec(doc, uc, product):
    scope = "FIFA Ticket Confidence Layer" if product == "FIFA" else "Chess Guided Beginner Review"
    tp = "TP-FIFA" if product == "FIFA" else "TP-CHESS"
    doc.add_heading(f"{uc['id']} {uc['name']}",level=2)
    add_table(doc,["Field","Specification"],[
        ["Scope / level",f"{scope} / user goal"],
        ["Primary actor",uc["actor"]],
        ["Supporting actors",uc["support"]],
        ["Stakeholders and interests",uc["stakeholders"]],
        ["Trigger",uc["trigger"]],
        ["Preconditions",uc["pre"]],
        ["Minimal guarantee","The originating artifact/context remains recoverable; no unsupported state or external success is invented."],
        ["Success guarantee",uc["output"]],
        ["Frequency",uc["frequency"]],
        ["Priority",uc["priority"]],
        ["Links",f"Proto-persona {uc['persona']}; {tp}; concepts {uc['concept']}; evidence {uc['evidence']}"],
        ["Open issue",uc["open"]],
    ],[4.0,11.6])
    doc.add_heading("Main success scenario",level=3)
    add_numbered_steps(doc, uc["main"])
    doc.add_heading("Alternative and exception flows",level=3)
    add_bullets(doc, uc["alts"])
    doc.add_heading("Use-case-specific data, rules, and special requirements",level=3)
    add_table(doc,["Category","Requirement"],[
        ["Required data",uc["data"]],
        ["Rules",uc["rules"]],
        ["Output artifact",uc["output"]],
        ["Special requirement",uc["special"]],
    ],[4.0,11.6])


def build_use_case_document(diagrams):
    doc=base_document("Use Case Document")
    add_cover(doc,"Group10-PA2 Use Case Document","Twelve user-goal use cases, UML system diagrams, and six alternative flow models")
    add_toc_placeholder(doc)
    doc.add_heading("1. Scope and notation",level=1)
    doc.add_paragraph("This document specifies two proposed responsive web modules. UML use-case diagrams show actors outside named system boundaries and user-goal ellipses inside. The six concept diagrams in Section 8 are flow models, not use-case diagrams. Task specifications follow goal, precondition, subtasks, exceptions, context, and involved-party guidance (LN04 - Task Analysis, pp. 6, 39-40, 48-49).")
    doc.add_heading("2. Actors and external systems",level=1)
    add_table(doc,["Actor/system","Role","Trust boundary"],[
        ["Ticket Buyer","Selects context, interprets governed status, saves a plan, and decides on an external handoff.","May cross partner, identity, and notification boundaries."],
        ["Official Ticketing Partner","Provides the external official action.","External domain/account/payment context; purchase success is never inferred."],
        ["Notification/Identity Services","Persist optional plan/consent and deliver alerts.","Personal data, consent, authentication, and retention."],
        ["Beginner / Returning Player","Loads an artifact, requests review, practices, and chooses learning continuation.","May be anonymous or authenticated."],
        ["Analysis/Explanation Services","Validate positions and supply analysis/explanation artifacts.","Uncertainty, entitlement, and explanation quality."],
        ["Learning Catalog","Supplies lessons/puzzles and entitlement metadata.","Relevance, availability, progress, and return context."],
    ],[4.1,7.5,4.0])
    doc.add_heading("3. Concept-to-use-case mapping",level=1)
    add_table(doc,["Concept","Primary use cases"],[
        ["F-A1","F-UC01, F-UC02, F-UC03, F-UC05, F-UC06"],["F-A2","F-UC01, F-UC02, F-UC05, F-UC06"],["F-A3","F-UC01, F-UC04, F-UC05, F-UC06"],
        ["C-A1","C-UC01-C-UC06"],["C-A2","C-UC01-C-UC04, C-UC06"],["C-A3","C-UC01-C-UC03, C-UC05-C-UC06"],
    ],[3.5,12.1])
    doc.add_heading("4. FIFA UML use-case diagram",level=1)
    add_figure(doc,diagrams["fifa_uc"],"UC-F-SYS","UML use-case diagram: named boundary, Ticket Buyer and external-system actors, use-case ellipses, and associations.",f"generated-diagrams/{diagrams['fifa_uc'].name}","F-P3; TP-FIFA")
    doc.add_heading("5. FIFA use case specifications",level=1)
    for uc in FIFA_UCS:
        add_use_case_spec(doc,uc,"FIFA")
    doc.add_heading("6. Chess UML use-case diagram",level=1)
    add_figure(doc,diagrams["chess_uc"],"UC-C-SYS","UML use-case diagram: named boundary, player and external-system actors, use-case ellipses, and associations.",f"generated-diagrams/{diagrams['chess_uc'].name}","C-P1/C-P3; TP-CHESS")
    doc.add_heading("7. Chess use case specifications",level=1)
    for uc in CHESS_UCS:
        add_use_case_spec(doc,uc,"Chess")
    doc.add_heading("8. Alternative-concept flow models",level=1)
    for cid,name,*_ in CONCEPTS:
        key=cid.lower()+"-flow"
        add_figure(doc,diagrams[key],f"UC-{cid}",f"{name} flow model with numbered exact steps and labeled branches/breakdowns; it is not a use-case diagram.",f"generated-diagrams/{diagrams[key].name}",f"{'F-P3' if cid.startswith('F') else 'C-P1/C-P3'}; {'TP-FIFA' if cid.startswith('F') else 'TP-CHESS'}")
    doc.add_heading("9. Shared system-level nonfunctional requirements",level=1)
    add_table(doc,["Category","System-level requirement"],[
        ["Usability","Plain task language, visible system status, low-memory choices, explicit closure, and recoverable reversal (LN02 - Fundamental Concepts, p. 116)."],
        ["Accessibility","Keyboard operation, visible focus, logical reading order, text alternatives, labeled controls/states/errors, no color-only meaning, and scalable layouts."],
        ["Security / trust","Allow-listed external destinations, provenance, calibrated uncertainty, protected tokens/preferences, and no inferred external success."],
        ["Privacy","Data minimization, purpose-specific consent, retention disclosure, edit/delete controls, and no bundled marketing consent."],
        ["Data freshness / provenance","Every governed state exposes source and timestamp; stale, unknown, not provided, closed, and unavailable remain distinct."],
        ["Error recovery","Preserve valid input/context, focus the exact error, provide retry/safe exit, and avoid destructive silent correction."],
        ["Localization","Localize dates/times/language while preserving canonical identifiers, chess notation, partner domains, and status semantics."],
    ],[3.6,12.0])
    doc.add_heading("10. Traceability and open issues",level=1)
    add_table(doc,["Tough problem","Concepts","Use cases","Evidence"],[
        ["TP-FIFA","F-A1/F-A2/F-A3","F-UC01-F-UC06","F2-E09-E11/E18; F-D2/F-D4"],
        ["TP-CHESS","C-A1/C-A2/C-A3","C-UC01-C-UC06","C2-E07/E08/E10/E15; C-D1/C-D2 (entry-only PA2 limitation)"],
    ],[3.0,3.6,3.6,5.4])
    add_bullets(doc,["Which official system owns each ticket state and freshness field?","Which markets and partner disclosures must each FIFA path express?","Which explanations can be produced reliably with calibrated uncertainty?","Which entitlement alternatives and return contexts are supported?","Which PA3 measures will validate the stated hypotheses without treating them as prior results?"])
    add_header_footer(doc,"Use Case Document")
    path=SOURCE/"Group10-PA2-UseCaseDocument.docx"; doc.save(path); return path


def build_peer_review():
    doc=base_document("Peer Review")
    add_cover(doc,"Group10-PA2 Peer Review","Pre-review presentation and evidence-completion record","INCOMPLETE - verified PA2 presentation feedback is absent")
    doc.add_heading("1. Evidence status",level=1)
    callout(doc,"Evidence blocker","No verified PA2 presentation date, duration, presenter list, audience record, commenter identity, feedback/question, group response, decision, owner, or revision evidence was found. This file cannot be completed truthfully until that evidence is supplied.")
    doc.add_heading("2. Presentation plan (5-10 minutes)",level=1)
    add_table(doc,["Time","Slide","Speaker focus"],[
        ["0:00-0:45","1. Scope and evidence boundary","Explain PA1 continuity, 105 screenshots, and absence of external participants."],
        ["0:45-1:45","2. FIFA task findings","Show ticket landing and partner-boundary evidence without unsupported claims."],
        ["1:45-2:45","3. Chess task findings","Show learning/analysis entry and positive guided-learning baseline."],
        ["2:45-3:45","4. Affinity and prioritization","Explain 11 clusters, rubric, and two tough problems."],
        ["3:45-5:15","5. FIFA concepts","Contrast dashboard, concierge, and alert-first planner."],
        ["5:15-6:45","6. Chess concepts","Contrast preset, conversational coach, and visual story."],
        ["6:45-7:45","7. Use cases and PA3","Show system boundaries and six alternative flows."],
        ["7:45-9:00","8. Questions","Capture commenter name, exact question, response, and revision decision."],
    ],[2.0,4.2,9.4])
    doc.add_heading("3. Speaking outline",level=1)
    add_bullets(doc,["State what is observed, inherited from PA1, inferred, and proposed.","Do not present personas as interview-derived or outcomes as measured.","Point to exact ticket cards, destination branding, lesson lists, and analysis-entry controls.","Ask reviewers whether the two tough problems are severe, distinct, and suitable for PA3 alternatives.","Record feedback verbatim enough to preserve meaning, but do not invent or paraphrase absent comments."])
    doc.add_heading("4. Missing real feedback record",level=1)
    add_table(doc,["Required field","Verified value"],[
        ["Presentation date","Not available - evidence blocker."],["Duration","Not available - evidence blocker."],["Presenters","Not available - evidence blocker."],["Audience","Not available - evidence blocker."],["Commenter name","Not available - evidence blocker."],["Comment or question","Not available - evidence blocker."],["Group response","Not available - evidence blocker."],["Decision and rationale","Not available - evidence blocker."],["Owner","Not available - evidence blocker."],["Affected document/page","Not available - evidence blocker."],["Revision and revision evidence","Not available - evidence blocker."],
    ],[5.0,10.6])
    doc.add_heading("5. Feedback and revision log status",level=1)
    doc.add_paragraph("No commenter row or revision row is created because an empty or synthetic row could be mistaken for real feedback. Once evidence exists, each row must preserve the comment's meaning and include response, decision, rationale, owner, affected document/page, revision, and revision evidence.")
    doc.add_heading("6. Completion gate",level=1)
    add_bullets(doc,["Add verified presentation metadata and audience/presenter records.","Add each real commenter and feedback/question without changing its meaning.","Record group response, decision, rationale, owner, affected document/page, revision, and revision evidence.","Regenerate affected reports and the six-file ZIP.","Remove the incomplete label only after the record and revisions are validated."])
    add_header_footer(doc,"Peer Review")
    path=SOURCE/"Group10-PA2-PeerReview.docx"; doc.save(path); return path


def build_weekly_report():
    doc=base_document("Weekly Report")
    add_cover(doc,"Group10-PA2 Weekly Report","Verified artifact status and missing team-process evidence","INCOMPLETE - verified PA2 meeting, attendance, and ownership records are absent")
    doc.add_heading("1. Evidence status",level=1)
    callout(doc,"Evidence blocker","No verified PA2 sprint dates, meeting schedule, attendance, objectives/minutes, per-member completed/current/next tasks, assigned actions, or ownership record was found. File existence is not used to infer who performed work.")
    callout(doc,"Package objective","Maintain six evidence-traceable reports with two provisional tough problems, three alternatives per problem, and twelve distinct user-goal use cases.")
    add_table(doc,["Field","Verified status"],[
        ["Sprint dates","Not available - evidence blocker. Capture/build timestamps are not sprint dates."],
        ["Meeting schedule and objectives","Not available - evidence blocker."],
        ["Attendance and minutes","Not available - evidence blocker."],
        ["Per-member completed/current/next work","Not available - evidence blocker."],
        ["Official template","The assignment-linked template is not present locally and was not fetched because browser/network use was prohibited."],
    ],[4.5,11.1])
    doc.add_heading("2. Team roster",level=1)
    add_table(doc,["Member","Student ID","PA2 completion ownership"],[[n,s,"Not verified; planned assignment only."] for n,s in TEAM],[5.0,3.2,7.4])
    doc.add_heading("3. Member task status",level=1)
    add_table(doc,["Member","Completed","Current","Next","Blocker"],[
        [name,"Not verifiable","Not verifiable","Not assignable without evidence","No PA2 task/meeting record"] for name,_ in TEAM
    ],[3.0,3.0,3.0,4.0,3.0])
    doc.add_heading("4. Verified workspace artifacts (ownership not inferred)",level=1)
    add_table(doc,["Artifact / activity","Evidence","Status"],[
        ["Source and filesystem audit","tmp/source-audit/audit.json; file-inventory.csv","Verified artifact"],
        ["Evidence index and validation","evidence-index.csv; evidence-validation.md","Verified artifact"],
        ["Traceability matrix","traceability-matrix.csv","Verified artifact"],
        ["Six editable report sources","source/Group10-PA2-*.docx","Present; owner unverified"],
        ["Generated diagrams and crops","generated-diagrams/*.png/*.svg","Present; owner unverified"],
        ["PDF export and QA","final/*.pdf and tmp/qa/qa-results.json","Checked by the final package QA"],
        ["Peer-review feedback integration","No PA2 feedback source found","Blocked by missing external evidence"],
        ["Group brainstorming record","No session/board/vote source found","Blocked by missing real group evidence"],
    ],[5.0,6.2,4.4])
    doc.add_heading("5. Issues, obstacles, and assigned actions",level=1)
    add_table(doc,["Issue","Impact","Assigned action","Owner/status"],[
        ["No external user research","Personas and findings remain provisional.","Conduct interviews or direct observation and update UserResearch.","Unassigned; open"],
        ["No real group brainstorming record","Requirement 2 remains incomplete.","Supply session metadata, real board export, notes, votes/consensus, decisions, and owners.","Unassigned; open"],
        ["No PA2 peer feedback","PeerReview remains incomplete.","Capture real lecture feedback and revision evidence.","Presenting team; open"],
        ["No PA2 meetings/attendance","Weekly ownership and attendance cannot be verified.","Add real meeting minutes or task tracker export.","Team coordinator; open"],
        ["Evidence template absent","Cannot validate entries intended for PA2_REAL_EVIDENCE_TEMPLATE.md.","Supply the completed template and linked artifacts.","Unassigned; open"],
        ["Misleading screenshot filenames","Risk of unsupported claims.","Use evidence index and excluded-state rules.","Package QA; controlled"],
    ],[4.2,4.5,5.0,2.3])
    doc.add_heading("6. Evidence-governance decisions",level=1)
    add_bullets(doc,["Keep PA1 canonical persona and drawback IDs, but convert PA2 personas to provisional role-based profiles.","Keep TP-FIFA and TP-CHESS as provisional evidence-driven selections; do not call them group-voted.","Restrict chess-53 because visible pixels do not unambiguously show the named ad panel.","Treat every proposed outcome as a hypothesis or PA3 design goal.","No post-presentation changes are reported because verified PA2 feedback is absent."])
    doc.add_heading("7. Package checklist",level=1)
    add_table(doc,["Item","Status"],[
        ["Six exact PDF filenames","Pass in final package QA"],["Group10 identifier and four student IDs","Pass in generated reports"],["Three alternatives per tough problem","Pass"],["Twelve distinct user-goal use case specifications","Pass"],["Embedded screenshots, crops, UML, screen maps, and flow models","Pass in rendered sources"],["Selectable PDF text and page/render QA","Pass in final package QA"],["Peer review evidence","Missing; report remains incomplete"],["Group brainstorming evidence","Missing; User Analysis remains incomplete"],["Weekly meeting/attendance/ownership evidence","Missing; report remains incomplete"],["ZIP contains only six PDFs","Pass in final package QA"],
    ],[10.5,5.1])
    add_header_footer(doc,"Weekly Report")
    path=SOURCE/"Group10-PA2-WeeklyReport.docx"; doc.save(path); return path


def main():
    SOURCE.mkdir(exist_ok=True)
    DIAGRAMS.mkdir(exist_ok=True)
    FINAL.mkdir(exist_ok=True)
    diagrams=generate_diagrams()
    evidence=read_evidence()
    paths=[
        build_user_research(diagrams,evidence),
        build_user_analysis(diagrams),
        build_project_proposal(diagrams),
        build_use_case_document(diagrams),
        build_peer_review(),
        build_weekly_report(),
    ]
    print("\n".join(str(p) for p in paths))
    print(f"Diagrams: {len(list(DIAGRAMS.glob('*.svg')))} SVG + {len(list(DIAGRAMS.glob('*.png')))} PNG")


if __name__ == "__main__":
    main()
