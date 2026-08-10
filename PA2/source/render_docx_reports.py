from __future__ import annotations

import io
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    BaseDocTemplate, Flowable, Frame, Image, KeepTogether, NextPageTemplate,
    PageBreak, PageTemplate, Paragraph, Spacer, Table, TableStyle,
)
from reportlab.platypus.tableofcontents import TableOfContents
from reportlab.lib.utils import ImageReader
from xml.sax.saxutils import escape


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "source"
FINAL = ROOT / "final"
STEMS = [
    "Group10-PA2-UserResearch", "Group10-PA2-UserAnalysis",
    "Group10-PA2-ProjectProposal", "Group10-PA2-UseCaseDocument",
    "Group10-PA2-PeerReview", "Group10-PA2-WeeklyReport",
]
NAVY, BLUE, GREEN, INK, MUTED = colors.HexColor("#123B65"), colors.HexColor("#1D70A2"), colors.HexColor("#4F7F35"), colors.HexColor("#1F2937"), colors.HexColor("#64748B")


class PA2Doc(BaseDocTemplate):
    def __init__(self, filename: str, title: str):
        super().__init__(filename, pagesize=A4, title=title, author="Group10", leftMargin=1.7*cm, rightMargin=1.7*cm, topMargin=1.7*cm, bottomMargin=1.6*cm)
        pw, ph = A4
        lw, lh = landscape(A4)
        self.title_text = title.replace("Group10-PA2-", "")
        self.addPageTemplates([
            PageTemplate("portrait", [Frame(1.7*cm, 1.6*cm, pw-3.4*cm, ph-3.3*cm, id="p")], pagesize=A4, onPage=self.decorate),
            PageTemplate("landscape", [Frame(1.3*cm, 1.3*cm, lw-2.6*cm, lh-2.6*cm, id="l")], pagesize=landscape(A4), onPage=self.decorate),
        ])

    def decorate(self, canvas, doc):
        canvas.saveState()
        w, h = canvas._pagesize
        canvas.setFont("Helvetica-Bold", 14)
        canvas.setFillColor(NAVY)
        canvas.drawString(1.3*cm, h-0.85*cm, f"GROUP10  |  CSC13112 UI/UX DESIGN  |  {self.title_text.upper()}")
        canvas.setStrokeColor(colors.HexColor("#D8E2EA")); canvas.line(1.3*cm, h-1.0*cm, w-1.3*cm, h-1.0*cm)
        canvas.setFont("Helvetica", 14); canvas.setFillColor(MUTED)
        canvas.drawCentredString(w/2, 0.75*cm, f"{self.title_text}  |  Group10  |  {doc.page}")
        canvas.restoreState()

    def afterFlowable(self, flowable):
        if isinstance(flowable, Paragraph) and flowable.style.name in ("Heading1", "Heading2"):
            level = 0 if flowable.style.name == "Heading1" else 1
            text = flowable.getPlainText()
            key = f"h{level}-{self.page}-{abs(hash(text))}"
            self.canv.bookmarkPage(key)
            self.canv.addOutlineEntry(text, key, level=level, closed=False)
            self.notify("TOCEntry", (level, text, self.page, key))


def styles():
    s = getSampleStyleSheet()
    return {
        "Normal": ParagraphStyle("NormalPA2", parent=s["BodyText"], fontName="Helvetica", fontSize=14, leading=18, textColor=INK, spaceAfter=7),
        "Title": ParagraphStyle("TitlePA2", parent=s["Title"], fontName="Helvetica-Bold", fontSize=24, leading=28, textColor=NAVY, alignment=TA_CENTER, spaceAfter=9),
        "Subtitle": ParagraphStyle("SubtitlePA2", parent=s["BodyText"], fontName="Helvetica", fontSize=14, leading=18, textColor=MUTED, alignment=TA_CENTER, spaceAfter=9),
        "Heading1": ParagraphStyle("Heading1", parent=s["Heading1"], fontName="Helvetica-Bold", fontSize=15, leading=18, textColor=NAVY, spaceBefore=10, spaceAfter=6, keepWithNext=True),
        "Heading2": ParagraphStyle("Heading2", parent=s["Heading2"], fontName="Helvetica-Bold", fontSize=16, leading=20, textColor=BLUE, spaceBefore=8, spaceAfter=4, keepWithNext=True),
        "Heading3": ParagraphStyle("Heading3", parent=s["Heading3"], fontName="Helvetica-Bold", fontSize=14, leading=18, textColor=GREEN, spaceBefore=6, spaceAfter=3, keepWithNext=True),
        "Caption": ParagraphStyle("CaptionPA2", parent=s["BodyText"], fontName="Helvetica", fontSize=14, leading=18, textColor=MUTED, spaceAfter=8),
        "TOC0": ParagraphStyle("TOC0", parent=s["BodyText"], fontName="Helvetica-Bold", fontSize=14, leading=18, leftIndent=0, firstLineIndent=0, textColor=NAVY),
        "TOC1": ParagraphStyle("TOC1", parent=s["BodyText"], fontName="Helvetica", fontSize=14, leading=18, leftIndent=12, firstLineIndent=0, textColor=INK),
    }


def blocks(doc: Document):
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield DocxParagraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield DocxTable(child, doc)


def paragraph_markup(p: DocxParagraph) -> str:
    parts = []
    for run in p.runs:
        text = escape(run.text).replace("\n", "<br/>")
        if not text:
            continue
        if run.bold:
            text = f"<b>{text}</b>"
        if run.italic:
            text = f"<i>{text}</i>"
        parts.append(text)
    return "".join(parts) or escape(p.text)


def image_flowables(p: DocxParagraph):
    out = []
    for blip in p._p.xpath(".//a:blip"):
        rid = blip.get(qn("r:embed"))
        if not rid or rid not in p.part.related_parts:
            continue
        blob = p.part.related_parts[rid].blob
        stream = io.BytesIO(blob)
        reader = ImageReader(stream)
        iw, ih = reader.getSize()
        scale = min((16.0*cm)/iw, (14.5*cm)/ih, 1.0)
        stream.seek(0)
        out.append(Image(stream, width=iw*scale, height=ih*scale, hAlign="CENTER"))
    return out


def table_flowables(t: DocxTable, stylemap, landscape_table: bool):
    rows = []
    for ri, row in enumerate(t.rows):
        cells = []
        for cell in row.cells:
            text = "<br/>".join(escape(p.text.strip()) for p in cell.paragraphs if p.text.strip()) or " "
            ps = ParagraphStyle("cell", parent=stylemap["Normal"], fontSize=14, leading=18, spaceAfter=0, textColor=colors.white if ri == 0 else INK)
            cells.append(Paragraph(text, ps))
        rows.append(cells)
    cols = max(1, len(rows[0]))
    if cols >= 4:
        cards = []
        headers = [cell.text.strip() or f"Field {i+1}" for i, cell in enumerate(t.rows[0].cells)]
        key_style = ParagraphStyle("card-key", parent=stylemap["Normal"], fontName="Helvetica-Bold", fontSize=14, leading=18, textColor=colors.white)
        value_style = ParagraphStyle("card-value", parent=stylemap["Normal"], fontSize=14, leading=18, textColor=INK)
        for record_number, row in enumerate(t.rows[1:], start=1):
            pairs = []
            for header, cell in zip(headers, row.cells):
                value = "<br/>".join(escape(p.text.strip()) for p in cell.paragraphs if p.text.strip()) or " "
                pairs.append([Paragraph(escape(header), key_style), Paragraph(value, value_style)])
            card = Table(pairs, colWidths=[5.2*cm, 11.8*cm], splitByRow=1, hAlign="LEFT")
            card.setStyle(TableStyle([
                ("BACKGROUND", (0,0), (0,-1), NAVY), ("TEXTCOLOR", (0,0), (0,-1), colors.white),
                ("GRID", (0,0), (-1,-1), 0.5, colors.HexColor("#94A3B8")),
                ("VALIGN", (0,0), (-1,-1), "TOP"), ("LEFTPADDING", (0,0), (-1,-1), 6),
                ("RIGHTPADDING", (0,0), (-1,-1), 6), ("TOPPADDING", (0,0), (-1,-1), 5), ("BOTTOMPADDING", (0,0), (-1,-1), 5),
            ]))
            cards.extend([Paragraph(f"Record {record_number}", stylemap["Heading3"]), card, Spacer(1, 10)])
        return cards
    width = (27.0 if landscape_table else 17.0) * cm
    tbl = Table(rows, colWidths=[width/cols]*cols, repeatRows=1, hAlign="LEFT", splitByRow=1)
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), NAVY), ("TEXTCOLOR", (0,0), (-1,0), colors.white),
        ("GRID", (0,0), (-1,-1), 0.35, colors.HexColor("#94A3B8")),
        ("VALIGN", (0,0), (-1,-1), "TOP"), ("LEFTPADDING", (0,0), (-1,-1), 3),
        ("RIGHTPADDING", (0,0), (-1,-1), 3), ("TOPPADDING", (0,0), (-1,-1), 3), ("BOTTOMPADDING", (0,0), (-1,-1), 3),
        ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, colors.HexColor("#F7FAFC")]),
    ]))
    return [tbl]


def convert(stem: str):
    src, dst = SOURCE / f"{stem}.docx", FINAL / f"{stem}.pdf"
    docx = Document(src)
    sm = styles()
    story = []
    list_counter = 0
    for block in blocks(docx):
        if isinstance(block, DocxParagraph):
            p = block
            style_name = p.style.name if p.style else "Normal"
            if "[[TOC]]" in p.text:
                toc = TableOfContents(); toc.levelStyles = [sm["TOC0"], sm["TOC1"]]
                story.append(toc); continue
            imgs = image_flowables(p)
            if imgs:
                story.extend(imgs)
            text = p.text.strip()
            if text:
                mapped = "Normal"
                if style_name == "Title": mapped = "Title"
                elif style_name == "Subtitle": mapped = "Subtitle"
                elif style_name == "Heading 1": mapped = "Heading1"
                elif style_name == "Heading 2": mapped = "Heading2"
                elif style_name == "Heading 3": mapped = "Heading3"
                elif text.startswith("Figure "): mapped = "Caption"
                if style_name == "List Bullet":
                    text = "&#8226; " + escape(text)
                    story.append(Paragraph(text, sm["Normal"])); continue
                if style_name == "List Number":
                    list_counter += 1
                    story.append(Paragraph(f"{list_counter}. {escape(text)}", sm["Normal"])); continue
                list_counter = 0
                rendered = Paragraph(paragraph_markup(p), sm[mapped])
                if mapped == "Caption" and story and isinstance(story[-1], Image):
                    story.append(KeepTogether([story.pop(), rendered]))
                else:
                    story.append(rendered)
            if "w:type=\"page\"" in p._p.xml:
                story.append(PageBreak())
        else:
            story.extend([*table_flowables(block, sm, False), Spacer(1, 8)])
    pdf = PA2Doc(str(dst), stem)
    pdf.multiBuild(story)
    print(stem, dst.stat().st_size)


def main():
    FINAL.mkdir(exist_ok=True)
    for stem in STEMS:
        convert(stem)


if __name__ == "__main__":
    main()
