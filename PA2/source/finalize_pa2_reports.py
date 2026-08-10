from __future__ import annotations

import csv
import hashlib
import importlib.util
import math
import zipfile
from pathlib import Path

from PIL import Image, ImageDraw
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "source"
FINAL = ROOT / "final"
DIAGRAMS = ROOT / "generated-diagrams"

spec = importlib.util.spec_from_file_location("pa2_base", SOURCE / "build_pa2_reports.py")
m = importlib.util.module_from_spec(spec)
assert spec.loader
spec.loader.exec_module(m)
m.ROOT = ROOT
m.PA1 = ROOT.parent / "PA1"
m.SOURCE = SOURCE
m.FINAL = FINAL
m.DIAGRAMS = DIAGRAMS
m.CAP = ROOT / "capture-work"


TEAM = ["Le Minh", "Nguyen Vu Bach", "Pham Nguyen Gia Bao", "Trang Minh Nhut"]
COURSE = (
    "Course basis: LN01 - Introduction - v2.pdf, pp. 20, 22, 24; "
    "LN02 - Fundamental Concepts - Usability Dimensions_2.pdf, pp. 62-67, 74-76, 83-84, 101-115; "
    "LN03 - UI Design Process.pdf, pp. 16-19, 23-24, 28-30, 34-36; "
    "LN04 - Task Analysis.pdf, pp. 3, 6-15, 39-49."
)


def save(doc: Document, stem: str) -> Path:
    m.add_header_footer(doc, stem.replace("Group10-PA2-", ""))
    path = SOURCE / f"{stem}.docx"
    doc.save(path)
    return path


def page(doc: Document):
    doc.add_page_break()


def landscape(doc: Document):
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = Cm(29.7), Cm(21.0)
    sec.top_margin = sec.bottom_margin = Cm(1.3)
    sec.left_margin = sec.right_margin = Cm(1.3)
    return sec


def portrait(doc: Document):
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.top_margin, sec.bottom_margin = Cm(2.0), Cm(1.8)
    sec.left_margin = sec.right_margin = Cm(2.0)
    return sec


def source_line(doc: Document, source: str, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{source}: ")
    r.bold = True
    p.add_run(text)


def cover(doc: Document, title: str, subtitle: str):
    m.add_cover(doc, title, subtitle)
    m.add_toc_placeholder(doc)


def kv(doc: Document, items: list[tuple[str, str]], widths=(4.2, 11.2)):
    return m.add_table(doc, ["Field", "Record"], [[a, b] for a, b in items], list(widths))


PARTICIPANTS = [
    ("F-SIM-01", "Casual football fan", "18-24", "Android phone", "Follows major finals", "Occasional FIFA.com visitor", "Monthly in season", "Find ticket entry without planning expertise", "Commute; one-handed use", "Limited time and data", "Large tap targets", "F1-F6"),
    ("F-SIM-02", "Casual football fan", "25-34", "Windows laptop + phone", "Watches selected national-team matches", "Uses Match Centre occasionally", "Several times per tournament", "Compare match and ticket status", "Evening planning at home", "Avoids creating accounts early", "High contrast", "F1-F6"),
    ("F-SIM-03", "Tournament follower", "25-34", "iPhone", "Tracks brackets and schedules", "Returning FIFA.com visitor", "Weekly in tournament periods", "Coordinate a group trip", "Mobile-first, shared planning", "Needs fast resumption", "Text zoom", "F1-F6"),
    ("F-SIM-04", "Tournament follower", "35-44", "MacBook + Android phone", "Follows multiple competitions", "Regular Match Centre user", "Weekly", "Connect schedule context to ticket action", "Desk-to-mobile transition", "Many open tabs", "Keyboard access on laptop", "F1-F6"),
    ("F-SIM-05", "Trust-sensitive ticket planner", "35-44", "Windows laptop", "Has bought event tickets online", "Knows FIFA.com but not partner boundaries", "At key sale windows", "Avoid unofficial or stale destinations", "Family purchase planning", "Needs provenance before leaving site", "Clear status text, not color alone", "F1-F6"),
    ("F-SIM-06", "Trust-sensitive ticket planner", "45-54", "Android phone", "Plans international tournament travel", "Low-frequency FIFA.com visitor", "A few times per year", "Know whether to buy, wait, or register interest", "Mobile on variable network", "Low confidence after handoff", "Plain language and persistent focus", "F1-F6"),
    ("C-SIM-01", "First-time beginner", "18-24", "Android phone", "Knows piece movement", "First Chess.com week", "Daily trial week", "Choose a safe beginner route", "Short study breaks", "Unfamiliar analysis terms", "Plain language", "C1-C9"),
    ("C-SIM-02", "First-time beginner", "25-34", "Windows laptop", "Played fewer than ten games", "New account", "Twice weekly", "Understand one mistake", "Home learning session", "Fears breaking the position", "Keyboard-only path", "C1-C9"),
    ("C-SIM-03", "Returning learner", "18-24", "iPhone", "Completed basic lessons", "Returns after four months", "Weekly when active", "Resume without relearning navigation", "Mobile evening practice", "Forgets feature locations", "Text zoom", "C1-C9"),
    ("C-SIM-04", "Returning learner", "35-44", "MacBook", "Knows opening principles", "Intermittent lesson and puzzle use", "Monthly", "Reconnect review to practice", "Desktop at home", "Forgets analysis vocabulary", "Reduced motion", "C1-C9"),
    ("C-SIM-05", "Intermediate player with low analysis use", "25-34", "Windows laptop", "Club-level casual", "Plays games and puzzles", "Weekly", "Review without engine overload", "After-game desktop session", "Skips dense panels", "Keyboard shortcuts with labels", "C1-C9"),
    ("C-SIM-06", "Intermediate player with low analysis use", "35-44", "Android tablet", "Understands tactics", "Rarely opens analysis", "Several times per month", "Find one actionable improvement", "Tablet practice session", "Terminology slows scanning", "Color-independent indicators", "C1-C9"),
]


SESSIONS = [
    ("F-SES-01", "16 Jul 2026", "28 min", "F-SIM-01", "Android phone", "F1, F2, F4, F5", "Nguyen Vu Bach", "Trang Minh Nhut"),
    ("F-SES-02", "16 Jul 2026", "31 min", "F-SIM-02", "Windows laptop + phone", "F1-F6", "Nguyen Vu Bach", "Le Minh"),
    ("C-SES-01", "16 Jul 2026", "29 min", "C-SIM-01", "Android phone", "C1-C3, C7-C9", "Pham Nguyen Gia Bao", "Trang Minh Nhut"),
    ("C-SES-02", "16 Jul 2026", "32 min", "C-SIM-02", "Windows laptop", "C1-C9", "Pham Nguyen Gia Bao", "Le Minh"),
    ("F-SES-03", "17 Jul 2026", "30 min", "F-SIM-03", "iPhone", "F1-F6", "Nguyen Vu Bach", "Trang Minh Nhut"),
    ("F-SES-04", "17 Jul 2026", "34 min", "F-SIM-04", "MacBook + Android phone", "F1-F6", "Nguyen Vu Bach", "Le Minh"),
    ("C-SES-03", "17 Jul 2026", "27 min", "C-SIM-03", "iPhone", "C1-C3, C5-C9", "Pham Nguyen Gia Bao", "Trang Minh Nhut"),
    ("C-SES-04", "17 Jul 2026", "35 min", "C-SIM-04", "MacBook", "C1-C9", "Pham Nguyen Gia Bao", "Le Minh"),
    ("F-SES-05", "18 Jul 2026", "33 min", "F-SIM-05", "Windows laptop", "F1-F6", "Nguyen Vu Bach", "Trang Minh Nhut"),
    ("F-SES-06", "18 Jul 2026", "35 min", "F-SIM-06", "Android phone", "F1-F6", "Nguyen Vu Bach", "Le Minh"),
    ("C-SES-05", "18 Jul 2026", "30 min", "C-SIM-05", "Windows laptop", "C1-C9", "Pham Nguyen Gia Bao", "Trang Minh Nhut"),
    ("C-SES-06", "18 Jul 2026", "34 min", "C-SIM-06", "Android tablet", "C1-C9", "Pham Nguyen Gia Bao", "Le Minh"),
]


NOTES = {
    "F-SIM-01": ("Opens menu, scrolls to Tickets & Hospitality, returns to Match Centre, then reopens ticket page.", "Pauses at tournament cards with different CTA wording.", "Tickets & Hospitality card and tournament status text.", "Ticket page should give one current action.", "I can find tickets, but I still do not know what is open.", "Partial success", "Uses Match Centre as a comparison anchor.", "Show state, freshness, and next action together.", "Scenario-based; no live sale data."),
    "F-SIM-02": ("Finds Match Centre quickly, compares schedule, opens ticket entry, then checks destination preview.", "Hesitates before the external handoff.", "Match Centre filters and partner CTA.", "An official domain preview should precede departure.", "The schedule is clear; the handoff is the risky part.", "Success with hesitation", "Opens destination in a second tab.", "Keep tournament context and return route.", "Scenario-based; partner response not tested."),
    "F-SIM-03": ("Selects tournament, saves context mentally, moves to phone view, and searches for the same state.", "Loses orientation after the mobile stack changes order.", "Long mobile tournament-card stack.", "Mobile should restore the selected tournament first.", "I want the same tournament to stay at the top.", "Partial success", "Uses page search and browser history.", "Persist context across devices.", "Scenario-based; account synchronization assumed."),
    "F-SIM-04": ("Uses Match Centre filters, opens ticket card, checks freshness, then returns.", "Questions whether the status changed since the last visit.", "Status label without a prominent timestamp.", "Current state requires visible provenance.", "Is this from today or from the last sale window?", "Failure on freshness", "Refreshes and compares wording.", "Display updated time and governed source.", "Scenario-based; timestamp policy unknown."),
    "F-SIM-05": ("Compares buy, wait, resale, and hospitality paths before choosing to stay.", "Stops when official destination is not explained.", "Multiple commercial paths and outbound CTA.", "Departure requires trust evidence and safe cancellation.", "I will not leave until I know where this goes.", "Success by choosing stay", "Copies the domain name for separate checking.", "Preview destination, purpose, and return behavior.", "Scenario-based; no transaction attempted."),
    "F-SIM-06": ("Finds ticket page, encounters missing state, retries on variable network, then subscribes conceptually.", "Cannot distinguish missing from sold out.", "Empty or stale status region.", "Missing data is not a ticket outcome.", "No status should not look like no tickets.", "Recovered after retry", "Returns to homepage and tries again.", "Differentiate missing, stale, sold out, and waiting room.", "Scenario-based; notification permission simulated."),
    "C-SIM-01": ("Chooses Learn, continues a lesson, solves a puzzle, then looks for review.", "Opens several analysis entry choices.", "Learn cards, puzzle board, analysis entry menu.", "Review should start with a beginner preset.", "Which button tells me what I did wrong?", "Partial success", "Returns to the completed game screen.", "Offer one labeled beginner route.", "Scenario-based; no engine output claimed."),
    "C-SIM-02": ("Uses keyboard to open a game, selects beginner review, tries a move, and reads help.", "Focus order reaches advanced controls before explanation.", "Analysis controls and terminology links.", "The main mistake should precede advanced depth.", "I found analysis, but not the first thing to learn.", "Success with wrong path", "Uses Tab and headings to relocate.", "Order content by learning priority.", "Scenario-based; keyboard path modeled."),
    "C-SIM-03": ("Recognizes puzzles, resumes a lesson, opens a completed game, then searches for analysis.", "Forgets the name of the review feature.", "Navigation labels and game actions.", "Returning learners recognize outcomes, not feature terms.", "I remember the result, not the menu name.", "Recovered", "Uses recent activity to reopen the game.", "Add review action at the completed game.", "Scenario-based; recency history assumed."),
    "C-SIM-04": ("Opens game, selects review, reads one explanation, then continues to a lesson.", "Stops at evaluation and engine terms.", "Evaluation label and advanced analysis controls.", "Terminology help should be available in place.", "Tell me why before showing the numbers.", "Success with hesitation", "Opens help text and ignores depth.", "Plain-language explanation with optional terms.", "Scenario-based; explanation quality not measured."),
    "C-SIM-05": ("Finds analysis, follows one candidate line, returns, and selects a puzzle.", "Takes a wrong path into full analysis first.", "Multiple setup paths and engine panel.", "Low-analysis users need a preset at entry.", "I opened the expert view by accident.", "Recovered", "Backs out and uses review link.", "Default to beginner review for the scenario.", "Scenario-based; recommendation accuracy not claimed."),
    "C-SIM-06": ("Opens recent game, identifies one turning point, tries a move, and reveals advanced detail.", "Scans slowly on tablet because labels compete.", "Board, move list, explanation, depth disclosure.", "A single learning card should lead the visual hierarchy.", "One mistake and one next exercise is enough.", "Success", "Collapses side panels.", "Progressive disclosure and continuation action.", "Scenario-based; engine availability assumed."),
}


RESULTS = [
    ("F-SIM-01 / F-SES-01", "3/4", "3", "1", "1", "3/5", "3/5"),
    ("F-SIM-02 / F-SES-02", "5/6", "2", "1", "1", "3/5", "4/5"),
    ("F-SIM-03 / F-SES-03", "4/6", "3", "1", "1", "3/5", "3/5"),
    ("F-SIM-04 / F-SES-04", "4/6", "3", "0", "1", "2/5", "3/5"),
    ("F-SIM-05 / F-SES-05", "4/6", "4", "1", "1", "2/5", "3/5"),
    ("F-SIM-06 / F-SES-06", "3/6", "4", "2", "2", "2/5", "2/5"),
    ("C-SIM-01 / C-SES-01", "5/6", "3", "2", "1", "3/5", "3/5"),
    ("C-SIM-02 / C-SES-02", "7/9", "4", "2", "1", "3/5", "3/5"),
    ("C-SIM-03 / C-SES-03", "6/8", "3", "2", "2", "3/5", "3/5"),
    ("C-SIM-04 / C-SES-04", "7/9", "3", "1", "1", "3/5", "4/5"),
    ("C-SIM-05 / C-SES-05", "7/9", "3", "2", "1", "3/5", "3/5"),
    ("C-SIM-06 / C-SES-06", "8/9", "2", "1", "1", "4/5", "4/5"),
]


def make_models():
    legend = ["Captured screen evidence", "Simulated participant behavior", "PA1 inherited finding", "Proposed response"]
    return {
        "f-flow": m.flow_diagram("final-fifa-flow", "UR-WM-F1 FIFA Flow Model", ["Select tournament", "Read ticket state", "Inspect freshness", "Preview destination", "Act, wait, or stay", "Return with context"], legend, "#1D70A2"),
        "f-seq": m.flow_diagram("final-fifa-sequence", "UR-WM-F2 FIFA Sequence Model", ["Intent", "Find entry", "Compare status", "Resolve missing/stale", "Review handoff", "Resume mobile"], legend, "#1D70A2"),
        "f-art": m.flow_diagram("final-fifa-artifact", "UR-WM-F3 FIFA Artifact Model", ["Tournament card", "Status + timestamp", "Official destination preview", "Saved plan", "Alert preference"], legend, "#1D70A2"),
        "f-info": m.flow_diagram("final-fifa-information", "UR-WM-F4 FIFA Information Model", ["Tournament identity", "Ticket state", "Freshness", "Channel", "Official destination", "Return state"], legend, "#1D70A2"),
        "c-flow": m.flow_diagram("final-chess-flow", "UR-WM-C1 Chess Flow Model", ["Choose beginner route", "Continue lesson", "Solve puzzle", "Open review", "Try move", "Continue practice"], legend, "#4F7F35"),
        "c-seq": m.flow_diagram("final-chess-sequence", "UR-WM-C2 Chess Sequence Model", ["Open game", "Select beginner review", "Identify mistake", "Try better move", "Read explanation", "Reveal depth"], legend, "#4F7F35"),
        "c-art": m.flow_diagram("final-chess-artifact", "UR-WM-C3 Chess Artifact Model", ["Game artifact", "Mistake card", "Board trial", "Plain-language explanation", "Practice link"], legend, "#4F7F35"),
        "c-learn": m.flow_diagram("final-chess-continuation", "UR-WM-C4 Chess Learning Continuation Model", ["Review outcome", "Learning priority", "Recommended lesson", "Relevant puzzle", "Advanced depth on demand"], legend, "#4F7F35"),
    }


def build_user_research() -> Path:
    models = make_models()
    doc = m.base_document("User Research")
    cover(doc, "Group10-PA2 User Research", "Source analysis, try-it-yourself walkthrough, and clearly labeled simulated study")
    doc.add_heading("1. Research scope, methods, and validity boundary", 1)
    m.callout(doc, "Simulated-study disclosure", "All participant profiles, sessions, quotations, observations, task outcomes, and ratings in this report are scenario-based synthetic evidence created to test design hypotheses; they are not presented as research with real users.")
    m.add_table(doc, ["Method", "What it supports", "Boundary", "Source label"], [
        ["Source analysis", "Interface inventory, documented states, PA1 continuity", "A file or screen does not prove behavior", "Screen evidence / PA1 inherited finding"],
        ["Try-it-yourself walkthrough", "Task path, transition, recovery, and handoff inspection", "Shows the path available to the evaluator, not population performance", "Screen evidence / Design inference"],
        ["Simulated user sessions", "Scenario coverage, test cases, hypotheses, and comparison logic", "Does not represent real-user measurement", "Simulated study / PA3 validation target"],
    ], [3.5, 5.0, 5.2, 3.0])
    source_line(doc, "Course source", COURSE)
    source_line(doc, "PA1 inherited finding", "FIFA separates useful tournament content from ticket decision confidence; Chess.com provides strong play and learning functions but exposes analysis choices with high recall demand.")
    source_line(doc, "Design inference", "Visibility of system status, recognition over recall, progressive disclosure, user control, and error prevention organize the PA2 questions.")

    doc.add_heading("2. Tasks and protocol", 1)
    doc.add_heading("2.1 FIFA task set", 2)
    m.add_numbered_steps(doc, ["Find ticket entry.", "Compare tournament status.", "Decide buy, wait, register interest, resale, or hospitality.", "Identify destination before leaving FIFA.com.", "Resume on mobile.", "Recover from missing or stale status."])
    doc.add_heading("2.2 Chess task set", 2)
    m.add_numbered_steps(doc, ["Choose a beginner route.", "Continue a lesson.", "Solve a puzzle.", "Open analysis or review.", "Identify one important mistake.", "Try a better move.", "Read a plain-language explanation.", "Continue to practice.", "Reveal advanced detail on demand."])
    m.callout(doc, "Simulated research session protocol", "Each 25-35 minute record uses 5 minutes of introduction, 15-20 minutes of task work, and 5 minutes of follow-up. Notes are structured; no audio or video is recorded.")

    doc.add_heading("3. Simulated participants", 1)
    landscape(doc)
    m.add_table(doc, ["ID", "Role", "Age", "Device", "Domain / application experience", "Frequency", "Motivation and context", "Constraints / accessibility", "Tasks"], [
        [p[0], p[1], p[2], p[3], f"{p[4]}; {p[5]}", p[6], f"{p[7]}; {p[8]}", f"{p[9]}; {p[10]}", p[11]] for p in PARTICIPANTS
    ], [2.0, 3.0, 1.4, 2.6, 4.4, 2.2, 4.5, 4.2, 1.7])
    source_line(doc, "Scenario-based synthetic evidence", "IDs are fictional study roles; no names identify participants.")
    portrait(doc)

    doc.add_heading("4. Three-day simulated session plan", 1)
    m.add_table(doc, ["Session ID", "Date", "Duration", "Participant", "Device", "Tasks", "Facilitator", "Note taker", "Recording method"], [
        [*s, "Structured notes only; no audio/video"] for s in SESSIONS
    ], [2.0, 2.1, 1.7, 2.0, 2.7, 2.0, 2.8, 2.8, 4.4])
    source_line(doc, "Simulated research session", "This is a simulated session record used for scenario coverage and planning, not an attendance or research claim.")

    doc.add_heading("5. Twelve simulated note sheets", 1)
    profiles = {p[0]: p for p in PARTICIPANTS}
    sessions = {s[3]: s for s in SESSIONS}
    for i, pid in enumerate(profiles):
        if i:
            page(doc)
        p = profiles[pid]
        s = sessions[pid]
        action, hesitation, trigger, mental, quote, outcome, workaround, opportunity, limitation = NOTES[pid]
        doc.add_heading(f"5.{i+1} Note sheet {pid}", 2)
        m.callout(doc, "Simulated observation", f"{s[0]} - {s[1]} - {s[2]}; structured notes only; no audio/video.", m.FIFA_LIGHT if pid.startswith("F") else m.CHESS_LIGHT)
        kv(doc, [
            ("Participant profile", f"Simulated participant {pid}; {p[1]}; {p[2]}; {p[3]}; motivation: {p[7]}; constraint: {p[9]}; accessibility: {p[10]}"),
            ("Task", s[5]),
            ("Action sequence", action),
            ("Hesitation point", hesitation),
            ("Visible interface trigger", trigger),
            ("Inferred mental model", mental),
            ("Simulated quote", f"\u201c{quote}\u201d"),
            ("Simulated task result", outcome),
            ("Workaround", workaround),
            ("Design opportunity", opportunity),
            ("Limitation", limitation),
        ])
        source_line(doc, "PA3 validation target", f"Test whether the proposed response reduces the hesitation in {s[0]} without hiding needed control.")

    doc.add_heading("6. Simulated task results", 1)
    m.add_table(doc, ["Trace", "Task completion", "Hesitations", "Wrong paths", "Recoveries", "Confidence", "Clarity"], RESULTS, [4.0, 2.5, 2.2, 2.2, 2.0, 2.1, 2.1])
    m.callout(doc, "Simulated task result", "The FIFA status/handoff records contain lower confidence and more hesitation than schedule-oriented work. Chess analysis entry has more wrong paths than puzzles; guided lesson work is clearer than analysis entry. No row is perfect, and every number traces to one note sheet.")

    doc.add_heading("7. Personas", 1)
    personas = [
        ("UR-P-F1", "FIFA trust-sensitive ticket planner", "Compares state and official destination before departure", "Checks labels, timestamp, domain, and return route", "Family purchase at a sale window", "High ticket-domain experience; moderate FIFA.com experience", "Laptop first; phone resume", "Ambiguous status and unexplained handoff", "Careful comparison and risk detection", "Will wait rather than accept uncertainty", "I need the state and destination before I leave.", "Compare sale, resale, hospitality, then choose act or stay.", "F-SIM-05/06; F2-E09; PA1 inherited ticket drawback", "Whether destination preview improves calibrated trust"),
        ("UR-P-F2", "FIFA returning tournament follower", "Reconnect schedule context to ticket action", "Starts in Match Centre and expects tournament continuity", "Weekly tournament planning", "High competition knowledge; regular FIFA.com use", "Phone first; laptop comparison", "Mobile stack and lost filters", "Fast recognition of tournaments and schedules", "Does not remember ticket channel wording", "Keep my tournament context when I switch devices.", "Follow a tournament, inspect status, resume on mobile.", "F-SIM-03/04; F2-E03/E04; PA1 continuity", "Whether persisted context reduces recovery work"),
        ("UR-P-C1", "Chess first-time beginner", "Understand one important mistake and next action", "Uses Learn and Puzzles; avoids engine terminology", "Short mobile learning session", "Basic rules; first week on Chess.com", "Phone; occasional laptop", "Too many entry choices", "Willing to try a move and read a short explanation", "Cannot interpret evaluation vocabulary", "Show me what to learn first.", "Complete a lesson, puzzle, and beginner review.", "C-SIM-01/02; C2-E05-E10; PA1 learning strength", "Whether one preset improves entry selection"),
        ("UR-P-C2", "Chess returning learner who forgets analysis terms", "Resume review and practice after a gap", "Recognizes recent game and lesson outcomes over feature names", "Intermittent desktop session", "Moderate chess knowledge; low analysis use", "Laptop; tablet backup", "Recall-heavy labels and dense advanced panels", "Can compare moves with plain explanations", "Forgets analysis terminology", "I remember the game, not the analysis label.", "Open recent game, identify one mistake, continue practice.", "C-SIM-03/04/05; C2-E07-E10", "Whether outcome-oriented labels improve memorability"),
    ]
    for j, x in enumerate(personas):
        doc.add_heading(f"7.{j+1} {x[1]}", 2)
        kv(doc, [("Persona ID", x[0]), ("Summary", x[1]), ("Goals", x[2]), ("Behaviors", x[3]), ("Context", x[4]), ("Domain / application experience", x[5]), ("Device pattern", x[6]), ("Frustrations", x[7]), ("Capabilities", x[8]), ("Limitations", x[9]), ("Simulated quote", f"\u201c{x[10]}\u201d"), ("Scenario", x[11]), ("Evidence basis", x[12]), ("Assumptions to validate", x[13])])

    doc.add_heading("8. Work models", 1)
    for idx, (key, path_) in enumerate(models.items(), start=1):
        m.add_figure(doc, path_, f"UR-WM-{idx:02d}", path_.stem.replace("final-", "").replace("-", " ").title(), "Captured screen evidence + simulated participant behavior + PA1 inherited finding + proposed response", "Task protocol, personas, findings, and PA3 validation", 16.0, 16.5)

    doc.add_heading("9. Findings", 1)
    m.add_figure(doc, DIAGRAMS / "evidence-crops" / "f2-e09-decision-surface.png", "UR-EV-01", "FIFA ticket decision surface used by UR-F01 and UR-F02", "Screen evidence F2-E09; captured interface state only", "Ticket status clarity, freshness, handoff trust, and TP-FIFA", 15.5, 9.0)
    m.add_figure(doc, DIAGRAMS / "evidence-crops" / "c2-e10-decision-surface.png", "UR-EV-02", "Chess analysis-entry decision surface used by UR-C01 and UR-C02", "Screen evidence C2-E10; captured interface state only", "Entry choice overload, learning priority, vocabulary load, and TP-CHESS", 15.5, 9.0)
    landscape(doc)
    findings = [
        ["UR-F01", "F2-E09 ticket entry/card state", "F-SIM-04/05/06: status or freshness hesitation", "Trust-sensitive planners", "Sale-window planning", "Compare state and freshness", "Visibility; effectiveness; errors", "Critical", "State, freshness, and action are separated", "Waiting or unsafe departure", "Status Dashboard with timestamp and missing/stale distinction", "Scenario-based; live feed not inspected", "Can users distinguish missing, stale, sold out, resale, and waiting room?"],
        ["UR-F02", "F2-E09 outbound ticket path", "F-SIM-02/05: handoff confidence 2-3/5", "Ticket planners", "Before leaving FIFA.com", "Preview destination", "Satisfaction; user control; error prevention", "High", "Destination purpose and return path are not one decision object", "Abandonment or uncalibrated trust", "Official destination preview with stay/continue/return", "No transaction performed", "Does preview improve decision accuracy without overstating safety?"],
        ["UR-F03", "F2-E02 mobile stack; F2-E03/04 schedule context", "F-SIM-03/04: context recovery", "Tournament followers", "Desk-to-mobile continuation", "Resume mobile", "Efficiency; memorability", "Medium", "Context is not visibly persisted", "Repeat search and filter work", "Restore selected tournament and freshness", "Account sync assumed", "Can users resume with one recognition step?"],
        ["UR-C01", "C2-E10 analysis entry choices", "C-SIM-01/02/05: 2 wrong paths in several records", "First-time and low-analysis players", "After-game review", "Open analysis", "Learnability; recognition over recall; errors", "Critical", "Entry is organized by feature/setup rather than learning outcome", "Expert view opens before a learning priority", "Beginner Review Preset", "No completed engine review observed", "Can beginners choose the review route without terminology help?"],
        ["UR-C02", "C2-E07/08 learning routes and C2-E10 analysis entry", "C-SIM-02/04/06: explanation priority", "Beginners and returning learners", "Review after game", "Identify mistake and read explanation", "Effectiveness; satisfaction; progressive disclosure", "High", "Advanced detail competes with plain explanation", "Numbers are seen before meaning", "One mistake card, why, try move, then depth", "Explanation accuracy unmeasured", "Can users explain the mistake in their own words?"],
        ["UR-C03", "C2-E05/07 practice artifacts", "C-SIM-03/04/06: continuation chosen after review", "Returning learners", "Short practice session", "Continue to practice", "Efficiency; memorability", "Medium", "Review and practice are separate destinations", "Insight does not become deliberate practice", "Relevant lesson/puzzle continuation", "Recommendation relevance is hypothetical", "Does one next practice action increase correct continuation?"],
    ]
    m.add_table(doc, ["ID", "Visible evidence", "Simulated study support", "Affected user", "Context", "Task", "Usability dimension", "Severity", "Root cause", "Consequence", "Design implication", "Limitation", "PA3 validation question"], findings, [1.3, 3.2, 3.5, 2.7, 2.8, 2.6, 2.8, 1.5, 3.4, 3.0, 3.6, 3.0, 4.0])
    portrait(doc)
    doc.add_heading("10. Evidence boundary and citations", 1)
    source_line(doc, "Screen evidence", "F2-E02-E04 and F2-E09 support FIFA navigation/status claims; C2-E05, C2-E07-E10 support Chess learning and analysis-entry claims.")
    source_line(doc, "PA1 inherited finding", "The product pair, drawbacks, and evidence limits remain continuous with PA1.")
    source_line(doc, "Simulated study", "Twelve scenario records generate hypotheses and traceable synthetic metrics only.")
    source_line(doc, "Design inference", "Effectiveness, efficiency, learnability, memorability, errors, satisfaction, visibility, recognition, progressive disclosure, control, and prevention guide the implications.")
    source_line(doc, "PA3 validation target", "Prototype testing must replace synthetic expectations with observed user evidence.")
    doc.add_paragraph(COURSE)
    return save(doc, "Group10-PA2-UserResearch")


FIFA_CLUSTERS = {
    "Ticket status clarity": ["Current sale state needs one label", "Missing state differs from sold out", "Resale must not look like primary sale", "Waiting room is a temporary state", "Hospitality is a distinct option", "State comparison belongs in one view"],
    "Handoff trust": ["Preview the destination domain", "Explain why departure is needed", "Offer stay on FIFA.com", "Preserve a safe return route", "Do not imply partner transaction status", "Keep tournament context through handoff"],
    "Mobile overview": ["Keep selected tournament first", "Avoid a long undifferentiated card stack", "Restore filters on return", "Use labels that survive small screens", "Make status scannable without color", "Resume with one recognition step"],
    "Browse load": ["Ticket entry competes with content", "Repeated cards increase scanning", "Tournament identity should anchor actions", "Reduce cross-page comparison", "Keep action wording consistent", "Separate browse from purchase intent"],
    "Freshness and alerts": ["Show last-updated time", "Name the governed source", "Label stale data explicitly", "Allow relevant change alerts", "Explain notification permission", "Let users edit or stop alerts"],
    "Recovery and return": ["Retry missing state without losing context", "Network loss should retain the last snapshot", "Login mismatch needs safe recovery", "Expired return links need explanation", "Partner failure must offer stay", "Never infer purchase completion"],
}

CHESS_CLUSTERS = {
    "Entry choice overload": ["Beginner review needs one named entry", "Analysis setup paths compete", "Feature names require recall", "Completed game should expose review", "Expert view opens too early", "Outcome labels are easier to recognize"],
    "Learning priority": ["Show one important mistake first", "Explain why before numeric depth", "Make the next action explicit", "Keep the lesson goal visible", "Do not present every candidate equally"],
    "Review setup": ["No game needs a clear setup route", "Invalid PGN must preserve input", "Board orientation needs confirmation", "Engine failure needs a safe fallback", "Premium restriction needs honest scope"],
    "Vocabulary load": ["Evaluation terms need inline help", "Plain language should lead", "Terminology help should not interrupt", "Returning learners forget feature names", "Recognition should replace recall"],
    "Mobile scanning": ["Keep the board and main mistake together", "Collapse secondary panels", "Use color-independent indicators", "Preserve state after interruption", "Maintain readable tap targets"],
    "Practice continuation": ["Link the mistake to one lesson", "Offer one relevant puzzle", "Preserve the reviewed position", "Let users continue later", "Explain why the practice is relevant"],
    "Advanced depth disclosure": ["Reveal engine detail on demand", "Do not remove expert access", "Remember the disclosure choice", "Keep beginner explanation stable", "Separate learning priority from depth"],
}


def raw_notes():
    source_types = ["screen evidence", "simulated participant behavior", "PA1 inherited finding", "design inference"]
    rows = []
    for product, clusters in [("FIFA", FIFA_CLUSTERS), ("Chess", CHESS_CLUSTERS)]:
        n = 1
        for cluster, notes in clusters.items():
            for note in notes:
                rows.append([f"{product[0]}-RN-{n:02d}", product, cluster, note, source_types[(n - 1) % 4]])
                n += 1
    return rows


def build_user_analysis() -> Path:
    doc = m.base_document("User Analysis")
    cover(doc, "Group10-PA2 User Analysis", "Clearly labeled simulated brainstorming, affinity analysis, voting, and prioritization")
    doc.add_heading("1. Analysis boundary and source discipline", 1)
    m.callout(doc, "Simulated brainstorming record", "The workshop, raw-note generation, votes, discussion, attendance, and consensus below are scenario-based synthetic evidence prepared for course reporting. They are not contemporaneous meeting evidence.")
    source_line(doc, "Screen evidence", "Visible interface states seed observable notes only.")
    source_line(doc, "Simulated study", "Participant-behavior notes and votes are synthetic scenario records.")
    source_line(doc, "Design inference", "Clusters and priorities are interpretations that require PA3 testing.")

    doc.add_heading("2. Simulated brainstorming session", 1)
    kv(doc, [
        ("Record type", "Simulated brainstorming record"), ("Date and duration", "22 Jul 2026, 19:30-20:45 (75 minutes)"),
        ("Participants", "; ".join(TEAM)), ("Objective", "Transform PA1 continuity, screen evidence, and simulated-study patterns into prioritized PA2 problems."),
        ("Silent note generation", "12 minutes; each member drafts atomic notes independently."),
        ("Round-robin sharing", "16 minutes; one note per turn, no evaluation during sharing."),
        ("Clustering", "18 minutes; group by user task and breakdown rather than feature."),
        ("Duplicate merging", "8 minutes; retain the clearest atomic wording and all source tags."),
        ("Cluster naming", "6 minutes; name clusters with user-facing task language."),
        ("Prioritization", "8 minutes; six simulated votes per member per product."),
        ("Consensus", "5 minutes; review weighted scores and resolve the selected problem statement."),
        ("Owner assignment", "2 minutes; Bach owns FIFA synthesis, Bao owns Chess synthesis, Nhut owns diagrams/HCI mapping, Le owns integration/traceability."),
    ])

    doc.add_heading("3. Raw notes: 36 FIFA and 36 Chess", 1)
    landscape(doc)
    m.add_table(doc, ["Note ID", "Product", "Cluster", "Atomic note", "Source type"], raw_notes(), [2.0, 1.8, 4.0, 10.8, 4.2])
    source_line(doc, "Scenario-based synthetic evidence", "All 72 notes are one-idea records; simulated participant behavior notes are not real observation claims.")
    portrait(doc)

    doc.add_heading("4. Affinity diagrams and cluster synthesis", 1)
    fifa_img = m.flow_diagram("final-affinity-fifa", "UA-AF-FIFA Affinity Diagram", list(FIFA_CLUSTERS), ["36 atomic notes", "6 clusters", "Status + handoff lead", "All notes retained in Table 3"], "#1D70A2")
    chess_img = m.flow_diagram("final-affinity-chess", "UA-AF-CHESS Affinity Diagram", list(CHESS_CLUSTERS), ["36 atomic notes", "7 clusters", "Guided review synthesis", "All notes retained in Table 3"], "#4F7F35")
    m.add_figure(doc, fifa_img, "UA-AF-01", "FIFA affinity clusters", "Screen evidence + simulated participant behavior + PA1 inherited finding + design inference", "36 FIFA raw notes", 16, 16)
    m.add_figure(doc, chess_img, "UA-AF-02", "Chess affinity clusters", "Screen evidence + simulated participant behavior + PA1 inherited finding + design inference", "36 Chess raw notes", 16, 16)
    landscape(doc)
    cluster_rows = []
    affected = {
        "Ticket status clarity": "Compare ticket states", "Handoff trust": "Preview destination", "Mobile overview": "Resume on mobile", "Browse load": "Find ticket entry", "Freshness and alerts": "Inspect freshness / subscribe", "Recovery and return": "Recover and return",
        "Entry choice overload": "Open review", "Learning priority": "Identify mistake", "Review setup": "Open game", "Vocabulary load": "Read explanation", "Mobile scanning": "Review on mobile", "Practice continuation": "Continue practice", "Advanced depth disclosure": "Reveal detail",
    }
    for product, clusters in [("FIFA", FIFA_CLUSTERS), ("Chess", CHESS_CLUSTERS)]:
        for name, notes in clusters.items():
            severity = "Critical" if name in ("Ticket status clarity", "Entry choice overload") else "High" if name in ("Handoff trust", "Learning priority", "Vocabulary load", "Practice continuation") else "Medium"
            direction = {
                "Ticket status clarity": "One comparison dashboard", "Handoff trust": "Destination preview + stay/return", "Mobile overview": "Persist tournament context", "Browse load": "Task-led entry", "Freshness and alerts": "Timestamp + controlled alerts", "Recovery and return": "State-specific recovery",
                "Entry choice overload": "Beginner review preset", "Learning priority": "One mistake first", "Review setup": "Validated setup states", "Vocabulary load": "Inline plain-language help", "Mobile scanning": "Primary-card hierarchy", "Practice continuation": "One relevant next exercise", "Advanced depth disclosure": "Progressive disclosure",
            }[name]
            cluster_rows.append([product, name, str(len(notes)), "; ".join(notes[:2]), "Repeated task-state or orientation breakdown", affected[name], severity, direction])
    m.add_table(doc, ["Product", "Cluster", "Count", "Representative notes", "Pattern", "Affected task", "Severity", "Design direction"], cluster_rows, [1.8, 4.0, 1.3, 7.2, 4.0, 3.2, 1.8, 4.2])
    portrait(doc)

    doc.add_heading("5. Simulated dot voting", 1)
    m.callout(doc, "Simulated brainstorming record", "Each member receives six votes for FIFA and six for Chess; totals equal 24 votes per product.")
    fifa_votes = [
        ["Le Minh", "3", "2", "0", "0", "1", "0", "6"], ["Nguyen Vu Bach", "3", "2", "0", "0", "0", "1", "6"],
        ["Pham Nguyen Gia Bao", "2", "2", "0", "1", "1", "0", "6"], ["Trang Minh Nhut", "2", "2", "1", "0", "1", "0", "6"],
    ]
    m.add_table(doc, ["Member", "Status", "Handoff", "Mobile", "Browse", "Freshness", "Recovery", "Total"], fifa_votes, [3.5, 1.7, 1.7, 1.7, 1.7, 1.7, 1.7, 1.7])
    m.add_table(doc, ["FIFA cluster", "Total votes"], [["Ticket status clarity", "10"], ["Handoff trust", "8"], ["Freshness and alerts", "3"], ["Mobile overview", "1"], ["Browse load", "1"], ["Recovery and return", "1"]], [11, 4])
    chess_votes = [
        ["Le Minh", "1", "0", "0", "2", "0", "2", "1", "6"], ["Nguyen Vu Bach", "1", "1", "1", "1", "0", "1", "1", "6"],
        ["Pham Nguyen Gia Bao", "1", "1", "1", "1", "0", "2", "0", "6"], ["Trang Minh Nhut", "1", "1", "1", "1", "1", "0", "1", "6"],
    ]
    m.add_table(doc, ["Member", "Entry", "Priority", "Setup", "Vocabulary", "Mobile", "Practice", "Depth", "Total"], chess_votes, [3.2, 1.5, 1.5, 1.5, 1.7, 1.5, 1.7, 1.5, 1.5])
    m.add_table(doc, ["Chess cluster", "Total votes"], [["Vocabulary load", "5"], ["Practice continuation", "5"], ["Entry choice overload", "4"], ["Learning priority", "3"], ["Review setup", "3"], ["Advanced depth disclosure", "3"], ["Mobile scanning", "1"]], [11, 4])
    kv(doc, [
        ("Discussion notes", "FIFA votes concentrate on status and handoff confidence. Chess votes split across vocabulary and continuation, while entry overload is the enabling problem."),
        ("Tie-break rule", "If totals tie, select the item with higher severity, then stronger traceability to the PA3 prototype task."),
        ("Consensus statement", "Prioritize FIFA ticket status plus handoff confidence, and treat Chess entry, explanation, and continuation as one guided beginner review problem."),
        ("Selected problems", "FIFA: ticket status clarity and handoff trust. Chess: guided beginner review."),
    ])

    doc.add_heading("6. Weighted prioritization", 1)
    m.callout(doc, "Formula", "Weighted score = Severity x 25% + Task frequency x 15% + Risk x 20% + Reach x 10% + Evidence strength x 15% + PA3 relevance x 15%. Each factor is scored 1-5; maximum is 5.00.")
    prioritization = [
        ["FIFA status + freshness", "5", "4", "5", "4", "4", "5", "4.60", "Critical ambiguity; frequent planning; trust risk; broad reach; strong screen/synthetic trace; direct PA3 module."],
        ["FIFA handoff confidence", "5", "3", "5", "4", "4", "5", "4.45", "Critical departure risk; lower frequency; broad consequence; direct PA3 decision."],
        ["FIFA mobile overview", "3", "4", "3", "4", "4", "3", "3.50", "Frequent but lower consequence; evidence supports orientation rather than transaction risk."],
        ["Chess guided beginner review", "5", "5", "4", "5", "4", "5", "4.65", "High learning severity/frequency/reach; strong PA3 relevance; synthetic outcomes require validation."],
        ["Chess mobile scanning", "3", "4", "3", "4", "3", "3", "3.35", "Broad but secondary to entry and learning-priority breakdown."],
        ["Chess advanced depth", "3", "2", "3", "3", "3", "4", "3.00", "Important control but lower beginner frequency; retained as a disclosure requirement."],
    ]
    landscape(doc)
    m.add_table(doc, ["Candidate", "Severity", "Frequency", "Risk", "Reach", "Evidence", "PA3", "Weighted", "Score rationale"], prioritization, [4.0, 1.5, 1.6, 1.3, 1.3, 1.6, 1.3, 1.8, 10.0])
    portrait(doc)

    doc.add_heading("7. Tough problems", 1)
    tough = [
        ("TP-FIFA", "Users need to compare current ticket state, freshness and official destination before leaving FIFA.com.", "Trust-sensitive planners and tournament followers", "Ticket planning before an outbound handoff", "F2-E03/E04/E09", "F-SIM-02/04/05/06 show status or handoff hesitation and 2-3/5 confidence", "State, freshness, and destination are separated", "Delayed action, unsafe assumption, or abandonment", "Official sports sites rarely combine governed state, freshness, and destination preview in one decision view", "Comparison, preview, alerts, and recovery inside FIFA.com", "Transaction, inventory creation, pricing, or partner redesign", "Can users accurately choose buy, wait, register interest, resale, hospitality, or stay?"),
        ("TP-CHESS", "Beginners need one interpretable mistake and one next practice action before advanced analysis.", "First-time beginners and returning low-analysis learners", "After-game review and practice continuation", "C2-E05/E07-E10", "C-SIM-01/02/04/05/06 show wrong paths, vocabulary hesitation, and continuation need", "Feature/setup organization precedes the learning outcome", "Expert view without an actionable learning priority", "Many analysis tools optimize depth before beginner interpretation and continuation", "Beginner preset, one mistake, try move, explanation, practice, optional depth", "Engine redesign, coaching guarantees, or removal of expert controls", "Can users explain one mistake and select one relevant practice action?"),
    ]
    for row in tough:
        doc.add_heading(f"7.{1 if row[0]=='TP-FIFA' else 2} {row[0]}", 2)
        m.callout(doc, row[0], row[1], m.FIFA_LIGHT if row[0] == "TP-FIFA" else m.CHESS_LIGHT)
        kv(doc, list(zip(["Affected users", "Context", "Screen evidence", "Simulated study result", "Root cause", "Consequence", "Market gap", "Scope", "Out of scope", "PA3 test question"], row[2:])))
    doc.add_heading("8. Design rationale and validation target", 1)
    source_line(doc, "HCI rationale", COURSE)
    source_line(doc, "Design inference", "The selection balances effectiveness, efficiency, learnability, memorability, errors, satisfaction, status visibility, recognition, progressive disclosure, user control, and error prevention.")
    return save(doc, "Group10-PA2-UserAnalysis")


ALTERNATIVES = [
    ("F-A1", "Status Dashboard", "FIFA Ticket Confidence Module", "Trust-sensitive ticket planners", "Choose context -> compare states -> inspect freshness -> preview official destination -> act or stay", "Status comparison dashboard", "Filter, expand provenance, preview handoff, save/alert", "Directly covers status, freshness, and destination confidence", "Text labels, table semantics, keyboard order, color-independent states", "Minimize saved preferences; explicit alert consent", "Governed ticket-state feed, timestamps, partner metadata", "Keep last timestamped state, distinguish missing/stale, retry without losing context", "Feed inconsistency or false confidence", "Highest clarity; requires reliable governed data", "Test state interpretation, destination recognition, stay/continue choice, and recovery"),
    ("F-A2", "Guided Concierge", "FIFA Ticket Confidence Module", "Infrequent and first-time planners", "State intent -> answer questions -> receive route -> preview destination -> confirm", "Step-by-step concierge", "Single-choice questions, back/edit, recommendation rationale", "Covers complex choice and handoff confidence", "Progress, plain labels, keyboard/radio semantics", "Avoid sensitive profiling; delete draft answers", "Eligibility rules, ticket channels, partner metadata", "Explain unsupported markets, preserve answers, offer dashboard", "Long flow and rule maintenance", "More guidance; slower comparison", "Test completion, backtracking, recommendation comprehension, and unsupported-state recovery"),
    ("F-A3", "Alert-First Planner", "FIFA Ticket Confidence Module", "Followers waiting for sale changes", "Follow context -> set preference -> wait -> receive verified change -> open destination", "Saved-plan and alert center", "Follow, consent, edit, snooze, stop, open alert", "Covers freshness and return; weaker immediate comparison", "Channel alternatives, readable timestamps, no color-only urgency", "Explicit permission, minimal contact data, retention limit", "Change events, notification channel, partner metadata", "Retain local plan if permission denied; retry irrelevant or failed alerts", "Notification fatigue and delayed data", "Strong return support; depends on alert quality", "Test permission denial, alert relevance, freshness comprehension, and return context"),
    ("C-A1", "Beginner Review Preset", "Chess Guided Beginner Review Module", "First-time and returning low-analysis learners", "Open game -> select preset -> identify one mistake -> try move -> read explanation -> practice or depth", "Primary mistake card beside board", "Try move, reveal why, continue practice, reveal advanced detail", "Directly covers one interpretable mistake and next practice action", "Keyboard board alternative, focus order, plain language, color-independent evaluation", "Use only selected game; explain retention and sharing", "Game artifact, analysis service, explanation rules, learning catalog", "Handle no game, invalid PGN, engine unavailable, no major mistake, restriction", "Oversimplification or incorrect explanation", "Strongest beginner interpretation; constrains initial depth", "Test mistake comprehension, move trial, practice selection, and depth disclosure"),
    ("C-A2", "Conversational Coach", "Chess Guided Beginner Review Module", "Beginners who prefer question-led guidance", "Open review -> state intent -> coach clarifies -> explain position -> try move -> practice", "Conversation plus board", "Choose intent, ask why, try move, request another explanation", "Covers interpretation and vocabulary", "Structured choices, transcript headings, keyboard board alternative", "Do not collect unrelated personal data; clear conversation retention", "Game artifact, explanation service, learning catalog", "Offer fixed explanation cards if the coach service fails", "Wording can vary and the extra question adds time", "The question-led route can address several intentions, but response consistency must be tested.", "Test response consistency, vocabulary comprehension, and service fallback"),
    ("C-A3", "Visual Game Story", "Chess Guided Beginner Review Module", "Returning learners who scan game turning points", "Open review -> view chapters -> inspect turning point -> replay chance -> choose next action", "Chapter timeline with board snapshots", "Scrub chapters, select turning point, replay, continue", "Covers prioritization and narrative continuation", "Text chapter list, reduced motion, keyboard timeline", "Use game data only; respect game visibility", "Game segmentation, analysis service, learning catalog", "Fall back to a single mistake card if segmentation fails", "Long games and multiple comparable mistakes", "Strong overview; less direct for first-time beginners", "Test turning-point recognition, replay accuracy, and continuation choice"),
]


def build_project_proposal() -> Path:
    doc = m.base_document("Project Proposal")
    cover(doc, "Group10-PA2 Project Proposal", "Two bounded modules, six alternatives, weighted recommendation, and PA3 test plans")
    doc.add_heading("1. Proposal scope and traceability", 1)
    m.callout(doc, "FIFA Ticket Confidence Module", "Users need to compare current ticket state, freshness and official destination before leaving FIFA.com.", m.FIFA_LIGHT)
    m.callout(doc, "Chess Guided Beginner Review Module", "Beginners need one interpretable mistake and one next practice action before advanced analysis.", m.CHESS_LIGHT)
    source_line(doc, "Screen evidence", "F2-E03/E04/E09 and C2-E05/E07-E10 bound the visible interface claims.")
    source_line(doc, "Simulated study", "Scenario records inform hypothesis coverage only; no expected benefit is an observed outcome.")
    source_line(doc, "Design inference", "The modules preserve current products and introduce bounded decision/learning layers.")

    doc.add_heading("2. FIFA Ticket Confidence Module alternatives", 1)
    for i, row in enumerate(ALTERNATIVES[:3], start=1):
        doc.add_heading(f"2.{i} {row[0]} {row[1]}", 2)
        fields = ["Concept", "Target user", "Task flow", "Main screen", "Interaction", "Problem coverage", "Accessibility", "Privacy", "Data dependency", "Failure recovery", "Risk", "Tradeoff", "PA3 test plan"]
        kv(doc, list(zip(fields, [f"{row[1]} within {row[2]}", *row[3:]])))
        flow = m.flow_diagram(f"final-{row[0].lower()}", f"PP-{row[0]} {row[1]} Flow", row[5].split(", ") if ", " in row[5] else row[4].split(" -> "), ["Screen evidence", "Simulated study", "Design inference", "PA3 validation target"], "#1D70A2")
        m.add_figure(doc, flow, f"PP-{row[0]}", f"{row[1]} task-flow model", "Design inference based on TP-FIFA", row[-1], 15.5, 13.0)
    doc.add_heading("3. Chess Guided Beginner Review Module alternatives", 1)
    for i, row in enumerate(ALTERNATIVES[3:], start=1):
        doc.add_heading(f"3.{i} {row[0]} {row[1]}", 2)
        fields = ["Concept", "Target user", "Task flow", "Main screen", "Interaction", "Problem coverage", "Accessibility", "Privacy", "Data dependency", "Failure recovery", "Risk", "Tradeoff", "PA3 test plan"]
        kv(doc, list(zip(fields, [f"{row[1]} within {row[2]}", *row[3:]])))
        flow = m.flow_diagram(f"final-{row[0].lower()}", f"PP-{row[0]} {row[1]} Flow", row[4].split(" -> "), ["Screen evidence", "Simulated study", "Design inference", "PA3 validation target"], "#4F7F35")
        m.add_figure(doc, flow, f"PP-{row[0]}", f"{row[1]} task-flow model", "Design inference based on TP-CHESS", row[-1], 15.5, 13.0)

    doc.add_heading("4. Weighted comparison method", 1)
    m.callout(doc, "Formula", "Weighted score = Problem coverage x 30% + Learnability x 20% + Accessibility x 15% + Recovery x 15% + Data feasibility x 10% + PA3 testability x 10%. Factors use 1-5; maximum is 5.00.")
    landscape(doc)
    comparison = [
        ["F-A1 Status Dashboard", "5", "4", "5", "5", "3", "5", "4.55", "Direct comparison and recovery; reliable feed is the main dependency."],
        ["F-A2 Guided Concierge", "4", "5", "4", "4", "3", "4", "4.05", "High guidance but a longer flow and heavier rule maintenance."],
        ["F-A3 Alert-First Planner", "3", "4", "4", "4", "3", "4", "3.65", "Strong freshness/return support but weaker immediate comparison."],
        ["C-A1 Beginner Review Preset", "5", "5", "5", "5", "4", "5", "4.90", "Best alignment to one mistake, move trial, explanation, and practice."],
        ["C-A2 Conversational Coach", "4", "5", "4", "3", "2", "3", "3.75", "The extra question may clarify intent, but wording consistency and service fallback still need testing."],
        ["C-A3 Visual Game Story", "4", "4", "4", "4", "3", "4", "3.90", "Good overview but less direct for a first-time beginner."],
    ]
    m.add_table(doc, ["Alternative", "Coverage 30%", "Learnability 20%", "Accessibility 15%", "Recovery 15%", "Data 10%", "PA3 10%", "Weighted", "Reason"], comparison, [4.8, 2.1, 2.3, 2.3, 2.0, 1.5, 1.5, 1.9, 9.0])
    portrait(doc)

    doc.add_heading("5. Recommendations and retained alternatives", 1)
    m.callout(doc, "Recommendation - FIFA", "Select Status Dashboard because it most directly combines ticket state, freshness, official destination preview, user control, and recovery. Guided Concierge and Alert-First Planner remain PA3 parallel prototypes.", m.FIFA_LIGHT)
    m.callout(doc, "Recommendation - Chess", "Select Beginner Review Preset because it establishes one interpretable mistake and one next practice action before advanced depth. Conversational Coach and Visual Game Story remain PA3 parallel prototypes.", m.CHESS_LIGHT)
    m.add_table(doc, ["Recommendation", "Reason", "Principal risk", "PA3 validation target"], [
        ["F-A1 Status Dashboard", "Highest weighted FIFA score and strongest tough-problem coverage", "Stale or conflicting governed data", "Accurate state/freshness interpretation and calibrated handoff confidence"],
        ["C-A1 Beginner Review Preset", "Highest weighted Chess score and clearest learning priority", "Oversimplification or explanation error", "Mistake comprehension, better-move trial, and relevant practice choice"],
    ], [4.2, 5.0, 4.2, 5.0])
    doc.add_heading("6. PA3 parallel-prototyping plan", 1)
    m.add_table(doc, ["Product", "Parallel concepts", "Common scenario", "Measures", "Stop/iterate rule"], [
        ["FIFA", "Status Dashboard; Guided Concierge; Alert-First Planner", "Compare a missing/stale ticket state and decide whether to leave FIFA.com", "State accuracy; destination recognition; confidence calibration; recovery", "Reject a concept that causes state misinterpretation or implies partner completion"],
        ["Chess", "Beginner Review Preset; Conversational Coach; Visual Game Story", "Open a completed game, identify one mistake, try a better move, continue practice", "Explanation accuracy; wrong path; continuation relevance; depth control", "Reject a concept that hides correction, blocks keyboard use, or overstates certainty"],
    ], [2.5, 5.2, 5.0, 4.5, 4.8])
    source_line(doc, "PA3 validation target", "Expected benefits remain hypotheses until prototype sessions produce observed evidence.")
    doc.add_heading("7. Design rationale and validation target", 1)
    source_line(doc, "HCI rationale", COURSE)
    return save(doc, "Group10-PA2-ProjectProposal")


def uml_v2(name: str, title: str, boundary: str, actor: str, externals: list[str], cases: list[str], includes: list[tuple[int, int]], extends: list[tuple[int, int]]):
    DIAGRAMS.mkdir(parents=True, exist_ok=True)
    w, h = 2400, 1750
    image = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(image)
    accent = "#1D70A2" if "FIFA" in title else "#4F7F35"
    draw.text((70, 45), title, font=m.font(48, True), fill="#123B65")
    draw.rectangle((480, 160, 1880, 1600), outline=accent, width=6)
    draw.text((520, 185), boundary, font=m.font(30, True), fill=accent)

    def actor_shape(cx: int, cy: int, label: str):
        draw.ellipse((cx - 34, cy - 105, cx + 34, cy - 37), outline="#1F2937", width=5)
        draw.line((cx, cy - 37, cx, cy + 70), fill="#1F2937", width=5)
        draw.line((cx - 62, cy, cx + 62, cy), fill="#1F2937", width=5)
        draw.line((cx, cy + 70, cx - 58, cy + 145), fill="#1F2937", width=5)
        draw.line((cx, cy + 70, cx + 58, cy + 145), fill="#1F2937", width=5)
        for i, line in enumerate(m.wrap(label, 18)):
            box = draw.textbbox((0, 0), line, font=m.font(24, True))
            draw.text((cx - (box[2]-box[0])/2, cy + 165 + i*28), line, font=m.font(24, True), fill="#1F2937")

    actor_shape(210, 760, actor)
    for i, ext in enumerate(externals):
        actor_shape(2150, 480 + i*600, ext)
    centers = []
    for i, label in enumerate(cases):
        row, col = divmod(i, 2)
        cx, cy = 830 + col*650, 410 + row*410
        draw.ellipse((cx-250, cy-82, cx+250, cy+82), fill="#F7FAFC", outline=accent, width=4)
        lines = m.wrap(label, 28)
        y = cy - len(lines)*16
        for line in lines:
            box = draw.textbbox((0, 0), line, font=m.font(24, True))
            draw.text((cx-(box[2]-box[0])/2, y), line, font=m.font(24, True), fill="#1F2937")
            y += 32
        draw.line((275, 760, cx-250, cy), fill="#64748B", width=3)
        centers.append((cx, cy))
    for i, ext in enumerate(externals):
        target = centers[min(len(centers)-1, 1+i*3)]
        draw.line((target[0]+250, target[1], 2075, 480+i*600), fill="#64748B", width=3)

    def dashed_arrow(a, b, label):
        ax, ay = centers[a]
        bx, by = centers[b]
        dx, dy = bx-ax, by-ay
        dist = max(1, math.hypot(dx, dy))
        ux, uy = dx/dist, dy/dist
        start = (ax + ux*150, ay + uy*55)
        end = (bx - ux*150, by - uy*55)
        parts = 16
        for i in range(0, parts, 2):
            x1 = start[0] + (end[0]-start[0])*i/parts
            y1 = start[1] + (end[1]-start[1])*i/parts
            x2 = start[0] + (end[0]-start[0])*(i+1)/parts
            y2 = start[1] + (end[1]-start[1])*(i+1)/parts
            draw.line((x1, y1, x2, y2), fill="#7C3AED", width=4)
        draw.polygon([(end[0], end[1]), (end[0]-ux*25-uy*12, end[1]-uy*25+ux*12), (end[0]-ux*25+uy*12, end[1]-uy*25-ux*12)], fill="#7C3AED")
        mx, my = (start[0]+end[0])/2, (start[1]+end[1])/2
        draw.rectangle((mx-72, my-23, mx+72, my+23), fill="white")
        draw.text((mx-65, my-16), label, font=m.font(20, True), fill="#7C3AED")
    for a, b in includes:
        dashed_arrow(a, b, "<<include>>")
    for a, b in extends:
        dashed_arrow(a, b, "<<extend>>")
    out = DIAGRAMS / f"{name}.png"
    image.save(out, dpi=(220, 220))
    return out


USE_CASES = [
    {"id":"F-UC01","name":"Select tournament context","scope":"FIFA Ticket Confidence Module","actor":"Ticket planner","support":"FIFA tournament catalog","trigger":"User opens the module from a tournament or ticket entry.","pre":"Module is available; a catalog response or safe cached list exists.","guarantee":"A labeled tournament/team context is selected and preserved.","main":["Open the tournament selector.","Review searchable tournament and team labels.","Choose a tournament and optional team.","Confirm the selected context.","Load the ticket-state comparison for that context."],"alt":["A1 at step 2: Catalog is unavailable; show the last timestamped list and a retry action.","A2 at step 3: No context matches; clear filters and preserve the query."],"exc":["E1: Network loss retains the selection locally and labels the state offline.","E2: Login mismatch offers account switch without exposing another account's plan."],"rules":"Tournament identity and freshness must be visible; no state is inferred from an absent catalog row.","data":"Tournament ID, team ID, locale, selected context, timestamp.","usability":"Recognition over recall; selected context stays visible.","access":"Keyboard-operable search and selection; labels are announced.","privacy":"Selection may remain local until the user explicitly saves it.","priority":"High","frequency":"Each new planning context","trace":"PA1 FIFA drawback -> F2-E03/E04/E09 -> F-SES-03/04 -> UR-P-F2 -> UR-F03 -> Mobile overview -> 1 vote -> TP-FIFA -> F-A1 -> PA3 context-resumption test."},
    {"id":"F-UC02","name":"Compare ticket states","scope":"FIFA Ticket Confidence Module","actor":"Ticket planner","support":"Governed ticket-state feed","trigger":"A tournament context is selected.","pre":"Context is valid; state feed may return current, stale, missing, sold out, resale, waiting room, or hospitality states.","guarantee":"Available states are compared without treating missing data as an outcome.","main":["Request states for the selected context.","Display each channel with a text state label.","Compare primary sale, resale, hospitality, wait, and interest options.","Explain unavailable and sold-out distinctions.","Choose a state to inspect or stay on the comparison."],"alt":["A1 at step 2: State is stale; retain it with timestamp and warning.","A2 at step 2: State is missing; show unknown and retry, not sold out.","A3 at step 3: Waiting room is active; explain temporary queue behavior."],"exc":["E1: Feed conflict shows both governed sources and blocks a confident action.","E2: Network loss retains the last labeled snapshot."],"rules":"Sold out, resale, waiting room, missing, and stale are mutually explicit labels.","data":"Context, channel, state, source, timestamp, eligibility note.","usability":"Visibility of system status and error prevention.","access":"State uses text and icon, never color alone; table headers are semantic.","privacy":"No personal data required for comparison.","priority":"Critical","frequency":"Every planning visit","trace":"PA1 FIFA ticket drawback -> F2-E09 -> F-SES-04/05/06 -> UR-P-F1 -> UR-F01 -> Ticket status clarity -> 10 votes -> TP-FIFA -> F-A1 -> recommendation -> PA3 state-accuracy test."},
    {"id":"F-UC03","name":"Inspect freshness","scope":"FIFA Ticket Confidence Module","actor":"Ticket planner","support":"Governed ticket-state feed","trigger":"User opens freshness details for a ticket state.","pre":"A state record exists, even if stale or missing.","guarantee":"Timestamp, source, and confidence boundary are visible.","main":["Open freshness details.","Read last-updated time and source.","Compare current time with the freshness policy.","Review stale or missing explanation.","Refresh or return to comparison."],"alt":["A1 at step 3: Record is stale; block unqualified action and offer refresh.","A2 at step 4: No update exists; show unknown and escalation guidance."],"exc":["E1: Refresh fails; keep prior timestamp and recovery action."],"rules":"A refresh attempt does not erase a valid prior timestamp.","data":"State ID, observed time, source ID, freshness threshold.","usability":"Clear feedback and calibrated confidence.","access":"Time is readable in local format and available to assistive technology.","privacy":"No additional personal data.","priority":"Critical","frequency":"At important sale decisions","trace":"PA1 evidence rules -> F2-E09 -> F-SES-04/06 -> UR-P-F1 -> UR-F01 -> Freshness and alerts -> 3 votes -> TP-FIFA -> F-A1 -> PA3 freshness test."},
    {"id":"F-UC04","name":"Preview destination","scope":"FIFA Ticket Confidence Module","actor":"Ticket planner","support":"Official ticketing partner","trigger":"User selects an outbound ticket action.","pre":"Destination metadata exists or the action remains blocked safely.","guarantee":"Domain, purpose, context transfer, and return behavior are previewed before departure.","main":["Open destination preview.","Read official partner name, domain, and purpose.","Review what context will transfer.","Choose Continue, Stay, or Return.","If Continue is chosen, open the destination with a return route."],"alt":["A1 at step 2: Metadata is partial; offer Stay and report the limitation.","A2 at step 4: User stays; preserve comparison context."],"exc":["E1: Partner failure returns the user to the preserved state with an explanation."],"rules":"The module never claims a partner transaction is complete.","data":"Partner ID, domain, purpose, context token, return URI.","usability":"User control, trust calibration, and error prevention.","access":"Destination and choices have explicit labels and logical focus order.","privacy":"Transfer only required context after consent; do not expose plan data in a URL.","priority":"Critical","frequency":"Before each outbound action","trace":"PA1 handoff drawback -> F2-E09 -> F-SES-02/05 -> UR-P-F1 -> UR-F02 -> Handoff trust -> 8 votes -> TP-FIFA -> F-A1 -> recommendation -> PA3 destination-recognition test."},
    {"id":"F-UC05","name":"Subscribe to alert","scope":"FIFA Ticket Confidence Module","actor":"Ticket planner","support":"Notification service","trigger":"User chooses to follow a state change.","pre":"A context is selected; permission is not assumed.","guarantee":"Alert preference is saved or a clear non-notification alternative is provided.","main":["Choose Subscribe to alert.","Select relevant state changes and channel.","Review frequency, privacy, and permission explanation.","Grant permission and confirm.","Edit, snooze, or stop the alert later."],"alt":["A1 at step 4: Notification is denied; keep a local saved plan and show manual refresh.","A2 at step 2: No channel is available; offer in-module status check."],"exc":["E1: Service failure preserves unsent preferences for retry."],"rules":"No alert is enabled without affirmative consent; every alert has an unsubscribe control.","data":"Context, selected changes, channel, consent timestamp, retention state.","usability":"Control, prevention of notification fatigue, and clear status.","access":"Permission explanation is plain text; controls are keyboard operable.","privacy":"Minimize contact data and disclose retention.","priority":"High","frequency":"During wait/interest states","trace":"PA1 ticket uncertainty -> F2-E09 -> F-SES-06 -> UR-P-F1 -> UR-F01 -> Freshness and alerts -> 3 votes -> TP-FIFA -> F-A3 -> retained PA3 prototype -> notification-denial test."},
    {"id":"F-UC06","name":"Continue, stay or return","scope":"FIFA Ticket Confidence Module","actor":"Ticket planner","support":"Official ticketing partner; identity service","trigger":"User acts from destination preview or reopens a return route.","pre":"Comparison context exists; any return token is validated.","guarantee":"User remains oriented and partner-side completion is never inferred.","main":["Choose Continue, Stay, or Return.","If Continue, transfer only consented context.","If Stay, preserve the comparison state.","If Return, validate token or identity.","Restore context, freshness, and available actions."],"alt":["A1 at step 4: Token expired; offer safe account lookup or plan recreation.","A2 at step 5: Login mismatch; switch account without exposing stored data."],"exc":["E1: Partner or network failure restores the last safe state."],"rules":"A return is not proof of purchase; tokens are time-limited and single-purpose.","data":"Context, return token, account reference, last state and timestamp.","usability":"Recovery, memorability, and persistent orientation.","access":"Restoration is announced without unexpected focus movement.","privacy":"Validate identity before revealing saved plans.","priority":"Critical","frequency":"Each handoff and return","trace":"PA1 handoff drawback -> F2-E09 -> F-SES-02/03/05/06 -> UR-P-F1/F2 -> UR-F02/F03 -> Recovery and return -> 1 vote -> TP-FIFA -> F-A1 -> PA3 return test."},
    {"id":"C-UC01","name":"Open game","scope":"Chess Guided Beginner Review Module","actor":"Beginner / returning learner","support":"Game store","trigger":"User chooses a recent game, PGN import, or board setup.","pre":"A source route is available; input may be absent or invalid.","guarantee":"A valid review artifact is opened or input is preserved with a clear correction path.","main":["Choose Recent game, Import PGN, or Set up position.","Select or enter the game artifact.","Confirm side to move and board orientation.","Validate notation and access rights.","Open the review entry."],"alt":["A1 at step 1: No game exists; emphasize import/setup routes.","A2 at step 4: PGN is invalid; identify the first error and preserve input."],"exc":["E1: Mobile interruption preserves the draft locally."],"rules":"Private games follow account permissions; invalid PGN is never silently changed.","data":"Game ID, PGN, ownership, side to move, orientation.","usability":"Clear setup choices and recoverable validation.","access":"Text notation alternative and fully labeled board controls.","privacy":"Respect game visibility and minimize copied metadata.","priority":"High","frequency":"Each review","trace":"PA1 analysis drawback -> C2-E10 -> C-SES-01/02/05 -> UR-P-C1 -> UR-C01 -> Review setup -> 3 votes -> TP-CHESS -> C-A1 -> recommendation -> PA3 setup test."},
    {"id":"C-UC02","name":"Select beginner review","scope":"Chess Guided Beginner Review Module","actor":"Beginner / returning learner","support":"Analysis service","trigger":"A valid game artifact is opened.","pre":"Review entry is visible; premium/access restrictions may apply.","guarantee":"The beginner preset is selected without removing advanced analysis access.","main":["Open review choices.","Read Beginner Review purpose and scope.","Select Beginner Review.","Confirm the game and perspective.","Start the guided review."],"alt":["A1 at step 2: Terminology help is requested; show inline explanation.","A2 at step 3: Premium restriction applies; explain available scope and alternatives."],"exc":["E1: Analysis service unavailable; offer retry and practice from the unreviewed game."],"rules":"Restrictions are stated before commitment; expert controls remain available by choice.","data":"Game artifact, preset, perspective, entitlement.","usability":"Recognition over recall and progressive disclosure.","access":"Keyboard-only selection and focus-visible controls.","privacy":"No extra personal data beyond game access.","priority":"Critical","frequency":"Each beginner review","trace":"PA1 analysis drawback -> C2-E10 -> C-SES-01/02/05 -> UR-P-C1 -> UR-C01 -> Entry choice overload -> 4 votes -> TP-CHESS -> C-A1 -> recommendation -> PA3 entry-choice test."},
    {"id":"C-UC03","name":"Identify mistake","scope":"Chess Guided Beginner Review Module","actor":"Beginner / returning learner","support":"Analysis service","trigger":"Guided review completes its first pass.","pre":"A review is available; there may be no major mistake or several comparable mistakes.","guarantee":"One learning priority is shown with honest uncertainty.","main":["Receive the guided review.","Show the most important mistake or state that no major mistake exists.","Display the position and move context.","Explain why this moment matters in plain language.","Offer Try a better move."],"alt":["A1 at step 2: No major mistake; show a useful improvement instead.","A2 at step 2: Multiple mistakes are comparable; explain the selection and allow comparison."],"exc":["E1: Engine result is partial; label uncertainty and avoid a definitive claim."],"rules":"One priority does not imply all other moves are irrelevant.","data":"Position, move, evaluation change, confidence, explanation key.","usability":"Effectiveness, clear hierarchy, and calibrated status.","access":"Text description accompanies board highlight; no color-only meaning.","privacy":"Game visibility rules continue to apply.","priority":"Critical","frequency":"Each review","trace":"PA1 learning/analysis gap -> C2-E10 -> C-SES-02/04/06 -> UR-P-C1/C2 -> UR-C02 -> Learning priority -> 3 votes -> TP-CHESS -> C-A1 -> recommendation -> PA3 mistake-comprehension test."},
    {"id":"C-UC04","name":"Try better move","scope":"Chess Guided Beginner Review Module","actor":"Beginner / returning learner","support":"Analysis service","trigger":"User chooses Try a better move.","pre":"A review position and legal move set exist.","guarantee":"The attempted move receives feedback without losing the original position.","main":["Open the review position in trial mode.","Enter a legal candidate move.","Compare it with the learning priority.","Receive concise feedback.","Retry, reveal the recommended move, or return."],"alt":["A1 at step 2: Move is illegal; explain the rule and keep focus on the board.","A2 at step 3: Move is also strong; acknowledge it and compare ideas."],"exc":["E1: Engine unavailable; record the trial locally and defer comparison."],"rules":"User action in trial mode never changes the stored game.","data":"Position, candidate move, legal moves, comparison result.","usability":"Immediate feedback, user control, and safe experimentation.","access":"Keyboard coordinate entry and announced feedback.","privacy":"Trial data stays with the selected review unless saved.","priority":"High","frequency":"At each selected mistake","trace":"PA1 analysis opportunity -> C2-E10 -> C-SES-02/06 -> UR-P-C1/C2 -> UR-C02 -> Learning priority -> 3 votes -> TP-CHESS -> C-A1 -> PA3 move-trial test."},
    {"id":"C-UC05","name":"Read explanation","scope":"Chess Guided Beginner Review Module","actor":"Beginner / returning learner","support":"Explanation service","trigger":"A mistake or better move is selected.","pre":"An explanation exists or a safe fallback can describe the position.","guarantee":"Plain language leads and terminology help is available on demand.","main":["Open the explanation card.","Read the one-sentence learning point.","Review board evidence and move comparison.","Open terminology help if needed.","Choose Try again, Continue practice, or Advanced detail."],"alt":["A1 at step 2: Explanation is unavailable; show factual board changes and avoid invented rationale.","A2 at step 4: Terminology help opens inline and returns focus to the term."],"exc":["E1: Service failure keeps the move trial and offers retry."],"rules":"Explanations distinguish fact, recommendation, and uncertainty.","data":"Learning point, board evidence, glossary terms, explanation version.","usability":"Plain language, satisfaction, and recognition over recall.","access":"Readable text, glossary semantics, and no forced motion.","privacy":"No unrelated conversation or profile data.","priority":"Critical","frequency":"Each selected mistake","trace":"PA1 terminology drawback -> C2-E10 -> C-SES-01/02/04/06 -> UR-P-C1/C2 -> UR-C02 -> Vocabulary load -> 5 votes -> TP-CHESS -> C-A1 -> recommendation -> PA3 explanation test."},
    {"id":"C-UC06","name":"Continue to practice or advanced depth","scope":"Chess Guided Beginner Review Module","actor":"Beginner / returning learner","support":"Learning catalog; analysis service","trigger":"User completes the explanation.","pre":"A learning point exists; practice or depth may be unavailable.","guarantee":"One relevant practice action is offered before optional advanced depth.","main":["Review the recommended lesson or puzzle.","Read why it matches the mistake.","Choose Continue to practice or Reveal advanced detail.","Open the chosen destination with review context.","Return to the review when desired."],"alt":["A1 at step 1: No exact practice exists; offer a broader topic and explain the gap.","A2 at step 3: Advanced detail is restricted; explain scope without blocking practice."],"exc":["E1: Mobile interruption preserves the review and selected next step."],"rules":"Practice relevance is explainable; advanced depth is opt-in, not removed.","data":"Learning point, catalog item, relevance reason, disclosure state.","usability":"Continuation, memorability, and progressive disclosure.","access":"Destination choice and return route are keyboard/screen-reader operable.","privacy":"Share only the minimum review context with the learning catalog.","priority":"Critical","frequency":"Each completed explanation","trace":"PA1 learning strength + analysis drawback -> C2-E05/E07/E10 -> C-SES-03/04/06 -> UR-P-C2 -> UR-C03 -> Practice continuation -> 5 votes -> TP-CHESS -> C-A1 -> recommendation -> PA3 continuation test."},
]


def build_use_cases() -> Path:
    doc = m.base_document("Use Case Document")
    cover(doc, "Group10-PA2 Use Case Document", "Twelve detailed use cases with UML, edge cases, and end-to-end traceability")
    doc.add_heading("1. Scope, actors, and notation", 1)
    m.add_table(doc, ["System", "Primary actor", "External systems", "Use cases"], [
        ["FIFA Ticket Confidence Module", "Ticket planner", "FIFA tournament catalog; governed ticket-state feed; official ticketing partner; notification and identity services", "F-UC01-F-UC06"],
        ["Chess Guided Beginner Review Module", "Beginner / returning learner", "Game store; analysis service; explanation service; learning catalog", "C-UC01-C-UC06"],
    ], [4.2, 3.5, 6.8, 2.2])
    m.callout(doc, "Notation", "Solid lines are actor associations. Dashed arrows labeled <<include>> point to required reused behavior. Dashed arrows labeled <<extend>> point from optional/conditional behavior to its base use case. External systems remain outside the system boundary.")
    fifa_uml = uml_v2("final-fifa-uml", "UC-UML-F FIFA Ticket Confidence Module", "FIFA Ticket Confidence Module", "Ticket planner", ["Ticket systems", "Notification / identity"], [u["name"] for u in USE_CASES[:6]], [(1,2),(3,2)], [(4,1),(5,3)])
    chess_uml = uml_v2("final-chess-uml", "UC-UML-C Chess Guided Beginner Review Module", "Chess Guided Beginner Review Module", "Beginner / returning learner", ["Game / analysis", "Learning catalog"], [u["name"] for u in USE_CASES[6:]], [(2,4),(3,4)], [(5,4),(1,0)])
    m.add_figure(doc, fifa_uml, "UC-UML-01", "FIFA UML use-case diagram", "Design inference; actors and systems derived from F-A1 and the traceability matrix", "F-UC01-F-UC06", 16, 16)
    m.add_figure(doc, chess_uml, "UC-UML-02", "Chess UML use-case diagram", "Design inference; actors and systems derived from C-A1 and the traceability matrix", "C-UC01-C-UC06", 16, 16)
    doc.add_heading("2. Detailed use cases", 1)
    for i, u in enumerate(USE_CASES):
        if i:
            page(doc)
        doc.add_heading(f"2.{i+1} {u['id']} - {u['name']}", 2)
        kv(doc, [("ID", u["id"]), ("Goal", u["name"]), ("Scope", u["scope"]), ("Primary actor", u["actor"]), ("Supporting actor", u["support"]), ("Trigger", u["trigger"]), ("Preconditions", u["pre"]), ("Success guarantee", u["guarantee"]), ("Priority", u["priority"]), ("Frequency", u["frequency"])])
        doc.add_heading("Main flow", 3)
        m.add_numbered_steps(doc, u["main"])
        doc.add_heading("Alternate flows", 3)
        m.add_bullets(doc, u["alt"])
        doc.add_heading("Exception flows", 3)
        m.add_bullets(doc, u["exc"])
        kv(doc, [("Business rules", u["rules"]), ("Data", u["data"]), ("Usability", u["usability"]), ("Accessibility", u["access"]), ("Privacy", u["privacy"]), ("Traceability", u["trace"])])
    doc.add_heading("3. Edge-case coverage", 1)
    m.add_table(doc, ["Product", "Required edge case", "Covered by"], [
        ["FIFA", "Stale state; missing state; sold out; resale; waiting room", "F-UC02-F-UC03"], ["FIFA", "Partner failure; login mismatch; notification denied; network loss", "F-UC01, F-UC04-F-UC06"],
        ["Chess", "No game; invalid PGN; engine unavailable", "C-UC01-C-UC04"], ["Chess", "No major mistake; multiple mistakes; premium restriction", "C-UC02-C-UC03"],
        ["Chess", "Mobile interruption; keyboard-only; terminology help", "C-UC01-C-UC02, C-UC05-C-UC06"],
    ], [3.0, 8.0, 5.0])
    doc.add_heading("4. Design rationale and validation target", 1)
    source_line(doc, "HCI rationale", COURSE)
    source_line(doc, "PA3 validation target", "Test the six core steps per product and the stated alternate/exception flows without assuming implementation success.")
    return save(doc, "Group10-PA2-UseCaseDocument")


def meeting(doc: Document, title: str, date: str, time: str, objective: str, topics: str, decisions: str, actions: str, summary: str):
    doc.add_heading(title, 2)
    kv(doc, [
        ("Record type", "Reconstructed project record / Simulated meeting record"), ("Date", date), ("Time", time),
        ("Present", "; ".join(TEAM)), ("Absent", "None"), ("Objective", objective), ("Topics", topics),
        ("Decisions", decisions), ("Actions", actions), ("Summary", summary),
    ])


def build_weekly_report() -> Path:
    doc = m.base_document("Weekly Report")
    cover(doc, "Group10-PA2 Weekly Report", "Reconstructed three-week RUP + Scrum project process")
    doc.add_heading("1. Process Overview", 1)
    m.callout(doc, "Disclosure", "This report documents a reconstructed three-week project process prepared for course reporting.")
    doc.add_paragraph("The process combines RUP phases with Scrum cadence: Inception establishes PA1 continuity and evidence rules; Elaboration defines simulated-study protocol, synthesis, and alternatives; Construction produces proposal, UML, use cases, traceability, and document builds; Transition performs visual QA and packaging.")
    source_line(doc, "Reconstructed project record", "Dates, meetings, attendance, workload, and progress below are reconstructed or simulated project records; they are not presented as contemporaneous records.")

    doc.add_heading("2. Meeting Conduct and Cross-Report Decisions", 1)
    m.add_table(doc, ["Rule", "Application"], [
        ["Evidence before claim", "Screen states prove interface state; PA1 findings remain inherited; simulated study remains labeled; design implications remain inference."],
        ["Cross-report IDs", "Participant, session, finding, cluster, problem, concept, recommendation, use case, and PA3 target IDs stay stable."],
        ["No unrecorded media", "Structured notes only; no external communication references."],
        ["Decision control", "Status Dashboard and Beginner Review Preset are recommendations; remaining alternatives stay in PA3 parallel prototyping."],
        ["Evidence validation", "The PA1 lesson on wrong interface states becomes a PA2 validation gate for every figure and claim."],
    ], [5.0, 11.0])

    doc.add_heading("3. Team Roster and Three-Week Schedule", 1)
    m.add_table(doc, ["Member", "Primary ownership", "Research", "Writing", "Review", "Presentation / packaging"], [
        ["Le Minh", "Coordination, integration, traceability, packaging, Weekly Report", "Protocol integration", "Weekly and cross-report sections", "Traceability and acceptance", "Package and submission"],
        ["Nguyen Vu Bach", "FIFA research, simulated sessions, findings, proposal, use cases", "FIFA tasks/sessions", "FIFA modules", "FIFA source/claim review", "FIFA presentation"],
        ["Pham Nguyen Gia Bao", "Chess research, simulated sessions, findings, proposal, use cases", "Chess tasks/sessions", "Chess modules", "Chess source/claim review", "Chess presentation"],
        ["Trang Minh Nhut", "Brainstorming, affinity, prioritization, HCI mapping, diagrams, visual QA", "Synthesis protocol", "Analysis and HCI mapping", "Diagram and layout review", "Visual consistency"],
    ], [3.3, 5.5, 3.0, 3.3, 3.2, 3.5])
    m.add_table(doc, ["Week", "Dates", "RUP focus", "Scrum goal", "Main outputs"], [
        ["1", "14-18 Jul 2026", "Inception / early Elaboration", "Scope, PA1 continuity, evidence audit, research plan, simulated session design, walkthrough", "Protocol, participants, session records, screen-evidence boundary"],
        ["2", "20-24 Jul 2026", "Elaboration", "Synthetic synthesis, brainstorming, affinity, voting, prioritization, tough problems, alternatives", "72 notes, clusters, votes, weighted problems, six alternatives"],
        ["3", "27-30 Jul 2026", "Construction / Transition", "Proposal, UML, use cases, traceability, QA, packaging", "Five revised reports, unchanged Peer Review, QA and six-PDF package"],
    ], [1.2, 3.1, 4.0, 7.4, 7.0])

    doc.add_heading("4. Sprint Planning", 1)
    meeting(doc, "4.1 Sprint Planning Meeting", "14 Jul 2026", "20:00-20:50", "Translate PA2 rubric into a three-week backlog while preserving PA1 continuity.", "Two products; PA1 drawbacks; evidence boundary; simulated-study labels; report ownership; definition of done.", "Keep FIFA and Chess; use three methods; label every synthetic record; do not change Peer Review; require traceability and page-render QA.", "Le builds traceability skeleton; Bach audits FIFA; Bao audits Chess; Nhut maps lecture concepts and visual rules.", "The team accepts one integrated sprint with weekly checkpoints and a final transition gate.")
    m.add_table(doc, ["Backlog item", "Owner", "Acceptance condition", "Week"], [
        ["PA1 continuity and screen-evidence audit", "Le + Bach + Bao", "Every high-level claim has a source label", "1"],
        ["12 simulated profiles/sessions/note sheets", "Bach + Bao", "IDs and metrics trace one-to-one", "1"],
        ["72 notes, affinity, votes, prioritization", "Nhut", "36 notes per product; 24 votes per product", "2"],
        ["Six alternatives and weighted comparison", "Bach + Bao", "All required evaluation fields and PA3 plan", "2"],
        ["12 detailed use cases and UML", "Bach + Bao + Nhut", "Six per product; include/extend; edge cases", "3"],
        ["Build, render, traceability, QA, ZIP", "Le + Nhut", "Six PDFs only; Peer Review byte-identical", "3"],
    ], [6.0, 3.6, 6.2, 1.5])

    doc.add_heading("5. Week 1 Scrum", 1)
    meeting(doc, "5.1 Weekly Scrum - Week 1", "17 Jul 2026", "20:00-20:30", "Check scope, evidence audit, protocol, simulated-session design, and walkthrough progress.", "PA1 continuity; screen inventory; task lists; participant balance; evidence-state mistakes.", "Retain 12 scenario profiles across two products; every note sheet must contain limitation and PA3 question.", "Finish final two sessions on 18 Jul; map each result to one note sheet; flag partner/engine claims as untested.", "Week 1 establishes a bounded research artifact without representing synthetic sessions as real research.")
    m.add_table(doc, ["Member", "Completed", "Next", "Obstacle"], [
        ["Le Minh", "Rubric matrix, continuity chain, source-label rules", "Integrate task/result IDs", "Legacy files use provisional wording"],
        ["Nguyen Vu Bach", "FIFA walkthrough, six profiles, first four session records", "Finish FIFA notes and findings", "Ticket feed and partner state unavailable"],
        ["Pham Nguyen Gia Bao", "Chess walkthrough, six profiles, first four session records", "Finish Chess notes and findings", "No completed engine review evidence"],
        ["Trang Minh Nhut", "Lecture mapping and evidence visual audit", "Model legends and layout rules", "Long captures require readable splits"],
    ], [3.5, 6.0, 5.0, 5.2])

    doc.add_heading("6. Week 2 Scrum", 1)
    meeting(doc, "6.1 Simulated Brainstorming Session", "22 Jul 2026", "19:30-20:45", "Cluster 72 atomic notes and establish product priorities.", "Silent generation; round-robin; duplicate merging; cluster naming; simulated voting; weighted scoring.", "FIFA prioritizes status and handoff confidence; Chess consolidates entry, explanation, and continuation into guided beginner review.", "Nhut maintains clusters/votes; Bach and Bao write problem profiles; Le checks traceability.", "The synthetic workshop produces design hypotheses and is not contemporaneous meeting evidence.")
    meeting(doc, "6.2 Weekly Scrum - Week 2", "24 Jul 2026", "20:00-20:30", "Review synthesis, tough problems, and six alternatives.", "Affinity counts; vote totals; weighted rationale; product boundary; PA3 parallel prototypes.", "Select TP-FIFA and TP-CHESS; recommend F-A1 and C-A1 while retaining four alternatives.", "Draft comparison formula, accessibility/privacy/data/recovery fields, and PA3 test measures.", "Week 2 converts evidence and synthetic patterns into explicit, testable design directions.")
    m.add_table(doc, ["Member", "Completed", "Next", "Obstacle"], [
        ["Le Minh", "Traceability through tough problems", "Integrate recommendations and use cases", "Cross-report terminology drift"],
        ["Nguyen Vu Bach", "FIFA findings, clusters, alternatives", "FIFA proposal and six use cases", "State taxonomy depends on governed data"],
        ["Pham Nguyen Gia Bao", "Chess findings, clusters, alternatives", "Chess proposal and six use cases", "Explanation quality requires PA3 validation"],
        ["Trang Minh Nhut", "72-note affinity, vote tables, HCI dimensions", "UML and visual QA", "Wide tables require landscape sections"],
    ], [3.5, 6.0, 5.0, 5.2])

    doc.add_heading("7. Week 3 Scrum", 1)
    meeting(doc, "7.1 Weekly Scrum - Week 3", "29 Jul 2026", "20:00-20:30", "Close proposal, UML, detailed use cases, traceability, and build defects.", "Weighted recommendations; include/extend direction; use-case numbering; edge cases; source labels; layout.", "Use 12 exact use cases; restart every main flow at 1; preserve advanced Chess depth by disclosure; never infer FIFA partner completion.", "Build revised DOCX/PDF, render every page, correct layout and banned text, prepare six-PDF ZIP.", "Week 3 moves from content completion to evidence-backed transition QA.")
    m.add_table(doc, ["Member", "Completed", "Next", "Obstacle"], [
        ["Le Minh", "Proposal integration, traceability CSV, weekly draft", "Package and final acceptance", "TOC page numbers require final render"],
        ["Nguyen Vu Bach", "FIFA proposal and use cases", "Cross-check FIFA edge cases", "Partner failure cannot be live-tested"],
        ["Pham Nguyen Gia Bao", "Chess proposal and use cases", "Cross-check Chess edge cases", "Engine unavailable path is scenario-based"],
        ["Trang Minh Nhut", "UML, diagrams, landscape tables", "All-page visual QA", "Dense note and traceability tables"],
    ], [3.5, 6.0, 5.0, 5.2])

    doc.add_heading("8. Sprint Review and Retrospective", 1)
    meeting(doc, "8.1 Sprint Review and Retrospective", "30 Jul 2026", "20:00-20:50", "Inspect final artifacts against rubric and record improvements for PA3.", "Research completeness; analysis trace; recommendation logic; UML/use cases; simulated labels; page render; package contents.", "Accept five revised reports when QA passes; preserve Peer Review unchanged; carry all unverified benefits to PA3 targets.", "Le completes package; Nhut records QA; Bach/Bao verify product-specific claims and edge cases.", "The sprint review closes PA2 reporting and leaves PA3 with explicit validation questions rather than synthetic certainty.")
    kv(doc, [
        ("What worked", "Stable IDs, source labels, product separation, and one evidence-to-use-case chain reduced ambiguity."),
        ("What did not work", "Legacy documents were too provisional, too short in analysis/weekly reporting, and lacked traceable synthetic records."),
        ("Improvement", "Validate state before capture, keep synthetic/observed evidence distinct, and prototype alternatives in parallel before commitment."),
        ("PA3 carryover", "Observe users, test prototypes, replace synthetic expectations with measured task evidence, and retain accessibility/privacy/recovery gates."),
    ])

    doc.add_heading("9. Workload Summary", 1)
    m.add_table(doc, ["Member", "Research", "Writing", "Review", "Presentation / packaging", "Reconstructed hours", "Status"], [
        ["Le Minh", "Protocol integration", "Weekly, traceability, integration", "Acceptance and cross-report consistency", "Packaging", "27", "Complete"],
        ["Nguyen Vu Bach", "FIFA walkthrough/sessions", "FIFA findings, proposal, use cases", "FIFA claims and edge cases", "FIFA report support", "26", "Complete"],
        ["Pham Nguyen Gia Bao", "Chess walkthrough/sessions", "Chess findings, proposal, use cases", "Chess claims and edge cases", "Chess report support", "26", "Complete"],
        ["Trang Minh Nhut", "Synthesis protocol", "Affinity, HCI mapping", "Diagram/layout consistency", "Visual QA", "25", "Complete"],
    ], [3.3, 4.3, 5.0, 5.0, 4.0, 2.4, 2.2])
    source_line(doc, "Reconstructed project record", "Hours are reconstructed workload history for balanced course reporting, not time-tracker evidence.")

    doc.add_heading("10. Acceptance Checklist", 1)
    m.add_table(doc, ["Acceptance item", "Evidence", "Status"], [
        ["User Research", "12 profiles, 12 sessions, 12 note sheets, metrics, 4 personas, 8 models, findings", "Complete"],
        ["User Analysis", "72 notes, 13 clusters, voting, weighted prioritization, two tough problems", "Complete"],
        ["Project Proposal", "Six alternatives, formula, two recommendations, PA3 parallel plan", "Complete"],
        ["Use Case Document", "Two UML diagrams, 12 use cases, required edge cases", "Complete"],
        ["Peer Review", "Existing PDF preserved without modification", "Unchanged"],
        ["Weekly Report", "Three weeks, meetings, scrums, workload, disclosure", "Complete"],
        ["QA and package", "Selectable text, TOC, renders, CSVs, labels, traceability, six-PDF ZIP", "Transition gate"],
    ], [7.0, 8.5, 2.5])

    doc.add_heading("11. Submission Status", 1)
    m.add_table(doc, ["Artifact", "Owner", "Submission state"], [
        ["Group10-PA2-UserResearch.pdf", "Bach + Bao + Le", "Ready after final QA"], ["Group10-PA2-UserAnalysis.pdf", "Nhut + Le", "Ready after final QA"],
        ["Group10-PA2-ProjectProposal.pdf", "Bach + Bao", "Ready after final QA"], ["Group10-PA2-UseCaseDocument.pdf", "Bach + Bao + Nhut", "Ready after final QA"],
        ["Group10-PA2-PeerReview.pdf", "Existing artifact", "Unchanged"], ["Group10-PA2-WeeklyReport.pdf", "Le Minh", "Ready after final QA"],
    ], [9.0, 4.0, 4.0])
    doc.add_heading("12. Continuity with PA1", 1)
    m.add_bullets(doc, ["The same FIFA.com and Chess.com product pair continues.", "The same four members and responsibilities continue.", "PA1 drawbacks become PA2 tough problems.", "PA1 evidence rules become the PA2 research protocol.", "The PA1 lesson about screenshots showing the wrong state becomes the PA2 evidence-validation gate."])
    return save(doc, "Group10-PA2-WeeklyReport")


def write_traceability():
    headers = ["pa1_finding", "screen_evidence", "simulated_session", "persona", "finding", "affinity_cluster", "vote", "tough_problem", "concept", "recommendation", "use_case", "pa3_test", "report_section"]
    rows = [
        ["FIFA ticket action lacks consolidated decision confidence", "F2-E09", "F-SES-04/05/06", "UR-P-F1", "UR-F01", "Ticket status clarity", "10", "TP-FIFA", "F-A1", "Status Dashboard", "F-UC02/F-UC03", "State/freshness accuracy", "UR 9; UA 7; PP 5; UC 2"],
        ["FIFA outbound route needs trust context", "F2-E09", "F-SES-02/05", "UR-P-F1", "UR-F02", "Handoff trust", "8", "TP-FIFA", "F-A1", "Status Dashboard", "F-UC04/F-UC06", "Destination recognition and return", "UR 9; UA 7; PP 5; UC 2"],
        ["FIFA mobile browsing loses overview", "F2-E02/E03/E04", "F-SES-03/04", "UR-P-F2", "UR-F03", "Mobile overview", "1", "TP-FIFA", "F-A1/F-A3", "Status Dashboard", "F-UC01/F-UC06", "Cross-device resumption", "UR 9; UA 7; PP 6; UC 2"],
        ["Chess analysis entry has high recall demand", "C2-E10", "C-SES-01/02/05", "UR-P-C1", "UR-C01", "Entry choice overload", "4", "TP-CHESS", "C-A1", "Beginner Review Preset", "C-UC01/C-UC02", "Entry-choice accuracy", "UR 9; UA 7; PP 5; UC 2"],
        ["Chess learning and analysis lack a beginner bridge", "C2-E07/E08/E10", "C-SES-02/04/06", "UR-P-C1/C2", "UR-C02", "Vocabulary load / Learning priority", "5 / 3", "TP-CHESS", "C-A1", "Beginner Review Preset", "C-UC03/C-UC04/C-UC05", "Mistake explanation and move trial", "UR 9; UA 7; PP 5; UC 2"],
        ["Chess practice is strong but disconnected from review", "C2-E05/E07/E10", "C-SES-03/04/06", "UR-P-C2", "UR-C03", "Practice continuation", "5", "TP-CHESS", "C-A1/C-A3", "Beginner Review Preset", "C-UC06", "Relevant continuation choice", "UR 9; UA 7; PP 6; UC 2"],
    ]
    with (ROOT / "traceability-matrix.csv").open("w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f); w.writerow(headers); w.writerows(rows)


def normalize_evidence_index():
    path = ROOT / "evidence-index.csv"
    with path.open(encoding="utf-8-sig", newline="") as f:
        reader = csv.reader(f); rows = list(reader)
    width = len(rows[0])
    fixed = [r[:width] + [""] * max(0, width-len(r)) for r in rows]
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        csv.writer(f).writerows(fixed)


def write_missing_evidence():
    text = """# Evidence boundary after PA2 finalization

- Screen evidence remains the only evidence of captured interface state.
- Simulated participant, session, quotation, observation, task result, brainstorming, meeting, attendance, and workload records are clearly labeled scenario-based synthetic evidence or reconstructed project records.
- The synthetic records support test-case design and hypotheses only; they are not research with real users.
- Verified presentation feedback remains unavailable. The Peer Review artifact is therefore labeled as internal/simulated design QA and does not attribute comments to external peers.
- PA3 must validate the recommendation, usability outcomes, explanation quality, ticket-state accuracy, external-service behavior, and real-user performance.
"""
    (ROOT / "MISSING-EVIDENCE.md").write_text(text, encoding="utf-8")


def main():
    normalize_evidence_index()
    write_traceability()
    write_missing_evidence()
    paths = [build_user_research(), build_user_analysis(), build_project_proposal(), build_use_cases(), build_weekly_report()]
    for path in paths:
        print(path.name, path.stat().st_size)


if __name__ == "__main__":
    main()
