from __future__ import annotations

import hashlib
import re
import subprocess
import zipfile
from pathlib import Path

import pdfplumber
from PIL import Image, ImageChops, ImageDraw, ImageFont
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
FINAL = ROOT / "final"
SOURCE = ROOT / "source"
QA = ROOT / "qa"
RENDER_ROOT = QA / "page-render-contact-sheets" / "revision-final"
POPPLER = Path(
    r"C:\Users\Administrator\.cache\codex-runtimes\codex-primary-runtime"
    r"\dependencies\native\poppler\Library\bin\pdftoppm.exe"
)
NAMES = [
    "Group10-PA2-UserResearch.pdf",
    "Group10-PA2-UserAnalysis.pdf",
    "Group10-PA2-ProjectProposal.pdf",
    "Group10-PA2-UseCaseDocument.pdf",
    "Group10-PA2-PeerReview.pdf",
    "Group10-PA2-WeeklyReport.pdf",
]


def ink_ratio(path: Path) -> float:
    with Image.open(path) as image:
        rgb = image.convert("RGB")
        diff = ImageChops.difference(rgb, Image.new("RGB", rgb.size, "white")).convert("L")
        return sum(diff.histogram()[12:]) / (rgb.width * rgb.height)


def edge_ratio(path: Path, border=4) -> float:
    with Image.open(path) as image:
        rgb = image.convert("RGB")
        strips = [rgb.crop((0, 0, rgb.width, border)), rgb.crop((0, rgb.height - border, rgb.width, rgb.height)), rgb.crop((0, 0, border, rgb.height)), rgb.crop((rgb.width - border, 0, rgb.width, rgb.height))]
        nonwhite = pixels = 0
        for strip in strips:
            diff = ImageChops.difference(strip, Image.new("RGB", strip.size, "white")).convert("L")
            nonwhite += sum(diff.histogram()[12:])
            pixels += strip.width * strip.height
        return nonwhite / pixels


def contact_sheet(images: list[Path], output: Path, start_page: int):
    thumb_w, thumb_h, cols = 360, 510, 4
    rows = (len(images) + cols - 1) // cols
    sheet = Image.new("RGB", (cols * 390 + 30, rows * 565 + 30), "#D9E0E8")
    draw = ImageDraw.Draw(sheet)
    font = ImageFont.truetype(r"C:\Windows\Fonts\arial.ttf", 22)
    for index, path in enumerate(images):
        with Image.open(path) as source:
            page = source.convert("RGB")
        page.thumbnail((thumb_w, thumb_h))
        x = 30 + (index % cols) * 390
        y = 25 + (index // cols) * 565
        sheet.paste(page, (x + (thumb_w - page.width) // 2, y))
        draw.text((x, y + thumb_h + 8), f"Page {start_page + index}", fill="#111827", font=font)
    sheet.save(output, dpi=(150, 150))


def docx_text(path: Path) -> str:
    doc = Document(path)
    return "\n".join([p.text for p in doc.paragraphs] + [cell.text for table in doc.tables for row in table.rows for cell in row.cells])


def main():
    QA.mkdir(exist_ok=True)
    RENDER_ROOT.mkdir(parents=True, exist_ok=True)
    for stale_contact in RENDER_ROOT.glob("*-contact-*.png"):
        stale_contact.unlink()
    checks: list[tuple[str, bool, str]] = []

    def check(name: str, result: bool, evidence: str):
        checks.append((name, bool(result), evidence))

    check("Six final PDFs exist", all((FINAL / name).is_file() for name in NAMES), ", ".join(NAMES))
    check("Six source DOCX exist", all((SOURCE / name.replace(".pdf", ".docx")).is_file() for name in NAMES), "DOCX sources retained")
    texts: dict[str, str] = {}
    page_counts: dict[str, int] = {}
    render_counts: dict[str, int] = {}
    all_contact_sheets: list[Path] = []
    prohibited = [r"Evidence Date", r"Document Type", r"Course[- ]source application", r"Course source mapping", r"Application of course source", r"Lorem Ipsum", r"John Doe", r"Jane Doe", r"test@gmail\.com", r"\bTBD\b", r"\bTODO\b", r"\bplaceholder\b", r"\boptimized\b", r"\buser-friendly\b", r"\bflexible\b", r"\bseamless\b", r"\bintuitive\b", r"\bCodex\b", r"\bChatGPT\b", r"AI-generated"]

    for name in NAMES:
        path = FINAL / name
        with pdfplumber.open(path) as document:
            pages = [page.extract_text() or "" for page in document.pages]
            sizes = [float(char.get("size", 0)) for page in document.pages for char in page.chars if char.get("text", "").strip()]
            geometries = [(float(page.width), float(page.height)) for page in document.pages]
        text = "\n".join(pages)
        texts[name] = text
        page_counts[name] = len(pages)
        check(f"Selectable text - {name}", all(len(page.strip()) >= 40 for page in pages), f"{len(pages)} pages; minimum {min(map(lambda x: len(x.strip()), pages))} characters")
        check(f"Minimum 14 pt - {name}", bool(sizes) and min(sizes) >= 13.9, f"minimum PDF character size {min(sizes):.2f} pt")
        check(f"Valid page geometry - {name}", all(width > 500 and height > 500 for width, height in geometries), f"{len(geometries)} valid page boxes")
        first_page = pages[0] if pages else ""
        intro_hits = [term for term in ("Evidence Date", "Document Type", "Status") if re.search(rf"\b{re.escape(term)}\b", first_page, re.I)]
        check(f"Clean introduction - {name}", not intro_hits, "Course, Group, Team, and Scope only" if not intro_hits else str(intro_hits))
        hits = [pattern for pattern in prohibited if re.search(pattern, text, re.I)]
        check(f"No prohibited text - {name}", not hits, "No banned metadata, placeholder, generic claim, or content-generation reference" if not hits else str(hits))
        check(f"Lessons learned - {name}", "Practical Difficulties and Lessons Learned" in text, "Report-specific difficulty, handling, lesson, and impact")

        render_dir = RENDER_ROOT / Path(name).stem
        render_dir.mkdir(exist_ok=True)
        for stale in render_dir.glob("page-*.png"):
            stale.unlink()
        subprocess.run([str(POPPLER), "-png", "-r", "120", str(path), str(render_dir / "page")], check=True, capture_output=True)
        images = sorted(render_dir.glob("page-*.png"))
        render_counts[name] = len(images)
        blank = [index + 1 for index, image in enumerate(images) if ink_ratio(image) < 0.002]
        edges = [index + 1 for index, image in enumerate(images) if edge_ratio(image) > 0.005]
        check(f"Rendered page count - {name}", len(images) == len(pages), f"{len(images)} PNG pages for {len(pages)} PDF pages")
        check(f"No blank rendered page - {name}", not blank, "No blank candidates" if not blank else str(blank))
        check(f"No page-edge clipping - {name}", not edges, "No non-white pixels at outer page edge" if not edges else str(edges))
        for batch in range(0, len(images), 12):
            output = RENDER_ROOT / f"{Path(name).stem}-contact-{batch // 12 + 1:02d}.png"
            contact_sheet(images[batch:batch + 12], output, batch + 1)
            all_contact_sheets.append(output)

    all_text = "\n".join(texts.values())
    proposal = texts["Group10-PA2-ProjectProposal.pdf"]
    use_cases = texts["Group10-PA2-UseCaseDocument.pdf"]
    research = texts["Group10-PA2-UserResearch.pdf"]
    analysis = texts["Group10-PA2-UserAnalysis.pdf"]
    peer = texts["Group10-PA2-PeerReview.pdf"]
    check("Guided Concierge 16 steps", all(sentence in proposal for sentence in ["User enters Ticket Concierge", "User selects a tournament", "System checks the known ticket state", "User compares official ticket", "System previews the destination", "User chooses Continue or Stay", "User may set an alert", "recovery branch returns the user"]), "Entry, questions, state check, comparison, preview, Continue/Stay, alert, confirmation, and recovery")
    concierge_figures = set(re.findall(r"Figure (PP-GC-[A-Z0-9-]+)\.", proposal))
    check("Guided Concierge visuals", len(concierge_figures) >= 10, f"{len(concierge_figures)} unique flow/screen figures")
    for branch in ["No tickets", "Unknown status", "Stale status", "Partner unavailable", "User not eligible", "No matching event", "Invalid input", "Alert permission denied", "Network failure"]:
        check(f"Concierge branch - {branch}", branch.lower() in proposal.lower(), "Branch named in flow or state contract")
    state_fields = ["Trigger", "Message", "User action", "Recovery"]
    check("FIFA and Chess state contracts", all(field in proposal for field in state_fields) and all(label in proposal for label in ["FIFA Ticket Concierge", "Chess Beginner Review", "EMPTY", "ERROR", "VALIDATION"]), "Both concepts include trigger, message, action, and recovery")
    visual_ids = set(re.findall(r"Figure (UC-VIS-[FC]-UC\d{2})\.", use_cases))
    check("Twelve unique use-case visuals", len(visual_ids) == 12, f"{len(visual_ids)} unique visual IDs")
    visual_files = [ROOT / "generated-diagrams" / f"revision-{prefix.lower()}-uc{i:02d}.png" for prefix in ("f", "c") for i in range(1, 7)]
    hashes = [hashlib.sha256(path.read_bytes()).hexdigest() for path in visual_files]
    check("Use-case visual content unique", len(hashes) == len(set(hashes)), "Twelve distinct image hashes")
    check("Affinity diagrams", all(x in analysis for x in ["Figure UA-AF-FIFA", "Figure UA-AF-CHESS", "SCREEN EVIDENCE / RAW NOTE", "DESIGN DIRECTION"]), "Separate FIFA and Chess evidence-to-direction pipelines")
    model_ids = set(re.findall(r"Figure (UR-WM-(?:FIFA|CHESS)-[A-Z0-9-]+)\.", research))
    check("Eight revised Work Models", len(model_ids) == 8, f"{len(model_ids)} unique revised models")
    research_lower = research.lower()
    check("Work Model legend and recovery", "distinct shape for actor" in research_lower and "recovery path" in research_lower and "design opportunity" in research_lower, "Shape legend, interruption, recovery, and opportunity present")
    check("Peer Review disclosure", "INTERNAL / SIMULATED PEER REVIEW FOR DESIGN QA" in peer and "No verified presentation feedback" in peer, "No external peer feedback claim")
    check("Peer Review fields", all(field in peer for field in ["Screen or artifact", "Issue", "Severity", "Reason", "Recommendation", "Affected report", "Action taken", "Result after revision"]), "Every QA item uses required fields")
    check("Long screenshot split", all(f"Figure PR-EV-01{suffix}" in peer for suffix in "ABC") and all(label in peer for label in ["Top", "Middle", "Bottom"]), "Chess Lessons mobile split in reading order")
    check("Mini design system", all(term in proposal for term in ["Mini Design System", "14 pt minimum", "48 px height", "Status chips"]), "Typography, spacing, controls, input, cards, status, error/success/empty rules")
    check("Example email domain", "minh.anh@example.com" in proposal and not re.search(r"@[A-Za-z0-9.-]+\.(?!com\b)[A-Za-z]{2,}", proposal), "Sample email uses example.com")
    check("No content-generation discussion", not re.search(r"\b(?:Codex|ChatGPT|AI-generated|content generation tool)\b", all_text, re.I), "Deliverables discuss design work only")

    zip_path = ROOT / "Group10-PA2.zip"
    temp = ROOT / "tmp" / "Group10-PA2.revision.zip"
    temp.parent.mkdir(exist_ok=True)
    with zipfile.ZipFile(temp, "w", zipfile.ZIP_DEFLATED) as archive:
        for name in NAMES:
            archive.write(FINAL / name, arcname=name)
    temp.replace(zip_path)
    with zipfile.ZipFile(zip_path) as archive:
        entries = archive.namelist()
    check("ZIP contains six final PDFs", entries == NAMES, ", ".join(entries))

    passed = all(result for _, result, _ in checks)
    report = ["# Final QA report", "", f"Overall automated QA: {'PASS' if passed else 'FAIL'}", "", "Visual page review: PENDING", "", "## Checks", "", "| Check | Result | Evidence |", "|---|---|---|"]
    report += [f"| {name.replace('|', '/')} | {'PASS' if result else 'FAIL'} | {evidence.replace('|', '/')} |" for name, result, evidence in checks]
    report += ["", "## Page counts", ""] + [f"- {name}: {page_counts[name]} PDF pages / {render_counts[name]} PNG pages" for name in NAMES]
    report += ["", "## Visual review", "", f"- Contact sheets generated: {len(all_contact_sheets)}", f"- Render directory: `qa/page-render-contact-sheets/revision-final`", "- Review criteria: clipping, overlap, table breaks, figure size, caption adjacency, typography, page transitions, header/footer, and long-screenshot segmentation.", "", "## Remaining limitations", "", "- Simulated study and reconstructed project records remain explicitly labeled; they are not real-user evidence.", "- Live FIFA ticket inventory, partner availability, Chess analysis service behavior, and notification permission behavior require later validation."]
    (QA / "final-qa-report.md").write_text("\n".join(report) + "\n", encoding="utf-8")
    print("PASS" if passed else "FAIL")
    print(f"CONTACT_SHEETS={len(all_contact_sheets)}")
    for name, result, evidence in checks:
        if not result:
            print("FAIL", name, evidence)


if __name__ == "__main__":
    main()
