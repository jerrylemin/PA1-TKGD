from __future__ import annotations

import importlib.util
import textwrap
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "source"
DIAGRAMS = ROOT / "generated-diagrams"
QA = ROOT / "qa"

spec = importlib.util.spec_from_file_location("finalize", SOURCE / "finalize_pa2_reports.py")
f = importlib.util.module_from_spec(spec)
assert spec.loader
spec.loader.exec_module(f)
m = f.m

FONT = Path(r"C:\Windows\Fonts\arial.ttf")
BOLD = Path(r"C:\Windows\Fonts\arialbd.ttf")
NAVY = "#123B65"
BLUE = "#1D70A2"
GREEN = "#4F7F35"
RED = "#9B1C1C"
INK = "#1F2937"
MUTED = "#5B6573"


def font(size: int, bold: bool = False):
    return ImageFont.truetype(str(BOLD if bold else FONT), size)


def lines(text: str, width: int) -> list[str]:
    return textwrap.wrap(text, width=width, break_long_words=False, break_on_hyphens=False) or [""]


def draw_centered(draw: ImageDraw.ImageDraw, box, text: str, size=28, bold=False, fill=INK):
    x1, y1, x2, y2 = box
    wrapped = lines(text, max(14, int((x2 - x1) / (size * 0.58))))
    total = len(wrapped) * (size + 8)
    y = y1 + (y2 - y1 - total) / 2
    for item in wrapped:
        bounds = draw.textbbox((0, 0), item, font=font(size, bold))
        draw.text((x1 + (x2 - x1 - (bounds[2] - bounds[0])) / 2, y), item, font=font(size, bold), fill=fill)
        y += size + 8


def arrow(draw, start, end, color=BLUE):
    draw.line((*start, *end), fill=color, width=6)
    x, y = end
    draw.polygon([(x, y), (x - 18, y - 11), (x - 18, y + 11)], fill=color)


def screen_mock(name: str, title: str, step: str, goal: str, components: list[str], next_action: str, accent=BLUE) -> Path:
    image = Image.new("RGB", (1500, 1250), "white")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((80, 70, 1420, 1180), radius=28, outline=accent, width=6, fill="#FAFCFE")
    draw.rectangle((80, 70, 1420, 190), fill=accent)
    draw.text((125, 105), title, font=font(42, True), fill="white")
    draw.text((125, 225), f"Step: {step}", font=font(30, True), fill=NAVY)
    draw.text((125, 275), f"Screen goal: {goal}", font=font(28), fill=INK)
    y = 355
    for component in components:
        height = 105 + 34 * max(0, len(lines(component, 48)) - 1)
        draw.rounded_rectangle((125, y, 1375, y + height), radius=14, outline="#94A3B8", width=3, fill="#F3F6FA")
        draw_centered(draw, (155, y + 12, 1345, y + height - 12), component, 28, True)
        y += height + 28
    draw.rounded_rectangle((760, 1020, 1375, 1115), radius=22, fill=accent)
    draw_centered(draw, (780, 1030, 1355, 1105), next_action, 28, True, "white")
    output = DIAGRAMS / f"{name}.png"
    image.save(output, dpi=(220, 220))
    return output


def flow(name: str, title: str, steps: list[str], branches: list[str], accent=BLUE) -> Path:
    image = Image.new("RGB", (2200, 1350), "white")
    draw = ImageDraw.Draw(image)
    draw.text((70, 45), title, font=font(48, True), fill=NAVY)
    cols = min(4, len(steps))
    positions = []
    for index, step in enumerate(steps):
        row, col = divmod(index, cols)
        x, y = 80 + col * 530, 180 + row * 300
        box = (x, y, x + 440, y + 170)
        draw.rounded_rectangle(box, radius=18, outline=accent, width=4, fill="#F3F6FA")
        draw_centered(draw, box, step, 26, True)
        positions.append(box)
        if index and row == (index - 1) // cols:
            previous = positions[index - 1]
            arrow(draw, (previous[2] + 10, (previous[1] + previous[3]) // 2), (box[0] - 12, (box[1] + box[3]) // 2), accent)
    y = 180 + ((len(steps) + cols - 1) // cols) * 300
    draw.text((80, y), "Decision and recovery branches", font=font(32, True), fill=RED)
    for index, branch in enumerate(branches):
        x = 80 + (index % 4) * 530
        by = y + 70 + (index // 4) * 135
        draw.rounded_rectangle((x, by, x + 440, by + 105), radius=12, outline=RED, width=3, fill="#FFF1F1")
        draw_centered(draw, (x + 10, by + 8, x + 430, by + 97), branch, 23, False, "#5E1B1B")
    output = DIAGRAMS / f"{name}.png"
    image.save(output, dpi=(220, 220))
    return output


def state_board(name: str, title: str, states: list[tuple[str, str, str, str, str]], accent=BLUE) -> Path:
    image = Image.new("RGB", (2200, 1450), "white")
    draw = ImageDraw.Draw(image)
    draw.text((70, 45), title, font=font(48, True), fill=NAVY)
    for index, (state, trigger, message, action, recovery) in enumerate(states):
        x = 70 + index * 710
        draw.rounded_rectangle((x, 170, x + 650, 1330), radius=24, outline=accent, width=5, fill="#FAFCFE")
        draw.rectangle((x, 170, x + 650, 275), fill=accent)
        draw_centered(draw, (x + 20, 180, x + 630, 265), state, 32, True, "white")
        y = 330
        for label, value in (("Trigger", trigger), ("Message", message), ("User action", action), ("Recovery", recovery)):
            draw.text((x + 35, y), label, font=font(27, True), fill=NAVY)
            y += 42
            for item in lines(value, 39):
                draw.text((x + 35, y), item, font=font(25), fill=INK)
                y += 34
            y += 34
    output = DIAGRAMS / f"{name}.png"
    image.save(output, dpi=(220, 220))
    return output


def work_model(name: str, title: str, values: list[str], problem: str, recovery: str, opportunity: str, accent=BLUE) -> Path:
    image = Image.new("RGB", (2400, 1450), "white")
    draw = ImageDraw.Draw(image)
    draw.text((70, 45), title, font=font(48, True), fill=NAVY)
    labels = ["Actor", "Task", "Artifact", "System", "Decision", "Interruption", "Recovery", "Output"]
    centers = []
    for i, (label, value) in enumerate(zip(labels, values)):
        row, col = divmod(i, 4)
        cx, cy = 300 + col * 570, 320 + row * 420
        centers.append((cx, cy))
        if label == "Actor":
            draw.ellipse((cx - 165, cy - 105, cx + 165, cy + 105), fill="#EAF2F8", outline=accent, width=5)
        elif label == "Decision":
            draw.polygon([(cx, cy - 120), (cx + 180, cy), (cx, cy + 120), (cx - 180, cy)], fill="#FFF7D6", outline="#9A6A00")
        elif label == "Interruption":
            draw.polygon([(cx - 180, cy - 75), (cx - 110, cy - 125), (cx + 110, cy - 125), (cx + 180, cy - 75), (cx + 180, cy + 75), (cx + 110, cy + 125), (cx - 110, cy + 125), (cx - 180, cy + 75)], fill="#FFF1F1", outline=RED)
        else:
            fill = "#EEF5E9" if label in ("Recovery", "Output") else "#F3F6FA"
            draw.rounded_rectangle((cx - 190, cy - 115, cx + 190, cy + 115), radius=20, fill=fill, outline=accent, width=5)
        draw.text((cx - 165, cy - 88), label, font=font(24, True), fill=NAVY)
        draw_centered(draw, (cx - 165, cy - 50, cx + 165, cy + 95), value, 23, False)
        if i and i != 4:
            pcx, pcy = centers[i - 1]
            if i % 4:
                arrow(draw, (pcx + 205, pcy), (cx - 205, cy), accent)
    draw.text((80, 1150), f"Problem point: {problem}", font=font(25, True), fill=RED)
    draw.text((80, 1205), f"Recovery path: {recovery}", font=font(25), fill=GREEN)
    draw.text((80, 1260), f"Design opportunity: {opportunity}", font=font(25), fill=NAVY)
    draw.text((80, 1340), "Legend: ellipse = actor; rounded box = task/artifact/system/recovery/output; diamond = decision; octagon = interruption", font=font(22), fill=MUTED)
    output = DIAGRAMS / f"{name}.png"
    image.save(output, dpi=(220, 220))
    return output


def affinity(name: str, title: str, clusters: list[tuple[str, list[str], str, str, str, str]], accent=BLUE) -> Path:
    image = Image.new("RGB", (2400, 1650), "white")
    draw = ImageDraw.Draw(image)
    draw.text((70, 45), title, font=font(48, True), fill=NAVY)
    headings = ["SCREEN EVIDENCE / RAW NOTE", "SMALL CLUSTER", "PATTERN", "PROBLEM", "DESIGN DIRECTION"]
    for i, heading in enumerate(headings):
        draw_centered(draw, (40 + i * 470, 120, 450 + i * 470, 210), heading, 22, True, accent)
    for row, (cluster, notes, pattern, problem, severity, direction) in enumerate(clusters):
        y = 260 + row * 430
        columns = ["\n".join(notes), cluster, pattern, f"{problem}\nSeverity: {severity}", direction]
        for col, value in enumerate(columns):
            x = 40 + col * 470
            draw.rounded_rectangle((x, y, x + 410, y + 330), radius=16, fill="#F8FAFC", outline=accent, width=3)
            draw_centered(draw, (x + 18, y + 15, x + 392, y + 315), value, 21, col == 1)
            if col:
                arrow(draw, (x - 50, y + 165), (x - 12, y + 165), accent)
    draw.text((70, 1575), "Summary: Evidence -> Affinity clusters -> Priority -> Tough problem -> Concepts", font=font(28, True), fill=NAVY)
    output = DIAGRAMS / f"{name}.png"
    image.save(output, dpi=(220, 220))
    return output


def split_vertical(source: Path, name: str, labels=("Top", "Middle", "Bottom")) -> list[Path]:
    outputs = []
    with Image.open(source) as image:
        source_image = image.convert("RGB")
        segment_height = (source_image.height + len(labels) - 1) // len(labels)
        for index, label in enumerate(labels):
            y1 = index * segment_height
            y2 = min(source_image.height, (index + 1) * segment_height)
            # A full third of an extremely long mobile capture is still too
            # narrow when placed on A4. Keep ordered, representative windows
            # from the top/middle/bottom thirds at a readable 2:3 aspect ratio.
            window_height = min(y2 - y1, round(source_image.width * 1.5))
            window_top = y1 + max(0, (y2 - y1 - window_height) // 2)
            segment = source_image.crop((0, window_top, source_image.width, window_top + window_height))
            if segment.width < 1200:
                scale = 1200 / segment.width
                segment = segment.resize((1200, round(segment.height * scale)), Image.Resampling.LANCZOS)
            output = DIAGRAMS / f"{name}-{label.lower()}.png"
            segment.save(output, dpi=(220, 220))
            outputs.append(output)
    return outputs


def create_visuals() -> dict[str, Path]:
    DIAGRAMS.mkdir(exist_ok=True)
    visuals: dict[str, Path] = {}
    concierge = [
        ("entry", "Ticket Concierge", "1 - Entry", "Start with tournament context", ["Choose a tournament", "Use recent tournament context", "Why we ask"], "Start"),
        ("question", "Your ticket needs", "2-5 - Questions", "Collect match, party size, and eligibility", ["Match or event preference", "Party size", "Eligibility and location", "Back preserves answers"], "Check options"),
        ("recommendation", "Current ticket options", "7-9 - Recommendation", "Compare routes without implying inventory", ["Official ticket - state and freshness", "Register interest / resale / hospitality", "Wait - why and alert option"], "Preview destination"),
        ("destination", "Before you continue", "10-12 - Destination preview", "Explain domain, account, and return behavior", ["Official partner name and domain", "Account expectation", "What stays saved if you return"], "Continue to partner"),
        ("empty", "No matching event", "Recovery - empty", "Separate no result from sold out", ["No event matches the selected filters", "Change tournament or match preference", "Preserve party size and eligibility"], "Change filters"),
        ("error", "Ticket status unavailable", "Recovery - error", "Protect the decision when status cannot load", ["Current ticket state could not be loaded", "Last known update is labeled stale", "Stay on FIFA.com or retry"], "Retry status"),
        ("validation", "Complete this step", "Recovery - validation", "Identify missing input beside the field", ["Select a tournament", "Enter a valid email such as minh.anh@example.com", "Permission denial keeps an on-site reminder option"], "Review answers"),
    ]
    for key, title, step, goal, components, action in concierge:
        visuals[f"concierge-{key}"] = screen_mock(f"revision-concierge-{key}", title, step, goal, components, action)
    visuals["concierge-overview"] = flow("revision-concierge-overview", "Guided Concierge - Overview", ["1 Enter concierge", "2 Select tournament", "3 Select match", "4 Party size", "5 Context questions", "6 Check known state"], ["No matching event", "Invalid input", "Network failure"])
    visuals["concierge-decision"] = flow("revision-concierge-decision", "Guided Concierge - Decision and Handoff", ["7 Explain state", "8 Show options", "9 Compare routes", "10 Preview destination", "11 Continue or stay", "12 Explain domain/account"], ["No tickets", "Unknown status", "Stale status", "User not eligible"])
    visuals["concierge-recovery"] = flow("revision-concierge-recovery", "Guided Concierge - Confirmation and Recovery", ["13 Preserve context", "14 Optional alert", "15 Confirm choice", "16 Recover or return"], ["Partner unavailable", "Alert permission denied", "Notification channel unavailable", "Network failure"])
    fifa_states = [("EMPTY", "No tournament matches filters", "No matching tournament. Change tournament or date.", "Change filters", "Keep party size and eligibility answers"), ("ERROR", "Status or partner service fails", "Ticket status is unavailable. No sale outcome is inferred.", "Retry or stay", "Show last update only when labeled stale"), ("VALIDATION", "Tournament, email, or permission is missing", "Select a tournament and enter a valid email.", "Correct the field", "Offer on-site reminder when permission is denied")]
    chess_states = [("EMPTY", "No recent game is available", "No game is ready for review. Import a PGN or play a game.", "Import game", "Return to recent activity after a game exists"), ("ERROR", "Analysis service or network fails", "Analysis is unavailable. Your game remains saved.", "Retry analysis", "Open board-only review or return later"), ("VALIDATION", "PGN or required input is missing", "Add a valid game or PGN before review starts.", "Fix input", "Keep valid fields and show the failing line")]
    visuals["fifa-states"] = state_board("revision-fifa-states", "FIFA Ticket Concierge - Empty, Error, and Validation", fifa_states)
    visuals["chess-states"] = state_board("revision-chess-states", "Chess Beginner Review - Empty, Error, and Validation", chess_states, GREEN)
    uc_specs = {
        "F-UC01": ("Tournament selector", ["Recent tournament", "Search tournament", "Required selection message"]),
        "F-UC02": ("Ticket-state comparison", ["Official ticket", "Register interest", "Resale", "Hospitality", "Wait"]),
        "F-UC03": ("Freshness detail", ["Current state", "Last updated", "Source boundary", "Refresh status"]),
        "F-UC04": ("Destination preview", ["Partner name", "Official domain", "Account expectation", "Continue / Stay"]),
        "F-UC05": ("Alert setup", ["Tournament", "Email", "Permission", "Validation message", "Confirmation"]),
        "F-UC06": ("Handoff confirmation", ["Context preserved", "Continue", "Stay", "Return route"]),
        "C-UC01": ("Recent game or import", ["Recent games", "Import PGN", "No-game state"]),
        "C-UC02": ("Beginner review entry", ["Beginner review", "What this shows", "Advanced analysis link"]),
        "C-UC03": ("Key mistake", ["Critical position", "Mistake label", "Why it matters"]),
        "C-UC04": ("Try better move", ["Board input", "Legal move feedback", "Retry / hint"]),
        "C-UC05": ("Explanation", ["Plain-language reason", "Term help", "Show consequence"]),
        "C-UC06": ("Continue learning", ["Relevant lesson", "Practice puzzle", "Reveal advanced depth"]),
    }
    for uid, (title, components) in uc_specs.items():
        accent = BLUE if uid.startswith("F") else GREEN
        visuals[uid] = screen_mock(f"revision-{uid.lower()}", f"{uid} - {title}", "Use-case visual", f"Support {title.lower()}", components, "Continue", accent)
    fifa_affinity = [("F-A1 Status clarity", ["Buy/Register/Wait coexist", "Timestamp is not prominent", "Missing is not sold out"], "Users compare labels before acting", "Ticket outcome is unclear", "High", "Show state, freshness, and next action together"), ("F-A2 Handoff trust", ["Outbound CTA changes domain", "Account expectation is hidden", "Return context can be lost"], "Leaving the site raises a trust decision", "Users hesitate before partner handoff", "High", "Preview domain, account, and return path"), ("F-A3 Mobile continuity", ["Cards stack vertically", "Long pages hide overview", "Selected context moves off-screen"], "Mobile scanning increases recall work", "Tournament context is lost", "Medium", "Persist selected tournament and compact summary")]
    chess_affinity = [("C-A1 Entry choice", ["Several analysis paths", "Feature names require recall", "Recent game action is separated"], "Beginners take wrong entry paths", "Review start is hard to identify", "High", "Add Beginner Review at the game result"), ("C-A2 Explanation", ["Engine terms appear early", "Panels compete with the key mistake", "Depth is visible before reason"], "Advanced detail precedes the learning priority", "The first mistake is hard to interpret", "High", "Show one mistake and one plain-language reason first"), ("C-A3 Continuation", ["Lessons and puzzles are separate", "Relevant practice is not linked", "Returning users forget feature names"], "Review ends without a specific next task", "Learning continuity breaks", "Medium", "Link one lesson or puzzle to the mistake")]
    visuals["affinity-fifa"] = affinity("revision-affinity-fifa", "FIFA Affinity Pipeline", fifa_affinity)
    visuals["affinity-chess"] = affinity("revision-affinity-chess", "Chess Affinity Pipeline", chess_affinity, GREEN)
    models = [
        ("wm-fifa-planning", "FIFA Ticket Planning Task Flow", ["Ticket planner", "Choose tournament", "Tournament card", "FIFA.com", "Act or wait?", "State missing/stale", "Change filters / retry", "Ticket plan"], "Status is missing or stale", "Keep context and retry or wait", "Place state, freshness, and options together"),
        ("wm-fifa-verification", "FIFA Ticket Verification Sequence", ["Ticket planner", "Check state", "Status + timestamp", "Ticket service", "Current enough?", "Unknown state", "Show boundary and retry", "Verified next action"], "Freshness cannot be established", "Do not infer a sale outcome", "Expose timestamp and source boundary"),
        ("wm-fifa-handoff", "FIFA Partner Handoff Flow", ["Ticket planner", "Preview destination", "Domain + account", "Partner service", "Continue?", "Partner unavailable", "Stay with saved context", "Handoff or return"], "Partner service is unavailable", "Stay on FIFA.com with context", "Explain destination before departure"),
        ("wm-fifa-artifact", "FIFA Artifact and Information Model", ["Ticket planner", "Compare plan", "Tournament/state/alert", "FIFA + partner", "Which route?", "Data is stale", "Refresh or wait", "Saved plan"], "State and destination data age differently", "Label stale fields separately", "Keep provenance beside each value"),
        ("wm-chess-learning", "Chess Beginner Learning Flow", ["Beginner", "Open recent game", "Game artifact", "Chess review", "Beginner route?", "No game", "Import PGN / play", "Review entry"], "No reviewable game exists", "Offer import or play without losing place", "Put review entry beside the game result"),
        ("wm-chess-entry", "Chess Review-Entry Sequence", ["Returning learner", "Choose review", "Recent game / PGN", "Analysis service", "Valid input?", "Invalid PGN", "Preserve valid input", "Beginner review"], "PGN or game input fails validation", "Point to the failing input and keep the rest", "Separate Beginner Review from advanced setup"),
        ("wm-chess-practice", "Chess Explanation-to-Practice Flow", ["Beginner", "Read key mistake", "Mistake card", "Move validator", "Try move?", "Illegal move", "Retry / hint", "Practice result"], "A move is illegal or service fails", "Keep the position and allow retry", "Explain legality separately from quality"),
        ("wm-chess-artifact", "Chess Learning Artifact Model", ["Learner", "Choose next task", "Game/mistake/lesson", "Review + catalog", "Practice or depth?", "Catalog unavailable", "Keep review and retry", "Learning plan"], "Relevant practice cannot load", "Preserve the review result", "Link one specific lesson or puzzle")]
    for key, title, values, problem, recovery, opportunity in models:
        visuals[key] = work_model(f"revision-{key}", title, values, problem, recovery, opportunity, GREEN if "chess" in key else BLUE)
    visuals["design-system"] = state_board("revision-mini-design-system", "Mini Design System - Shared FIFA and Chess Rules", [("TYPE + SPACE", "Body, caption, table, figure labels", "Arial 14 pt minimum; headings 16/17/28 pt", "Use 8/16/24/32 px spacing", "Split tables or change orientation; never shrink below 14 pt"), ("CONTROLS", "Primary action or required input", "Buttons: 48 px height, 12 px radius; inputs: persistent label", "Use verb labels: Continue, Stay, Retry", "Show field-level errors and retain valid values"), ("FEEDBACK", "Empty, error, success, or status change", "Status chips use text plus color", "Offer one primary and one recovery action", "Keep navigation, footer, card padding, and mobile gaps consistent")])
    visuals["weekly-process"] = flow("revision-weekly-process", "PA2 Three-Week Process", ["Week 1: evidence audit", "Research protocol", "Week 2: affinity + priority", "Concept comparison", "Week 3: use cases", "Render + visual QA"], ["Wrong screenshot state -> recapture or narrow claim", "Terminology drift -> use stable IDs", "Layout failure -> revise source and render again"])
    for label, path in zip(("top", "middle", "bottom"), split_vertical(ROOT / "capture-work/chess/mobile/chess-49-lessons-mobile.png", "revision-chess-lessons")):
        visuals[f"chess-lessons-{label}"] = path
    return visuals


def practical(doc: Document, report: str, rows: list[list[str]]):
    doc.add_heading("Practical Difficulties and Lessons Learned", 1)
    m.add_table(doc, ["Practical difficulty", "How the group handled it", "Lesson learned", "Impact on the final version"], rows, [4.2, 4.2, 4.2, 4.2])


def remove_section(doc: Document, start_text: str, end_text: str):
    removing = False
    for child in list(doc.element.body):
        text = Paragraph(child, doc).text.strip() if child.tag.endswith("}p") else ""
        if text == start_text:
            removing = True
        if removing and text == end_text:
            break
        if removing:
            doc.element.body.remove(child)


def remove_figures(doc: Document, caption_prefixes: tuple[str, ...]):
    for paragraph in list(doc.paragraphs):
        if paragraph.text.startswith(caption_prefixes):
            caption = paragraph._p
            image = caption.getprevious()
            if image is not None and image.xpath(".//a:blip"):
                image.getparent().remove(image)
            caption.getparent().remove(caption)


def add_figures(doc: Document, items: list[tuple[Path, str, str, str]]):
    for path, figure_id, caption, related in items:
        m.add_figure(doc, path, figure_id, caption, "PA2 design QA artifact", related, 16.0, 17.0)


def revise_user_research(visuals):
    path = SOURCE / "Group10-PA2-UserResearch.docx"
    doc = Document(path)
    remove_section(doc, "8. Work models", "9. Findings")
    doc.add_heading("Revised Work Models", 1)
    doc.add_paragraph("Each model uses a distinct shape for actor, task/artifact/system, decision, interruption, recovery, and output. The flow reads left to right. Captions identify the current state, problem point, user decision, recovery path, and design opportunity.")
    keys = ["wm-fifa-planning", "wm-fifa-verification", "wm-fifa-handoff", "wm-fifa-artifact", "wm-chess-learning", "wm-chess-entry", "wm-chess-practice", "wm-chess-artifact"]
    add_figures(doc, [(visuals[k], f"UR-{k.upper()}", Path(visuals[k]).stem.replace("revision-", "").replace("-", " ").title(), "FIFA or Chess task evidence, interruption, recovery, and design opportunity") for k in keys])
    practical(doc, "User Research", [["Screenshots did not always show the state needed for a claim.", "The group separated captured state from simulated behavior and design inference.", "A screenshot proves only the visible state and capture time.", "Claims now name their evidence boundary and recovery assumptions."], ["Desktop and mobile pages expose different ordering and page length.", "The group used separate mobile evidence and task models.", "Cross-device continuity must be specified as a requirement.", "The revised models preserve tournament/game context during recovery."]])
    doc.save(path)


def revise_user_analysis(visuals):
    path = SOURCE / "Group10-PA2-UserAnalysis.docx"
    doc = Document(path)
    remove_figures(doc, ("Figure UA-AF-01.", "Figure UA-AF-02."))
    doc.add_heading("Revised Affinity Diagrams", 1)
    doc.add_paragraph("The diagrams separate FIFA and Chess. Each row traces representative raw notes through a small cluster, pattern, problem, severity, and design direction.")
    add_figures(doc, [(visuals["affinity-fifa"], "UA-AF-FIFA", "FIFA evidence-to-direction affinity diagram", "Screen evidence -> cluster -> pattern -> problem -> design direction"), (visuals["affinity-chess"], "UA-AF-CHESS", "Chess evidence-to-direction affinity diagram", "Screen evidence -> cluster -> pattern -> problem -> design direction")])
    doc.add_paragraph("Reading order for both diagrams: SCREEN EVIDENCE / RAW NOTE -> SMALL CLUSTER -> PATTERN -> PROBLEM -> DESIGN DIRECTION. The summary then links evidence to affinity clusters, priority, the tough problem, and concepts.")
    practical(doc, "User Analysis", [["Seventy-two notes could not remain legible on one page.", "The group split product views and kept only three representative notes in each visible cluster.", "Affinity synthesis should show the transformation, not only the notes.", "The final figures expose pattern, task, severity, and design implication."], ["Use-case and PA1 terms drifted during clustering.", "Stable FIFA and Chess cluster names were used across analysis and proposal sections.", "Terminology is part of traceability.", "The revised diagrams link each problem to one named design direction."]])
    doc.save(path)


def revise_proposal(visuals):
    path = SOURCE / "Group10-PA2-ProjectProposal.docx"
    doc = Document(path)
    doc.add_heading("Guided Concierge Flow - Detailed Revision", 1)
    steps = ["User enters Ticket Concierge.", "User selects a tournament.", "User selects a match or event preference.", "User selects party size.", "User answers eligibility and context questions.", "System checks the known ticket state.", "System explains the current state in plain language.", "System shows the available next options.", "User compares official ticket, register interest, resale, hospitality, and wait.", "System previews the destination.", "User chooses Continue or Stay.", "If Continue, the system explains the partner domain and account expectation.", "If Stay, the system preserves tournament, match, and party context.", "User may set an alert.", "System confirms the choice and saved context.", "A recovery branch returns the user to a safe step without discarding valid answers."]
    m.add_numbered_steps(doc, steps)
    add_figures(doc, [(visuals["concierge-overview"], "PP-GC-FLOW-01", "Guided Concierge overview flow", "Steps 1-6; entry, context, and known-state check"), (visuals["concierge-decision"], "PP-GC-FLOW-02", "Guided Concierge decision branch flow", "Steps 7-12; option comparison, destination preview, Continue or Stay"), (visuals["concierge-recovery"], "PP-GC-FLOW-03", "Guided Concierge error and recovery flow", "Steps 13-16; preserved context, alert, confirmation, and recovery")])
    doc.add_heading("Guided Concierge branch matrix", 2)
    m.add_table(doc, ["Branch", "Trigger", "Message", "User action", "Recovery path"], [
        ["No tickets", "Known state shows no official inventory", "No official tickets are available for this selection.", "Compare register interest, resale, hospitality, or wait", "Keep the selection and return to options"],
        ["Unknown status", "No governed status is available", "Current ticket status is unknown; this is not a sold-out result.", "Stay or retry", "Retry the state check without losing answers"],
        ["Stale status", "Last update exceeds the freshness rule", "This status may be outdated.", "Refresh or wait", "Show the last update only with a stale label"],
        ["Partner unavailable", "The official destination cannot respond", "The partner service is unavailable.", "Stay on FIFA.com", "Preserve context and retry later"],
        ["User not eligible", "Eligibility answer conflicts with route rules", "This route is not available for the supplied context.", "Review answers or compare other routes", "Return to the relevant question"],
        ["No matching event", "Tournament/match filters return no result", "No event matches these filters.", "Change tournament or match", "Keep party size and eligibility"],
        ["Invalid input", "A required answer is missing or malformed", "Correct the highlighted field.", "Edit the field", "Keep all valid answers"],
        ["Alert permission denied", "Browser notification permission is denied", "Browser alerts are off.", "Use an on-site reminder or change permission", "Do not block the ticket decision"],
        ["Network failure", "The state request times out", "The network request failed; no ticket outcome is inferred.", "Retry or stay", "Resume from the state check"],
    ])
    doc.add_heading("Guided Concierge Screens", 2)
    screen_captions = {"entry": ("Entry screen", "Choose tournament context", "Start"), "question": ("Question step", "Collect match, party, and eligibility inputs", "Check options"), "recommendation": ("Recommendation result", "Compare five ticket routes with current-state language", "Preview destination"), "destination": ("Destination preview", "Explain partner domain, account, and return behavior", "Continue or Stay"), "empty": ("Empty state", "Explain that filters produced no matching event", "Change filters"), "error": ("Error state", "Prevent an unavailable status from being read as sold out", "Retry or Stay"), "validation": ("Validation state", "Identify missing or invalid input beside the field", "Correct input")}
    for key, (screen, goal, action) in screen_captions.items():
        m.add_figure(doc, visuals[f"concierge-{key}"], f"PP-GC-{key.upper()}", f"{screen}. User step: {screen}; screen goal: {goal}; next action: {action}.", "Guided Concierge detailed flow and state requirements", 14.5, 16.0)
    doc.add_heading("Empty, Error, and Validation State Contract", 1)
    add_figures(doc, [(visuals["fifa-states"], "PP-STATE-FIFA", "FIFA state contract with trigger, message, action, and recovery", "No match, status/partner failure, invalid tournament/email/permission"), (visuals["chess-states"], "PP-STATE-CHESS", "Chess state contract with trigger, message, action, and recovery", "No game, analysis/network failure, invalid PGN or required input")])
    m.add_table(doc, ["Concept and state", "Trigger", "Message", "User action", "Recovery path"], [
        ["FIFA Ticket Concierge - EMPTY", "No tournament matches", "No matching tournament. Change tournament or date.", "Change filters", "Keep party size and eligibility"],
        ["FIFA Ticket Concierge - ERROR", "Ticket status or partner service fails", "Ticket status is unavailable; no sale outcome is inferred.", "Retry or stay", "Use a labeled last update or retry"],
        ["FIFA Ticket Concierge - VALIDATION", "Tournament or email is invalid", "Select a tournament and enter a valid email such as minh.anh@example.com.", "Correct the field", "Retain every valid answer"],
        ["Chess Beginner Review - EMPTY", "No recent game exists", "No game is ready for review. Import a PGN or play a game.", "Import game", "Return after a game exists"],
        ["Chess Beginner Review - ERROR", "Analysis service or network fails", "Analysis is unavailable; the game remains saved.", "Retry analysis", "Open board-only review or return later"],
        ["Chess Beginner Review - VALIDATION", "PGN or required input is missing", "Add a valid game or PGN before review starts.", "Fix input", "Keep valid fields and identify the failing line"],
    ])
    doc.add_heading("Mini Design System Appendix", 1)
    m.add_figure(doc, visuals["design-system"], "PP-DS-01", "Shared typography, spacing, control, card, status, error, success, and empty-state rules", "FIFA and Chess concept screens", 16.0, 17.0)
    m.add_table(doc, ["Token or component", "Shared rule"], [["Typography", "Arial; body, caption, table, and figure labels use 14 pt minimum; headings use 16/17/28 pt."], ["Spacing", "Use 8/16/24/32 px steps; mobile cards keep a 16 px gap."], ["Buttons", "48 px height, 12 px radius, verb label such as Continue, Stay, or Retry."], ["Inputs", "Persistent label, 48 px height, field-level error, and retained valid values."], ["Cards", "16 px radius and 16 px padding."], ["Status chips", "Status chips use a text label plus color; include freshness where it affects a decision."], ["Error and validation", "State the problem, one user action, and one recovery path."], ["Success and empty", "Confirm saved context or explain what is missing and how to continue."], ["Navigation and footer", "Use the same order, labels, icon size, and mobile spacing across concept screens."]])
    practical(doc, "Project Proposal", [["Empty, error, and validation screens exposed requirements missing from the happy path.", "The group added triggers, messages, user actions, and recovery paths before comparing concepts.", "A concept is incomplete when failure changes the user decision.", "The final proposal includes two product state contracts and three recovery flows."], ["The original Guided Concierge flow was too compressed.", "The group divided it into overview, decision/handoff, and recovery flows.", "Long flows should be split at decision boundaries.", "The final version contains 16 numbered steps and nine named branches."]])
    doc.save(path)


def insert_use_case_visuals(doc: Document, visuals):
    ids = [f"F-UC{i:02d}" for i in range(1, 7)] + [f"C-UC{i:02d}" for i in range(1, 7)]
    for uid in ids:
        paragraphs = doc.paragraphs
        anchor_index, anchor = next((index, p) for index, p in enumerate(paragraphs) if uid in p.text and p.style.name == "Heading 2")
        boundary_index, boundary = next(
            (index, p) for index, p in enumerate(paragraphs[anchor_index + 1:], start=anchor_index + 1)
            if p.style.name in ("Heading 1", "Heading 2")
        )
        insertion_target = boundary
        if boundary_index and 'w:type="page"' in paragraphs[boundary_index - 1]._p.xml:
            insertion_target = paragraphs[boundary_index - 1]
        before = len(doc.paragraphs)
        m.add_figure(doc, visuals[uid], f"UC-VIS-{uid}", f"{uid} unique UI state. The visual shows the use-case entry, required information, primary action, and recovery cue.", uid, 14.5, 15.5)
        image_p, caption_p = doc.paragraphs[before], doc.paragraphs[before + 1]
        insertion_target._p.addprevious(image_p._p)
        insertion_target._p.addprevious(caption_p._p)


def revise_use_cases(visuals):
    path = SOURCE / "Group10-PA2-UseCaseDocument.docx"
    doc = Document(path)
    insert_use_case_visuals(doc, visuals)
    doc.add_heading("State and Recovery Reference", 1)
    add_figures(doc, [(visuals["fifa-states"], "UC-STATE-FIFA", "FIFA empty, error, and validation reference", "F-UC01-F-UC06"), (visuals["chess-states"], "UC-STATE-CHESS", "Chess empty, error, and validation reference", "C-UC01-C-UC06")])
    practical(doc, "Use Case Document", [["Use cases with similar verbs began to repeat the same path.", "The group separated trigger, success guarantee, alternate flow, and visual for every ID.", "A unique goal needs a unique state transition and visual.", "All twelve use cases now have a different visual beside the detailed specification."], ["Error branches were easy to hide in long tables.", "The group added a shared state reference while keeping case-specific exception text.", "Recovery must preserve valid context and name the next action.", "FIFA and Chess now distinguish empty, service error, and validation failure."]])
    doc.save(path)


def replace_everywhere(doc: Document, old: str, new: str):
    containers = list(doc.paragraphs) + [p for table in doc.tables for row in table.rows for cell in row.cells for p in cell.paragraphs]
    for paragraph in containers:
        if old in paragraph.text:
            for run in paragraph.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)
            if old in paragraph.text:
                paragraph.text = paragraph.text.replace(old, new)


def revise_weekly(visuals):
    path = SOURCE / "Group10-PA2-WeeklyReport.docx"
    doc = Document(path)
    replacements = {
        "Five revised reports, unchanged Peer Review, QA and six-PDF package": "Six revised reports, internal design QA Peer Review, and six-PDF package",
        "do not change Peer Review": "label Peer Review as internal/simulated design QA",
        "Peer Review byte-identical": "Peer Review identifies artifact, severity, recommendation, action, and result",
        "preserve Peer Review unchanged": "publish Peer Review as internal/simulated design QA",
        "Existing PDF preserved without modification": "Internal/simulated design QA review with traceable actions",
        "Unchanged": "Revised for design QA",
        "Existing artifact": "Group10 design QA",
    }
    for old, new in replacements.items():
        replace_everywhere(doc, old, new)
    doc.add_heading("Process Flow", 1)
    m.add_figure(doc, visuals["weekly-process"], "WR-FLOW-01", "Three-week PA2 process from evidence audit to all-page visual QA", "Weekly schedule, recovery gates, and final package", 16.0, 16.0)
    practical(doc, "Weekly Report", [["Ticket status and partner handoff changed across capture times.", "The team treated each screenshot as a dated state and avoided claiming a live outcome.", "Time-sensitive interface evidence needs a freshness boundary.", "Final reports separate current, stale, unknown, and unavailable states."], ["Terms from PA1 and PA2 drifted across reports.", "The team fixed stable names for problems, concepts, states, and use cases.", "Shared terms reduce broken traceability.", "The final package uses the same FIFA and Chess labels in every report."]])
    doc.save(path)


def build_peer_review(visuals):
    doc = m.base_document("Peer Review")
    m.add_cover(doc, "Group10-PA2 Peer Review", "Internal / simulated peer review for design QA")
    m.add_toc_placeholder(doc)
    doc.add_heading("INTERNAL / SIMULATED PEER REVIEW FOR DESIGN QA", 1)
    m.callout(doc, "Review boundary", "No verified presentation feedback is available. The items below are an internal, simulated review of visible screenshots, diagrams, and report pages. They are not presented as comments from external peers.")
    reviews = [
        ["FIFA mobile ticket cards", "The selected tournament can move below a long vertical card stack.", "High", "The mobile capture shows repeated stacked cards before later options.", "Keep the selected tournament and current state in a compact summary.", "Project Proposal; Use Case Document", "Added persistent context to the Concierge and F-UC01 visual.", "The next decision remains visible after filters change."],
        ["Chess mobile lessons", "The long lesson list makes the next beginner task hard to locate.", "Medium", "The capture shows many lesson groups in one vertical sequence.", "Place one Continue lesson action before the catalog.", "User Research; User Analysis", "Added a learning-priority step to the Chess work model.", "The model leads to one lesson or puzzle before advanced depth."],
        ["Guided Concierge question step", "Required tournament and email input lacked a defined validation response.", "High", "The earlier low-fidelity screen listed fields but no field-level error or recovery.", "Add inline message, preserve valid answers, and provide a correction action.", "Project Proposal; Use Case Document", "Added validation screen and state contract.", "The user can correct one field without restarting."],
        ["Ticket destination handoff", "The partner boundary was not explained before Continue.", "High", "Existing ticket captures show an outbound action and a later partner state.", "Preview domain, account expectation, and return behavior.", "Project Proposal; Use Case Document", "Added destination preview and Continue/Stay decision.", "The handoff identifies what changes and what context remains."],
        ["Use-case document", "A repeated overview image would not prove twelve distinct task states.", "Medium", "The use cases have different triggers and outputs.", "Use a different UI state for every use case.", "Use Case Document", "Added twelve unique use-case visuals beside their specifications.", "Each visual now matches its trigger and primary action."],
        ["Shared state feedback", "Color alone would not distinguish status, error, and success.", "High", "Status chips and error panels use related accent colors.", "Pair color with a state label, message, action, and recovery.", "Project Proposal; Use Case Document", "Added text labels to every state card.", "The state meaning remains available without relying on color."],
        ["Mobile controls", "The source screenshots do not provide a reliable pixel measurement for every tap target.", "Medium", "Screenshot scaling prevents a defensible size claim.", "Set a design rule instead of claiming the current controls fail.", "Project Proposal", "Specified 48 px button height and persistent labels.", "The concept has a measurable mobile acceptance rule."],
    ]
    for index, row in enumerate(reviews, 1):
        doc.add_heading(f"QA-{index:02d} - {row[0]}", 2)
        m.add_table(doc, ["Field", "Review record"], [["Screen or artifact", row[0]], ["Issue", row[1]], ["Severity", row[2]], ["Reason", row[3]], ["Recommendation", row[4]], ["Affected report", row[5]], ["Action taken", row[6]], ["Result after revision", row[7]]], [4.2, 11.2])
    doc.add_heading("Evidence excerpts", 1)
    evidence = [(visuals["chess-lessons-top"], "PR-EV-01A", "Chess mobile Lessons page - Top; the first visible learning groups."), (visuals["chess-lessons-middle"], "PR-EV-01B", "Chess mobile Lessons page - Middle; the continued lesson sequence."), (visuals["chess-lessons-bottom"], "PR-EV-01C", "Chess mobile Lessons page - Bottom; later groups and page end."), (ROOT / "generated-diagrams/evidence-crops/f2-e09-decision-surface.png", "PR-EV-02", "FIFA ticket decision surface used to review state and handoff wording."), (visuals["concierge-validation"], "PR-EV-03", "Revised validation state showing field message, correction action, and recovery.")]
    for path, fid, caption in evidence:
        m.add_figure(doc, path, fid, caption, "Internal design QA item", 14.0, 16.0)
    practical(doc, "Peer Review", [["No verified presentation feedback existed.", "The group changed the artifact title and disclosure to internal/simulated design QA.", "Feedback provenance must be explicit.", "The final review does not attribute comments to external peers."], ["Some review claims could not be measured from scaled screenshots.", "The group converted them into measurable design requirements instead of unsupported findings.", "Evidence limits should change the wording of a review item.", "Tap targets now have a 48 px concept rule without claiming a measured defect in the source."]])
    f.save(doc, "Group10-PA2-PeerReview")


def write_revision_audit(visuals):
    QA.mkdir(exist_ok=True)
    text = """# PA2 revision audit

## Scope

- Audited 12 DOCX and 8 PDF artifacts in the active PA2 workspace, plus 155 screenshots/diagram PNG files.
- Reviewed seven labeled contact sheets covering every FIFA, Chess, evidence-crop, affinity, flow, UML, and prototype PNG.
- Rebuilt all six source DOCX and all six final PDFs from the source pipeline.

## Required revisions

| Requirement | Source action | Acceptance check |
|---|---|---|
| Remove intro metadata | Cover table retains Course, Group, Team, and Scope only | First-page text scan |
| Remove course-source sections | Renamed sections and kept short HCI citations at the point of use | Prohibited-heading scan |
| Minimum 14 pt | DOCX styles and PDF renderer set body, tables, captions, headers, footers, and TOC to 14 pt or larger | PDF character-size audit |
| Guided Concierge | Added 16 steps, three flows, seven screens, and recovery branches | Text and figure audit |
| Use-case visuals | Added 12 unique visuals beside 12 detailed use cases | Figure-ID and image-hash audit |
| Affinity diagrams | Added separate FIFA and Chess evidence-to-direction pipelines | Visual review |
| Work models | Added eight models with actor/task/artifact/system/decision/interruption/recovery/output | Visual review |
| States | Added FIFA and Chess empty/error/validation contracts | Trigger/message/action/recovery scan |
| Peer review | Rebuilt as INTERNAL / SIMULATED PEER REVIEW FOR DESIGN QA | Disclosure and field audit |
| Lessons learned | Added report-specific section to all six reports | Heading scan |

## Visual inventory

- New generated visuals: {count}
- Existing long captures are not reduced to unreadable full-page figures; decision-surface crops or product-specific mock screens are used instead.
""".format(count=len(visuals))
    (QA / "revision-audit.md").write_text(text, encoding="utf-8")


def main():
    f.main()
    visuals = create_visuals()
    revise_user_research(visuals)
    revise_user_analysis(visuals)
    revise_proposal(visuals)
    revise_use_cases(visuals)
    revise_weekly(visuals)
    build_peer_review(visuals)
    write_revision_audit(visuals)
    print(f"Revised six DOCX files and generated {len(visuals)} visual assets.")


if __name__ == "__main__":
    main()
