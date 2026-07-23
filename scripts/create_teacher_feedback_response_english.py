from __future__ import annotations

import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

import create_teacher_feedback_response as base


ROOT = Path(__file__).resolve().parents[1]
OUT_PDF = ROOT / "Group10-PA1-PeerReview-Revised-English.pdf"


def add_title_block(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = base.Pt(8)
    p.paragraph_format.space_after = base.Pt(2)
    base.set_run_font(p.add_run("TEACHER FEEDBACK RESPONSE"), size=22, bold=True, color=base.NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = base.Pt(3)
    base.set_run_font(p.add_run("PA1 - HCI Analysis of FIFA.com and Chess.com"), size=13, bold=True, color=base.BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = base.Pt(8)
    base.set_run_font(p.add_run("Group10 | Content baseline: PA1.pptx"), size=9.5, color=base.MUTED)


def add_scope(doc):
    base.add_section_bar(doc, "01 Scope update", page_break=False)
    base.add_body(
        doc,
        "This document is the formal response to the points that required clarification after the PA1 presentation. The product scope, user groups, use cases, benefits, drawbacks, and solutions remain consistent with PA1.pptx; the new analysis goes further by explaining how specific UI decisions affect user behavior.",
    )
    base.add_callout(
        doc,
        "Scope principle",
        "FIFA.com is treated as a browse-first ecosystem for information, viewing content, and ticketing, while Chess.com is an action-first platform for playing, learning, and reviewing. Every proposal identifies the exact screen area that would change.",
    )
    base.add_subheading(doc, "FIFA.com - PA1 baseline")
    base.add_table(
        doc,
        ["User groups", "Four use cases", "Benefits", "Drawbacks", "Solutions"],
        [[
            "Football fan; football viewer; ticket buyer.",
            "Fixtures and scores; match and player information; ticket booking control; live viewing and highlights.",
            "Task-clear navigation; date-based Match Centre; official source; tournament hub.",
            "Ecosystem sprawl; FIFA+ handoff; unclear ticket status; browse-first friction for quick tasks.",
            "Task-first navigation; audience intent chips; FIFA+ handoff explainer; ticket status dashboard; article utility rail.",
        ]],
        [2200, 3400, 3100, 3300, 3534],
        font_size=8.4,
    )
    base.add_subheading(doc, "Chess.com - PA1 baseline")
    base.add_table(
        doc,
        ["User groups", "Four use cases", "Benefits", "Drawbacks", "Solutions"],
        [[
            "Beginner learner; competitive player; returning casual player.",
            "Start a chess game; solve a puzzle; take a beginner lesson; review game analysis.",
            "Fast play flow; direct manipulation; Game Review feedback; puzzles and study path.",
            "Feature overload; dense analysis screen; premium gating; premove risk; hard-to-find Focus Mode; empty space.",
            "Goal-based onboarding; personal dashboard; beginner analysis preset; premove queue preview; Focus Mode coachmark.",
        ]],
        [2200, 3300, 3100, 3434, 3500],
        font_size=8.4,
    )


def add_fifa_response(doc):
    base.add_section_bar(doc, "02 FIFA.com teacher feedback response")
    base.add_subheading(doc, "A. Current ticket seating flow")
    base.add_numbered(doc, [
        "The user opens Tickets & Hospitality from the FIFA.com header.",
        "The user selects a tournament from the tournament-logo row.",
        "The page displays the Tickets or Hospitality card for the selected tournament.",
        "The CTA changes with the sales state and package type: Explore details, Buy Packages Now, Register your interest, or Buy now.",
        "After selecting a CTA, the user usually continues on a separate ticketing or hospitality channel. The seat map and seat selection appear later, but FIFA.com does not clearly explain when that step occurs.",
    ])
    base.add_body(
        doc,
        "The current flow tells users where to go next, but it does not answer four questions at the decision point: Are tickets on sale or is the page only collecting expressions of interest? At which step does the seat map appear? Which ticket categories or seating areas remain available? Does the CTA take the user away from FIFA.com? This is a visibility-of-system-status gap, not merely a missing flow step.",
    )
    base.add_callout(
        doc,
        "Ticket-buyer use case",
        "Before paying, the user must verify the official ticket source, seat location, ticket type, seat status, sale date, resale options, and hospitality options. This is a high-risk task because the decision involves money, a physical seat, and a real event that is difficult to reverse.",
    )
    base.add_subheading(doc, "B. Why third-party redirection feels uncomfortable")
    base.add_bullets(doc, [
        "The journey begins on FIFA.com, so the initial mental model is tied to an official source. When the header, logo, colors, account system, and CTA wording all change, the user must rebuild that mental model in the middle of the task.",
        "DAZN, ticketing vendors, hospitality providers, Store, Collect, and Rewards are different destinations within the ecosystem. Without a clear handoff signal, moving between them creates a continuity break rather than feeling like an ordinary page transition.",
        "Users must decide whether the destination is still official and safe, whether it could be a fraudulent page, and whether they can return to FIFA.com. This creates trust friction and mode-boundary friction.",
        "The friction is more serious in a ticket flow because the next step may request sign-in credentials, personal information, and payment.",
    ])

    doc.add_page_break()
    base.add_subheading(doc, "C. Confirmation before the handoff")
    base.add_body(doc, "The confirmation is not intended to slow the user down. It announces the system boundary and lets the user decide before the context, domain, and account system change.")
    base.add_table(
        doc,
        ["Component", "Displayed content"],
        [
            ["Primary message", "You are leaving FIFA.com to continue on the official ticketing partner site."],
            ["Context", "Tournament: FIFA World Cup 2026. | Task: Continue ticket purchase."],
            ["Trust information", "Official partner name; destination domain; official-partner indicator; return link to FIFA.com."],
            ["User control", "Buttons: Continue | Stay on FIFA.com."],
        ],
        [3100, 12434],
        font_size=8.8,
    )
    base.add_callout(
        doc,
        "HCI value",
        "The dialog improves visibility of system status, reinforces trust, increases user control, and reduces confusion about which website the user is visiting.",
    )
    base.add_subheading(doc, "D. Deeper UI design analysis")
    base.add_bullets(doc, [
        "The header should distinguish tasks performed on FIFA.com from links to sibling properties. A quick bar should prioritize Scores, Tickets, and Watch instead of making every destination compete at the same level.",
        "The tournament-logo row needs a clear selected state, and the breadcrumb should retain the tournament name and current location so that users do not lose context when switching tournaments.",
        "Each ticket card needs a state-specific CTA. Explore details is too vague for a purchase task; the card should also show a status label, sale date, ticket type, and seat-availability summary.",
        "The color hierarchy should separate the primary CTA from secondary links. The confirmation dialog should show the partner, domain, and option to stay. The dashboard table should support scanning by tournament, sale phase, seat category, and last update.",
        "The FIFA+ handoff needs an explainer card before the DAZN-branded surface opens, together with a return link to the originating FIFA.com content.",
    ])
    base.add_subheading(doc, "E. FIFA.com color analysis")
    base.add_body(
        doc,
        "Navy and blue support a sense of official authority, stability, and trust; white keeps news cards and content readable; tournament-specific colors create emotion and event recognition. However, when the colors of events, Store, Collect, FIFA+, Tickets, and News appear together, they compete for visual priority. In the ticket flow, color should communicate status rather than serve only as decoration.",
    )
    base.add_table(
        doc,
        ["Proposed color", "Status", "UI effect"],
        [
            ["Green", "Open / Available", "Immediately indicates that a ticket or seating area can be selected."],
            ["Yellow", "Register interest / Waiting", "Shows that immediate purchase is unavailable but a next action still exists."],
            ["Red", "Sold out / Closed", "Prevents CTA misunderstanding and repeated unsuccessful clicks."],
            ["Blue", "Official information", "Separates official information from sale status and partner actions."],
        ],
        [2200, 4100, 9234],
        font_size=8.6,
    )


FIFA_UI_ROWS = [
    ["Header navigation", "The top-level menu mixes news, tournaments, and sibling destinations.", "The browse-first structure makes quick-task users scan too many items.", "Add a task-first quick bar: Scores, Tickets, Watch, Rankings.", "Reduce cognitive load and shorten the route to key tasks."],
    ["Home hero", "The hero prioritizes featured content but does not clearly indicate a goal-based starting point.", "Weak information scent for users who need an immediate task.", "Add intent chips below the hero: Scores, Tickets, Watch.", "Users recognize a goal-based entry point without opening the menu."],
    ["Match Centre", "The date rail, match rows, and filters support scanning but disappear while scrolling.", "Users must return to the top to change scope.", "Add sticky quick filters: Today, Live, Results, Team, Competition.", "Keep controls visible and improve efficiency."],
    ["Tickets page", "Tournament logos and Tickets/Hospitality cards identify destinations but not sale or seat status.", "Low visibility of system status in a high-risk task.", "Add a ticket status dashboard: sale phase, seat availability, ticket type, date, resale, hospitality, and last update.", "Users know whether to buy, wait, register interest, or use resale."],
    ["Third-party handoff", "A CTA may open a different domain, brand, or account system.", "Continuity break and trust friction.", "Use a confirmation popup with official partner, domain preview, Continue, Stay, and a return link.", "Users understand the boundary crossing and retain control."],
    ["FIFA+ handoff", "The DAZN-branded surface uses Log in, Sign up, and Watch free controls that differ from FIFA.com.", "Mode-boundary friction; users may not understand the relationship between the two surfaces.", "Show an explainer card before handoff; preserve tournament/task context and a return link.", "Reduce hesitation and strengthen continuity in the viewing flow."],
    ["Article page", "After reading an article, users must return to the menu to find scores, tickets, or video.", "Browse-first friction interrupts the next task.", "Add an article utility rail: Scores, Tickets, Watch, Rankings.", "Move from content to a task without losing context."],
]


def add_fifa_ui_analysis(doc):
    base.add_section_bar(doc, "03 FIFA.com deeper UI design analysis")
    base.add_body(doc, "The table below moves each observation from flow-level description to the specific component, state, and on-screen behavior involved.")
    base.add_subheading(doc, "Entry points and primary tasks")
    base.add_table(
        doc,
        ["UI area", "Current design observation", "HCI issue", "Proposed UI change", "Expected result"],
        FIFA_UI_ROWS[:4],
        [1900, 3500, 3000, 4000, 3134],
        font_size=7.8,
    )
    base.add_subheading(doc, "Continuity across surfaces")
    base.add_table(
        doc,
        ["UI area", "Current design observation", "HCI issue", "Proposed UI change", "Expected result"],
        FIFA_UI_ROWS[4:],
        [1900, 3500, 3000, 4000, 3134],
        font_size=8.0,
    )
    base.add_callout(doc, "Priority", "The ticket status dashboard and handoff confirmation should be implemented together: the dashboard answers the ticket-status question, while the dialog explains where the user is going.")


def add_chess_response(doc):
    base.add_section_bar(doc, "04 Chess.com teacher feedback response")
    base.add_subheading(doc, "A. Effects of empty space")
    base.add_body(
        doc,
        "Whitespace around the chessboard can be purposeful: it separates the board from secondary content, reduces glare, and protects attention during a long game. The problem begins when empty space sits between modules or pushes important controls away from the main viewing area. In that case, information density decreases without providing additional focus.",
    )
    base.add_table(
        doc,
        ["User group", "Effect in a specific UI area", "HCI impact"],
        [
            ["Beginner learner", "Empty space between Play, Learn, Puzzles, and Review does not reveal the next action; priority cards are not placed near the starting point.", "Weak visual hierarchy, reduced affordance, higher scan load, and longer task-completion time."],
            ["Competitive player", "Whitespace around the board is acceptable if clocks, draw/resign controls, premove state, and settings remain close to the board.", "Hidden or distant controls increase motor and visual search under time pressure."],
            ["Returning casual player", "Large gaps between modules make Review, Learn, Puzzle, or Focus Mode harder to relocate.", "Recognition becomes less effective because the user must remember feature locations."],
        ],
        [2600, 7900, 5034],
        font_size=8.5,
    )
    base.add_callout(
        doc,
        "Conclusion",
        "The interface should not fill every empty area. It should preserve whitespace around the board for focus, while using priority cards and compact task grouping on the homepage and in module areas so that the remaining space is intentional.",
    )
    base.add_subheading(doc, "B. Chess.com color analysis")
    base.add_body(
        doc,
        "The dark background reduces glare during long sessions; green connects visually to the chessboard and highlights primary CTAs; white and gray separate text, panels, and the move list. However, green appears across many modules, so the primary CTA can compete with other components. On the analysis screen, evaluation colors, move labels, charts, and icons appear together, forcing beginners to decode both the content and the color semantics.",
    )
    base.add_table(
        doc,
        ["Color in beginner mode", "Single meaning", "Usage"],
        [
            ["Green", "Good move or primary action", "Use only for the best move, correct state, and continue CTA."],
            ["Yellow", "Requires attention", "Use for an inaccuracy or a prompt to review a move."],
            ["Red", "Blunder or error", "Use sparingly for the main error, not across several panels at once."],
            ["Gray", "Secondary information", "De-emphasize engine detail, advanced options, and metadata."],
        ],
        [3100, 4200, 8234],
        font_size=8.5,
    )


CHESS_UI_ROWS = [
    ["Home / feature menu", "Play, Learn, Puzzles, Train, and Watch compete at the entry point.", "Feature overload; beginners do not know where to start.", "Use goal-based onboarding cards: Play a game, Learn basics, Solve a puzzle.", "Improve information scent and reduce initial choices."],
    ["Empty-space areas", "Space between modules is not always connected to a focus goal.", "Low information density and increased scan load.", "Use priority cards for the three most relevant tasks and compact grouping for secondary tasks.", "Use the space without turning the homepage into a dense dashboard."],
    ["Game board", "Whitespace around the board supports focus, but some controls are subtle or depend on hover.", "Reduced affordance and discoverability when controls are hidden.", "Keep breathing room around the board; place clocks, draw/resign, settings, and premove state in a stable viewing area.", "Preserve focus without sacrificing user control."],
    ["Analysis screen", "Charts, evaluation, engine lines, move labels, and coach feedback appear simultaneously.", "High cognitive load and weak progressive disclosure.", "Use a beginner analysis preset: main error, best move, one short explanation; reveal advanced detail on demand.", "Beginners understand feedback before reading engine detail."],
    ["Color system", "Green, red, yellow, chart colors, and icons can communicate several meanings.", "Feedback ambiguity and visual noise.", "Standardize Green/Yellow/Red/Gray so each color has one role in beginner mode.", "Reduce status-decoding time and improve consistency."],
    ["Focus Mode", "The control is hard to notice when it depends on hovering near the board/sidebar boundary.", "Low discoverability for a feature intended to reduce distraction.", "Show a coachmark after several games; add a persistent Settings shortcut and labeled tooltip.", "Returning users can find and enable the mode more easily."],
    ["Premove", "A queued move can execute after an unexpected reply.", "Insufficient error prevention and feedback under time pressure.", "Show a queue preview on the board, a clear state, and a one-action clear shortcut.", "Reduce unintended premove blunders."],
]


def add_chess_ui_analysis(doc):
    base.add_section_bar(doc, "05 Chess.com deeper UI design analysis")
    base.add_subheading(doc, "Entry points, empty space, and the board")
    base.add_table(
        doc,
        ["UI area", "Current design observation", "HCI issue", "Proposed UI change", "Expected result"],
        CHESS_UI_ROWS[:4],
        [1900, 3500, 3000, 4000, 3134],
        font_size=7.8,
    )
    base.add_subheading(doc, "Feedback, focus, and error prevention")
    base.add_table(
        doc,
        ["UI area", "Current design observation", "HCI issue", "Proposed UI change", "Expected result"],
        CHESS_UI_ROWS[4:],
        [1900, 3500, 3000, 4000, 3134],
        font_size=8.0,
    )
    base.add_callout(doc, "Principle", "The beginner preset reduces the number of signals that must be read, while the competitive view retains advanced data. Both modes use the same color semantics so that feedback does not change meaning.")


MAPPING_ROWS = [
    ["FIFA.com", "Ticket seating flow", "Ticket status uncertainty", "Visibility of system status; trust", "Ticket status dashboard + seat availability summary + official partner confirmation", "Cards and dashboard show sale phase, seat category, availability, and last update; the dialog shows partner/domain before the CTA.", "Users know whether to buy, wait, or use resale; risk and repeated clicks decrease."],
    ["FIFA.com", "Third-party redirection", "Ecosystem sprawl; FIFA+ handoff", "Continuity break; trust friction", "Handoff explainer popup + destination preview + return link", "The modal states brand, domain, tournament/task, Continue, and Stay; the destination retains a return link.", "Reduce hesitation while preserving the mental model and user control."],
    ["FIFA.com", "Design depth", "Browse-first friction", "Cognitive load", "Task-first navigation + quick task bar + article utility rail", "The header adds Scores/Tickets/Watch/Rankings; intent chips appear below the hero; a utility rail sits beside articles.", "Quick tasks require less scanning and fewer returns to the menu."],
    ["FIFA.com", "Color", "Competing visual hierarchy", "Weak visual priority", "Color-coded status labels + consistent CTA hierarchy", "Green/Yellow/Red indicate sale state; Blue marks official information; each ticket card has one primary CTA.", "Users recognize status quickly and distinguish primary from secondary actions."],
    ["Chess.com", "Empty space", "Empty space / inefficient screen use", "Visual hierarchy; information density; scan load", "Purposeful whitespace + priority cards + compact task grouping", "Preserve breathing room around the board; place three priority cards in homepage space; group secondary items compactly.", "Maintain focus while reducing the time needed to find Learn, Puzzle, and Review."],
    ["Chess.com", "Color", "Analysis density; feature overload", "Cognitive load; feedback ambiguity", "Simplified beginner color system + clearer feedback states", "Beginner view uses Green/Yellow/Red/Gray consistently; engine detail and metadata move to a gray advanced layer.", "Make feedback easier to understand and reduce visual noise."],
]


def add_mapping(doc):
    base.add_section_bar(doc, "06 Revised solution mapping")
    base.add_body(doc, "Each concern is linked directly to a PA1 drawback, its HCI cause, and a visible change on the screen.")
    base.add_subheading(doc, "FIFA.com")
    base.add_table(
        doc,
        ["Website", "Teacher concern", "Linked PA1 drawback", "HCI cause", "UI-level solution", "What exactly changes on screen", "Expected improvement"],
        MAPPING_ROWS[:4],
        [1000, 1800, 2100, 2100, 2950, 3300, 2284],
        font_size=7.15,
    )
    base.add_subheading(doc, "Chess.com")
    base.add_table(
        doc,
        ["Website", "Teacher concern", "Linked PA1 drawback", "HCI cause", "UI-level solution", "What exactly changes on screen", "Expected improvement"],
        MAPPING_ROWS[4:],
        [1000, 1800, 2100, 2100, 2950, 3300, 2284],
        font_size=7.5,
    )
    base.add_callout(doc, "Evidence boundary", "The seat-availability dashboard, handoff confirmation, and status color system are proposed designs. They are not described as existing features in the current interface.")


def add_presentation_notes(doc):
    base.add_section_bar(doc, "07 Final presentation notes")
    base.add_body(doc, "During the presentation, every claim should be tied to a specific UI area rather than reading the table row by row.")
    base.add_bullets(doc, [
        "FIFA ticketing: point to Tickets & Hospitality in the header, the tournament-logo row, the Tickets/Hospitality cards, and the CTA. Explain that the card identifies a destination but does not yet show sale state, seat-map stage, or seat availability.",
        "FIFA handoff: point to FIFA+ or the DAZN-branded surface and the Log in, Sign up, and Watch free controls. Explain how the change in brand, account system, and CTA breaks continuity.",
        "FIFA color: point to the navy header, white content cards, tournament colors, and primary CTA. Then compare them with Green/Yellow/Red status labels and Blue official information.",
        "Chess.com empty space: point to the homepage/module area or another region with substantial whitespace. Distinguish purposeful whitespace around the board from gaps that separate priority modules.",
        "Chess.com color: point to the board, CTA, analysis feedback, and move labels. Explain why beginner mode needs fewer colors and why each color should retain one meaning.",
    ])
    base.add_callout(doc, "Closing statement", "The group is not proposing that every screen be filled. The goal is to clarify status, preserve continuity, and place important actions where users make decisions.")


def build():
    doc = Document()
    base.configure_document(doc)
    add_title_block(doc)
    add_scope(doc)
    add_fifa_response(doc)
    add_fifa_ui_analysis(doc)
    add_chess_response(doc)
    add_chess_ui_analysis(doc)
    add_mapping(doc)
    add_presentation_notes(doc)
    doc.core_properties.title = "Group10-PA1 Teacher Feedback Response - English"
    doc.core_properties.subject = "HCI analysis of FIFA.com and Chess.com"
    doc.core_properties.author = "Group10"

    with tempfile.TemporaryDirectory(prefix="pa1-english-") as temp_dir:
        temp_docx = Path(temp_dir) / "teacher-feedback-response-english.docx"
        doc.save(temp_docx)
        base.OUT_DOCX = temp_docx
        base.OUT_PDF = OUT_PDF
        base.PDF_TITLE = "Group10-PA1 Teacher Feedback Response - English"
        base.PDF_PAGE_LABEL = "Page"
        base.build_pdf_from_docx()


if __name__ == "__main__":
    build()
