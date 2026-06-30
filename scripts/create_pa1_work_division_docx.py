from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "output"
DOCS_DIR = ROOT / "docs"
OUT_DOCX = OUTPUT_DIR / "GroupID-PA1-WorkDivision.docx"
ROOT_DOCX = ROOT / "GroupID-PA1-WorkDivision.docx"
LOG_PATH = DOCS_DIR / "pa1_work_division_generation_log.md"
ROOT_COPY_STATUS = {"copied": False, "error": ""}

MEMBERS = [
    ("Le Minh", "21127645"),
    ("Nguyen Vu Bach", "21127224"),
    ("Pham Nguyen Gia Bao", "20127119"),
    ("Trang Minh Nhut", "22127318"),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="8A8F98", size="6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_cell_text(cell, bold=False, color="000000", size=9) -> None:
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.05
        for run in paragraph.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.color.rgb = RGBColor.from_string(color)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_borders(table)
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
        set_cell_shading(hdr[i], "F2F4F7")
        style_cell_text(hdr[i], bold=True, color="0B2545", size=9)
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            cells[i].text = str(value)
            style_cell_text(cells[i], size=8.6)
    set_table_width(table, widths)
    doc.add_paragraph()
    return table


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run(text)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_page_break_heading(doc: Document, text: str) -> None:
    doc.add_page_break()
    add_heading(doc, text)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(13)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(46, 116, 181)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.10

    header = section.header.paragraphs[0]
    header.text = "Phan chia cong viec PA1"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(85, 85, 85)

    footer = section.footer.paragraphs[0]
    footer.text = "Du an HCI cho FIFA.com va Chess.com"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(85, 85, 85)


def build_docx() -> None:
    OUTPUT_DIR.mkdir(exist_ok=True)
    DOCS_DIR.mkdir(exist_ok=True)

    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("PHÂN CHIA CÔNG VIỆC PA1")
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.font.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("Nghiên cứu HCI cho FIFA.com và Chess.com")
    run.font.name = "Calibri"
    run.font.size = Pt(12)

    add_heading(doc, "1. Mục tiêu đồ án")
    add_body(
        doc,
        "Nhóm hoàn thiện gói PA1 cho cặp sản phẩm FIFA.com và Chess.com theo góc nhìn HCI. "
        "Các sản phẩm nộp gồm ProductResearch, PotentialSolutions, PeerReview, WeeklyReport và gói nén cuối cùng. "
        "Công việc nhấn mạnh bằng chứng ảnh chụp màn hình, chú thích hình, nguồn chính thức, phân tích HCI, giải pháp cải thiện và kiểm tra chất lượng trước khi nộp.",
    )

    add_heading(doc, "2. Thành viên nhóm")
    add_table(
        doc,
        ["STT", "Thành viên", "MSSV", "Vai trò chính", "Vai trò hỗ trợ"],
        [
            ["1", "Le Minh", "21127645", "Điều phối dự án, tích hợp, PeerReview, WeeklyReport, đóng gói cuối", "Kiểm tra nguồn, rà soát PDF, kiểm tra zip"],
            ["2", "Nguyen Vu Bach", "21127224", "Trưởng nghiên cứu FIFA.com, phụ trách bằng chứng ảnh FIFA.com", "Đồng dẫn ProductResearch và rà soát phát hiện HCI"],
            ["3", "Pham Nguyen Gia Bao", "20127119", "Trưởng nghiên cứu Chess.com, phụ trách bằng chứng ảnh Chess.com", "Đồng dẫn ProductResearch và rà soát phát hiện HCI"],
            ["4", "Trang Minh Nhut", "22127318", "Trưởng phân tích HCI, PotentialSolutions, kiểm tra hình ảnh", "Rà soát chú thích, ma trận drawback-solution và nhất quán báo cáo"],
        ],
        [600, 1900, 1200, 3000, 2660],
    )

    add_heading(doc, "3. Phân chia trách nhiệm tổng quan")
    add_table(
        doc,
        ["Thành viên", "Trách nhiệm chính", "Trách nhiệm hỗ trợ", "Đầu ra dự kiến"],
        [
            [
                "Le Minh",
                "Khóa phạm vi, điều phối tiến độ, tích hợp nội dung, viết PeerReview và WeeklyReport, xuất PDF và đóng gói zip.",
                "Hỗ trợ rà soát nguồn, tên file, checklist nộp bài và kiểm tra văn bản trích xuất từ PDF.",
                "PeerReview, WeeklyReport, checklist cuối, PDF và GroupID-PA1.zip.",
            ],
            [
                "Nguyen Vu Bach",
                "Nghiên cứu FIFA.com, thu thập nguồn chính thức, chụp và chú thích ảnh, viết persona, use case, lợi ích, hạn chế và phát hiện HCI cho FIFA.com.",
                "Hỗ trợ đối chiếu ProductResearch với hình ảnh, caption và nguồn FIFA.",
                "Phần FIFA.com trong ProductResearch và bằng chứng ảnh FIFA.com.",
            ],
            [
                "Pham Nguyen Gia Bao",
                "Nghiên cứu Chess.com, thu thập nguồn chính thức, chụp và chú thích ảnh, viết persona, use case, lợi ích, hạn chế và phát hiện HCI cho Chess.com.",
                "Hỗ trợ kiểm tra phạm vi Chess.com để báo cáo không quá rộng.",
                "Phần Chess.com trong ProductResearch và bằng chứng ảnh Chess.com.",
            ],
            [
                "Trang Minh Nhut",
                "Ánh xạ khái niệm HCI, dẫn PotentialSolutions, kiểm tra caption, hình minh họa và tính nhất quán giữa drawback và solution.",
                "Hỗ trợ kiểm tra trực quan cho ProductResearch, PotentialSolutions và PeerReview.",
                "PotentialSolutions, ma trận RACI/QA trực quan và rà soát hình ảnh cuối.",
            ],
        ],
        [1550, 3200, 2500, 2110],
    )

    add_page_break_heading(doc, "4. Phân công theo sản phẩm nộp")
    add_table(
        doc,
        ["Sản phẩm", "Người phụ trách chính", "Người phối hợp", "Nội dung công việc"],
        [
            ["GroupID-PA1-ProductResearch.pdf", "Nguyen Vu Bach, Pham Nguyen Gia Bao", "Le Minh, Trang Minh Nhut", "Nghiên cứu FIFA.com và Chess.com, persona, use case, HCI findings, hình ảnh, nguồn tham khảo."],
            ["GroupID-PA1-PotentialSolutions.pdf", "Trang Minh Nhut", "Nguyen Vu Bach, Pham Nguyen Gia Bao", "Chuyển drawback thành giải pháp, ưu tiên impact-effort, mô tả UI và kiểm tra mapping."],
            ["GroupID-PA1-PeerReview.pdf", "Le Minh", "Cả nhóm", "Kịch bản 7 phút, slide outline, câu hỏi dự kiến, phản hồi mock/internal rehearsal và owner thật."],
            ["GroupID-PA1-WeeklyReport.pdf", "Le Minh", "Cả nhóm", "Lịch họp, scrum theo từng thành viên, sprint review, workload matrix và checklist cuối."],
            ["GroupID-PA1.zip", "Le Minh", "Trang Minh Nhut", "Đóng gói đúng bốn PDF ở cấp cao nhất và kiểm tra bằng zip listing."],
        ],
        [2200, 2300, 2100, 2760],
    )

    add_heading(doc, "5. Tóm tắt cân bằng khối lượng công việc")
    add_table(
        doc,
        ["Hạng mục", "Le Minh", "Nguyen Vu Bach", "Pham Nguyen Gia Bao", "Trang Minh Nhut"],
        [
            ["Điều phối và tích hợp", "Chính", "Hỗ trợ", "Hỗ trợ", "Hỗ trợ"],
            ["Nghiên cứu FIFA.com", "Rà soát", "Chính", "Tham khảo", "HCI review"],
            ["Nghiên cứu Chess.com", "Rà soát", "Tham khảo", "Chính", "HCI review"],
            ["Bằng chứng ảnh", "Kiểm tra cuối", "FIFA.com", "Chess.com", "Visual QA"],
            ["ProductResearch", "Tích hợp", "FIFA.com", "Chess.com", "HCI consistency"],
            ["PotentialSolutions", "Rà soát", "Input FIFA", "Input Chess.com", "Chính"],
            ["PeerReview và WeeklyReport", "Chính", "Input FIFA", "Input Chess.com", "Visual/HCI input"],
            ["Đóng gói cuối", "Chính", "Kiểm tra PDF", "Kiểm tra PDF", "Kiểm tra hình/caption"],
        ],
        [2050, 1800, 1900, 1900, 1710],
    )

    add_page_break_heading(doc, "6. Ma trận RACI")
    add_body(doc, "Chú thích: R = Responsible, A = Accountable, C = Consulted, I = Informed.")
    add_table(
        doc,
        ["Công việc", "Le Minh", "Nguyen Vu Bach", "Pham Nguyen Gia Bao", "Trang Minh Nhut"],
        [
            ["Scope và checklist", "A/R", "C", "C", "C"],
            ["FIFA.com research", "C", "A/R", "I", "C"],
            ["Chess.com research", "C", "I", "A/R", "C"],
            ["HCI analysis", "C", "C", "C", "A/R"],
            ["PotentialSolutions", "C", "C", "C", "A/R"],
            ["PeerReview", "A/R", "C", "C", "C"],
            ["WeeklyReport", "A/R", "C", "C", "C"],
            ["Visual QA", "C", "C", "C", "A/R"],
            ["Final export và zip packaging", "A/R", "C", "C", "C"],
        ],
        [2600, 1500, 1900, 1900, 1460],
    )

    add_heading(doc, "7. Đóng góp dự kiến của từng thành viên")
    contributions = {
        "Le Minh": "Điều phối phạm vi, tích hợp nội dung, viết PeerReview và WeeklyReport, chạy xuất PDF, kiểm tra văn bản PDF và đóng gói zip cuối.",
        "Nguyen Vu Bach": "Phụ trách nghiên cứu FIFA.com, nguồn chính thức, ảnh chụp và chú thích FIFA.com, persona/use case, HCI findings, lợi ích và hạn chế của FIFA.com.",
        "Pham Nguyen Gia Bao": "Phụ trách nghiên cứu Chess.com, nguồn chính thức, ảnh chụp và chú thích Chess.com, persona/use case, HCI findings, lợi ích và hạn chế của Chess.com.",
        "Trang Minh Nhut": "Phụ trách ánh xạ HCI, viết PotentialSolutions, kiểm tra giải pháp với drawback, rà soát hình ảnh, caption và tính nhất quán trình bày.",
    }
    for member, text in contributions.items():
        p = doc.add_paragraph()
        p.add_run(f"{member}: ").bold = True
        p.add_run(text)

    add_page_break_heading(doc, "8. Checklist chất lượng")
    checklist = [
        "Cặp sản phẩm là FIFA.com và Chess.com.",
        "Tất cả bốn thành viên và MSSV xuất hiện đúng.",
        "ProductResearch có persona, use case, ảnh chụp, HCI findings, lợi ích, hạn chế, so sánh và nguồn.",
        "PotentialSolutions map từng drawback sang giải pháp và có ưu tiên impact-effort.",
        "PeerReview có kịch bản 7 phút, slide outline, owner thật và phản hồi mock/internal rehearsal được ghi nhãn rõ.",
        "WeeklyReport có sprint planning, hai weekly scrum theo từng thành viên, sprint review, workload matrix và checklist.",
        "Zip cuối chỉ chứa bốn PDF ở cấp cao nhất.",
        "Tài liệu phân công chỉ mô tả trách nhiệm tổng quan, không mô tả lịch triển khai chi tiết.",
        "Mỗi thành viên có đóng góp nghiên cứu, viết, rà soát hoặc QA có ý nghĩa.",
        "RACI matrix, quality checklist và signature table đều có mặt.",
    ]
    for item in checklist:
        p = doc.add_paragraph()
        p.style = doc.styles["Normal"]
        p.add_run("\u2610 ").font.size = Pt(11)
        p.add_run(item)

    add_heading(doc, "9. Bảng xác nhận")
    add_table(
        doc,
        ["Thành viên", "MSSV", "Xác nhận trách nhiệm", "Chữ ký"],
        [[name, sid, "Đã xác nhận", ""] for name, sid in MEMBERS],
        [2300, 1500, 3100, 2460],
    )

    doc.save(OUT_DOCX)
    try:
        shutil.copy2(OUT_DOCX, ROOT_DOCX)
        ROOT_COPY_STATUS["copied"] = True
    except PermissionError as exc:
        ROOT_COPY_STATUS["copied"] = False
        ROOT_COPY_STATUS["error"] = str(exc)


def validate_docx(path: Path) -> dict[str, object]:
    if not path.exists():
        return {
            "path": str(path),
            "exists": False,
            "size": 0,
            "size_gt_10kb": False,
            "required_terms_ok": False,
            "balanced_research_ok": False,
            "no_detailed_timeline": False,
            "table_count": 0,
        }
    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            text += "\n" + "\t".join(cell.text for cell in row.cells)
    required = [
        "PHÂN CHIA CÔNG VIỆC PA1",
        "Nghiên cứu HCI cho FIFA.com và Chess.com",
        "Mục tiêu đồ án",
        "Thành viên nhóm",
        "Phân chia trách nhiệm tổng quan",
        "Phân công theo sản phẩm nộp",
        "Tóm tắt cân bằng khối lượng công việc",
        "Ma trận RACI",
        "Đóng góp dự kiến của từng thành viên",
        "Checklist chất lượng",
        "Bảng xác nhận",
        "Le Minh",
        "Nguyen Vu Bach",
        "Pham Nguyen Gia Bao",
        "Trang Minh Nhut",
        "21127645",
        "21127224",
        "20127119",
        "22127318",
        "GroupID-PA1-ProductResearch.pdf",
        "GroupID-PA1-PotentialSolutions.pdf",
        "GroupID-PA1-PeerReview.pdf",
        "GroupID-PA1-WeeklyReport.pdf",
        "GroupID-PA1.zip",
    ]
    forbidden = ["14-day", "day-by-day", "Day 1", "Day 14", "14 ngày", "từng ngày"]
    return {
        "path": str(path),
        "exists": path.exists(),
        "size": path.stat().st_size if path.exists() else 0,
        "size_gt_10kb": path.exists() and path.stat().st_size > 10_000,
        "required_terms_ok": all(term in text for term in required),
        "balanced_research_ok": all(
            term in text
            for term in [
                "FIFA.com",
                "Chess.com",
                "PotentialSolutions",
                "PeerReview",
                "WeeklyReport",
                "A/R",
            ]
        ),
        "no_detailed_timeline": not any(term in text for term in forbidden),
        "table_count": len(doc.tables),
    }


def write_log(results: list[dict[str, object]]) -> None:
    cwd = ROOT
    lines = [
        "# PA1 Work Division Generation Log",
        "",
        f"Generated at: {datetime.now().isoformat(timespec='seconds')}",
        "",
        "## Phase 0 inspection",
        f"- Current working directory: `{cwd}`",
        "- Repository tree listed to depth 3.",
        "- Existing folders checked: `docs` exists; `output` created if missing; no `reports` or `deliverables` folder required for this task.",
        "- Python check: bundled Codex Python 3.12.13.",
        "- python-docx check: OK.",
        "- Design preset: `standard_business_brief`, with A4 page-size override retained from the existing script.",
        "",
        "## Generation actions",
        "- Created `output/GroupID-PA1-WorkDivision.docx`.",
        f"- Root copy status for `GroupID-PA1-WorkDivision.docx`: {ROOT_COPY_STATUS['copied']}.",
        f"- Root copy error: {ROOT_COPY_STATUS['error'] or 'None'}.",
        "- Applied A4 page size, normal margins, centered title/subtitle, bold section headings, visible table borders, header, footer, and required page breaks.",
        "- Updated assignment model: this Vietnamese document covers ProductResearch, PotentialSolutions, PeerReview, WeeklyReport, and final zip packaging.",
        "- RACI model: Le Minh leads scope, PeerReview, WeeklyReport, final export, and zip packaging; Nguyen Vu Bach leads FIFA.com research; Pham Nguyen Gia Bao leads Chess.com research; Trang Minh Nhut leads HCI analysis, PotentialSolutions, and visual QA.",
        "- Each member has meaningful research, writing, review, or QA responsibilities.",
        "",
        "## Validation results",
    ]
    for result in results:
        lines.extend(
            [
                f"### {result['path']}",
                f"- Exists: {result['exists']}",
                f"- Size: {result['size']} bytes",
                f"- Size > 10 KB: {result['size_gt_10kb']}",
                f"- Required title/member/RACI/checklist terms present: {result['required_terms_ok']}",
                f"- Balanced all-deliverable responsibility model present: {result['balanced_research_ok']}",
                f"- No detailed day-by-day or 14-day plan terms: {result['no_detailed_timeline']}",
                f"- Table count: {result['table_count']}",
                "",
            ]
        )
    LOG_PATH.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    build_docx()
    results = [validate_docx(OUT_DOCX), validate_docx(ROOT_DOCX)]
    write_log(results)
    for result in results:
        print(result)
    output_result = results[0]
    if not (
        output_result["exists"]
        and output_result["size_gt_10kb"]
        and output_result["required_terms_ok"]
        and output_result["balanced_research_ok"]
        and output_result["no_detailed_timeline"]
    ):
        raise SystemExit(1)


if __name__ == "__main__":
    main()
