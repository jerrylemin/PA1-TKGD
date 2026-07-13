from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt
import re


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "archive/potential_solutions_before_live_ui_revision_20260707/Group10-PA1-PotentialSolutions_VisualReport (1).docx"
OUTPUT = ROOT / "Group10-PA1-PotentialSolutions_VisualReport_links_updated.docx"


CARDS = {
    4: {
        "image": ROOT / "assets/screenshots/solution-references/fifa_footer_clean.png",
        "evidence": "FIFA.com footer/ecosystem navigation. Corresponds to ProductResearch F-09B and F-HCI6.",
        "bullets": [
            "Global header: keep Match Centre, News, Rankings, Tickets, and Watch visible as the five primary task destinations (F-S1).",
            "Header overflow: move Store, Collect, Rewards, and other corporate properties into a clearly labelled More FIFA menu.",
            "Directly below the homepage hero: add intent chips for Scores, News, Rankings, Tickets, and Watch (F-S2).",
            "Cross-property links: show destination brand and a Back to FIFA.com path before users leave the current property.",
        ],
    },
    5: {
        "image": ROOT / "assets/screenshots/solution-references/fifa_plus_boundary_clean.png",
        "evidence": "FIFA+ landing surface after the DAZN handoff. Corresponds to ProductResearch F-10B and F-HCI7.",
        "bullets": [
            "Before the FIFA+ outbound action: insert an explainer card stating that FIFA+ is powered by DAZN and whether sign-in is required (F-S3).",
            "FIFA+ top bar: add a persistent Back to FIFA.com control and retain FIFA visual identity during the handoff (F-S4).",
            "Primary CTA area: state what opens next—live stream, replay, or sign-in—before navigation occurs.",
            "Account boundary: distinguish FIFA.com and DAZN account state with explicit labels instead of relying on branding alone.",
        ],
    },
    6: {
        "image": ROOT / "assets/screenshots/solution-references/fifa_plus_boundary_clean.png",
        "evidence": "FIFA+ live, archive, originals, schedule, and media rails. Corresponds to ProductResearch F-10B/F-HCI8 and source [6].",
        "bullets": [
            "Above the first media rail: add filters for Live, Highlights, Documentaries, Competitions, and Archive (F-S5).",
            "Each rail header: show item count and a concise See all action so users can predict the content scope.",
            "Media results area: add Compact scan mode with smaller cards and metadata-first rows (F-S6).",
            "Nonmatching rails: collapse them to labelled summaries after a filter is selected, with an option to restore all content.",
        ],
    },
    7: {
        "image": ROOT / "assets/screenshots/solution-references/fifa_tickets_clean.png",
        "evidence": "Official FIFA Tickets entry page. Corresponds to ProductResearch F-HCI9 and sources [5][9].",
        "bullets": [
            "Tickets landing page, above tournament cards: add a status dashboard for Official sale, Resale, Waiting room, and Sold out (F-S7).",
            "Every ticket card: display last-updated time, next known milestone, and one official action button.",
            "Beside each tournament status: add opt-in email/browser availability alerts by date and team (F-S8).",
            "Waiting-room and unavailable states: preserve the selected tournament and explain whether users should act now or return later.",
        ],
    },
    8: {
        "image": ROOT / "assets/screenshots/solution-references/fifa_article_clean.png",
        "evidence": "FIFA.com news article surface. Corresponds to ProductResearch F-06 and F-HCI10.",
        "bullets": [
            "Article desktop layout: add a sticky utility rail for Scores, Tickets, Watch, and related tournament information (F-S9).",
            "Article header, below title metadata: add context chips such as Open Match Centre, Ticket status, and Watch highlights (F-S10).",
            "On mobile: convert the utility rail into a bottom sheet that does not cover the article text.",
            "Action labels: derive the linked competition/team from article metadata so links remain contextual rather than generic.",
        ],
    },
    10: {
        "image": ROOT / "assets/screenshots/solution-references/chess_navigation_ui.png",
        "evidence": "Chess.com home navigation and task surface. Corresponds to ProductResearch C-01/C-06 and C-HCI7.",
        "bullets": [
            "First-run home area: replace product taxonomy with goal choices—Play now, Review, Learn basics, and Train puzzles (C-S1).",
            "After goal selection: reveal only the next two or three relevant actions, with a View all features escape hatch.",
            "Returning-user home: add a personal dashboard where users pin three frequent tasks and hide unused modules (C-S2).",
            "Left navigation: keep stable labels but visually separate beginner goals from advanced tools and community content.",
        ],
    },
    11: {
        "image": ROOT / "assets/screenshots/solution-references/chess_analysis_ui.png",
        "evidence": "Chess.com Analysis Board UI. Corresponds to ProductResearch C-03, C-HCI4/C-HCI10, and sources [15][16][17].",
        "bullets": [
            "Game Review toolbar: add a Beginner preset that keeps evaluation, best move, and one plain-language explanation visible (C-S3).",
            "Advanced charts, engine lines, classifications, and toggles: place them behind a clearly labelled Advanced controls disclosure.",
            "Beside unfamiliar labels: add tap/hover glossary definitions in plain language (C-S4).",
            "Move explanation panel: present one recommended action first, then progressively reveal engine depth and alternative lines.",
        ],
    },
    12: {
        "image": ROOT / "assets/screenshots/solution-references/chess_lessons_ui.png",
        "evidence": "Chess.com Lessons UI with learning routes and access prompts. Corresponds to ProductResearch C-05/C-HCI9 and source [19].",
        "bullets": [
            "Lesson, puzzle, and review cards: label Free, Daily limit, or Premium before users start (C-S5).",
            "Progress area: show remaining free attempts and reset time without forcing users into the activity first.",
            "When a limit is reached: offer a free alternative, resume-later action, and a saved-progress confirmation (C-S6).",
            "Upgrade prompt: keep the user's current learning context visible and avoid replacing the entire task surface.",
        ],
    },
    13: {
        "image": ROOT / "assets/screenshots/solution-references/chess_board_ui.png",
        "evidence": "Chess.com live board and Options area where premove behavior is configured. Corresponds to ProductResearch C-03/C-HCI8 and source [12].",
        "bullets": [
            "Board edge/status area: show every queued premove as a visible pending sequence (C-S7).",
            "Queued move indicator: add a warning when the opponent's reply could invalidate the user's tactical assumption.",
            "Near the board and in keyboard help: provide a one-click Clear action and Esc shortcut (C-S8).",
            "After clearing: show immediate confirmation and restore normal legal-move highlighting without interrupting the clock.",
        ],
    },
    14: {
        "image": ROOT / "assets/screenshots/solution-references/chess_board_ui.png",
        "evidence": "Chess.com board UI and Options area where Focus Mode must be discoverable. Corresponds to ProductResearch C-03/C-HCI3 and sources [13][14].",
        "bullets": [
            "After two completed games: show a one-time coachmark beside the board—Need fewer distractions? Try Focus Mode (C-S9).",
            "Board settings: add a persistent Focus Mode toggle instead of requiring hover discovery (C-S10).",
            "In-game top bar: show the active Focus Mode state and provide a reversible exit action.",
            "Mobile board menu: place the same toggle under Display/Board settings so the control is consistent across viewports.",
        ],
    },
}

EVIDENCE_REFS = {
    4: "F-09B + F-HCI6 + [1][6][9]",
    5: "F-10B + F-HCI7 + [6]",
    6: "F-10B + F-HCI8 + [6]",
    7: "F-HCI9 + [5][9]",
    8: "F-06 + F-HCI10 + [1][2][3][7]",
    10: "C-01/C-06 + C-HCI7 + [10][18][19][20]",
    11: "C-03 + C-HCI4/C-HCI10 + [15][16][17]",
    12: "C-05 + C-HCI9 + [19]",
    13: "C-03 + C-HCI8 + [12]",
    14: "C-03 + C-HCI3 + [13][14]",
}


def clear_cell(cell):
    tc = cell._tc
    for child in list(tc):
        if child.tag != qn("w:tcPr"):
            tc.remove(child)


def add_text(cell, text, bold=False, size=8.5, color=None):
    paragraph = cell.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    paragraph.paragraph_format.space_after = Pt(3)
    return paragraph


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_hyperlink(paragraph, text, url):
    relation_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    run_style = OxmlElement("w:rStyle")
    run_style.set(qn("w:val"), "Hyperlink")
    run_props.append(run_style)
    run.append(run_props)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_reference(paragraph, number, title, description, url):
    paragraph.clear()
    title_suffix = "" if title.endswith((".", "?", "!")) else "."
    paragraph.add_run(f"[{number}] {title}{title_suffix} {description} ")
    add_hyperlink(paragraph, url, url)
    paragraph.add_run(". Accessed 2026-07-07.")


def insert_paragraph_after(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, paragraph._parent)


doc = Document(SOURCE)

for paragraph in doc.paragraphs:
    if "Solution visuals are labeled as low-fidelity UI proposals, not captured live UI." in paragraph.text:
        paragraph.text = paragraph.text.replace(
            "Solution visuals are labeled as low-fidelity UI proposals, not captured live UI.",
            "Each solution card uses a relevant live website baseline on the left and explicit text-only UI changes on the right.",
        )

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if "Solution visuals are labelled as low-fidelity UI proposals, not captured live UI." in paragraph.text:
                    paragraph.text = paragraph.text.replace(
                        "Solution visuals are labelled as low-fidelity UI proposals, not captured live UI.",
                        "Each solution card uses a relevant live website baseline on the left and explicit text-only UI changes on the right.",
                    )

for table_index, spec in CARDS.items():
    table = doc.tables[table_index]
    table.cell(3, 0).text = "Live web evidence"
    table.cell(3, 4).text = "Required UI changes"

    evidence_cell = table.cell(4, 0)
    clear_cell(evidence_cell)
    image_p = evidence_cell.add_paragraph()
    image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_p.add_run().add_picture(str(spec["image"]), width=Inches(3.05))
    caption = evidence_cell.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
    caption_run = caption.add_run(spec["evidence"])
    caption_run.font.size = Pt(8)
    caption_run.italic = True
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(2)

    proposal_cell = table.cell(4, 4)
    clear_cell(proposal_cell)
    intro = proposal_cell.add_paragraph()
    intro_run = intro.add_run("Change the following UI locations:")
    intro_run.bold = True
    intro_run.font.size = Pt(9)
    intro.paragraph_format.space_after = Pt(3)
    for bullet in spec["bullets"]:
        paragraph = proposal_cell.add_paragraph(style="List Bullet")
        paragraph.add_run(bullet).font.size = Pt(8.5)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.keep_together = True
    shade_cell(proposal_cell, "F4F7FA")

    table.cell(2, 0).text = f"Evidence {EVIDENCE_REFS[table_index]}"

# Add an explicit cross-report contract near the beginning without changing the visual system.
overview = doc.tables[2].cell(0, 0)
for paragraph in list(overview.paragraphs):
    if paragraph.text.startswith("Cross-report linkage:"):
        paragraph._element.getparent().remove(paragraph._element)
link_p = overview.add_paragraph()
link_run = link_p.add_run(
    "Cross-report linkage: drawback IDs, persona IDs, HCI finding IDs, figure IDs, and source numbers in this report are inherited from Group10-PA1-ProductResearch_VisualReport.docx. Live captures provide the observed baseline; the right-hand bullets specify proposed UI changes and are not presented as existing UI."
)
link_run.bold = True
link_run.font.size = Pt(9)
link_p.paragraph_format.space_before = Pt(5)
link_p.paragraph_format.space_after = Pt(2)

REFERENCES = {
    1: ("Inside FIFA", "Official FIFA news and ecosystem navigation.", "https://inside.fifa.com/"),
    2: ("All stories and topics", "Official FIFA topic index.", "https://inside.fifa.com/all-stories"),
    3: ("FIFA World Cup 2026 Blog", "Official tournament story hub.", "https://inside.fifa.com/blogs/fwc-2026"),
    4: ("FIFA World Cup 26 Ticketing Programme launches this September", "Official FIFA media release used for the article baseline.", "https://inside.fifa.com/tournament-organisation/commercial/media-releases/world-cup-26-ticketing-programme-launch-september"),
    5: ("FIFA World Cup 2026 Last-Minute Sales Phase", "Official ticket-status media release.", "https://inside.fifa.com/media-releases/last-minute-ticket-sales-phase-fifa-world-cup-2026"),
    6: ("FIFA+", "Live FIFA+ interface containing live, schedule, originals, and archive routes.", "https://www.plus.fifa.com/"),
    7: ("FIFA Match Centre", "Official fixtures and results interface.", "https://www.fifa.com/en/match-centre"),
    8: ("FIFA/Coca-Cola Men's World Ranking", "Official FIFA ranking interface.", "https://inside.fifa.com/fifa-world-ranking/men"),
    9: ("FIFA Tickets and Hospitality", "Official FIFA ticket entry interface.", "https://www.fifa.com/en/tickets"),
    10: ("Chess.com homepage", "Live navigation and task-entry interface.", "https://www.chess.com/"),
    11: ("Play Online Chess", "Live Chess.com game-entry interface.", "https://www.chess.com/play/online"),
    12: ("What are pre-moves and how do they work?", "Official behavior reference supporting the premove-risk finding.", "https://support.chess.com/en/articles/8562432-what-are-pre-moves-and-how-do-they-work"),
    13: ("What is Focus Mode?", "Official behavior reference supporting the Focus Mode discoverability finding.", "https://support.chess.com/en/articles/8588088-what-is-focus-mode-how-do-i-turn-it-on"),
    14: ("Play Computer", "Live Chess.com board and Options interface used in the solution card.", "https://www.chess.com/play/computer"),
    15: ("How does Game Review work?", "Official behavior reference for post-game review.", "https://support.chess.com/en/articles/8584089-how-does-game-review-work"),
    16: ("Chess Analysis Board", "Live Chess.com analysis interface used in the solution card.", "https://www.chess.com/analysis"),
    17: ("How do I use the Analysis Board?", "Official behavior reference for analysis controls.", "https://support.chess.com/en/articles/8583825-how-do-i-use-the-analysis-board"),
    18: ("Chess Puzzles", "Live Chess.com puzzle interface used as a learning-flow baseline.", "https://www.chess.com/puzzles"),
    19: ("Chess Lessons", "Live Chess.com lessons interface used in the solution card.", "https://www.chess.com/lessons"),
    20: ("Chess Study Plans for All Levels", "Official Chess.com learning-path article.", "https://www.chess.com/article/view/study-plan-directory"),
}

reference_paragraphs = {}
for paragraph in doc.paragraphs:
    match = re.match(r"^\[(\d+)\]", paragraph.text.strip())
    if match:
        reference_paragraphs[int(match.group(1))] = paragraph

for missing, previous in ((8, 7), (16, 15), (18, 17)):
    if missing not in reference_paragraphs:
        reference_paragraphs[missing] = insert_paragraph_after(reference_paragraphs[previous])

for number, (title, description, url) in REFERENCES.items():
    set_reference(reference_paragraphs[number], number, title, description, url)

doc.save(OUTPUT)
print(f"Saved {OUTPUT}")
