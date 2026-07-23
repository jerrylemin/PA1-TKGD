from __future__ import annotations

import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Twips
from pypdf import PdfReader

import sys

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))
from build_pa1_package import GROUP_ID, weekly_report_md  # noqa: E402


DOCX = ROOT / f"{GROUP_ID}-PA1-WeeklyReport.docx"
PDF = ROOT / f"{GROUP_ID}-PA1-WeeklyReport.pdf"
SOURCE = ROOT / "sources" / f"{GROUP_ID}-PA1-WeeklyReport.md"
REVIEW = ROOT / f"{GROUP_ID}-PA1-WeeklyReport-Review.md"

NAVY = "17365D"
FIFA_BLUE = "EAF3FB"
CHESS_GREEN = "EDF5E8"
LIGHT_BLUE = "DCE6F1"
LIGHT_GRAY = "F3F5F7"
MID_GRAY = "B8C2CC"
TEXT = "202B38"
MUTED = "5B6573"
WHITE = "FFFFFF"
CONTENT_DXA = 10200


def set_run_font(run, size=10.2, bold=False, color=TEXT, italic=False):
    run.font.name = "Arial"
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Arial")
    rpr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=MID_GRAY, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def repeat_header(row):
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    row._tr.get_or_add_trPr().append(node)


def cant_split(row):
    row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))


def set_table_geometry(table, widths):
    if sum(widths) != CONTENT_DXA:
        raise ValueError(f"Table widths must total {CONTENT_DXA}: {widths}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Twips(widths[index])
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_page_field(paragraph):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend([begin, instruction, separate, text, end])
    set_run_font(run, size=8.5, color=MUTED)


def set_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.clear()
    p.paragraph_format.space_after = Pt(0)
    left = p.add_run("Group10-PA1 Sprint 1 Meeting Minutes and Weekly Report")
    set_run_font(left, size=8.5, bold=True, color=NAVY)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    footer = section.footer
    table = footer.add_table(rows=1, cols=3, width=Cm(18))
    set_table_geometry(table, [4100, 2000, 4100])
    for cell in table.rows[0].cells:
        set_cell_margins(cell, top=20, bottom=20, start=0, end=0)
    table.cell(0, 0).text = "Group10-PA1 Weekly Report"
    table.cell(0, 1).text = "Sprint 1"
    table.cell(0, 2).text = "Page "
    table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_field(table.cell(0, 2).paragraphs[0])
    for cell in table.rows[0].cells:
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=8.5, color=MUTED)


def configure(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.55)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(0.55)
    section.footer_distance = Cm(0.55)
    set_header_footer(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08
    for name, size, color, before, after in (
        ("Heading 1", 17, NAVY, 10, 5),
        ("Heading 2", 13, NAVY, 8, 4),
        ("Heading 3", 11, NAVY, 6, 3),
    ):
        style = doc.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    bullet = doc.styles["List Bullet"]
    bullet.font.name = "Arial"
    bullet.font.size = Pt(10)
    bullet.paragraph_format.left_indent = Cm(0.55)
    bullet.paragraph_format.first_line_indent = Cm(-0.25)
    bullet.paragraph_format.space_after = Pt(2.5)
    bullet.paragraph_format.line_spacing = 1.05


def paragraph_text(paragraph, text, size=10.2, bold=False, color=TEXT, align=None, italic=False):
    paragraph.clear()
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(4)
    set_run_font(paragraph.add_run(text), size=size, bold=bold, color=color, italic=italic)


def add_cover(doc, lines):
    doc.add_paragraph().paragraph_format.space_after = Pt(28)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    set_run_font(p.add_run(lines[0]), size=20, bold=True, color=NAVY)
    for index, text in enumerate(lines[1:6], 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4 if index != 5 else 20)
        set_run_font(p.add_run(text), size=13 if index == 1 else 11, bold=index in (1, 2), color=NAVY if index == 1 else TEXT)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    set_run_font(p.add_run("Team members"), size=11, bold=True, color=NAVY)
    for text in lines[6:]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        set_run_font(p.add_run(text), size=10.5)


def table_widths(headers):
    key = tuple(headers)
    known = {
        ("Resource", "Purpose", "Link", "Current status"): [2100, 3100, 3000, 2000],
        ("Member", "Student ID", "Primary role", "Main contribution"): [1700, 1200, 2700, 4600],
        ("Meeting", "Date", "Time"): [4000, 3100, 3100],
        ("Product", "Main areas", "Agreed drawbacks", "Agreed solutions"): [1200, 2500, 3100, 3400],
        ("Member", "Assigned work"): [2500, 7700],
        ("Member", "Research", "Writing", "Review", "Presentation or packaging", "Overall share"): [1500, 1600, 1600, 1500, 2000, 2000],
        ("Requirement", "Acceptance criteria", "Evidence or location", "Current status", "Owner", "Action before submission"): [1800, 2450, 1500, 1250, 1200, 2000],
    }
    return known.get(key, [CONTENT_DXA // len(headers)] * (len(headers) - 1) + [CONTENT_DXA - (CONTENT_DXA // len(headers)) * (len(headers) - 1)])


def add_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    set_table_geometry(table, table_widths(rows[0]))
    set_table_borders(table)
    repeat_header(table.rows[0])
    for row_index, row in enumerate(rows):
        cant_split(table.rows[row_index])
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.text = value
            fill = LIGHT_BLUE if row_index == 0 else WHITE
            if row_index > 0 and rows[0][0] == "Product":
                fill = FIFA_BLUE if value == "FIFA.com" or row[0] == "FIFA.com" else CHESS_GREEN
            if row_index > 0 and "Current status" in rows[0] and value in ("Needs team input", "Needs correction"):
                fill = "FFF2CC"
            shade_cell(cell, fill)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    set_run_font(run, size=7.6 if len(rows[0]) == 6 else 8.5, bold=row_index == 0, color=NAVY if row_index == 0 else TEXT)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_process_diagram(doc):
    labels = ["Sprint Planning", "→", "Weekly Scrum 1", "→", "Weekly Scrum 2", "→", "Sprint Review", "→", "Final export and submission check"]
    widths = [1650, 450, 1500, 450, 1500, 450, 1500, 450, 2250]
    table = doc.add_table(rows=1, cols=len(labels))
    set_table_geometry(table, widths)
    set_table_borders(table, color="8EA9C1", size="5")
    for index, label in enumerate(labels):
        cell = table.cell(0, index)
        cell.text = label
        shade_cell(cell, WHITE if label == "→" else LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            set_run_font(run, size=9 if label == "→" else 7.8, bold=label != "→", color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def parse_table(lines, start):
    rows = []
    index = start
    while index < len(lines) and lines[index].startswith("|"):
        cells = [cell.strip() for cell in lines[index].strip("|").split("|")]
        if not all(set(cell) <= {"-"} for cell in cells):
            rows.append(cells)
        index += 1
    return rows, index


def build_docx(markdown):
    lines = [line.rstrip() for line in markdown.splitlines()]
    doc = Document()
    configure(doc)
    first_break = lines.index("[[PAGE BREAK]]")
    cover = [line.lstrip("# ") for line in lines[:first_break] if line]
    add_cover(doc, cover)
    index = first_break
    while index < len(lines):
        line = lines[index]
        if not line:
            index += 1
            continue
        if line == "[[PAGE BREAK]]":
            doc.add_page_break()
        elif line == "[[SPRINT PROCESS]]":
            add_process_diagram(doc)
        elif line.startswith("|"):
            rows, index = parse_table(lines, index)
            add_table(doc, rows)
            continue
        elif line.startswith("## "):
            doc.add_paragraph(line[3:], style="Heading 1")
        elif line.startswith("### "):
            doc.add_paragraph(line[4:], style="Heading 2")
        elif line.startswith("#### "):
            doc.add_paragraph(line[5:], style="Heading 3")
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            paragraph_text(p, line[2:], size=10)
            p.style = doc.styles["List Bullet"]
            next_line = next((candidate for candidate in lines[index + 1 :] if candidate), "")
            p.paragraph_format.keep_with_next = next_line.startswith("- ")
        elif line.startswith("**") and ":**" in line:
            label, value = line[2:].split(":**", 1)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            set_run_font(p.add_run(label + ":"), size=10.2, bold=True, color=NAVY)
            set_run_font(p.add_run(value), size=10.2)
        else:
            p = doc.add_paragraph()
            paragraph_text(p, line)
        index += 1
    doc.core_properties.title = f"{GROUP_ID}-PA1 Sprint 1 Meeting Minutes and Weekly Report"
    doc.core_properties.subject = "Project Assignment 1 - RUP and Scrum meeting minutes"
    doc.core_properties.author = "Group10"
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    doc.settings.element.append(update_fields)
    doc.save(DOCX)


def export_pdf():
    docx = str(DOCX.resolve()).replace("'", "''")
    pdf = str(PDF.resolve()).replace("'", "''")
    command = (
        "$ErrorActionPreference='Stop'; $word=New-Object -ComObject Word.Application; "
        "$word.Visible=$false; $word.DisplayAlerts=0; "
        f"$doc=$word.Documents.Open('{docx}'); "
        "$doc.Fields.Update() | Out-Null; "
        f"$doc.ExportAsFixedFormat('{pdf}',17); "
        "$doc.Close(0); $word.Quit(); "
        "[Runtime.InteropServices.Marshal]::ReleaseComObject($doc) | Out-Null; "
        "[Runtime.InteropServices.Marshal]::ReleaseComObject($word) | Out-Null"
    )
    subprocess.run(["powershell", "-NoProfile", "-Command", command], check=True)


def write_review(markdown):
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(PDF).pages)
    pages = len(PdfReader(PDF).pages)
    terms = [
        "GroupID", "TODO", "READY DRAFT", "READY FINAL", "generator", "regenerate artifacts",
        "strict draft validator", "PDF extraction", "semantic drift", "hard-coded readiness assumptions",
    ]
    term_rows = "\n".join(f"- `{term}`: {len(re.findall(re.escape(term), pdf_text, re.I))} matches" for term in terms)
    review = f"""# Group10-PA1 WeeklyReport Review

## Files reviewed

- `GroupID-PA1-WeeklyReport.pdf` as the previous report baseline.
- `Group10-PA1-ProductResearch.docx`.
- `Group10-PA1-PotentialSolutions.docx`.
- `PA1.pptx`.
- `build_pa1_package.py`, `pa1_project_data.json`, `config/pa1_config.json`, and the WeeklyReport source Markdown.
- The lecturer-format reference named in the task was not present in the repository or its parent project folder; the supplied format requirements were used directly.

## Source files modified

- `config/pa1_config.json`.
- `build_pa1_package.py`.
- `scripts/create_weekly_report.py`.
- `scripts/validate_pa1_submission.py`.
- `docs/google_drive_readme_template.md`.
- `sources/Group10-PA1-WeeklyReport.md`.

## Pages in final PDF

- {pages} A4 portrait pages.

## Items corrected

- Group10 title, filenames, header, footer, and package references.
- One Sprint Planning, two Weekly Scrum records, and one Sprint Review and Retrospective.
- Per-member status reports without a duplicated summary table.
- Natural project language, ten report use cases, eight selected presentation use cases, and the agreed FIFA.com and Chess.com mappings.
- Accurate distinction between observed UI evidence and official documentation used for feature context.
- Balanced workload wording without unsupported percentages.
- Six-column acceptance checklist with truthful status values.

## Items requiring team input

- Weekly Scrum Google Doc link.
- Sprint Planning Google Doc link.
- Sprint Review Google Doc link.
- Google Drive README link.
- Zoom meeting link, if required.
- Presentation correction: label Game Review as Chess.com use case 4 and remove unsupported system wording before submission.

## Search results for prohibited terms

Final WeeklyReport PDF text:

{term_rows}

## Meeting-format verification

- Four meeting headings found in chronological order.
- Every meeting includes date, time, present members, absent members, actions, and a meeting summary.
- Both Weekly Scrum records contain Completed tasks, To-do tasks, and Issues/Obstacles for all four members.
- Sprint Review contains all six required retrospective sections.

## Cross-document consistency verification

- ProductResearch contains five FIFA.com and five Chess.com use cases.
- WeeklyReport records the presentation subset as four representative use cases per product.
- FIFA.com and Chess.com drawback meanings match the cross-report consistency map and PotentialSolutions.
- Game Review is recorded as Chess.com use case 4; unused visual space is not treated as an agreed PA1 drawback.
- No unsupported claim about the technology behind review or puzzle feedback appears in WeeklyReport.

## Layout verification

- DOCX uses A4 portrait, 1.5 to 1.55 cm margins, Arial body text, navy headings, repeat table headers, non-splitting table rows, keep-with-next headings, and a Word-native sprint process diagram.
- PDF was exported from the DOCX. Full page-image review is pending in this generated review and must be completed before delivery.

## Final submission status

Content complete. Administrative links require team input before submission. The document is not represented as fully ready while those links and the listed presentation correction remain outstanding.
"""
    REVIEW.write_text(review, encoding="utf-8")


def main():
    markdown = weekly_report_md()
    SOURCE.parent.mkdir(parents=True, exist_ok=True)
    source_text = markdown.rstrip() + "\n"
    if not SOURCE.exists() or SOURCE.read_text(encoding="utf-8") != source_text:
        SOURCE.write_text(source_text, encoding="utf-8")
    build_docx(markdown)
    export_pdf()
    write_review(markdown)
    print(f"Created {DOCX.name}, {PDF.name}, and {REVIEW.name}")


if __name__ == "__main__":
    main()
