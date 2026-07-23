from __future__ import annotations

import re
import zipfile
from pathlib import Path

from docx import Document
from pypdf import PdfReader

ROOT = Path(__file__).resolve().parents[1]
PDFS = [ROOT / f"GroupID-PA1-{name}.pdf" for name in ("ProductResearch", "PotentialSolutions", "PeerReview", "WeeklyReport")]
REPORTS = {p.stem.removeprefix("GroupID-PA1-"): ROOT / "sources" / f"{p.stem}.md" for p in PDFS}
DIAGRAMS = ("pa1_productresearch_task_flow", "pa1_productresearch_navigation_map", "pa1_potentialsolutions_traceability", "pa1_peerreview_traceability", "pa1_weeklyreport_workflow", "pa1_workdivision_raci")
MEMBERS = (("Le Minh", "21127645"), ("Nguyen Vu Bach", "21127224"), ("Pham Nguyen Gia Bao", "20127119"), ("Trang Minh Nhut", "22127318"))
FORBIDDEN = ("Member1", "Member2", "Member3", "Member4", "Member5", "placeholder", "TODO", "TBD", "FIXME", "lorem ipsum", "insert screenshot here", "Strava", "Nike Run Club", "NRC", "Garmin", "Garmin Connect", "Forerunner", "smartwatch")


def main() -> None:
    checks: list[tuple[str, bool, str]] = []
    source_text = "\n".join(path.read_text(encoding="utf-8") for path in REPORTS.values())
    pdf_text = "\n".join("\n".join(page.extract_text() or "" for page in PdfReader(path).pages) for path in PDFS)
    final_text = source_text + "\n" + pdf_text
    checks.append(("Four PDFs exist and exceed 10 KB", all(p.exists() and p.stat().st_size > 10_000 for p in PDFS), ""))
    checks.append(("Zip exists", (ROOT / "GroupID-PA1.zip").exists(), ""))
    with zipfile.ZipFile(ROOT / "GroupID-PA1.zip") as zf:
        names = zf.namelist()
    checks.append(("Zip contains exactly four top-level PDFs", names == [p.name for p in PDFS], str(names)))
    checks.append(("WorkDivision copies exist", all((ROOT / rel).exists() for rel in ("GroupID-PA1-WorkDivision.docx", "output/GroupID-PA1-WorkDivision.docx")), ""))
    checks.append(("Six Mermaid sources exist", all((ROOT / "assets/diagrams/mermaid" / f"{d}.mmd").exists() for d in DIAGRAMS), ""))
    checks.append(("Six rendered PNGs exceed 5 KB", all((ROOT / "assets/diagrams/rendered" / f"{d}.png").stat().st_size > 5_000 for d in DIAGRAMS), ""))
    matches = [term for term in FORBIDDEN if re.search(rf"\b{re.escape(term)}\b", final_text, re.I)]
    checks.append(("Final source and PDF prohibited-term scan", not matches, str(matches)))
    weekly = REPORTS["WeeklyReport"].read_text(encoding="utf-8")
    peer = REPORTS["PeerReview"].read_text(encoding="utf-8")
    checks.append(("WeeklyReport contains all members and IDs", all(name in weekly and sid in weekly for name, sid in MEMBERS), ""))
    checks.append(("PeerReview contains all real owners/speakers", all(name in peer for name, _ in MEMBERS), ""))
    checks.append(("Canonical website teams stated", "FIFA.com team: Le Minh and Nguyen Vu Bach" in final_text and "Chess.com team: Pham Nguyen Gia Bao and Trang Minh Nhut" in final_text, ""))
    checks.append(("Four balanced 25% rows", weekly.count("25%") == 4, f"count={weekly.count('25%')}"))
    product = REPORTS["ProductResearch"].read_text(encoding="utf-8")
    solutions = REPORTS["PotentialSolutions"].read_text(encoding="utf-8")
    id_ok = all(f"F-HCI{i}" in product and f"C-HCI{i}" in product and f"F-S{i}" in solutions and f"C-S{i}" in solutions for i in range(1, 11)) and all(f"F-D{i}" in product and f"F-D{i}" in solutions and f"C-D{i}" in product and f"C-D{i}" in solutions for i in range(1, 6))
    checks.append(("Canonical IDs present", id_ok, ""))
    refs = {"ProductResearch": ("pa1_productresearch_task_flow.png", "pa1_productresearch_navigation_map.png"), "PotentialSolutions": ("pa1_potentialsolutions_traceability.png",), "PeerReview": ("pa1_peerreview_traceability.png",), "WeeklyReport": ("pa1_weeklyreport_workflow.png",)}
    checks.append(("All report diagram references present", all(all(ref in REPORTS[name].read_text(encoding="utf-8") for ref in expected) for name, expected in refs.items()), ""))
    doc = Document(ROOT / "GroupID-PA1-WorkDivision.docx")
    doc_text = "\n".join(p.text for p in doc.paragraphs) + "\n" + "\n".join(c.text for t in doc.tables for row in t.rows for c in row.cells)
    checks.append(("WorkDivision contains roster, IDs, ownership, 25%", all(name in doc_text and sid in doc_text for name, sid in MEMBERS) and "FIFA.com: Le Minh" in doc_text and "Chess.com: Pham Nguyen Gia Bao" in doc_text and "25%" in doc_text, ""))
    checks.append(("WorkDivision embeds diagram", bool(doc.inline_shapes), f"inline_shapes={len(doc.inline_shapes)}"))
    passed = all(ok for _, ok, _ in checks)
    report = ["# PA1 Balance and Diagram Validation", "", f"Status: {'PASS' if passed else 'FAIL'}", "", "| Check | Result | Detail |", "| --- | --- | --- |"]
    report += [f"| {name} | {'PASS' if ok else 'FAIL'} | {detail} |" for name, ok, detail in checks]
    report += ["", "PDF first-page rendering: not performed by this validator; PDF extraction and size sanity completed."]
    (ROOT / "docs/pa1_balance_diagrams_validation.md").write_text("\n".join(report) + "\n", encoding="utf-8")
    if not passed:
        raise SystemExit("Validation failed: " + ", ".join(name for name, ok, _ in checks if not ok))
    print("PASS: PA1 balance, diagrams, PDFs, DOCX, and zip")


if __name__ == "__main__":
    main()
