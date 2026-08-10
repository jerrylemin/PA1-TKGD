from __future__ import annotations

from pathlib import Path
import re
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Twips
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "Group10-PA1-PeerReview-Revised.docx"
OUT_PDF = ROOT / "Group10-PA1-PeerReview-Revised.pdf"
PDF_TITLE = "Group10-PA1 Teacher Feedback Response"
PDF_PAGE_LABEL = "Trang"

NAVY = "12263F"
BLUE = "1769AA"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F4F6F8"
MID_GRAY = "D7DEE7"
TEXT = "172033"
MUTED = "526173"
WHITE = "FFFFFF"
GREEN = "2E7D32"
YELLOW = "9A6700"
RED = "B42318"

# A4 landscape with 1.15 cm left/right margins.
CONTENT_DXA = 15534


def set_run_font(run, size=9.5, bold=False, color=TEXT, italic=False):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=110):
    if sum(widths) != CONTENT_DXA:
        raise ValueError(f"Table widths must total {CONTENT_DXA}, got {sum(widths)}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, width in enumerate(widths):
        table.columns[i].width = Twips(width)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Twips(widths[i])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color=MID_GRAY, size="5"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:cantSplit")
    tr_pr.append(node)


def add_page_field(paragraph):
    paragraph.add_run("Teacher Feedback Response")


def style_paragraph(paragraph, size=9.5, bold=False, color=TEXT, align=None, italic=False):
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.06
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold, color=color, italic=italic)


def add_header_footer(section):
    header = section.header
    header.is_linked_to_previous = False
    # Word requires a paragraph in the header part; keep the default empty one.
    if header.paragraphs:
        header.paragraphs[0].clear()
        header.paragraphs[0].paragraph_format.space_after = Pt(0)
    table = header.add_table(rows=1, cols=1, width=Cm(27.4))
    table.cell(0, 0).text = "Group10-PA1  |  TEACHER FEEDBACK RESPONSE"
    shade_cell(table.cell(0, 0), NAVY)
    set_cell_margins(table.cell(0, 0), top=75, bottom=75, start=130, end=130)
    style_paragraph(table.cell(0, 0).paragraphs[0], size=10.5, bold=True, color=WHITE)

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("FIFA.com x Chess.com  |  ")
    add_page_field(p)
    style_paragraph(p, size=8, color=MUTED)


def configure_document(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.15)
    section.left_margin = Cm(1.15)
    section.right_margin = Cm(1.15)
    section.header_distance = Cm(0.35)
    section.footer_distance = Cm(0.35)
    add_header_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color, before, after in (
        ("Heading 1", 15, NAVY, 9, 5),
        ("Heading 2", 11.5, NAVY, 7, 3),
        ("Heading 3", 10.2, BLUE, 5, 2),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(9.3)
        style.paragraph_format.left_indent = Cm(0.55)
        style.paragraph_format.first_line_indent = Cm(-0.25)
        style.paragraph_format.space_after = Pt(2.5)
        style.paragraph_format.line_spacing = 1.05


def add_title_block(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    set_run_font(p.add_run("TEACHER FEEDBACK RESPONSE"), size=22, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    set_run_font(p.add_run("PA1 - Phân tích HCI cho FIFA.com và Chess.com"), size=13, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    set_run_font(p.add_run("Group10  |  Căn cứ nội dung: PA1.pptx"), size=9.5, color=MUTED)


def add_section_bar(doc, title, page_break=True):
    if page_break:
        doc.add_page_break()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = title
    shade_cell(cell, NAVY)
    set_cell_margins(cell, top=90, bottom=90, start=140, end=140)
    style_paragraph(cell.paragraphs[0], size=14, bold=True, color=WHITE)
    set_table_geometry(table, [CONTENT_DXA])
    set_table_borders(table, color=NAVY, size="4")
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_subheading(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    p.add_run(text)
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        set_run_font(p.add_run(bold_lead), size=9.5, bold=True, color=NAVY)
        set_run_font(p.add_run(text[len(bold_lead):]), size=9.5)
    else:
        set_run_font(p.add_run(text), size=9.5)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.08
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_run_font(p.add_run(item), size=9.3)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_run_font(p.add_run(item), size=9.3)


def add_callout(doc, label, text, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    set_run_font(p.add_run(f"{label}: "), size=9.5, bold=True, color=NAVY)
    set_run_font(p.add_run(text), size=9.5, color=TEXT)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    shade_cell(cell, fill)
    set_cell_margins(cell, top=110, bottom=110, start=140, end=140)
    set_table_geometry(table, [CONTENT_DXA])
    set_table_borders(table, color="B9CBE0", size="5")
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths, font_size=8.2):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_borders(table)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        shade_cell(cell, NAVY)
        style_paragraph(cell.paragraphs[0], size=8.4, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row_index, row_data in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            cells[i].text = str(value)
            if row_index % 2:
                shade_cell(cells[i], LIGHT_GRAY)
            style_paragraph(cells[i].paragraphs[0], size=font_size, color=TEXT)
    set_table_geometry(table, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)
    return table


def add_scope(doc):
    add_section_bar(doc, "01 Scope update", page_break=False)
    add_body(
        doc,
        "Tài liệu này là bản phản hồi chính thức cho các điểm cần làm rõ sau phần trình bày PA1. Phạm vi sản phẩm, nhóm người dùng, use case, lợi ích, hạn chế và giải pháp dưới đây được giữ đúng theo PA1.pptx; phần phân tích mới chỉ đào sâu cách các quyết định UI tác động tới hành vi người dùng.",
    )
    add_callout(
        doc,
        "Nguyên tắc phạm vi",
        "FIFA.com được xem như một hệ sinh thái browse-first phục vụ thông tin, xem nội dung và vé; Chess.com là nền tảng action-first phục vụ chơi, học và review. Mọi đề xuất đều chỉ rõ vị trí thay đổi trên màn hình.",
    )
    add_subheading(doc, "FIFA.com - baseline của PA1")
    add_table(
        doc,
        ["Nhóm user", "4 use case", "Benefits", "Drawbacks", "Solutions"],
        [[
            "Fan bóng đá; người xem bóng đá; người đặt vé.",
            "Lịch thi đấu và tỉ số; thông tin trận đấu và tuyển thủ; kiểm soát đặt vé; xem trực tiếp và highlight.",
            "Navigation rõ tác vụ; Match Centre theo ngày; nguồn chính thức; tournament hub.",
            "Ecosystem sprawl; FIFA+ handoff; trạng thái vé chưa rõ; browse-first chậm tác vụ nhanh.",
            "Task-first navigation; audience intent chips; FIFA+ handoff explainer; ticket status dashboard; article utility rail.",
        ]],
        [2200, 3400, 3100, 3300, 3534],
        font_size=8.4,
    )
    add_subheading(doc, "Chess.com - baseline của PA1")
    add_table(
        doc,
        ["Nhóm user", "4 use case", "Benefits", "Drawbacks", "Solutions"],
        [[
            "Beginner learner; competitive player; returning casual player.",
            "Bắt đầu ván cờ; giải puzzle; học lesson beginner; review game analysis.",
            "Play flow nhanh; direct manipulation; Game Review feedback; puzzles và study path.",
            "Feature overload; analysis screen dày đặc; premium gating; premove risk; Focus Mode khó tìm; empty space.",
            "Goal-based onboarding; personal dashboard; beginner analysis preset; premove queue preview; Focus Mode coachmark.",
        ]],
        [2200, 3300, 3100, 3434, 3500],
        font_size=8.4,
    )


def add_fifa_response(doc):
    add_section_bar(doc, "02 FIFA.com teacher feedback response")
    add_subheading(doc, "A. Ticket seating flow hiện tại")
    add_numbered(doc, [
        "User mở Tickets & Hospitality từ header của FIFA.com.",
        "User chọn tournament qua hàng logo giải đấu.",
        "Trang hiển thị card Tickets hoặc Hospitality tương ứng với giải đã chọn.",
        "CTA thay đổi theo trạng thái và gói dịch vụ: Explore details, Buy Packages Now, Register your interest hoặc Buy now.",
        "Sau CTA, user thường tiếp tục trên kênh ticketing hoặc hospitality riêng; seat map và lựa chọn chỗ ngồi xuất hiện ở giai đoạn sau, không được giải thích rõ ngay trên FIFA.com.",
    ])
    add_body(
        doc,
        "Flow hiện tại cho biết người dùng nên đi đâu tiếp, nhưng không trả lời đủ bốn câu hỏi tại điểm quyết định: vé đang mở bán hay chỉ nhận đăng ký quan tâm; seat map nằm ở bước nào; hạng vé hoặc khu vực ghế còn chỗ không; và CTA có đưa người dùng ra khỏi FIFA.com hay không. Đây là khoảng trống về visibility of system status, không chỉ là thiếu một bước trong flow.",
    )
    add_callout(
        doc,
        "Use case người đặt vé",
        "Người dùng cần xác nhận nguồn vé chính thức, vị trí ghế, loại vé, trạng thái chỗ ngồi, ngày mở bán, resale và hospitality trước khi trả tiền. Tác vụ có rủi ro cao vì quyết định gắn với tiền, chỗ ngồi và một sự kiện thật không thể hoàn tác dễ dàng.",
    )
    add_subheading(doc, "B. Vì sao chuyển hướng sang bên thứ ba gây khó chịu")
    add_bullets(doc, [
        "Điểm bắt đầu là FIFA.com nên mental model ban đầu gắn với official source. Khi header, logo, màu sắc, account system và wording của CTA cùng đổi, người dùng phải xây lại mental model ở giữa tác vụ.",
        "DAZN, ticketing vendor, hospitality provider, Store, Collect và Rewards là các destination khác nhau trong hệ sinh thái; việc thiếu dấu hiệu chuyển hệ tạo continuity break thay vì một chuyển trang bình thường.",
        "Người dùng phải tự kiểm tra trang mới có chính thức, có an toàn, có phải trang giả, và có thể quay lại FIFA.com hay không. Đây là trust friction và mode-boundary friction.",
        "Trong ticket flow, ma sát này nặng hơn vì CTA dẫn tới bước có thể yêu cầu đăng nhập, thông tin cá nhân và thanh toán.",
    ])

    doc.add_page_break()
    add_subheading(doc, "C. Confirmation trước khi handoff")
    add_body(doc, "Confirmation không nhằm làm chậm user. Nó báo trước điểm chuyển hệ và cho phép người dùng quyết định trước khi ngữ cảnh, domain và account system thay đổi.")
    add_table(
        doc,
        ["Thành phần", "Nội dung hiển thị"],
        [
            ["Thông báo chính", "You are leaving FIFA.com to continue on the official ticketing partner site."],
            ["Ngữ cảnh", "Tournament: FIFA World Cup 2026. | Task: Continue ticket purchase."],
            ["Thông tin tin cậy", "Official partner name; destination domain; dấu xác nhận official partner; return link về FIFA.com."],
            ["Quyền kiểm soát", "Buttons: Continue | Stay on FIFA.com."],
        ],
        [3100, 12434],
        font_size=8.8,
    )
    add_callout(
        doc,
        "HCI value",
        "Dialog làm tăng visibility of system status, củng cố trust, tăng user control và giảm nhầm lẫn về việc người dùng đang ở website nào.",
    )
    add_subheading(doc, "D. Đi sâu vào thiết kế UI")
    add_bullets(doc, [
        "Header cần phân biệt rõ task trên FIFA.com và link sang property khác; quick bar ưu tiên Scores, Tickets và Watch thay vì để tất cả destination cạnh tranh ngang nhau.",
        "Tournament logo row phải có selected state rõ, breadcrumb phải giữ tên giải và vị trí hiện tại, để user không mất context khi đổi tournament.",
        "Ticket card cần CTA cụ thể theo trạng thái. Explore details không đủ cho tác vụ mua; card phải kèm status label, ngày mở bán, loại vé và seat availability summary.",
        "Color hierarchy phải tách CTA chính khỏi link phụ; confirmation dialog phải hiển thị partner, domain và lựa chọn ở lại; dashboard table phải cho phép scan trạng thái theo tournament, sale phase, seat category và last update.",
        "FIFA+ handoff cần explainer card trước khi mở bề mặt DAZN-branded, đồng thời giữ return link về nội dung gốc trên FIFA.com.",
    ])
    add_subheading(doc, "E. Màu sắc FIFA.com")
    add_body(
        doc,
        "Navy và xanh dương hỗ trợ cảm nhận official, ổn định và đáng tin; nền trắng giữ card tin tức và nội dung dễ đọc; màu riêng của tournament tạo cảm xúc và nhận diện giải. Tuy nhiên, khi màu của event, Store, Collect, FIFA+, Tickets và News xuất hiện gần nhau, visual hierarchy bị cạnh tranh. Trong ticket flow, màu cần truyền trạng thái thay vì chỉ trang trí.",
    )
    add_table(
        doc,
        ["Màu đề xuất", "Trạng thái", "Tác dụng trên UI"],
        [
            ["Green", "Open / Available", "Nhận biết ngay vé hoặc khu vực ghế có thể tiếp tục."],
            ["Yellow", "Register interest / Waiting", "Báo trạng thái chưa mua ngay được nhưng vẫn có hành động tiếp theo."],
            ["Red", "Sold out / Closed", "Ngăn user hiểu nhầm CTA và giảm click lặp lại."],
            ["Blue", "Official information", "Phân biệt thông tin chính thức với sale status và partner action."],
        ],
        [2200, 4100, 9234],
        font_size=8.6,
    )


FIFA_UI_ROWS = [
    ["Header navigation", "Top-level menu trộn news, tournament và sibling destination.", "Browse-first; task nhanh phải scan nhiều mục.", "Thêm task-first quick bar: Scores, Tickets, Watch, Rankings.", "Giảm cognitive load và rút ngắn đường vào tác vụ."],
    ["Home hero", "Hero ưu tiên nội dung nổi bật, chưa chỉ rõ điểm bắt đầu theo mục tiêu.", "Information scent yếu với user cần task ngay.", "Thêm intent chips dưới hero: Scores, Tickets, Watch.", "User nhận ra điểm vào theo mục tiêu mà không cần mở menu."],
    ["Match Centre", "Date rail, match rows và filter hỗ trợ scan nhưng mất khi cuộn.", "User phải quay lại vùng đầu để đổi phạm vi.", "Sticky quick filter: Today, Live, Results, Team, Competition.", "Giữ control trong tầm nhìn và tăng efficiency."],
    ["Tickets page", "Tournament logos và Tickets/Hospitality cards cho biết destination, chưa cho biết sale state hoặc seat state.", "Visibility of system status thấp trong tác vụ rủi ro cao.", "Ticket status dashboard: sale phase, seat availability, ticket type, date, resale, hospitality, last update.", "User biết mua, chờ, đăng ký quan tâm hay chuyển resale."],
    ["Third-party handoff", "CTA có thể mở một domain, brand hoặc account system khác.", "Continuity break và trust friction.", "Confirmation popup với official partner, domain preview, Continue, Stay và return link.", "User hiểu điểm chuyển hệ và giữ quyền kiểm soát."],
    ["FIFA+ handoff", "Bề mặt DAZN-branded có Log in, Sign up và Watch free khác với FIFA.com.", "Mode-boundary friction; user không rõ quan hệ hai bề mặt.", "Explainer card trước handoff; giữ tournament/task context và return link.", "Giảm nghi ngại, tăng continuity cho flow xem."],
    ["Article page", "Sau khi đọc tin, user phải quay lại menu để xem tỉ số, vé hoặc video.", "Browse-first friction làm đứt task tiếp theo.", "Utility rail cạnh bài viết: Scores, Tickets, Watch, Rankings.", "Chuyển từ content sang task mà không mất context."],
]


def add_fifa_ui_analysis(doc):
    add_section_bar(doc, "03 FIFA.com deeper UI design analysis")
    add_body(doc, "Phân tích dưới đây chuyển từng nhận định từ mức flow sang mức component, trạng thái và hành vi hiển thị trên màn hình.")
    add_subheading(doc, "Điểm vào và tác vụ chính")
    add_table(
        doc,
        ["UI area", "Current design observation", "HCI issue", "Proposed UI change", "Expected result"],
        FIFA_UI_ROWS[:4],
        [1900, 3500, 3000, 4000, 3134],
        font_size=7.8,
    )
    add_subheading(doc, "Continuity giữa các bề mặt")
    add_table(
        doc,
        ["UI area", "Current design observation", "HCI issue", "Proposed UI change", "Expected result"],
        FIFA_UI_ROWS[4:],
        [1900, 3500, 3000, 4000, 3134],
        font_size=8.0,
    )
    add_callout(doc, "Ưu tiên", "Ticket status dashboard và handoff confirmation cần đi cùng nhau: dashboard trả lời trạng thái vé, còn dialog trả lời người dùng đang chuyển sang đâu.")


def add_chess_response(doc):
    add_section_bar(doc, "04 Chess.com teacher feedback response")
    add_subheading(doc, "A. Empty space ảnh hưởng thế nào")
    add_body(
        doc,
        "Khoảng trống quanh bàn cờ có thể là purposeful whitespace: nó tách board khỏi nội dung phụ, giảm chói mắt và giữ attention trong một ván dài. Vấn đề xuất hiện khi khoảng trống nằm giữa các module hoặc đẩy control quan trọng ra xa vùng nhìn chính. Khi đó information density giảm nhưng người dùng không nhận được thêm focus.",
    )
    add_table(
        doc,
        ["Nhóm user", "Tác động tại UI cụ thể", "HCI impact"],
        [
            ["Beginner learner", "Vùng trống giữa Play, Learn, Puzzles và Review không cho biết hành động tiếp theo; card ưu tiên không hiện gần điểm bắt đầu.", "Visual hierarchy yếu, affordance giảm, scan load và task completion time tăng."],
            ["Competitive player", "Khoảng trống quanh board không phải vấn đề lớn nếu clocks, resign/draw, premove state và settings vẫn nằm gần board.", "Control bị ẩn hoặc nằm xa làm tăng motor/visual search dưới áp lực giờ."],
            ["Returning casual player", "Khoảng cách lớn giữa module khiến người dùng mất thời gian tìm lại Review, Learn, Puzzle hoặc Focus Mode.", "Recognition kém hiệu quả hơn; user phải nhớ vị trí tính năng."],
        ],
        [2600, 7900, 5034],
        font_size=8.5,
    )
    add_callout(
        doc,
        "Kết luận",
        "Không nên lấp kín mọi khoảng trống. Cần giữ whitespace quanh board để tập trung, nhưng dùng priority cards và compact task grouping tại homepage/module area để khoảng trống có chủ đích.",
    )
    add_subheading(doc, "B. Màu sắc Chess.com")
    add_body(
        doc,
        "Nền tối giúp giảm chói khi chơi lâu; xanh lá gợi liên tưởng tới chess board và làm nổi CTA chính; trắng và xám phân tách text, panel và move list. Tuy nhiên, xanh lá xuất hiện ở nhiều module khiến CTA chính có thể cạnh tranh với thành phần khác. Trên analysis screen, evaluation, move labels, chart và icons dùng nhiều màu cùng lúc nên beginner phải giải mã cả nội dung lẫn color semantics.",
    )
    add_table(
        doc,
        ["Màu trong beginner mode", "Ý nghĩa duy nhất", "Cách dùng"],
        [
            ["Green", "Move tốt hoặc action chính", "Chỉ dùng cho best move, correct state và CTA tiếp tục."],
            ["Yellow", "Cần chú ý", "Dùng cho inaccuracy hoặc gợi ý cần xem lại."],
            ["Red", "Blunder hoặc lỗi", "Dùng tiết chế cho lỗi chính, không phủ nhiều panel cùng lúc."],
            ["Gray", "Thông tin phụ", "Engine detail, option nâng cao và metadata được hạ ưu tiên."],
        ],
        [3100, 4200, 8234],
        font_size=8.5,
    )


CHESS_UI_ROWS = [
    ["Home / feature menu", "Play, Learn, Puzzles, Train và Watch cùng cạnh tranh ở điểm vào.", "Feature overload; beginner không biết bắt đầu.", "Goal-based onboarding với card Play a game, Learn basics, Solve a puzzle.", "Tăng information scent và giảm lựa chọn ban đầu."],
    ["Empty space areas", "Khoảng trống giữa module không luôn gắn với mục tiêu focus.", "Information density thấp, scan load tăng.", "Dùng priority cards cho 3 tác vụ gần nhất; compact task grouping cho mục phụ.", "Tận dụng không gian mà không biến homepage thành dashboard dày đặc."],
    ["Game board", "Whitespace quanh board hỗ trợ tập trung nhưng một số control ít nổi hoặc phụ thuộc hover.", "Affordance và discoverability giảm khi control bị ẩn.", "Giữ vùng thở quanh board; đặt clocks, draw/resign, settings và premove state trong vùng nhìn ổn định.", "Giữ focus nhưng không đánh đổi user control."],
    ["Analysis screen", "Chart, evaluation, engine lines, move labels và coach feedback xuất hiện cùng lúc.", "Cognitive load cao; progressive disclosure yếu.", "Beginner analysis preset: lỗi chính, best move, một giải thích ngắn; phần nâng cao mở theo nhu cầu.", "Beginner hiểu feedback trước khi đọc engine detail."],
    ["Color system", "Green, red, yellow, chart colors và icons có thể mang nhiều ý nghĩa.", "Feedback ambiguity và visual noise.", "Chuẩn hóa Green/Yellow/Red/Gray; một màu một vai trò trong beginner mode.", "Giảm thời gian giải mã trạng thái và tăng consistency."],
    ["Focus Mode", "Control khó nhận ra khi phụ thuộc hover gần board/sidebar boundary.", "Discoverability thấp đối với tính năng giảm phân tán.", "Coachmark sau vài ván; shortcut cố định trong Settings và tooltip có nhãn.", "Returning user tự tìm và bật mode dễ hơn."],
    ["Premove", "Nước đã queue có thể thực thi sau một reply ngoài dự đoán.", "Error prevention và feedback chưa đủ rõ dưới áp lực giờ.", "Queue preview trên board, trạng thái rõ và clear shortcut một thao tác.", "Giảm blunder do premove ngoài ý muốn."],
]


def add_chess_ui_analysis(doc):
    add_section_bar(doc, "05 Chess.com deeper UI design analysis")
    add_subheading(doc, "Điểm vào, khoảng trống và bàn cờ")
    add_table(
        doc,
        ["UI area", "Current design observation", "HCI issue", "Proposed UI change", "Expected result"],
        CHESS_UI_ROWS[:4],
        [1900, 3500, 3000, 4000, 3134],
        font_size=7.8,
    )
    add_subheading(doc, "Feedback, focus và error prevention")
    add_table(
        doc,
        ["UI area", "Current design observation", "HCI issue", "Proposed UI change", "Expected result"],
        CHESS_UI_ROWS[4:],
        [1900, 3500, 3000, 4000, 3134],
        font_size=8.0,
    )
    add_callout(doc, "Nguyên tắc", "Beginner preset giảm số tín hiệu phải đọc; competitive view vẫn giữ dữ liệu nâng cao. Hai mode dùng cùng semantics màu để feedback không đổi nghĩa.")


MAPPING_ROWS = [
    ["FIFA.com", "Ticket seating flow", "Ticket status uncertainty", "Visibility of system status; trust", "Ticket status dashboard + seat availability summary + official partner confirmation", "Card và dashboard hiển thị sale phase, seat category, availability, last update; dialog hiện partner/domain trước CTA.", "Biết mua, chờ hoặc resale; giảm rủi ro và click lặp lại."],
    ["FIFA.com", "Third-party redirection", "Ecosystem sprawl; FIFA+ handoff", "Continuity break; trust friction", "Handoff explainer popup + destination preview + return link", "Modal nêu brand, domain, tournament/task; Continue và Stay; return link giữ trong trang đích.", "Giảm nghi ngại, giữ mental model và user control."],
    ["FIFA.com", "Design depth", "Browse-first friction", "Cognitive load", "Task-first navigation + quick task bar + article utility rail", "Header có Scores/Tickets/Watch/Rankings; intent chips dưới hero; utility rail cạnh article.", "Tác vụ nhanh cần ít scan và ít quay lại menu."],
    ["FIFA.com", "Color", "Visual hierarchy cạnh tranh", "Weak visual priority", "Color-coded status labels + consistent CTA hierarchy", "Green/Yellow/Red cho sale state; Blue cho official info; một primary CTA mỗi ticket card.", "Nhìn trạng thái nhanh và phân biệt action chính/phụ."],
    ["Chess.com", "Empty space", "Empty space / inefficient screen use", "Visual hierarchy; information density; scan load", "Purposeful whitespace + priority cards + compact task grouping", "Giữ khoảng thở quanh board; homepage đặt 3 priority cards vào vùng trống; nhóm mục phụ gọn hơn.", "Giữ focus nhưng giảm thời gian tìm Learn, Puzzle, Review."],
    ["Chess.com", "Color", "Analysis density; feature overload", "Cognitive load; feedback ambiguity", "Simplified beginner color system + clearer feedback states", "Beginner view dùng Green/Yellow/Red/Gray nhất quán; engine detail và metadata chuyển sang Gray/advanced layer.", "Feedback dễ hiểu hơn và giảm visual noise."],
]


def add_mapping(doc):
    add_section_bar(doc, "06 Revised solution mapping")
    add_body(doc, "Mỗi concern được nối trực tiếp với drawback của PA1, nguyên nhân HCI và thay đổi nhìn thấy được trên màn hình.")
    add_subheading(doc, "FIFA.com")
    add_table(
        doc,
        ["Website", "Teacher concern", "Linked PA1 drawback", "HCI cause", "UI-level solution", "What exactly changes on screen", "Expected improvement"],
        MAPPING_ROWS[:4],
        [1000, 1800, 2100, 2100, 2950, 3300, 2284],
        font_size=7.15,
    )
    add_subheading(doc, "Chess.com")
    add_table(
        doc,
        ["Website", "Teacher concern", "Linked PA1 drawback", "HCI cause", "UI-level solution", "What exactly changes on screen", "Expected improvement"],
        MAPPING_ROWS[4:],
        [1000, 1800, 2100, 2100, 2950, 3300, 2284],
        font_size=7.5,
    )
    add_callout(doc, "Ranh giới bằng chứng", "Seat availability dashboard, handoff confirmation và hệ màu trạng thái là thiết kế đề xuất. Chúng không được mô tả như chức năng đã có trên màn hình hiện tại.")


def add_presentation_notes(doc):
    add_section_bar(doc, "07 Final presentation notes")
    add_body(doc, "Khi trình bày, mỗi nhận định phải đi cùng một vùng UI cụ thể; không đọc lại bảng theo hàng.")
    add_bullets(doc, [
        "FIFA ticket: chỉ vào Tickets & Hospitality trên header, hàng tournament logos, card Tickets/Hospitality và CTA. Nêu rõ card đang chỉ destination nhưng chưa cho thấy sale state, seat-map stage và seat availability.",
        "FIFA handoff: chỉ vào FIFA+ hoặc bề mặt DAZN-branded, các control Log in, Sign up và Watch free. Giải thích sự đổi brand, account system và CTA làm đứt continuity.",
        "Màu FIFA: chỉ navy header, white content cards, tournament colors và primary CTA. Sau đó đối chiếu với status labels Green/Yellow/Red và Blue cho official information.",
        "Chess empty space: chỉ homepage/module area hoặc vùng có nhiều khoảng trắng. Phân biệt purposeful whitespace quanh board với khoảng trống làm module ưu tiên bị xa nhau.",
        "Màu Chess: chỉ board, CTA, analysis feedback và move labels. Giải thích vì sao beginner mode cần ít màu hơn và một màu chỉ giữ một nghĩa.",
    ])
    add_callout(doc, "Câu chốt", "Nhóm không đề xuất lấp đầy mọi màn hình. Mục tiêu là làm rõ trạng thái, giữ continuity và đặt action quan trọng đúng nơi người dùng đang ra quyết định.")


def _pdf_text(text):
    return escape(text).replace("\n", "<br/>")


def _register_pdf_fonts():
    font_dir = Path(r"C:\Windows\Fonts")
    pdfmetrics.registerFont(TTFont("ArialPA1", str(font_dir / "arial.ttf")))
    pdfmetrics.registerFont(TTFont("ArialPA1-Bold", str(font_dir / "arialbd.ttf")))


def build_pdf_from_docx():
    """Render the same DOCX content into a deterministic A4-landscape PDF."""
    _register_pdf_fonts()
    page_size = landscape(A4)
    page_width, page_height = page_size
    left = right = 1.15 * cm
    usable_width = page_width - left - right

    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "PA1Body", parent=styles["BodyText"], fontName="ArialPA1", fontSize=8.8,
        leading=10.6, textColor=colors.HexColor(f"#{TEXT}"), spaceAfter=4,
    )
    bullet = ParagraphStyle(
        "PA1Bullet", parent=body, leftIndent=15, firstLineIndent=-8,
        bulletIndent=2, spaceAfter=2.5,
    )
    h2 = ParagraphStyle(
        "PA1H2", parent=body, fontName="ArialPA1-Bold", fontSize=11.2,
        leading=13.5, textColor=colors.HexColor(f"#{NAVY}"), spaceBefore=5, spaceAfter=4,
    )
    title = ParagraphStyle(
        "PA1Title", parent=body, fontName="ArialPA1-Bold", fontSize=20,
        leading=23, alignment=TA_CENTER, textColor=colors.HexColor(f"#{NAVY}"), spaceBefore=8, spaceAfter=2,
    )
    subtitle = ParagraphStyle(
        "PA1Subtitle", parent=body, fontName="ArialPA1-Bold", fontSize=12,
        leading=14, alignment=TA_CENTER, textColor=colors.HexColor(f"#{BLUE}"), spaceAfter=2,
    )
    metadata = ParagraphStyle(
        "PA1Meta", parent=body, fontSize=8.8, leading=10.5, alignment=TA_CENTER,
        textColor=colors.HexColor(f"#{MUTED}"), spaceAfter=7,
    )
    table_body = ParagraphStyle(
        "PA1Table", parent=body, fontSize=7.25, leading=8.7, spaceAfter=0,
    )
    table_head = ParagraphStyle(
        "PA1TableHead", parent=table_body, fontName="ArialPA1-Bold",
        fontSize=7.35, leading=8.8, alignment=TA_CENTER, textColor=colors.white,
    )
    section_style = ParagraphStyle(
        "PA1Section", parent=body, fontName="ArialPA1-Bold", fontSize=13.2,
        leading=15, textColor=colors.white, spaceAfter=0,
    )
    callout = ParagraphStyle(
        "PA1Callout", parent=body, fontSize=8.5, leading=10.2, spaceAfter=0,
    )

    def on_page(canvas, _doc):
        canvas.saveState()
        canvas.setTitle(PDF_TITLE)
        canvas.setAuthor("Group10")
        canvas.setFillColor(colors.HexColor(f"#{NAVY}"))
        canvas.rect(left, page_height - 29, usable_width, 21, fill=1, stroke=0)
        canvas.setFillColor(colors.white)
        canvas.setFont("ArialPA1-Bold", 9.3)
        canvas.drawString(left + 8, page_height - 22, "Group10-PA1  |  TEACHER FEEDBACK RESPONSE")
        canvas.setFillColor(colors.HexColor(f"#{MUTED}"))
        canvas.setFont("ArialPA1", 7.4)
        canvas.drawRightString(page_width - right, 13, f"FIFA.com x Chess.com  |  {PDF_PAGE_LABEL} {canvas.getPageNumber()}")
        canvas.restoreState()

    story = []
    source = Document(OUT_DOCX)
    number_index = 0
    for item in source.iter_inner_content():
        if hasattr(item, "_p"):
            if item._p.xpath('.//w:br[@w:type="page"]'):
                story.append(PageBreak())
                number_index = 0
                continue
            text = item.text.strip()
            if not text:
                continue
            if text == "TEACHER FEEDBACK RESPONSE":
                story.append(Paragraph(_pdf_text(text), title))
            elif text.startswith("PA1 - Phân tích HCI") or text.startswith("PA1 - HCI Analysis"):
                story.append(Paragraph(_pdf_text(text), subtitle))
            elif text.startswith("Group10  |"):
                story.append(Paragraph(_pdf_text(text), metadata))
            elif item.style and item.style.name == "Heading 2":
                story.append(Paragraph(_pdf_text(text), h2))
                number_index = 0
            elif item.style and item.style.name == "List Bullet":
                story.append(Paragraph(_pdf_text(text), bullet, bulletText="•"))
            elif item.style and item.style.name == "List Number":
                number_index += 1
                story.append(Paragraph(_pdf_text(text), bullet, bulletText=f"{number_index}."))
            else:
                number_index = 0
                story.append(Paragraph(_pdf_text(text), body))
            continue

        rows = [[cell.text.strip() for cell in row.cells] for row in item.rows]
        if not rows:
            continue
        if len(rows) == 1 and len(rows[0]) == 1:
            cell_text = rows[0][0]
            is_section = bool(re.match(r"^\d{2} ", cell_text))
            p_style = section_style if is_section else callout
            t = Table([[Paragraph(_pdf_text(cell_text), p_style)]], colWidths=[usable_width])
            fill = colors.HexColor(f"#{NAVY if is_section else LIGHT_BLUE}")
            border = colors.HexColor(f"#{NAVY if is_section else 'B9CBE0'}")
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), fill),
                ("BOX", (0, 0), (-1, -1), 0.55, border),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]))
            story.extend([t, Spacer(1, 4)])
            continue

        first_widths = []
        for cell in item.rows[0].cells:
            first_widths.append(cell.width.pt if cell.width is not None else 1)
        total = sum(first_widths) or len(first_widths)
        col_widths = [usable_width * width / total for width in first_widths]
        data = []
        for row_index, row in enumerate(rows):
            p_style = table_head if row_index == 0 else table_body
            data.append([Paragraph(_pdf_text(value), p_style) for value in row])
        t = Table(data, colWidths=col_widths, repeatRows=1, splitByRow=1, hAlign="LEFT")
        commands = [
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{NAVY}")),
            ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor(f"#{MID_GRAY}")),
            ("LEFTPADDING", (0, 0), (-1, -1), 4.5),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4.5),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]
        for row_index in range(2, len(data), 2):
            commands.append(("BACKGROUND", (0, row_index), (-1, row_index), colors.HexColor(f"#{LIGHT_GRAY}")))
        t.setStyle(TableStyle(commands))
        story.extend([t, Spacer(1, 4)])

    pdf = SimpleDocTemplate(
        str(OUT_PDF), pagesize=page_size, leftMargin=left, rightMargin=right,
        topMargin=1.25 * cm, bottomMargin=1.05 * cm, title=PDF_TITLE,
        author="Group10", subject="HCI analysis of FIFA.com and Chess.com",
    )
    pdf.build(story, onFirstPage=on_page, onLaterPages=on_page)
    print(OUT_PDF)


def build():
    doc = Document()
    configure_document(doc)
    add_title_block(doc)
    add_scope(doc)
    add_fifa_response(doc)
    add_fifa_ui_analysis(doc)
    add_chess_response(doc)
    add_chess_ui_analysis(doc)
    add_mapping(doc)
    add_presentation_notes(doc)
    doc.core_properties.title = "Group10-PA1 Teacher Feedback Response"
    doc.core_properties.subject = "HCI analysis of FIFA.com and Chess.com"
    doc.core_properties.author = "Group10"
    doc.core_properties.keywords = "PA1, HCI, FIFA.com, Chess.com, Teacher Feedback Response"
    doc.save(OUT_DOCX)
    print(OUT_DOCX)
    build_pdf_from_docx()


if __name__ == "__main__":
    build()
