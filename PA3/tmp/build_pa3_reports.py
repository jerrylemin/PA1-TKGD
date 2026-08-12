from __future__ import annotations

from copy import deepcopy
from datetime import datetime
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


PA = Path(r"C:\Users\Administrator\Documents\MEGA\tkgd\PA")
ROOT = PA / "PA3"
FINAL = ROOT / "final"
BACKUP = ROOT / "backup"
PROTOS = ROOT / "paper prototypes"
TEMPLATE_REPORT = PA / "PA2" / "source" / "Group10-PA2-ProjectProposal.docx"
TEMPLATE_WEEKLY = PA / "PA2" / "source" / "Group10-PA2-WeeklyReport.docx"

TEAM = [
    ("Le Minh", "21127645", "Project Coordinator and Integration Lead"),
    ("Nguyen Vu Bach", "21127224", "FIFA.com Research Lead"),
    ("Pham Nguyen Gia Bao", "20127119", "Chess.com Research Lead"),
    ("Trang Minh Nhut", "22127318", "HCI Analysis and Visual Review Lead"),
]

NAVY = "123B65"
BLUE = "1D70A2"
GREEN = "4F7F35"
INK = "1F2937"
MUTED = "5B6573"
LIGHT = "F3F6FA"
FIFA_LIGHT = "EAF2F8"
CHESS_LIGHT = "EEF5E9"
ORANGE = "C66A1B"
RED = "9B1C1C"


def set_font(run, size=14, color=INK, bold=None, italic=None):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def clear_body(doc):
    body = doc._element.body
    sect = body.sectPr
    for child in list(body):
        if child is not sect:
            body.remove(child)


def field(paragraph, instruction, shown="1"):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = shown
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, sep, txt, end])
    return run


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def keep_table_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    tr_pr.append(cant)


def set_table_widths(table, widths_cm):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    repeat_header(table.rows[0])
    for ri, row in enumerate(table.rows):
        keep_table_row(row)
        for ci, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[ci])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.05
                for r in p.runs:
                    set_font(r, 12 if len(widths_cm) >= 5 else 13)
            if ri == 0:
                shade(cell, NAVY)
                for p in cell.paragraphs:
                    for r in p.runs:
                        set_font(r, 13, "FFFFFF", True)


def configure_doc(template, title, short_title, landscape=False):
    doc = Document(template)
    clear_body(doc)
    sec = doc.sections[0]
    if landscape:
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width = Cm(29.7)
        sec.page_height = Cm(21.0)
        sec.top_margin = Cm(1.7)
        sec.bottom_margin = sec.left_margin = sec.right_margin = Cm(1.3)
    else:
        sec.orientation = WD_ORIENT.PORTRAIT
        sec.page_width = Cm(21.0)
        sec.page_height = Cm(29.7)
        sec.top_margin = sec.left_margin = sec.right_margin = Cm(2.0)
        sec.bottom_margin = Cm(1.8)
    sec.header_distance = Cm(0.8)
    sec.footer_distance = Cm(0.8)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(14)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, color, before, after in [
        ("Title", 28, NAVY, 0, 8),
        ("Subtitle", 14, MUTED, 0, 12),
        ("Heading 1", 26, NAVY, 14, 7),
        ("Heading 2", 20, BLUE, 10, 5),
        ("Heading 3", 17, GREEN, 8, 4),
    ]:
        s = doc.styles[name]
        s.font.name = "Arial"
        s._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        s._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        s.font.size = Pt(size)
        s.font.bold = name != "Subtitle"
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True
    for style_name in ("List Bullet", "List Number"):
        s = doc.styles[style_name]
        s.font.name = "Arial"
        s.font.size = Pt(14)
        s.paragraph_format.space_after = Pt(3)
    sec.different_first_page_header_footer = False
    doc.settings.odd_and_even_pages_header_footer = False
    for hdr in (sec.header, sec.first_page_header, sec.even_page_header):
        hp = hdr.paragraphs[0]
        hp.clear()
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hr = hp.add_run(f"GROUP10  |  CSC13112 UI/UX DESIGN  |  {short_title.upper()}")
        set_font(hr, 12, NAVY, True)
    for ftr in (sec.footer, sec.first_page_footer, sec.even_page_footer):
        fp = ftr.paragraphs[0]
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = fp.add_run(f"{short_title}  |  Group10  |  ")
        set_font(fr, 11, MUTED)
        set_font(field(fp, "PAGE"), 11, MUTED)
    doc.core_properties.title = f"Group10 PA3 - {title}"
    doc.core_properties.subject = "CSC13112 UI/UX Design Project Assignment 3"
    doc.core_properties.author = "Group10"
    return doc


def para(doc, text="", bold_label=None, fill=None):
    if fill:
        t = doc.add_table(rows=1, cols=1)
        t.style = "Table Grid"
        shade(t.cell(0, 0), fill)
        cell_margins(t.cell(0, 0), 140, 170, 140, 170)
        p = t.cell(0, 0).paragraphs[0]
    else:
        p = doc.add_paragraph()
    if bold_label and text.startswith(bold_label):
        r = p.add_run(bold_label)
        set_font(r, 14, NAVY, True)
        r = p.add_run(text[len(bold_label):])
        set_font(r)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_font(p.add_run(item))


def numbered(doc, items):
    numbering = doc.part.numbering_part.element
    abstract_id = "0"
    for abstract in numbering.findall(qn("w:abstractNum")):
        styles = abstract.findall(".//" + qn("w:pStyle"))
        if any(s.get(qn("w:val")) == "ListNumber" for s in styles):
            abstract_id = abstract.get(qn("w:abstractNumId"))
            break
    ids = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    num_id = str(max(ids, default=0) + 1)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), num_id)
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), abstract_id)
    num.append(ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_font(p.add_run(item))
        num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
        num_pr.get_or_add_ilvl().set(qn("w:val"), "0")
        num_pr.get_or_add_numId().set(qn("w:val"), num_id)


def table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.cell(0, i).text = h
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    set_table_widths(t, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return t


def cover(doc, title, subtitle, status=None, links=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28 if links else 65)
    set_font(p.add_run("PROJECT ASSIGNMENT 3"), 14, ORANGE, True)
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(title), 28, NAVY, True)
    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(subtitle), 14, MUTED)
    if status:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(status), 14, RED, True)
    rows = [
        ["Course", "CSC13112 - UI/UX Design, FIT-HCMUS"],
        ["Group", "Group10"],
        ["Team", "; ".join(f"{n} ({sid})" for n, sid, _ in TEAM)],
        ["Scope", "FIFA.com and Chess.com"],
    ]
    table(doc, ["Field", "Record"], rows, [3.2, 12.4])
    if links:
        p = doc.add_paragraph()
        set_font(p.add_run("YouTube Demo Links"), 17, NAVY, True)
        rows = [
            ["FIFA - Status Dashboard", "[YOUTUBE LINK REQUIRED]"],
            ["FIFA - Timeline Tracker", "[YOUTUBE LINK REQUIRED]"],
            ["FIFA - Action Hub", "[YOUTUBE LINK REQUIRED]"],
            ["Chess - Beginner Review Flow", "[YOUTUBE LINK REQUIRED]"],
            ["Chess - Visual Card Dashboard", "[YOUTUBE LINK REQUIRED]"],
            ["Chess - Side-by-Side Assistant", "[YOUTUBE LINK REQUIRED]"],
        ]
        table(doc, ["Prototype", "Demo"], rows, [7.6, 8.0])
    doc.add_page_break()


def figure(doc, path, fid, caption, related, width_cm=16.4, max_height_cm=17.4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    with Image.open(path) as im:
        projected = width_cm * im.height / im.width
    run = p.add_run()
    if projected > max_height_cm:
        inline = run.add_picture(str(path), height=Cm(max_height_cm))
    else:
        inline = run.add_picture(str(path), width=Cm(width_cm))
    inline._inline.docPr.set("title", f"Figure {fid}")
    inline._inline.docPr.set("descr", f"{caption} Related use cases: {related}.")
    cp = doc.add_paragraph()
    cp.paragraph_format.keep_together = True
    cp.paragraph_format.space_after = Pt(8)
    set_font(cp.add_run(f"Figure {fid}. "), 12, NAVY, True)
    set_font(cp.add_run(caption), 12, INK)
    set_font(cp.add_run(f"\nSource: Group10 PA3 paper-prototype asset. Related: {related}"), 11, MUTED, italic=True)


FIFA_ALTS = [
    {
        "name": "Status Dashboard", "image": "alt1scenario1.png", "model": "Status-first, confidence-first dashboard.",
        "workflow": "Scan aggregate states -> select an event -> follow the status-specific action or explanation.",
        "problem": "Makes governed ticket state, action-needed status, and event-level next actions visible before users leave FIFA.com.",
        "motivation": "Direct evolution of PA2 F-A1 Status Dashboard and the PA1 F-S7 ticket-status concept.",
        "strengths": "Fast triage; explicit Confirmed/Pending/Action needed/Cancelled states; contextual primary actions; visible last update.",
        "weaknesses": "Dense summary plus side actions; a count does not explain the process; Pending still needs a reason.",
        "risks": "Color dependence; users may read zero action needed too broadly; confirmed allocation may be confused with a usable mobile ticket.",
        "hyp": "Users can identify the pending event, explain whether action is required, and choose the correct event action without detouring.",
        "ucs": "F-UC01, F-UC02, F-UC03, F-UC06",
    },
    {
        "name": "Timeline Tracker", "image": "alt2scenario1.png", "model": "Progress-first, freshness-first milestone tracker.",
        "workflow": "Select an event -> locate completed/current/future stages -> inspect what changed or official next-step guidance.",
        "problem": "Explains where an order stands, when the source changed, and what normally happens next.",
        "motivation": "Explores PA2's state/freshness/provenance principles as a process model rather than a status dashboard.",
        "strengths": "Strong process transparency; timestamps and update history; visible future delivery milestone and official source.",
        "weaknesses": "More interpretation than a summary; dense timelines; verification terminology is unexplained.",
        "risks": "A linear track may overpromise predictable fulfilment; future milestones may look like user tasks; small labels may be missed.",
        "hyp": "Users can identify the current stage, distinguish normal waiting from a problem, and explain the next expected event and freshness source.",
        "ucs": "F-UC02, F-UC03, F-UC04, F-UC06",
    },
    {
        "name": "Action Hub", "image": "alt3scenario1.png", "model": "Task-first shortcut hub with official-option and handoff cues.",
        "workflow": "Confirm the event -> select a management shortcut -> complete on FIFA or review the partner boundary before leaving.",
        "problem": "Reduces navigation effort for common ticket tasks and makes external destinations more explicit.",
        "motivation": "Tests one-clear-next-action and trust-through-provenance principles using shortcuts rather than guided sequencing.",
        "strengths": "Direct task access; broad ticket-management coverage; explicit official badges and pre-handoff reassurance.",
        "weaknesses": "Choice-heavy; critical management tasks compete with commercial options; little progress detail.",
        "risks": "Promotional cards may reduce trust; Official may not explain provider/data terms; equally weighted shortcuts slow routine work.",
        "hyp": "Users choose the correct shortcut, separate included ticket tasks from optional purchases, and recognize the partner handoff before clicking.",
        "ucs": "F-UC02, F-UC04, F-UC05, F-UC06",
    },
]

CHESS_ALTS = [
    {
        "name": "Beginner Review Flow", "image": "alt1scenario2.png", "model": "Linear guided wizard; the system controls review order.",
        "workflow": "Enter Beginner Review -> inspect one mistake -> show/try the recommended move -> practice or continue.",
        "problem": "Makes one learning priority and one next-learning action explicit before advanced analysis.",
        "motivation": "Direct evolution of PA2 C-A1 Beginner Analysis Preset and PA1 C-S3 Beginner analysis preset.",
        "strengths": "Clear progress; plain causal explanation; manageable chunking; direct practice bridge; advanced analysis remains available.",
        "weaknesses": "Forced order; notation and terms remain; leaving for practice may interrupt review context.",
        "risks": "Simplification can hide nuance; users may memorize Qe2 without understanding the idea; board/text split attention.",
        "hyp": "Beginners can explain the mistake in their own words, try the recommended move, and find relevant practice without facilitator help.",
        "ucs": "C-UC02, C-UC03, C-UC04, C-UC05, C-UC06",
    },
    {
        "name": "Visual Card Dashboard", "image": "alt2scenario2.png", "model": "Non-linear overview; users self-select review cards.",
        "workflow": "Scan category summary -> choose a moment card -> read/try/practice -> select another card.",
        "problem": "Supports scan-first review and user-chosen order while retaining direct learning actions.",
        "motivation": "Explores PA2 C-A3 Visual Game Story as a card-based overview rather than a fixed chapter timeline.",
        "strengths": "Strong scanability; visual previews; user control; direct links from each concept to review, trial, or puzzle.",
        "weaknesses": "High initial choice load; category terms are unexplained; checks may mean selection or completion.",
        "risks": "Beginners may skip the highest-value mistake; dense cards; positive/opening cards can dilute the learning priority.",
        "hyp": "Users understand card selection, choose a useful priority, relate the thumbnail to the board, and distinguish Review, Try, and Puzzle.",
        "ucs": "C-UC03, C-UC04, C-UC05, C-UC06",
    },
    {
        "name": "Side-by-Side Assistant", "image": "alt3scenario2.png", "model": "Conversational contextual review; users decide what to ask.",
        "workflow": "Select a key moment -> ask or choose a suggested question -> inspect board-linked explanation -> follow up.",
        "problem": "Lets learners request plain-language explanations without following a fixed review order.",
        "motivation": "Evolves PA2 C-A2 Conversational Coach with the board and contextual assistant visible together.",
        "strengths": "Flexible follow-up; contextual board highlights; suggested prompts reduce blank-page friction; explanation and position stay together.",
        "weaknesses": "No bounded completion path; text-heavy; user questions determine coverage.",
        "risks": "Answer correctness and inconsistency; important mistakes may be skipped; notation remains; conversation competes with board attention.",
        "hyp": "Users can start without composing a prompt, map an answer to the board, navigate key moments, and recognize when an explanation is uncertain.",
        "ucs": "C-UC03, C-UC04, C-UC05, C-UC06",
    },
]


def add_alt(doc, number, alt, scenario, fig_id):
    doc.add_heading(f"{number} {alt['name']}", level=2)
    figure(doc, PROTOS / alt["image"], fig_id, f"{scenario} - {alt['name']} full paper prototype.", alt["ucs"])
    if alt["name"] == "Side-by-Side Assistant":
        break_anchor = doc.add_paragraph()
        break_anchor.paragraph_format.page_break_before = True
        break_anchor.paragraph_format.space_before = Pt(0)
        break_anchor.paragraph_format.space_after = Pt(0)
        break_anchor.paragraph_format.line_spacing = Pt(1)
        set_font(break_anchor.add_run(" "), 1)
    table(doc, ["Dimension", "Evidence-grounded analysis"], [
        ["Concept / interaction model", alt["model"]],
        ["Primary workflow", alt["workflow"]],
        ["Problem addressed", alt["problem"]],
        ["Motivation", alt["motivation"]],
        ["Strengths", alt["strengths"]],
        ["Weaknesses", alt["weaknesses"]],
        ["Usability risks", alt["risks"]],
        ["Formative hypothesis", alt["hyp"]],
        ["Exact PA2 use-case coverage", alt["ucs"]],
    ], [4.3, 11.3])


def build_paper():
    doc = configure_doc(TEMPLATE_REPORT, "Paper Prototype and Storyboarding", "PaperProtype")
    cover(doc, "Group10-PA3 Paper Prototypes", "Parallel lo-fi alternatives for FIFA.com and Chess.com", links=True)
    doc.add_heading("1. Introduction and Project Continuity", level=1)
    para(doc, "PA1 identified product-level drawbacks, including FIFA ticket-state uncertainty (F-D4) and dense Chess analysis before novices know what matters (C-D2). PA2 converted those concerns into provisional tough problems, solution alternatives, and use cases. PA3 now compares behaviorally distinct lo-fi interaction models before any PA4 hi-fi commitment.")
    para(doc, "Evidence boundary: Expected benefits are hypotheses. No genuine PA3 participant evidence or YouTube demo links were found, so this document does not declare a winning design.", fill="FFF1F1")
    doc.add_heading("2. Problem and Traceability Overview", level=1)
    table(doc, ["Chain", "FIFA.com", "Chess.com"], [
        ["PA1", "F-D4; F-S7 Ticket status dashboard; F-S8 Official availability alerts", "C-D2; C-S3 Beginner analysis preset; C-S4 Inline analysis glossary"],
        ["PA2 tough problem", "TP-FIFA: compare governed visible states, freshness, and official destination before acting or leaving", "TP-CHESS: make one beginner path explicit, then connect a critical position to optional depth and practice"],
        ["PA2 alternatives", "F-A1 Status Dashboard; F-A2 Guided Ticket Concierge; F-A3 Alert-First Ticket Planner", "C-A1 Beginner Analysis Preset; C-A2 Conversational Coach; C-A3 Visual Game Story"],
        ["PA2 use cases", "F-UC01-F-UC06", "C-UC01-C-UC06"],
        ["PA3 scenarios", "Plan and Manage Tickets with Confidence", "Beginner Review After a Game"],
        ["PA3 testing target", "State, freshness, next action, official destination, and handoff understanding", "Mistake comprehension, recommended-move trial, control, and relevant practice continuation"],
    ], [3.2, 6.2, 6.2])
    para(doc, "Identifier note: The handwritten sketch annotations reuse some F-A/C-A labels inconsistently with final PA2. Formal mappings above preserve final-PA2 meanings; the six PA3 concepts are identified by descriptive names.", fill="FFF4E5")

    doc.add_heading("3. Scenario 1 - FIFA.com: Plan and Manage Tickets with Confidence", level=1)
    para(doc, "Target user and context: A trust-sensitive ticket planner or tournament follower needs to understand the current visible ticket situation, decide whether to act or wait, and recognize the official destination before an outbound handoff.")
    para(doc, "Design objective: Apply state before action, one clear next action, trust through provenance, and accessible redundancy while keeping transaction processing and partner redesign out of scope.")
    for i, alt in enumerate(FIFA_ALTS, 1):
        add_alt(doc, f"3.{i}", alt, "FIFA.com", f"PP-F-{i:02d}")
    doc.add_heading("3.4 FIFA Alternative Comparison", level=2)
    table(doc, ["Dimension", "Status Dashboard", "Timeline Tracker", "Action Hub"], [
        ["Information architecture", "Aggregate states + event cards", "Event milestones + update history", "Ticket tasks + official options"],
        ["Primary interaction", "Scan then select", "Trace progress then inspect", "Choose shortcut then act"],
        ["User control", "Moderate", "Moderate", "High"],
        ["Cognitive load", "Low-medium", "Medium-high", "Medium-high choice load"],
        ["Status visibility", "Highest", "Stage-based", "Compact current event"],
        ["Freshness visibility", "Timestamp", "Highest: source + history", "Timestamp + source badge"],
        ["Next-step clarity", "Status-specific action", "Expected future milestone", "Direct task shortcuts"],
        ["External-handoff trust", "Limited", "Official source/steps", "Highest pre-handoff cue"],
        ["Principal risk", "Counts oversimplify", "Linear model overpromises", "Commercial/choice distraction"],
    ], [3.2, 4.1, 4.1, 4.1])

    doc.add_heading("4. Scenario 2 - Chess.com: Beginner Review After a Game", level=1)
    para(doc, "Target user and context: A first-time beginner or returning low-analysis learner has completed a game and needs one interpretable mistake, a recommended move they can try safely, and a relevant next-learning action before optional advanced analysis.")
    para(doc, "Evidence caution: PA2 captured analysis-entry choices, not completed-review overload. These sketches test a proposed learning flow; they do not prove that current Chess.com users experience the depicted breakdowns.", fill="FFF4E5")
    for i, alt in enumerate(CHESS_ALTS, 1):
        add_alt(doc, f"4.{i}", alt, "Chess.com", f"PP-C-{i:02d}")
    doc.add_heading("4.4 Chess Alternative Comparison", level=2)
    table(doc, ["Dimension", "Beginner Review Flow", "Visual Card Dashboard", "Side-by-Side Assistant"], [
        ["Workflow authority", "System-led", "User-led", "User-led questions"],
        ["Navigation", "Linear", "Non-linear", "Conversational/contextual"],
        ["Information density", "Low-medium", "High overview", "High dialogue + board"],
        ["Learnability", "Highest initial guidance", "Requires category/card literacy", "Requires prompt/answer literacy"],
        ["User control", "Bounded", "High", "Highest"],
        ["Memory load", "Low", "Medium", "Medium-high across dialogue"],
        ["Explanation style", "Prepared plain-language steps", "Expandable card explanation", "Adaptive Q&A"],
        ["Practice bridge", "Direct recommended practice", "Per-card puzzle action", "Suggested; needs explicit practice action"],
        ["Likely beginner suitability", "Strong for first review", "Strong for self-directed learners", "Strong if answers are reliable"],
        ["Principal risk", "Passive compliance", "Choice overload/skipping", "Incorrect or unbounded dialogue"],
    ], [3.2, 4.1, 4.1, 4.1])

    doc.add_heading("5. Cross-Scenario Design Reflection", level=1)
    bullets(doc, [
        "The most guided concepts reduce immediate decision burden but can constrain exploration.",
        "The most user-controlled concepts accelerate expert or confident behavior but increase choice and interpretation demands.",
        "Freshness, provenance, plain-language explanation, and recovery must be observable rather than assumed.",
        "Color supports scanning but must not be the only carrier of status or severity.",
        "No alternative is proven best until genuine novice participants complete the core tasks.",
    ])
    doc.add_heading("6. Testing Hypotheses for Requirement 2", level=1)
    rows = []
    for product, alts in (("FIFA", FIFA_ALTS), ("Chess", CHESS_ALTS)):
        for alt in alts:
            rows.append([product, alt["name"], alt["hyp"]])
    table(doc, ["Scenario", "Alternative", "Observable hypothesis"], rows, [2.4, 4.3, 8.9])
    doc.add_heading("7. Strengths and Weaknesses Summary", level=1)
    para(doc, "The portfolio deliberately spans status, progress, task, guided, non-linear, and conversational models. This breadth enables formative testing to expose which representation best supports the target task without treating visual polish as evidence. Each concept retains a plausible strength and a falsifiable risk.")
    doc.add_heading("8. Conclusion", level=1)
    para(doc, "All six paper prototypes are documented as alternatives to evaluate. Final best-prototype selection pending real formative testing. Six YouTube demonstration links must also be inserted on the first page before submission.", fill="FFF1F1")
    doc.add_heading("Evidence References", level=1)
    bullets(doc, [
        "Group10-PA1-PotentialSolutions.pdf, especially F-D4/F-S7/F-S8 and C-D2/C-S3/C-S4.",
        "Group10-PA2-UserAnalysis.pdf, final tough-problem and recommendation wording.",
        "Group10-PA2-ProjectProposal.pdf, final alternatives and PA3 validation targets.",
        "Group10-PA2-UseCaseDocument.pdf, final F-UC01-F-UC06 and C-UC01-C-UC06 names.",
        "PA3-LKDuy-2026-Public.pdf, Requirements 1-2 and submission rules.",
    ])
    return doc


SIM_ROWS = {
    "FIFA-P1 - casual fan; moderate web experience; low FIFA ticketing familiarity": [
        ["Status Dashboard", "Starts with Pending/What does this mean; notices counts and validity; may confuse View Tickets with View Order.", "Add expected resolution and clarify pending reason.", "Preferred: answers whether tickets are okay with least interpretation."],
        ["Timeline Tracker", "Traces pending stages; may read future gray stages as required tasks.", "Label milestone owner and normal waiting clearly.", "Useful but busier."],
        ["Action Hub", "Starts with View Tickets; optional official packages may look required.", "Separate Manage my ticket from Optional extras.", "Concrete but choice-heavy."],
    ],
    "FIFA-P2 - online shopper/travel planner; strong transaction familiarity": [
        ["Status Dashboard", "Looks for pending order/payment detail and last update.", "Expose pending reason and transaction milestone inline.", "Familiar but not diagnostic enough."],
        ["Timeline Tracker", "Uses current milestone and update history to establish trust.", "Add owner, normal duration, and escalation threshold.", "Preferred for transaction trust."],
        ["Action Hub", "Uses ticket/calendar/travel shortcuts; scrutinizes partner policy.", "Show provider, inclusions, and policy before handoff.", "Strong for post-purchase planning."],
    ],
    "FIFA-P3 - experienced event attendee; high web fluency": [
        ["Status Dashboard", "Immediately opens confirmed tickets; asks whether confirmed means usable now.", "Separate confirmed allocation from mobile-ticket readiness.", "Fast triage."],
        ["Timeline Tracker", "Checks Ticket ready timing and exception handling.", "Add exception branches and action-required state.", "Preferred for troubleshooting."],
        ["Action Hub", "Uses View/Transfer Tickets; repeated/equal shortcuts slow routine action.", "Rank applicable actions and disable unavailable ones.", "Preferred for routine management."],
    ],
    "CHESS-P1 - basic rules; minimal Chess.com Analysis experience": [
        ["Beginner Review Flow", "Uses Show me on the board; notation and hanging pieces remain barriers.", "Pair notation with piece names and arrows.", "Preferred for explicit guidance."],
        ["Visual Card Dashboard", "Clicks highlighted card; may treat checks as completion and cannot choose priority.", "Recommend one starting card while preserving choice.", "Crowded for a first review."],
        ["Side-by-Side Assistant", "Uses suggested question; may not detect inconsistent pawn references.", "Provide Explain this move without notation and uncertainty cues.", "Approachable only with strong prompts."],
    ],
    "CHESS-P2 - casual beginner/intermediate; puzzle/lesson oriented": [
        ["Beginner Review Flow", "Tries move then practice; worries review position will be lost.", "Make practice inline or guarantee return state.", "Strong guided path."],
        ["Visual Card Dashboard", "Chooses tactic/puzzle cards and may skip strategic weaknesses.", "Add recommended priority and learning-goal tags.", "Preferred for self-directed learning."],
        ["Side-by-Side Assistant", "Uses similar-ideas prompt; dialogue may not create retained practice.", "Add Practice this idea after explanations.", "Feels like coaching."],
    ],
    "CHESS-P3 - low-rated casual; some Chess.com experience; low engine vocabulary": [
        ["Beginner Review Flow", "Reads plain explanation but may memorize Qe2 without principle.", "Visualize before/after defended squares and name the concept.", "Safest learning flow."],
        ["Visual Card Dashboard", "Follows red severity; taxonomy may discourage or distract.", "Lead with beginner learning priority, not raw severity.", "Useful overview after orientation."],
        ["Side-by-Side Assistant", "Asks about red key moment; broad prompts may yield shallow or contradictory answers.", "Add beginner-language mode, reviewed checklist, and confidence cue.", "Preferred only if answer quality is reliable."],
    ],
}


def build_testing():
    doc = configure_doc(TEMPLATE_REPORT, "Formative Testing", "FormativeTesting")
    cover(doc, "Group10-PA3 Formative Testing", "Executable plan, simulated pretest, and evidence-ready result framework", "REAL PARTICIPANT TESTING NOT YET EVIDENCED")
    para(doc, "Evidence status: Real participant testing has not yet been evidenced in the project files. Requirement 2 empirical testing remains incomplete. The simulated pretest below is design rehearsal only and cannot support final prototype selection.", fill="FFF1F1")
    doc.add_heading("1. Testing Plan", level=1)
    doc.add_heading("1.1 Objectives", level=2)
    table(doc, ["Scenario", "Evaluation targets"], [
        ["FIFA.com", "Ticket-state comprehension; freshness/source recognition; next-action recognition; official-option discovery; external-handoff understanding; hesitation and navigation breakdowns."],
        ["Chess.com", "Mistake recognition; explanation comprehension; recommended-move discovery and trial; practice continuation; entry-choice difficulty; perceived control; information overload."],
    ], [4.0, 11.6])
    doc.add_heading("1.2 Target Users", level=2)
    bullets(doc, [
        "FIFA: trust-sensitive ticket planners or tournament followers with realistic need to act, wait, compare, or review an official destination.",
        "Chess: first-time beginners or returning low-analysis learners who know basic rules but need a clear review and next-learning path.",
        "Do not recruit team members or anyone already briefed on the sketches; screen for no prior prototype knowledge.",
    ])
    doc.add_heading("1.3 Participant Requirement", level=2)
    para(doc, "Official PA3 requirement: conduct several testing sessions with 2-3 participants who do not have prior knowledge about the prototype being tested. The group must systematically document qualitative observations, feedback, encountered issues, and points of improvement.", fill=FIFA_LIGHT)
    doc.add_heading("1.4 Method", level=2)
    numbered(doc, [
        "Screen the participant and record only an anonymous participant ID plus relevant experience.",
        "Present the neutral scenario and paper prototype without explaining controls or intended workflow.",
        "Ask the participant to think aloud while completing the task; the moderator manipulates paper states only in response to a visible action.",
        "Observe first action, task path, hesitations, wrong paths, misunderstandings, interventions, and participant language.",
        "Run a short post-task interview, then compare all three alternatives for the scenario.",
        "Separate observed behavior, participant feedback, and team interpretation during synthesis.",
    ])
    doc.add_heading("1.5 Order and Bias Control", level=2)
    table(doc, ["Participant", "FIFA order", "Chess order"], [
        ["P01", "Status Dashboard -> Timeline Tracker -> Action Hub", "Beginner Review Flow -> Visual Card Dashboard -> Side-by-Side Assistant"],
        ["P02", "Timeline Tracker -> Action Hub -> Status Dashboard", "Visual Card Dashboard -> Side-by-Side Assistant -> Beginner Review Flow"],
        ["P03", "Action Hub -> Status Dashboard -> Timeline Tracker", "Side-by-Side Assistant -> Beginner Review Flow -> Visual Card Dashboard"],
    ], [2.6, 6.5, 6.5])
    para(doc, "If only two participants are recruited, use P01 and P02 orders and note the smaller counterbalancing coverage as a limitation.")
    doc.add_heading("1.6 Testing Tasks", level=2)
    table(doc, ["Scenario prompt", "Success conditions tied to final PA2"], [
        ["FIFA: You want to attend a FIFA event. Understand the current ticket situation and decide the correct next step without assuming that an external partner has completed the task.", "Select context (F-UC01); interpret official state/options (F-UC02-F-UC03); locate alert/plan support if relevant (F-UC04); identify destination and choose Continue or Stay (F-UC05); explain return/recovery (F-UC06)."],
        ["Chess: You completed a game. Find one important mistake, explain why it matters, try the recommended move, and choose useful practice.", "Open/start Beginner Review (C-UC01-C-UC02); explain the main mistake and recommended move (C-UC03); try it safely (C-UC04); find optional advanced depth (C-UC05); continue to a relevant lesson or puzzle (C-UC06)."],
    ], [7.2, 8.4])
    doc.add_heading("1.7 Success Criteria", level=2)
    table(doc, ["Measure", "Observable criterion", "Rationale"], [
        ["Task completion", "Completes the scenario without facilitator intervention; partial/failed is recorded with the stopping point.", "Direct effectiveness evidence."],
        ["State / mistake comprehension", "Explains the ticket state or chess mistake in their own words without merely repeating a label.", "Tests interpretation, not recognition alone."],
        ["Correct next action", "Chooses a context-appropriate action and explains why.", "Tests decision clarity."],
        ["Provenance / handoff", "Recognizes official source or partner boundary before leaving FIFA.com.", "Tests calibrated trust."],
        ["Move trial / practice", "Finds and tries the recommended move, then reaches a relevant practice path.", "Tests learning continuation."],
        ["Wrong paths", "Each unrelated selection or reversal is counted and annotated by cause.", "Locates navigation breakdowns."],
        ["Hesitation", "Moderator marks pauses accompanied by scanning, uncertainty, or a question; no arbitrary time threshold is imposed.", "Avoids unsupported numeric targets."],
        ["Perceived control", "Participant can state what they can do next and how to return or change direction.", "Tests user control/recovery."],
    ], [3.2, 7.3, 5.1])
    doc.add_heading("1.8 Moderator Script", level=2)
    numbered(doc, [
        "Thank you for helping us evaluate the design, not your ability. We will not collect your name in the report.",
        "You have not seen this prototype before. I will not explain the interface; please act on what you see.",
        "Read the scenario prompt. Tell me what you are thinking as you decide where to look or act.",
        "Neutral reminders: What are you looking for? What do you expect that to do? What makes you say that? Please keep thinking aloud.",
        "Avoid praise, correction, feature names not yet noticed, or questions that imply the intended answer.",
        "Post-task: What was clear? What was confusing? What did you expect next? What would you change? Which alternative best supports the task, and why?",
        "Close by explaining that the notes will be anonymized and used to revise the paper prototype.",
    ])
    doc.add_heading("1.9 Observation Sheet", level=2)
    table(doc, ["Field", "Record during session"], [[f, ""] for f in [
        "Participant ID / relevant profile", "Scenario / alternative / order", "Task completion and stopping point", "First action", "Hesitations", "Wrong paths", "Questions", "Misunderstandings", "Facilitator intervention", "Participant comments (verbatim only if genuinely recorded)", "Observed issue / severity", "Improvement opportunity"
    ]], [5.5, 10.1])

    doc.add_heading("2. SIMULATED PRETEST - NOT HUMAN PARTICIPANT EVIDENCE", level=1)
    para(doc, "Purpose: rehearse the protocol and expose likely ambiguities before recruitment. The following are independent role-play hypotheses. They are not participants, sessions, measured results, quotes, or evidence of Requirement 2 completion.", fill="FFF1F1")
    for persona, rows in SIM_ROWS.items():
        if persona.startswith("CHESS-P3"):
            doc.add_page_break()
        doc.add_heading(persona, level=2)
        table(doc, ["Alternative", "Likely reasoning / breakdown", "Improvement to test", "Provisional preference"], rows, [3.2, 4.8, 4.2, 3.4])
    doc.add_page_break()
    doc.add_heading("2.7 UX Auditor Synthesis", level=2)
    table(doc, ["Recurring provisional risk", "Affected alternatives", "Implication for real testing"], [
        ["Labels do not distinguish system work, user work, and normal waiting.", "FIFA Status Dashboard; Timeline Tracker", "Probe who owns the next step and whether waiting is interpreted as failure."],
        ["Commercial or equally weighted shortcuts compete with core ticket tasks.", "FIFA Action Hub", "Observe first action and whether optional extras are mistaken for requirements."],
        ["Official source/partner language may be over-trusted.", "All FIFA alternatives", "Ask the participant to explain provider, destination, data transfer, and return expectations."],
        ["Chess notation and engine vocabulary remain despite beginner framing.", "All Chess alternatives", "Ask for explanation in the participant's own words and note every term requiring help."],
        ["Practice can break review context or be skipped.", "Beginner Review Flow; Card Dashboard; Assistant", "Test return state and whether practice relevance is understood."],
        ["High user control can lead to skipping the most important learning moment.", "Card Dashboard; Assistant", "Record chosen order, ignored moments, and rationale."],
        ["Conversational answers can be inconsistent or unbounded.", "Side-by-Side Assistant", "Include a deliberately ambiguous follow-up and observe confidence calibration."],
    ], [5.5, 4.1, 6.0])

    doc.add_page_break()
    doc.add_heading("3. Real Formative Testing Results", level=1)
    para(doc, "Real participant testing has not yet been evidenced in the project files. This section must be completed using 2-3 genuine participants according to the PA3 brief.", fill="FFF1F1")
    table(doc, ["Participant", "Relevant profile", "Scenario", "Alternative order", "Task result", "Errors / hesitations", "Feedback / issues"], [
        ["P01", "", "FIFA / Chess", "", "", "", ""],
        ["P02", "", "FIFA / Chess", "", "", "", ""],
        ["P03 (if recruited)", "", "FIFA / Chess", "", "", "", ""],
    ], [2.0, 2.6, 2.1, 2.8, 2.1, 2.2, 2.8])
    doc.add_heading("4. Evaluation", level=1)
    para(doc, "PROVISIONAL ONLY: the simulated pretest suggests which risks to probe; it does not establish task performance or comparative superiority. After real sessions, compare each alternative by scenario and maintain three columns: observed evidence, participant feedback, and team interpretation.", fill="FFF4E5")
    table(doc, ["Alternative", "Observed evidence", "Participant feedback", "Team interpretation"], [[a["name"], "", "", ""] for a in FIFA_ALTS + CHESS_ALTS], [4.0, 4.0, 4.0, 3.6])
    doc.add_heading("5. Points of Improvement", level=1)
    para(doc, "All entries below are provisional because their only evidence is the simulated pretest. Each proposed change requires confirmation and retesting with genuine participants.", fill="FFF4E5")
    improvements = [
        ("PI-01 - Clarify waiting and ownership", "Pending/future states may be misread as user tasks.", "Milestone ownership is implicit.", "Provisional major", "Status Dashboard; Timeline Tracker", "Label system-owned versus user-owned steps and normal duration.", "Calmer, more accurate next-action choice."),
        ("PI-02 - Separate essential and optional ticket actions", "Core tasks compete with optional purchases.", "Equal visual weight and mixed grouping.", "Provisional major", "Action Hub", "Separate Manage my ticket from Optional extras; rank applicable tasks.", "Lower choice cost and fewer commercial detours."),
        ("PI-03 - Lead Chess explanations with plain concepts", "Notation and terminology can block explanation.", "Beginner mode still leads with chess codes.", "Provisional major", "All Chess alternatives", "Lead with piece names and plain concepts; retain notation secondarily.", "Improved comprehension without removing depth."),
        ("PI-04 - Preserve review-to-practice context", "Review context may be lost when practice starts.", "Practice is modeled as navigation away.", "Provisional moderate", "Beginner Review Flow; Visual Card Dashboard", "Use inline mini-practice or a guaranteed return state.", "Stronger review-to-practice continuity."),
        ("PI-05 - Calibrate assistant trust", "Assistant answers may be inconsistent or over-trusted.", "No provenance, confidence, or contradiction handling.", "Provisional critical", "Side-by-Side Assistant", "Add uncertainty cues, correction path, reviewed-moment checklist, and practice action.", "Calibrated trust and bounded learning progress."),
    ]
    for title, problem, cause, severity, affected, change, benefit in improvements:
        doc.add_heading(title, level=3)
        table(doc, ["Field", "Provisional record"], [
            ["Evidence", "SIMULATED PRETEST ONLY - NOT HUMAN EVIDENCE"],
            ["Observed problem", problem], ["Likely cause", cause], ["Severity", severity],
            ["Affected alternative", affected], ["Proposed change", change],
            ["Expected benefit", benefit], ["Retest requirement", "Required with genuine participants"],
        ], [4.1, 11.5])
    doc.add_heading("6. Best Prototype Selection", level=1)
    para(doc, "Final best-prototype selection pending real formative testing. Simulated-pretest preferences are mixed by profile and therefore cannot justify a final winner. Selection must be based on genuine task evidence, participant feedback, issue severity, and retest results.", fill="FFF1F1")
    doc.add_heading("7. Improved Prototype Specification", level=1)
    para(doc, "Provisional cross-cutting revision specification: clarify ownership and freshness for FIFA states; separate essential ticket management from optional offers; expose provider and return behavior before handoff; lead Chess explanations with piece names and principles; preserve review position through practice; provide a visible learning priority and a clear completion/recovery path. Apply only the changes supported by real session evidence, then retest the revised concept.")
    doc.add_heading("8. YouTube Link for Best Improved Paper Prototype", level=1)
    para(doc, "[YOUTUBE LINK REQUIRED]", fill="FFF1F1")
    doc.add_heading("Appendix A - Participant Screening Checklist", level=1)
    bullets(doc, [
        "Participant has no prior knowledge of the paper prototypes.",
        "Participant is not a Group10 member and has not participated in sketch creation.",
        "Relevant FIFA ticket-planning or beginner Chess profile is recorded without unnecessary personal data.",
        "Participant understands think-aloud and voluntary participation; local course consent practice is followed.",
        "Alternative order is assigned before the session.",
        "Recording, if any, is used only with explicit permission and stored according to course rules.",
    ])
    doc.add_heading("Appendix B - Post-Task Questions", level=1)
    bullets(doc, [
        "What did you think the current situation or main learning point was?",
        "What did you expect to happen after your first action?",
        "Which label, control, or explanation was hardest to interpret?",
        "How would you recover, go back, or choose another path?",
        "Which alternative best supported the task? What concrete behavior made it better?",
        "What is the single most important change before another person uses it?",
    ])
    return doc


def build_weekly():
    doc = configure_doc(TEMPLATE_WEEKLY, "Weekly Report Draft / Planned Two-Week Continuation", "WeeklyReport", landscape=True)
    cover(doc, "Group10-PA3 Weekly Report", "Weekly Report Draft / Planned Two-Week Continuation", "PLANNED RECORD - NO ACTUAL PA3 MEETING LOGS FOUND")
    doc.add_heading("1. Evidence Status and Reporting Boundary", level=1)
    para(doc, "This report is a planned two-week continuation after PA2 evidence audited through 30 July 2026. No contemporaneous PA3 meeting logs, attendance records, minutes, completed testing notes, or member completion confirmations were found. Every item below is therefore a plan, assignment, or acceptance target—not a historical claim.", fill="FFF1F1")
    doc.add_heading("2. Team Roster and Continuing Roles", level=1)
    table(doc, ["Member", "Student ID", "Continuing role", "Planned PA3 focus"], [
        [TEAM[0][0], TEAM[0][1], TEAM[0][2], "Evidence gate, cross-report traceability, final integration and submission QA"],
        [TEAM[1][0], TEAM[1][1], TEAM[1][2], "FIFA alternatives, FIFA task script, source/claim verification"],
        [TEAM[2][0], TEAM[2][1], TEAM[2][2], "Chess alternatives, Chess task script, explanation/practice review"],
        [TEAM[3][0], TEAM[3][1], TEAM[3][2], "Prototype distinctness, HCI synthesis, layout and visual QA"],
    ], [4.2, 3.0, 7.2, 10.2])
    doc.add_heading("3. Planned Two-Week Schedule", level=1)
    table(doc, ["Week", "Planned dates", "Goal", "Planned outputs", "Evidence gate"], [
        ["Week 1", "31 Jul-6 Aug 2026", "Carry PA2 problems into six distinct paper alternatives and prepare the formative protocol.", "Three FIFA sketches; three Chess sketches; storyboard/demo outline; test objectives, tasks, order control, instruments.", "No design benefit is marked proven; every PA2 ID is checked against final PDFs."],
        ["Week 2", "7-13 Aug 2026", "Recruit and run genuine testing, revise evidence-backed issues, complete videos and reports.", "2-3 genuine participant records; comparison; improvement decision; improved prototype/demo; final DOCX QA.", "If testing/videos/meetings remain unevidenced, mark them incomplete rather than backfilling history."],
    ], [2.3, 3.1, 6.2, 7.4, 6.0])
    doc.add_heading("4. Week 1 Planned Work", level=1)
    table(doc, ["Owner", "Planned task", "Acceptance condition", "Status at report creation"], [
        ["Le Minh", "Review final PA2 outcomes; lock terminology and evidence boundaries; structure three PA3 reports.", "Trace table uses final PA2 IDs and no unsupported empirical claim.", "Planned; artifact structure prepared, completion not asserted."],
        ["Nguyen Vu Bach", "Develop/review Status Dashboard, Timeline Tracker, and Action Hub; prepare FIFA task and success conditions.", "Three interaction-distinct FIFA models map to TP-FIFA/F-UC01-F-UC06.", "Planned; prototype assets exist, member ownership unverified."],
        ["Pham Nguyen Gia Bao", "Develop/review Beginner Flow, Card Dashboard, and Assistant; prepare Chess task and success conditions.", "Three interaction-distinct Chess models map to TP-CHESS/C-UC01-C-UC06.", "Planned; prototype assets exist, member ownership unverified."],
        ["Trang Minh Nhut", "Check prototype distinctness, readability, HCI terminology, storyboard and Word layout.", "Figures readable; status/progress/task and linear/non-linear/conversational differences remain explicit.", "Planned; completion not asserted."],
    ], [4.0, 8.2, 8.2, 5.4])
    doc.add_heading("4.1 Planned Week 1 Coordination Checkpoint", level=2)
    table(doc, ["Field", "Planned record"], [
        ["Record type", "Planned checkpoint - not an actual meeting minute"],
        ["Timing", "During Week 1; exact date/time to be recorded only if the checkpoint occurs"],
        ["Proposed attendees", "; ".join(n for n, _, _ in TEAM)],
        ["Agenda", "Confirm scenarios; compare three alternatives per product; test distinctness; assign storyboard/demo and report ownership; approve moderator materials."],
        ["Required evidence after occurrence", "Attendance, decisions, assigned actions, unresolved questions, and date/time recorded contemporaneously."],
    ], [5.4, 20.2])
    doc.add_heading("5. Week 2 Planned Work", level=1)
    table(doc, ["Owner", "Planned task", "Acceptance condition", "Current blocker / dependency"], [
        ["Le Minh", "Coordinate evidence ledger, consolidate genuine observations, maintain observed/feedback/interpretation separation.", "Every conclusion traces to participant IDs and observation records.", "No genuine PA3 session evidence currently exists."],
        ["Nguyen Vu Bach", "Recruit/schedule or facilitate FIFA task; compare ticket-state, freshness, action, and handoff understanding.", "Anonymous session notes with assigned order and no prior prototype knowledge.", "Participant recruitment and sessions required."],
        ["Pham Nguyen Gia Bao", "Recruit/schedule or facilitate Chess task; compare explanation, move trial, control, and practice continuation.", "Anonymous session notes with assigned order and no prior prototype knowledge.", "Participant recruitment and sessions required."],
        ["Trang Minh Nhut", "Synthesize severity/improvements, update prototype/storyboard, prepare demo materials, perform visual QA.", "Changes cite evidence; six demos and best-improved demo link are inserted.", "YouTube demonstrations and improvement evidence required."],
    ], [4.0, 8.2, 8.2, 5.4])
    doc.add_heading("5.1 Planned Week 2 Coordination Checkpoint", level=2)
    table(doc, ["Field", "Planned record"], [
        ["Record type", "Planned checkpoint - not an actual meeting minute"],
        ["Timing", "After genuine sessions; exact date/time recorded only after occurrence"],
        ["Agenda", "Review raw observations; separate evidence from interpretation; compare alternatives; prioritize issues; select only if real evidence supports it; assign retest/demo/report QA."],
        ["Decision rule", "No final winner, measured result, attendance, quote, or completed task is recorded without source evidence."],
        ["Required follow-up", "Attach anonymous result sheets, insert real links, update report status, and run document QA."],
    ], [5.4, 20.2])
    doc.add_heading("6. Planned Cross-Document Responsibilities", level=1)
    table(doc, ["Deliverable", "Planned lead", "Planned reviewers", "Acceptance target", "Current state"], [
        ["Paper Prototypes", "Nguyen Vu Bach / Pham Nguyen Gia Bao", "Le Minh / Trang Minh Nhut", "Six sketches, motivations, risks, comparisons, hypotheses, YouTube links on first page.", "Document prepared; YouTube links missing."],
        ["Formative Testing", "Le Minh", "All members", "Executable plan; 2-3 genuine novice participants; evidence-based evaluation and improvements.", "Plan prepared; empirical result criterion blocked."],
        ["Weekly Report", "Le Minh", "All members", "Roster continuity; accurate factual/planned status; task tracking and evidence gates.", "Planned/draft record only."],
        ["Peer Review", "Not assigned", "Not applicable", "Excluded by explicit user instruction.", "Out of scope."],
    ], [4.0, 4.5, 4.5, 8.2, 4.6])
    doc.add_heading("7. Planned Risk and Dependency Register", level=1)
    table(doc, ["Risk / dependency", "Impact", "Planned control", "Owner"], [
        ["No YouTube demo links", "Paper-prototype and improved-prototype video criteria fail.", "Record six prototype videos and the improved-prototype demo; replace markers before submission.", "Team / Le Minh integration"],
        ["No genuine participant evidence", "Requirement 2 testing result, evaluation, and final selection remain incomplete.", "Recruit 2-3 unbriefed participants; run counterbalanced sessions; retain anonymous notes.", "Testing leads"],
        ["No PA3 meeting log", "Weekly reporting cannot claim attendance or decisions.", "Record future checkpoints contemporaneously; keep this version labeled planned/draft.", "Le Minh"],
        ["Sketch IDs conflict with final PA2 labels", "Traceability could misrepresent project history.", "Use descriptive PA3 names; reserve F-A/C-A IDs for exact PA2 concepts.", "Le Minh / Trang Minh Nhut"],
        ["Peer Review excluded", "Known 5% rubric item remains absent.", "Report as intentional out-of-scope item; do not create the file.", "User instruction"],
    ], [6.4, 6.6, 9.2, 3.6])
    doc.add_heading("8. Planned QA and Submission Gate", level=1)
    table(doc, ["Gate", "Pass condition", "Current status"], [
        ["Prototype coverage", "Two scenarios, three distinct alternatives each, six readable embedded sketches.", "Ready for document QA."],
        ["Traceability", "Claims map to final PA1/PA2 evidence and exact final-PA2 identifiers.", "Ready; label conflict controlled."],
        ["YouTube demos", "Six prototype links on first page and best-improved link in testing report.", "Blocked."],
        ["Human testing", "2-3 genuine unbriefed participants with systematic qualitative records.", "Blocked."],
        ["Best selection", "Supported by real task evidence and improvement/retest reasoning.", "Blocked; provisional only."],
        ["Document package", "Three valid DOCX files, no final PDFs, no Peer Review, clean render and integrity checks.", "Pending final QA."],
    ], [5.0, 14.8, 6.0])
    doc.add_heading("9. Resume Instructions", level=1)
    numbered(doc, [
        "Record six prototype demos and insert their real YouTube links on the Paper Prototypes first page.",
        "Recruit 2-3 participants with no prior prototype knowledge and assign the counterbalanced orders.",
        "Run sessions using the moderator script and observation sheet; preserve anonymous raw notes.",
        "Replace the empty real-results tables, complete evidence/feedback/interpretation evaluation, and select only if supported.",
        "Revise and retest the selected concept; record the improved-prototype YouTube link.",
        "If real PA3 coordination occurs, replace planned checkpoint records with contemporaneous factual records and evidence.",
        "Run final DOCX structural and visual QA; submit only the instructor-required package format after user approval.",
    ])
    return doc


def backup_existing(path):
    if not path.exists():
        return
    stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    target_dir = BACKUP / stamp
    target_dir.mkdir(parents=True, exist_ok=True)
    target = target_dir / path.name
    target.write_bytes(path.read_bytes())


def save(doc, name):
    FINAL.mkdir(parents=True, exist_ok=True)
    BACKUP.mkdir(parents=True, exist_ok=True)
    path = FINAL / name
    backup_existing(path)
    doc.save(path)
    print(path)


def main():
    save(build_paper(), "Group10-PA3-PaperProtype.docx")
    save(build_testing(), "Group10-PA3-FormativeTesting.docx")
    save(build_weekly(), "Group10-PA3-WeeklyReport.docx")


if __name__ == "__main__":
    main()
