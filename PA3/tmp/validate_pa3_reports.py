from __future__ import annotations

import re
import sys
import zipfile
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
FINAL = ROOT / "final"
EXPECTED = {
    "Group10-PA3-PaperProtype.docx",
    "Group10-PA3-FormativeTesting.docx",
    "Group10-PA3-WeeklyReport.docx",
}
ALTERNATIVES = {
    "Status Dashboard",
    "Timeline Tracker",
    "Action Hub",
    "Beginner Review Flow",
    "Visual Card Dashboard",
    "Side-by-Side Assistant",
}
REQUIRED_PARTS = {
    "[Content_Types].xml",
    "word/document.xml",
    "word/styles.xml",
    "word/_rels/document.xml.rels",
}


def all_text(doc: Document) -> str:
    chunks: list[str] = []
    chunks.extend(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    for section in doc.sections:
        for container in (section.header, section.footer):
            chunks.extend(p.text for p in container.paragraphs)
            for table in container.tables:
                for row in table.rows:
                    chunks.extend(cell.text for cell in row.cells)
    return "\n".join(chunks)


def main() -> int:
    errors: list[str] = []
    actual = {p.name for p in FINAL.iterdir() if p.is_file()}
    if actual != EXPECTED:
        errors.append(f"final directory mismatch: {sorted(actual)}")

    texts: dict[str, str] = {}
    media_counts: dict[str, int] = {}
    inline_shape_counts: dict[str, int] = {}
    for name in sorted(EXPECTED):
        path = FINAL / name
        if not path.exists():
            errors.append(f"missing: {name}")
            continue
        if not zipfile.is_zipfile(path):
            errors.append(f"not a valid ZIP/DOCX: {name}")
            continue
        with zipfile.ZipFile(path) as archive:
            members = set(archive.namelist())
            missing_parts = REQUIRED_PARTS - members
            if missing_parts:
                errors.append(f"{name}: missing OOXML parts {sorted(missing_parts)}")
            media_counts[name] = len([m for m in members if m.startswith("word/media/") and not m.endswith("/")])
            bad = archive.testzip()
            if bad:
                errors.append(f"{name}: corrupt member {bad}")
        try:
            document = Document(path)
            texts[name] = all_text(document)
            inline_shape_counts[name] = len(document.inline_shapes)
        except Exception as exc:
            errors.append(f"{name}: python-docx open failed: {exc}")

    paper = texts.get("Group10-PA3-PaperProtype.docx", "")
    testing = texts.get("Group10-PA3-FormativeTesting.docx", "")
    weekly = texts.get("Group10-PA3-WeeklyReport.docx", "")

    if inline_shape_counts.get("Group10-PA3-PaperProtype.docx") != 6:
        errors.append(f"paper prototype inline-image count != 6: {inline_shape_counts.get('Group10-PA3-PaperProtype.docx')}")
    for alternative in sorted(ALTERNATIVES):
        if alternative not in paper:
            errors.append(f"paper missing alternative: {alternative}")
        if alternative not in testing:
            errors.append(f"testing missing alternative: {alternative}")
    for marker in ("[YOUTUBE LINK REQUIRED]", "Final best-prototype selection pending real formative testing"):
        if marker not in paper and marker not in testing:
            errors.append(f"missing evidence marker: {marker}")
    for phrase in (
        "SIMULATED PRETEST - NOT HUMAN PARTICIPANT EVIDENCE",
        "Real participant testing has not yet been evidenced",
        "2-3 genuine participants",
    ):
        if phrase not in testing:
            errors.append(f"testing missing evidence boundary: {phrase}")
    for phrase in (
        "Weekly Report Draft / Planned Two-Week Continuation",
        "No contemporaneous PA3 meeting logs",
        "planned/draft record only",
    ):
        if phrase.lower() not in weekly.lower():
            errors.append(f"weekly missing factual-status language: {phrase}")

    combined = "\n".join(texts.values())
    for pattern in (r"\bTODO\b", r"\bTBD\b", r"Lorem ipsum", r"\bplaceholder\b", r"\bfake\b"):
        if re.search(pattern, combined, flags=re.IGNORECASE):
            errors.append(f"prohibited token found: {pattern}")

    print(f"files={sorted(actual)}")
    print(f"media_counts={media_counts}")
    print(f"inline_shape_counts={inline_shape_counts}")
    print(f"youtube_markers={combined.count('[YOUTUBE LINK REQUIRED]')}")
    print(f"errors={len(errors)}")
    for error in errors:
        print(f"ERROR: {error}")
    return 1 if errors else 0


if __name__ == "__main__":
    sys.exit(main())
